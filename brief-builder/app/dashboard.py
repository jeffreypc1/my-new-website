"""Brief Builder -- Streamlit dashboard.

Production-quality UI for drafting immigration law briefs with section-based
editing, legal boilerplate insertion, live preview, draft persistence, and
Word/text export. Works entirely offline without the API server.

Part of the O'Brien Immigration Law tool suite.
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from datetime import date
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.drafts import delete_draft, list_drafts, load_draft, new_draft_id, save_draft
from app.sections import BRIEF_TYPES, get_brief_types, load_sections

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.google_upload import upload_to_google_docs
from shared.client_banner import render_client_banner
from shared.tool_notes import render_tool_notes
from shared.theme import render_theme_css, render_nav_bar
from shared.config_store import is_component_enabled
try:
    from shared.box_folder_browser import render_box_folder_browser
    from shared.box_client import parse_folder_id as _parse_folder_id
except ImportError:
    render_box_folder_browser = None
    _parse_folder_id = None
try:
    from shared.preview_modal import show_preview_modal
except ImportError:
    show_preview_modal = None
try:
    from shared.draft_box import render_draft_box
except ImportError:
    render_draft_box = None
try:
    from shared.tool_help import render_tool_help
except ImportError:
    render_tool_help = None
try:
    from shared.feedback_button import render_feedback_button
except ImportError:
    render_feedback_button = None

# -- Page config --------------------------------------------------------------

st.set_page_config(
    page_title="Brief Builder -- O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -- CSS ----------------------------------------------------------------------

render_theme_css()

from shared.auth import require_auth, render_logout
require_auth()

# -- Navigation bar -----------------------------------------------------------

render_nav_bar("Brief Builder")
render_logout()

render_client_banner()
if render_tool_help:
    render_tool_help("brief-builder")
if render_feedback_button:
    render_feedback_button("brief-builder")

# -- Session state defaults ---------------------------------------------------

_DEFAULTS: dict = {
    "draft_id": None,
    "last_saved_msg": "",
    "prev_brief_type": None,
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state.draft_id is None:
    st.session_state.draft_id = new_draft_id()

# -- SF auto-fill for client fields -------------------------------------------
_sf_client = st.session_state.get("sf_client")
if _sf_client:
    _sf_cid = _sf_client.get("Id", "")
    _prev_cid = st.session_state.get("_sf_autofill_cid", "")
    if _sf_cid and _sf_cid != _prev_cid:
        _sf_name = _sf_client.get("Name", "")
        if _sf_name:
            st.session_state["inp_client_name"] = _sf_name
        _sf_anum = _sf_client.get("A_Number__c", "")
        if _sf_anum:
            st.session_state["inp_a_number"] = _sf_anum
        _sf_court = _sf_client.get("Immigration_Court__c", "")
        if _sf_court:
            st.session_state["inp_court_or_office"] = _sf_court
        # Pull from active legal case
        _sf_legal_case = _sf_client.get("selected_legal_case", {})
        if _sf_legal_case:
            _sf_judge = _sf_legal_case.get("Immigration_Judge__c", "")
            if _sf_judge:
                st.session_state["inp_ij_name"] = _sf_judge
            _sf_hearing = _sf_legal_case.get("Next_Government_Date__c", "")
            if _sf_hearing:
                # Format as MM/DD/YYYY if it comes as YYYY-MM-DD
                try:
                    from datetime import datetime as _dt
                    _parsed = _dt.strptime(_sf_hearing, "%Y-%m-%d")
                    st.session_state["inp_hearing_date"] = _parsed.strftime("%m/%d/%Y")
                except (ValueError, TypeError):
                    st.session_state["inp_hearing_date"] = _sf_hearing
        st.session_state["_sf_autofill_cid"] = _sf_cid

# Roman numeral helper
_ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
           "XI", "XII", "XIII", "XIV", "XV"]


# -- Helpers ------------------------------------------------------------------

def _content_key(section_key: str, sub_key: str | None = None) -> str:
    """Build a session-state key for a content text_area."""
    if sub_key:
        return f"content_{section_key}_{sub_key}"
    return f"content_{section_key}"


def _collect_section_content(brief_type: str) -> dict[str, str]:
    """Gather all section text_area values from session state."""
    sections = load_sections(brief_type)
    content: dict[str, str] = {}
    for section in sections:
        subs = section.get("subsections", [])
        if subs:
            for sub in subs:
                ck = _content_key(section["key"], sub["key"])
                content[ck] = st.session_state.get(ck, "")
        else:
            ck = _content_key(section["key"])
            content[ck] = st.session_state.get(ck, "")
    return content


def _clear_all_content_keys() -> None:
    """Remove all content_* keys from session state."""
    to_delete = [k for k in st.session_state if k.startswith("content_")]
    for k in to_delete:
        del st.session_state[k]


def _do_save(brief_type: str) -> None:
    """Save the current state as a draft."""
    case_info = {
        "client_name": st.session_state.get("inp_client_name", ""),
        "a_number": st.session_state.get("inp_a_number", ""),
        "court_or_office": st.session_state.get("inp_court_or_office", ""),
        "ij_name": st.session_state.get("inp_ij_name", ""),
        "hearing_date": st.session_state.get("inp_hearing_date", ""),
    }
    section_content = _collect_section_content(brief_type)
    save_draft(st.session_state.draft_id, brief_type, case_info, section_content)
    name = case_info["client_name"] or "draft"
    st.session_state.last_saved_msg = f"Saved -- {name}"


def _do_load(draft_id: str) -> None:
    """Load a draft into session state."""
    draft = load_draft(draft_id)
    if not draft:
        return
    _clear_all_content_keys()
    st.session_state.draft_id = draft["id"]
    st.session_state.inp_brief_type = draft.get("brief_type", list(get_brief_types().keys())[0])
    ci = draft.get("case_info", {})
    st.session_state.inp_client_name = ci.get("client_name", "")
    st.session_state.inp_a_number = ci.get("a_number", "")
    st.session_state.inp_court_or_office = ci.get("court_or_office", "")
    st.session_state.inp_ij_name = ci.get("ij_name", "")
    st.session_state.inp_hearing_date = ci.get("hearing_date", "")
    st.session_state.prev_brief_type = draft.get("brief_type", "")
    for ck, val in draft.get("section_content", {}).items():
        st.session_state[ck] = val


def _do_new() -> None:
    """Start a fresh brief."""
    _clear_all_content_keys()
    st.session_state.draft_id = new_draft_id()
    st.session_state.last_saved_msg = ""
    st.session_state.prev_brief_type = None
    for k in (
        "inp_client_name", "inp_a_number", "inp_court_or_office",
        "inp_ij_name", "inp_hearing_date", "inp_brief_type",
        "_preview_edited_text", "_show_preview",
    ):
        if k in st.session_state:
            del st.session_state[k]


def _open_preview_modal():
    """Gather current brief data and open the shared preview modal."""
    brief_type = st.session_state.get("inp_brief_type", list(get_brief_types().keys())[0])
    case_info = {
        "client_name": st.session_state.get("inp_client_name", ""),
        "a_number": st.session_state.get("inp_a_number", ""),
        "court_or_office": st.session_state.get("inp_court_or_office", ""),
        "ij_name": st.session_state.get("inp_ij_name", ""),
        "hearing_date": st.session_state.get("inp_hearing_date", ""),
    }
    sections_content = _collect_section_content(brief_type)
    preview_html = _build_preview_html(brief_type, case_info, sections_content)
    plain_text = _build_plain_text(brief_type, case_info, sections_content)
    show_preview_modal(
        title="Brief Preview",
        preview_html=preview_html,
        plain_text=plain_text,
        tool_name="brief-builder",
    )


def _build_preview_html(
    brief_type: str,
    case_info: dict,
    sections_content: dict[str, str],
) -> str:
    """Render the brief as styled HTML for the live preview panel."""
    esc = html_mod.escape
    parts: list[str] = []

    # Title
    title_map = {
        "Asylum Merits Brief": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR ASYLUM",
        "Motion to Reopen": "MOTION TO REOPEN PROCEEDINGS",
        "Appeal Brief": "RESPONDENT'S BRIEF ON APPEAL",
        "Bond Brief": "RESPONDENT'S BRIEF IN SUPPORT OF BOND REDETERMINATION",
        "Cancellation of Removal": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR CANCELLATION OF REMOVAL",
        "Motion to Terminate": "MOTION TO TERMINATE PROCEEDINGS",
        "Motion for Continuance": "MOTION FOR CONTINUANCE",
        "Motion to Change Venue": "MOTION TO CHANGE VENUE",
        "Motion to Suppress": "MOTION TO SUPPRESS EVIDENCE",
    }
    parts.append(f'<div class="brief-title">{esc(title_map.get(brief_type, brief_type.upper()))}</div>')

    # Caption
    caption_lines = []
    client = case_info.get("client_name", "")
    a_num = case_info.get("a_number", "")
    court = case_info.get("court_or_office", "")
    ij = case_info.get("ij_name", "")
    hearing = case_info.get("hearing_date", "")
    if client:
        caption_lines.append(f"IN THE MATTER OF: {esc(client.upper())}")
    if a_num:
        caption_lines.append(f"A-Number: {esc(a_num)}")
    if court:
        caption_lines.append(f"Before the {esc(court)}")
    if ij:
        caption_lines.append(f"Immigration Judge {esc(ij)}")
    if hearing:
        caption_lines.append(f"Hearing Date: {esc(hearing)}")
    if caption_lines:
        parts.append(f'<div class="brief-caption">{"<br>".join(caption_lines)}</div>')

    # Sections
    sections = load_sections(brief_type)
    for idx, section in enumerate(sections):
        roman = _ROMAN[idx] if idx < len(_ROMAN) else str(idx + 1)
        heading = section["heading"]
        subs = section.get("subsections", [])

        parts.append(f'<div class="brief-heading">{roman}. {esc(heading)}</div>')

        if subs:
            for sub_idx, sub in enumerate(subs):
                sub_letter = chr(ord("A") + sub_idx)
                parts.append(f'<div class="brief-subheading">{sub_letter}. {esc(sub["heading"])}</div>')
                ck = _content_key(section["key"], sub["key"])
                body = sections_content.get(ck, "").strip()
                if body:
                    for para in body.split("\n\n"):
                        para = para.strip()
                        if para:
                            parts.append(f'<div class="brief-body">{esc(para)}</div>')
        else:
            ck = _content_key(section["key"])
            body = sections_content.get(ck, "").strip()
            if body:
                for para in body.split("\n\n"):
                    para = para.strip()
                    if para:
                        parts.append(f'<div class="brief-body">{esc(para)}</div>')

    # Signature block
    parts.append(
        '<div class="brief-sig">'
        "Respectfully submitted,<br><br>"
        "____________________________<br>"
        "Attorney for Respondent<br>"
        f"Date: {date.today().strftime('%m/%d/%Y')}"
        "</div>"
    )

    return "\n".join(parts)


def _build_plain_text(
    brief_type: str,
    case_info: dict,
    sections_content: dict[str, str],
) -> str:
    """Render the brief as plain text for .txt export."""
    lines: list[str] = []

    title_map = {
        "Asylum Merits Brief": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR ASYLUM",
        "Motion to Reopen": "MOTION TO REOPEN PROCEEDINGS",
        "Appeal Brief": "RESPONDENT'S BRIEF ON APPEAL",
        "Bond Brief": "RESPONDENT'S BRIEF IN SUPPORT OF BOND REDETERMINATION",
        "Cancellation of Removal": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR CANCELLATION OF REMOVAL",
        "Motion to Terminate": "MOTION TO TERMINATE PROCEEDINGS",
        "Motion for Continuance": "MOTION FOR CONTINUANCE",
        "Motion to Change Venue": "MOTION TO CHANGE VENUE",
        "Motion to Suppress": "MOTION TO SUPPRESS EVIDENCE",
    }
    lines.append(title_map.get(brief_type, brief_type.upper()))
    lines.append("")

    # Caption
    client = case_info.get("client_name", "")
    a_num = case_info.get("a_number", "")
    court = case_info.get("court_or_office", "")
    ij = case_info.get("ij_name", "")
    hearing = case_info.get("hearing_date", "")
    if client:
        lines.append(f"IN THE MATTER OF: {client.upper()}")
    if a_num:
        lines.append(f"A-Number: {a_num}")
    if court:
        lines.append(f"Before the {court}")
    if ij:
        lines.append(f"Immigration Judge {ij}")
    if hearing:
        lines.append(f"Hearing Date: {hearing}")
    lines.append("")
    lines.append("=" * 60)
    lines.append("")

    # Sections
    sections = load_sections(brief_type)
    for idx, section in enumerate(sections):
        roman = _ROMAN[idx] if idx < len(_ROMAN) else str(idx + 1)
        heading = section["heading"]
        subs = section.get("subsections", [])

        lines.append(f"{roman}. {heading}")
        lines.append("")

        if subs:
            for sub_idx, sub in enumerate(subs):
                sub_letter = chr(ord("A") + sub_idx)
                lines.append(f"    {sub_letter}. {sub['heading']}")
                lines.append("")
                ck = _content_key(section["key"], sub["key"])
                body = sections_content.get(ck, "").strip()
                if body:
                    lines.append(f"    {body}")
                    lines.append("")
        else:
            ck = _content_key(section["key"])
            body = sections_content.get(ck, "").strip()
            if body:
                lines.append(body)
                lines.append("")

    # Signature
    lines.append("")
    lines.append("Respectfully submitted,")
    lines.append("")
    lines.append("____________________________")
    lines.append("Attorney for Respondent")
    lines.append(f"Date: {date.today().strftime('%m/%d/%Y')}")

    return "\n".join(lines)


def _build_docx(
    brief_type: str,
    case_info: dict,
    sections_content: dict[str, str],
) -> bytes:
    """Build a properly formatted legal brief Word document."""
    doc = Document()

    # Page layout: 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, text, size=12, bold=False, italic=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return r

    # Title
    title_map = {
        "Asylum Merits Brief": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR ASYLUM",
        "Motion to Reopen": "MOTION TO REOPEN PROCEEDINGS",
        "Appeal Brief": "RESPONDENT'S BRIEF ON APPEAL",
        "Bond Brief": "RESPONDENT'S BRIEF IN SUPPORT OF BOND REDETERMINATION",
        "Cancellation of Removal": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR CANCELLATION OF REMOVAL",
        "Motion to Terminate": "MOTION TO TERMINATE PROCEEDINGS",
        "Motion for Continuance": "MOTION FOR CONTINUANCE",
        "Motion to Change Venue": "MOTION TO CHANGE VENUE",
        "Motion to Suppress": "MOTION TO SUPPRESS EVIDENCE",
    }
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, title_map.get(brief_type, brief_type.upper()), size=14, bold=True)

    # Caption block
    client = case_info.get("client_name", "")
    a_num = case_info.get("a_number", "")
    court = case_info.get("court_or_office", "")
    ij = case_info.get("ij_name", "")
    hearing = case_info.get("hearing_date", "")

    caption_lines = []
    if client:
        caption_lines.append(f"IN THE MATTER OF: {client.upper()}")
    if a_num:
        caption_lines.append(f"A-Number: {a_num}")
    if court:
        caption_lines.append(f"Before the {court}")
    if ij:
        caption_lines.append(f"Immigration Judge {ij}")
    if hearing:
        caption_lines.append(f"Hearing Date: {hearing}")

    if caption_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "\n".join(caption_lines), size=11)
        p.paragraph_format.space_after = Pt(12)

    # Spacer
    doc.add_paragraph()

    # Sections
    sections = load_sections(brief_type)
    for idx, section in enumerate(sections):
        roman = _ROMAN[idx] if idx < len(_ROMAN) else str(idx + 1)
        heading = section["heading"]
        subs = section.get("subsections", [])

        # Section heading
        p = doc.add_paragraph()
        _run(p, f"{roman}. {heading}", bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        if subs:
            for sub_idx, sub in enumerate(subs):
                sub_letter = chr(ord("A") + sub_idx)

                # Subsection heading
                p = doc.add_paragraph()
                _run(p, f"    {sub_letter}. {sub['heading']}", bold=True, italic=True)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)

                # Body text
                ck = _content_key(section["key"], sub["key"])
                body = sections_content.get(ck, "").strip()
                if body:
                    for para_text in body.split("\n\n"):
                        para_text = para_text.strip()
                        if para_text:
                            p = doc.add_paragraph()
                            _run(p, para_text)
                            p.paragraph_format.line_spacing = 2.0
                            p.paragraph_format.space_after = Pt(0)
        else:
            ck = _content_key(section["key"])
            body = sections_content.get(ck, "").strip()
            if body:
                for para_text in body.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph()
                        _run(p, para_text)
                        p.paragraph_format.line_spacing = 2.0
                        p.paragraph_format.space_after = Pt(0)

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, "Respectfully submitted,")
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, "____________________________")
    p = doc.add_paragraph()
    _run(p, "Attorney for Respondent")
    p = doc.add_paragraph()
    _run(p, f"Date: {date.today().strftime('%m/%d/%Y')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# -- Sidebar ------------------------------------------------------------------

with st.sidebar:
    # Draft management
    st.markdown("#### Drafts")
    btn_cols = st.columns(2)
    with btn_cols[0]:
        if st.button("New", use_container_width=True):
            _do_new()
            st.rerun()
    with btn_cols[1]:
        save_clicked = st.button("Save", use_container_width=True, type="primary")

    saved_drafts = list_drafts()
    if saved_drafts:
        labels_map = {
            d["id"]: f"{d['client_name']} -- {d['brief_type']}"
            for d in saved_drafts
        }
        draft_ids = list(labels_map.keys())
        selected_draft = st.selectbox(
            "Load a saved draft",
            options=[""] + draft_ids,
            format_func=lambda x: labels_map.get(x, "Select..."),
            label_visibility="collapsed",
        )
        load_cols = st.columns(2)
        with load_cols[0]:
            if selected_draft and st.button("Load", use_container_width=True):
                _do_load(selected_draft)
                st.rerun()
        with load_cols[1]:
            if selected_draft and st.button("Delete", use_container_width=True):
                delete_draft(selected_draft)
                st.rerun()

    if st.session_state.last_saved_msg:
        st.markdown(
            f'<div class="saved-toast">{html_mod.escape(st.session_state.last_saved_msg)}</div>',
            unsafe_allow_html=True,
        )

    st.divider()

    # Box folder browser
    if render_box_folder_browser and _parse_folder_id:
        _sf = st.session_state.get("sf_client")
        _box_raw = (_sf.get("Box_Folder_Id__c", "") or "") if _sf else ""
        if _box_raw:
            render_box_folder_browser(
                _parse_folder_id(_box_raw),
                mode="viewer",
                key_prefix="_bb_box",
                header_label="Client Documents",
            )
            st.divider()

    render_tool_notes("brief-builder")



# -- Brief Type (main area, top) -----------------------------------------------

_bt_col, _preview_btn_col = st.columns([3, 1])

with _bt_col:
    brief_type = st.selectbox(
        "Brief Type",
        options=list(get_brief_types().keys()),
        key="inp_brief_type",
    )

# If brief type changed, clear section content for fresh start
if st.session_state.prev_brief_type is not None and st.session_state.prev_brief_type != brief_type:
    _clear_all_content_keys()
    # Also clear any edited preview text
    if "_preview_edited_text" in st.session_state:
        del st.session_state["_preview_edited_text"]
st.session_state.prev_brief_type = brief_type

with _preview_btn_col:
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("View Full Preview", type="primary", use_container_width=True):
        st.session_state["_show_preview"] = True
        st.rerun()

# -- Handle save (after sidebar renders) --------------------------------------

if save_clicked:
    _do_save(brief_type)
    st.rerun()

# -- Preview dialog trigger ----------------------------------------------------

if st.session_state.get("_show_preview"):
    st.session_state["_show_preview"] = False
    if show_preview_modal and is_component_enabled("preview_modal", "brief-builder"):
        _open_preview_modal()

# -- Main area ----------------------------------------------------------------

sections = load_sections(brief_type)

edit_col, right_col = st.columns([3, 2], gap="large")

# -- Left column: Section editing ---------------------------------------------

with edit_col:
    st.markdown('<div class="section-label">Brief Sections</div>', unsafe_allow_html=True)

    for section in sections:
        section_key = section["key"]
        heading = section["heading"]
        subs = section.get("subsections", [])
        section_boilerplate = section.get("boilerplate", "")

        with st.expander(heading, expanded=False):

            # Section-level boilerplate
            if section_boilerplate:
                st.markdown(
                    f'<div class="boilerplate-block">'
                    f"<strong>Standard language:</strong><br>"
                    f"{html_mod.escape(section_boilerplate)}"
                    f"</div>",
                    unsafe_allow_html=True,
                )
                ck = _content_key(section_key)
                insert_key = f"insert_{section_key}"
                if st.button("Insert Boilerplate", key=insert_key):
                    st.session_state[ck] = section_boilerplate
                    st.rerun()

            if subs:
                for sub in subs:
                    sub_key = sub["key"]
                    sub_heading = sub["heading"]
                    sub_boilerplate = sub.get("boilerplate", "")

                    st.markdown(f"**{sub_heading}**")

                    if sub_boilerplate:
                        st.markdown(
                            f'<div class="boilerplate-block">'
                            f"<strong>Standard language:</strong><br>"
                            f"{html_mod.escape(sub_boilerplate)}"
                            f"</div>",
                            unsafe_allow_html=True,
                        )
                        ck = _content_key(section_key, sub_key)
                        insert_key = f"insert_{section_key}_{sub_key}"
                        if st.button("Insert Boilerplate", key=insert_key):
                            st.session_state[ck] = sub_boilerplate
                            st.rerun()

                    ck = _content_key(section_key, sub_key)
                    st.text_area(
                        f"Content for {sub_heading}",
                        height=150,
                        key=ck,
                        placeholder=f"Draft your {sub_heading.lower()} argument here...",
                        label_visibility="collapsed",
                    )
            else:
                # Single section without subsections
                ck = _content_key(section_key)
                height = 200 if section_key in ("statement_of_facts", "country_conditions") else 150
                st.text_area(
                    f"Content for {heading}",
                    height=height,
                    key=ck,
                    placeholder=f"Draft your {heading.lower()} here...",
                    label_visibility="collapsed",
                )

# -- Right column: Info summary + Draft Box + Export ---------------------------

with right_col:
    # Info summary (replaces the old static preview)
    st.markdown('<div class="section-label">Brief Info</div>', unsafe_allow_html=True)

    # Gather current case info
    case_info = {
        "client_name": st.session_state.get("inp_client_name", ""),
        "a_number": st.session_state.get("inp_a_number", ""),
        "court_or_office": st.session_state.get("inp_court_or_office", ""),
        "ij_name": st.session_state.get("inp_ij_name", ""),
        "hearing_date": st.session_state.get("inp_hearing_date", ""),
    }

    # Gather section content
    sections_content = _collect_section_content(brief_type)

    # Check if there is any content at all
    has_content = any(v.strip() for v in sections_content.values())

    _info_parts = [f"**Type:** {brief_type}"]
    if case_info["client_name"]:
        _info_parts.append(f"**Client:** {case_info['client_name']}")
    if case_info["a_number"]:
        _info_parts.append(f"**A-Number:** {case_info['a_number']}")
    if case_info["court_or_office"]:
        _info_parts.append(f"**Court:** {case_info['court_or_office']}")
    if case_info["ij_name"]:
        _info_parts.append(f"**Judge:** {case_info['ij_name']}")
    if case_info["hearing_date"]:
        _info_parts.append(f"**Hearing:** {case_info['hearing_date']}")
    st.markdown("  \n".join(_info_parts))

    if st.session_state.get("_preview_edited_text"):
        st.caption("Preview text has been edited. Exports will use your edits.")

    # Draft Box
    if render_draft_box is not None and is_component_enabled("draft_box", "brief-builder"):
        plain_for_draft = _build_plain_text(brief_type, case_info, sections_content) if has_content else ""
        render_draft_box("brief-builder", {
            "document_type": "legal brief",
            "client_name": case_info["client_name"],
            "case_id": st.session_state.get("draft_id", ""),
            "content": plain_for_draft,
        })

    # Export controls
    st.markdown("---")
    st.markdown('<div class="section-label">Export</div>', unsafe_allow_html=True)

    safe_name = (case_info["client_name"] or "Brief").replace(" ", "_")

    exp_cols = st.columns(2)
    with exp_cols[0]:
        if has_content and case_info["client_name"]:
            if st.session_state.get("_preview_edited_text"):
                plain_text = st.session_state["_preview_edited_text"]
            else:
                plain_text = _build_plain_text(brief_type, case_info, sections_content)
            st.download_button(
                "Download .txt",
                data=plain_text,
                file_name=f"{brief_type.replace(' ', '_')}_{safe_name}.txt",
                mime="text/plain",
                use_container_width=True,
            )
        else:
            st.button("Download .txt", disabled=True, use_container_width=True)

    with exp_cols[1]:
        if has_content and case_info["client_name"]:
            if st.session_state.get("_preview_edited_text"):
                # Build docx from edited plain text
                docx_bytes = _build_docx(brief_type, case_info, sections_content)
            else:
                docx_bytes = _build_docx(brief_type, case_info, sections_content)
            st.download_button(
                "Download .docx",
                data=docx_bytes,
                file_name=f"{brief_type.replace(' ', '_')}_{safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            if st.button("Upload to Google Docs", use_container_width=True):
                with st.spinner("Uploading to Google Docs..."):
                    try:
                        url = upload_to_google_docs(docx_bytes, f"{brief_type} - {case_info['client_name']}")
                        st.session_state.google_doc_url = url
                    except Exception as e:
                        st.error(f"Upload failed: {e}")
            if st.session_state.get("google_doc_url"):
                st.markdown(f"[Open Google Doc]({st.session_state.google_doc_url})")
        else:
            st.button("Download .docx", disabled=True, use_container_width=True)
