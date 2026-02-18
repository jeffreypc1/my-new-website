"""FastAPI backend for the Brief Builder tool.

Provides endpoints for listing brief types, retrieving standard sections,
assembling briefs from structured arguments and citations, draft CRUD,
and exporting to Word (.docx) format with proper legal brief formatting.

Part of the O'Brien Immigration Law tool suite.
"""

from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

from app.drafts import (
    delete_draft,
    list_drafts,
    load_draft,
    new_draft_id,
    save_draft,
)
from app.sections import BRIEF_TYPES, get_boilerplate, load_sections

app = FastAPI(title="Brief Builder API")

_ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
          "XI", "XII", "XIII", "XIV", "XV"]

_TITLE_MAP = {
    "Asylum Merits Brief": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR ASYLUM",
    "Motion to Reopen": "MOTION TO REOPEN PROCEEDINGS",
    "Appeal Brief": "RESPONDENT'S BRIEF ON APPEAL",
    "Bond Brief": "RESPONDENT'S BRIEF IN SUPPORT OF BOND REDETERMINATION",
    "Cancellation of Removal": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR CANCELLATION OF REMOVAL",
}


# ---------------------------------------------------------------------------
# Request / response models
# ---------------------------------------------------------------------------

class CaseInfo(BaseModel):
    """Basic case identification fields."""
    client_name: str = ""
    a_number: str = ""
    court_or_office: str = ""
    ij_name: str = ""
    hearing_date: str = ""


class SectionContent(BaseModel):
    """A single section of a brief with its content and optional citations."""
    section_key: str
    heading: str
    body: str
    citations: list[str] = []


class GenerateRequest(BaseModel):
    """Payload for assembling a brief from sections."""
    brief_type: str
    case_info: CaseInfo
    sections: list[SectionContent]


class ExportDocxRequest(BaseModel):
    """Payload for exporting an assembled brief to Word format."""
    brief_type: str
    case_info: CaseInfo
    sections: list[SectionContent]


class DraftSaveRequest(BaseModel):
    """Payload for saving a draft."""
    draft_id: str | None = None
    brief_type: str
    case_info: CaseInfo
    section_content: dict[str, str] = {}


# ---------------------------------------------------------------------------
# Docx builder (shared logic with dashboard)
# ---------------------------------------------------------------------------

def _build_docx_from_sections(
    brief_type: str,
    case_info: CaseInfo,
    section_payloads: list[SectionContent],
) -> bytes:
    """Build a properly formatted legal brief Word document from section payloads."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, text, size=12, bold=False, italic=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return r

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, _TITLE_MAP.get(brief_type, brief_type.upper()), size=14, bold=True)

    # Caption block
    caption_lines = []
    if case_info.client_name:
        caption_lines.append(f"IN THE MATTER OF: {case_info.client_name.upper()}")
    if case_info.a_number:
        caption_lines.append(f"A-Number: {case_info.a_number}")
    if case_info.court_or_office:
        caption_lines.append(f"Before the {case_info.court_or_office}")
    if case_info.ij_name:
        caption_lines.append(f"Immigration Judge {case_info.ij_name}")
    if case_info.hearing_date:
        caption_lines.append(f"Hearing Date: {case_info.hearing_date}")

    if caption_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "\n".join(caption_lines), size=11)
        p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    # Sections with proper heading formatting
    for idx, sec in enumerate(section_payloads):
        roman = _ROMAN[idx] if idx < len(_ROMAN) else str(idx + 1)

        p = doc.add_paragraph()
        _run(p, f"{roman}. {sec.heading}", bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        if sec.body.strip():
            for para_text in sec.body.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph()
                    _run(p, para_text)
                    p.paragraph_format.line_spacing = 2.0
                    p.paragraph_format.space_after = Pt(0)

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, "Respectfully submitted,")
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, "____________________________")
    p = doc.add_paragraph()
    _run(p, "Attorney for Respondent")
    p = doc.add_paragraph()
    _run(p, f"Date: {date.today().strftime('%m/%d/%Y')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Endpoints: brief types and sections
# ---------------------------------------------------------------------------

@app.get("/api/brief-types")
def list_brief_types() -> list[str]:
    """Return the list of supported brief types."""
    return list(BRIEF_TYPES.keys())


@app.get("/api/sections/{brief_type}")
def get_sections(brief_type: str) -> list[dict[str, Any]]:
    """Return the standard sections for a given brief type."""
    sections = load_sections(brief_type)
    return sections


# ---------------------------------------------------------------------------
# Endpoints: brief generation and export
# ---------------------------------------------------------------------------

@app.post("/api/generate")
def generate_brief(request: GenerateRequest) -> dict[str, Any]:
    """Assemble a brief from structured sections."""
    boilerplate = get_boilerplate(request.brief_type)
    assembled_sections: list[dict[str, Any]] = []

    for section in request.sections:
        assembled_sections.append({
            "heading": section.heading,
            "body": section.body,
            "citations": section.citations,
            "boilerplate_available": section.section_key in boilerplate,
        })

    return {
        "brief_type": request.brief_type,
        "case_info": request.case_info.model_dump(),
        "sections": assembled_sections,
        "status": "draft",
    }


@app.post("/api/export/docx")
def export_docx(request: ExportDocxRequest) -> Response:
    """Export the assembled brief to a Word (.docx) file."""
    docx_bytes = _build_docx_from_sections(
        request.brief_type,
        request.case_info,
        request.sections,
    )
    filename = f"{request.brief_type.replace(' ', '_')}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Endpoints: draft CRUD
# ---------------------------------------------------------------------------

@app.post("/api/drafts")
def create_or_update_draft(request: DraftSaveRequest) -> dict:
    """Create or update a draft."""
    draft_id = request.draft_id or new_draft_id()
    case_info = request.case_info.model_dump()
    result = save_draft(draft_id, request.brief_type, case_info, request.section_content)
    return result


@app.get("/api/drafts")
def get_drafts() -> list[dict]:
    """List all saved drafts."""
    return list_drafts()


@app.get("/api/drafts/{draft_id}")
def get_draft(draft_id: str) -> dict:
    """Load a single draft by ID."""
    draft = load_draft(draft_id)
    if draft is None:
        raise HTTPException(status_code=404, detail=f"Draft {draft_id} not found")
    return draft


@app.delete("/api/drafts/{draft_id}")
def remove_draft(draft_id: str) -> dict:
    """Delete a draft by ID."""
    deleted = delete_draft(draft_id)
    if not deleted:
        raise HTTPException(status_code=404, detail=f"Draft {draft_id} not found")
    return {"deleted": True, "id": draft_id}
