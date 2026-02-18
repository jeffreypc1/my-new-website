"""FastAPI backend for the Cover Letter Generator tool.

Provides endpoints for template listing, draft CRUD, cover letter
generation, and Word document export.
"""

from __future__ import annotations

import io
from datetime import date

from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.drafts import delete_draft, list_drafts, load_draft, save_draft
from app.templates import (
    CASE_TYPES,
    TEMPLATES,
    get_filing_office_address,
    render_cover_letter,
)

app = FastAPI(title="Cover Letter Generator API")


# ---------------------------------------------------------------------------
# Request / response models
# ---------------------------------------------------------------------------

class EnclosedDoc(BaseModel):
    name: str
    description: str = ""


class GenerateRequest(BaseModel):
    case_type: str
    client_name: str
    a_number: str = ""
    receipt_number: str = ""
    filing_office: str = ""
    enclosed_docs: list[EnclosedDoc] = []
    attorney_name: str = ""
    bar_number: str = ""
    firm_name: str = "O'Brien Immigration Law"
    firm_address: str = ""
    recipient_address: str = ""
    salutation: str = ""


class DraftSaveRequest(BaseModel):
    draft_id: str
    case_type: str
    client: dict
    attorney: dict
    filing_office: str = ""
    enclosed_docs: list[dict] = []
    recipient_type: str = "agency"
    recipient_address: str = ""
    salutation: str = "Dear Sir or Madam:"


# ---------------------------------------------------------------------------
# Template endpoints
# ---------------------------------------------------------------------------

@app.get("/api/templates")
def api_list_templates() -> list[dict]:
    """List all available case types and their template metadata."""
    result = []
    for case_type in CASE_TYPES:
        tpl = TEMPLATES[case_type]
        result.append({
            "case_type": case_type,
            "form_numbers": tpl.get("form_numbers", []),
            "filing_offices": tpl.get("filing_offices", []),
            "standard_enclosed_docs": tpl.get("standard_enclosed_docs", []),
            "required_fields": tpl.get("required_fields", []),
        })
    return result


@app.get("/api/templates/{case_type}")
def api_get_template(case_type: str) -> dict:
    """Retrieve a single template by case type."""
    tpl = TEMPLATES.get(case_type)
    if tpl is None:
        raise HTTPException(status_code=404, detail=f"Template '{case_type}' not found")
    return tpl


# ---------------------------------------------------------------------------
# Draft endpoints
# ---------------------------------------------------------------------------

@app.get("/api/drafts")
def api_list_drafts() -> list[dict]:
    """List all saved drafts."""
    return list_drafts()


@app.get("/api/drafts/{draft_id}")
def api_get_draft(draft_id: str) -> dict:
    """Load a specific draft."""
    draft = load_draft(draft_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="Draft not found")
    return draft


@app.post("/api/drafts")
def api_save_draft(req: DraftSaveRequest) -> dict:
    """Save or update a draft."""
    return save_draft(
        draft_id=req.draft_id,
        case_type=req.case_type,
        client=req.client,
        attorney=req.attorney,
        filing_office=req.filing_office,
        enclosed_docs=req.enclosed_docs,
        recipient_type=req.recipient_type,
        recipient_address=req.recipient_address,
        salutation=req.salutation,
    )


@app.delete("/api/drafts/{draft_id}")
def api_delete_draft(draft_id: str) -> dict:
    """Delete a draft."""
    if delete_draft(draft_id):
        return {"status": "deleted", "id": draft_id}
    raise HTTPException(status_code=404, detail="Draft not found")


# ---------------------------------------------------------------------------
# Generation endpoints
# ---------------------------------------------------------------------------

@app.post("/api/generate")
def api_generate(req: GenerateRequest) -> dict:
    """Generate a cover letter and return as plain text."""
    docs = [{"name": d.name, "description": d.description} for d in req.enclosed_docs]
    text = render_cover_letter(
        case_type=req.case_type,
        client_name=req.client_name,
        a_number=req.a_number,
        receipt_number=req.receipt_number,
        filing_office=req.filing_office,
        enclosed_docs=docs,
        attorney_name=req.attorney_name,
        bar_number=req.bar_number,
        firm_name=req.firm_name,
        firm_address=req.firm_address,
        recipient_address=req.recipient_address,
        salutation=req.salutation,
    )
    return {"text": text}


@app.post("/api/export/docx")
def api_export_docx(req: GenerateRequest) -> StreamingResponse:
    """Generate a cover letter and return as a .docx file."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    docs = [{"name": d.name, "description": d.description} for d in req.enclosed_docs]
    text = render_cover_letter(
        case_type=req.case_type,
        client_name=req.client_name,
        a_number=req.a_number,
        receipt_number=req.receipt_number,
        filing_office=req.filing_office,
        enclosed_docs=docs,
        attorney_name=req.attorney_name,
        bar_number=req.bar_number,
        firm_name=req.firm_name,
        firm_address=req.firm_address,
        recipient_address=req.recipient_address,
        salutation=req.salutation,
    )

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, txt, size=12, bold=False):
        r = para.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        return r

    for line in text.split("\n"):
        p = doc.add_paragraph()
        _run(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = req.client_name.replace(" ", "_") if req.client_name else "draft"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="Cover_Letter_{safe_name}.docx"'},
    )
