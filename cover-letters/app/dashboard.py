"""Filing Assembler — Streamlit dashboard.

Production-quality UI for generating immigration filing cover pages with
live preview, enclosed document management, draft persistence, and
Word/text export. Works entirely offline without the API server.
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from datetime import date
from pathlib import Path

import streamlit as st
from docx import Document
from docx.shared import Inches, Pt

from app.drafts import delete_draft, list_drafts, load_draft, new_draft_id, save_draft
import streamlit.components.v1 as st_components

from app.templates import (
    CASE_TYPES,
    DEFAULT_EOIR_COVER_TEMPLATE,
    RECIPIENT_CATEGORIES,
    TEMPLATES,
    get_filing_office_address,
    get_filing_offices,
    get_recipient_addresses,
    render_cover_letter,
    render_eoir_from_template,
    render_eoir_submission,
    save_recipient_addresses,
    split_eoir_into_blocks,
)

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.config_store import get_config_value, load_config, is_component_enabled
from shared.google_upload import copy_template_and_fill, upload_to_google_docs
from shared.client_banner import render_client_banner
from shared.tool_notes import render_tool_notes
from shared.theme import render_theme_css, render_nav_bar
try:
    from shared.preview_modal import show_preview_modal
except ImportError:
    show_preview_modal = None
try:
    from shared.draft_box import render_draft_box
except ImportError:
    render_draft_box = None
try:
    from shared.salesforce_client import upload_file_to_contact
except ImportError:
    upload_file_to_contact = None
try:
    from shared.salesforce_client import get_beneficiaries
except ImportError:
    get_beneficiaries = None
try:
    from shared.salesforce_client import update_legal_case
except ImportError:
    update_legal_case = None
try:
    from shared.salesforce_client import get_legal_case_field_metadata
except ImportError:
    get_legal_case_field_metadata = None
try:
    from shared.salesforce_client import get_case_beneficiaries, update_case_beneficiary, describe_case_contact_relationships
except ImportError:
    get_case_beneficiaries = None
    update_case_beneficiary = None
    describe_case_contact_relationships = None
try:
    from shared.google_doc_creator import render_google_doc_button
except ImportError:
    render_google_doc_button = None
try:
    from shared.tool_help import render_tool_help
except ImportError:
    render_tool_help = None
try:
    from shared.feedback_button import render_feedback_button
except ImportError:
    render_feedback_button = None

# -- EOIR Canvas custom Streamlit component -----------------------------------
_EOIR_CANVAS_DIR = Path(__file__).parent / "eoir_canvas"
_eoir_canvas_component = st_components.declare_component("eoir_canvas", path=str(_EOIR_CANVAS_DIR))

# -- Page config --------------------------------------------------------------

st.set_page_config(
    page_title="Filing Assembler — O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -- CSS ----------------------------------------------------------------------

_FILING_EXTRA_CSS = """\
/* Header */
.cl-header {
    color: #1a2744;
    font-size: 1.8rem;
    font-weight: 800;
    letter-spacing: -0.03em;
    margin-bottom: 0;
}
.cl-sub {
    color: #5a6a85;
    font-size: 0.95rem;
    margin-bottom: 0.5rem;
}

/* Document checklist item */
.doc-item {
    font-size: 0.88rem;
    padding: 4px 0;
    color: #1a2744;
}

/* SF task row marked for deletion */
.sf-task-deleted {
    background: #fef2f2;
    border: 1px solid #fca5a5;
    border-radius: 6px;
    padding: 6px 10px;
    margin: 2px 0;
}
.sf-task-deleted .doc-item {
    text-decoration: line-through;
    color: #991b1b;
}

/* SF task row — normal */
.sf-task-row {
    padding: 4px 10px;
    margin: 2px 0;
    border-radius: 6px;
    border: 1px solid transparent;
}
.sf-task-row:hover {
    background: #f8fafc;
}

/* Pending changes badge */
.pending-badge {
    display: inline-block;
    font-size: 0.72rem;
    font-weight: 600;
    padding: 2px 8px;
    border-radius: 10px;
    margin-left: 8px;
}
.pending-badge.edits {
    background: #fef3c7;
    color: #92400e;
}
.pending-badge.deletes {
    background: #fee2e2;
    color: #991b1b;
}

/* Filing Assembler letter preview overrides */
.preview-panel {
    white-space: pre-wrap;
    word-wrap: break-word;
}

/* Custom doc add row */
.add-doc-row {
    display: flex;
    gap: 8px;
    align-items: flex-end;
}
"""
render_theme_css(extra_css=_FILING_EXTRA_CSS)

from shared.auth import require_auth, render_logout
require_auth()

# -- Navigation bar -----------------------------------------------------------

render_nav_bar("Filing Assembler")
render_logout()

render_client_banner()
if render_tool_help:
    render_tool_help("cover-letters")
if render_feedback_button:
    render_feedback_button("cover-letters")

# -- Session state defaults ---------------------------------------------------

_DEFAULTS: dict = {
    "draft_id": None,
    "last_saved_msg": "",
    "custom_docs": [],
    "sf_task_docs": [],
    "doc_descriptions": {},
    "sf_tasks_pending_delete": [],
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state.draft_id is None:
    st.session_state.draft_id = new_draft_id()

# -- SF auto-fill for client fields -------------------------------------------
_sf_client = st.session_state.get("sf_client")
if _sf_client:
    _sf_cid = _sf_client.get("Id", "")
    _prev_cid = st.session_state.get("_sf_autofill_cid", "")
    if _sf_cid and _sf_cid != _prev_cid:
        # New client pulled — auto-fill fields
        _sf_name = _sf_client.get("Name", "")
        if _sf_name:
            st.session_state["inp_client_name"] = _sf_name
        _sf_anum = _sf_client.get("A_Number__c", "")
        if _sf_anum:
            st.session_state["inp_a_number"] = _sf_anum
        st.session_state["_sf_autofill_cid"] = _sf_cid


# -- Helpers ------------------------------------------------------------------

def _build_enclosed_docs_list() -> list[dict[str, str]]:
    """Collect the currently selected enclosed documents with descriptions."""
    docs: list[dict[str, str]] = []
    # SF tasks that are checked
    for doc_name in st.session_state.get("sf_task_docs", []):
        docs.append({"name": doc_name, "description": ""})
    # Custom docs
    for doc_name in st.session_state.get("custom_docs", []):
        desc = st.session_state.get("doc_descriptions", {}).get(doc_name, "")
        docs.append({"name": doc_name, "description": desc})
    return docs


def _substitute_placeholders(text: str) -> str:
    """Resolve {first_name}, {last_name}, {client_name}, {case_type}, {a_number} from session state."""
    client_name = st.session_state.get("inp_client_name", "")
    case_type = st.session_state.get("inp_case_type", "")
    a_number = st.session_state.get("inp_a_number", "")

    # Try SF client for first/last, fall back to splitting client_name
    _sf = st.session_state.get("sf_client")
    if _sf:
        first_name = _sf.get("FirstName", "")
        last_name = _sf.get("LastName", "")
    else:
        parts = client_name.rsplit(" ", 1)
        first_name = parts[0] if len(parts) > 1 else client_name
        last_name = parts[1] if len(parts) > 1 else ""

    return (
        text
        .replace("{first_name}", first_name)
        .replace("{last_name}", last_name)
        .replace("{client_name}", client_name)
        .replace("{case_type}", case_type)
        .replace("{a_number}", a_number)
    )


def _do_save(case_type: str) -> None:
    """Save the current state as a draft."""
    client = {
        "name": st.session_state.get("inp_client_name", ""),
        "a_number": st.session_state.get("inp_a_number", ""),
        "receipt_number": st.session_state.get("inp_receipt_number", ""),
    }
    # Resolve attorney from staff directory selection
    _staff_sel = st.session_state.get("inp_attorney_staff", "")
    _staff_list = load_config("staff-directory") or []
    _staff_match = next(
        (m for m in _staff_list
         if f"{m.get('first_name', '')} {m.get('last_name', '')}".strip() == _staff_sel),
        None,
    )
    _gs = load_config("global-settings") or {}
    attorney = {
        "name": _staff_sel,
        "bar_number": _staff_match.get("bar_number", "") if _staff_match else "",
        "firm_name": _gs.get("firm_name", "O'Brien Immigration Law"),
        "firm_address": _gs.get("firm_address", ""),
        "staff_id": _staff_match.get("id", "") if _staff_match else "",
    }
    filing_office = st.session_state.get("inp_filing_office", "")
    enclosed = _build_enclosed_docs_list()

    # Recipient info
    _rt = st.session_state.get("inp_recipient_type", "Government Agency")
    _rtype = "client" if _rt == "Client" else "agency"
    if _rtype == "agency":
        _all_addr = get_recipient_addresses()
        _sel_name = st.session_state.get("inp_recipient_name", "")
        _entry = next((a for a in _all_addr if a["name"] == _sel_name), None)
        _raddr = _entry["address"] if _entry else ""
        _rsal = _entry.get("salutation", "Dear Sir or Madam:") if _entry else "Dear Sir or Madam:"
    else:
        _raddr = st.session_state.get("inp_client_address", "")
        _rsal = st.session_state.get("inp_client_salutation", "Dear Sir or Madam:")

    _tmpl_sel = st.session_state.get("inp_template_selection", "Default (from Case Type)")
    _ltr_subject = st.session_state.get("inp_letter_subject", "")
    _ltr_body = st.session_state.get("inp_letter_body", "")

    # Backward-compat: also populate custom_purpose/custom_closing
    if _tmpl_sel == "Default (from Case Type)":
        _body_text = _ltr_body
        _split = _body_text.split("\n\n", 1)
        _cp = _split[0] if _split else ""
        _cc = _split[1] if len(_split) > 1 else ""
    else:
        _cp = ""
        _cc = ""

    save_draft(
        st.session_state.draft_id,
        case_type,
        client,
        attorney,
        filing_office,
        enclosed,
        recipient_type=_rtype,
        recipient_address=_raddr,
        salutation=_rsal,
        custom_purpose=_cp,
        custom_closing=_cc,
        template_selection=_tmpl_sel,
        letter_subject=_ltr_subject,
        letter_body=_ltr_body,
    )
    name = client["name"] or "draft"
    st.session_state.last_saved_msg = f"Saved -- {name}"


def _do_load(draft_id: str) -> None:
    """Load a draft into session state."""
    draft = load_draft(draft_id)
    if not draft:
        return
    st.session_state.draft_id = draft["id"]
    st.session_state.inp_case_type = draft.get("case_type", CASE_TYPES[0])

    c = draft.get("client", {})
    st.session_state.inp_client_name = c.get("name", "")
    st.session_state.inp_a_number = c.get("a_number", "")
    st.session_state.inp_receipt_number = c.get("receipt_number", "")

    a = draft.get("attorney", {})
    # Restore attorney selection — try staff_id first, fall back to name match
    _staff_id = a.get("staff_id", "")
    _atty_name = a.get("name", "")
    if _staff_id or _atty_name:
        _staff_list = load_config("staff-directory") or []
        _match = None
        if _staff_id:
            _match = next((m for m in _staff_list if m.get("id") == _staff_id), None)
        if not _match and _atty_name:
            _match = next(
                (m for m in _staff_list
                 if f"{m.get('first_name', '')} {m.get('last_name', '')}".strip() == _atty_name),
                None,
            )
        if _match:
            st.session_state.inp_attorney_staff = (
                f"{_match.get('first_name', '')} {_match.get('last_name', '')}".strip()
            )
        else:
            st.session_state.inp_attorney_staff = ""
    else:
        st.session_state.inp_attorney_staff = ""

    st.session_state.inp_filing_office = draft.get("filing_office", "")

    # Restore recipient info
    _rtype = draft.get("recipient_type", "")
    _raddr = draft.get("recipient_address", "")
    _rsal = draft.get("salutation", "Dear Sir or Madam:")

    if _rtype == "client":
        st.session_state.inp_recipient_type = "Client"
        st.session_state.inp_client_address = _raddr
        st.session_state.inp_client_salutation = _rsal
    elif _rtype == "agency" and _raddr:
        st.session_state.inp_recipient_type = "Government Agency"
        # Try to find the matching address entry to set category + name
        _all_addr = get_recipient_addresses()
        _match = next((a for a in _all_addr if a["address"] == _raddr), None)
        if _match:
            st.session_state.inp_recipient_category = _match["category"]
            st.session_state.inp_recipient_name = _match["name"]
    elif not _rtype and draft.get("filing_office"):
        # Old draft: map filing_office to an address entry
        st.session_state.inp_recipient_type = "Government Agency"
        _all_addr = get_recipient_addresses()
        _match = next((a for a in _all_addr if a["name"] == draft["filing_office"]), None)
        if _match:
            st.session_state.inp_recipient_category = _match["category"]
            st.session_state.inp_recipient_name = _match["name"]

    # Restore enclosed docs
    enclosed = draft.get("enclosed_docs", [])
    custom = []
    descs: dict[str, str] = {}
    for doc in enclosed:
        name = doc.get("name", "")
        desc = doc.get("description", "")
        custom.append(name)
        if desc:
            descs[name] = desc
    st.session_state.custom_docs = custom
    st.session_state.doc_descriptions = descs

    # Restore template / letter body fields
    _loaded_case_type = draft.get("case_type", CASE_TYPES[0])
    _loaded_tmpl = draft.get("template_selection", "")

    if _loaded_tmpl:
        # New-format draft — restore directly
        # Guard: if saved template name not in current template list, fall back
        _cl_templates = load_config("client-cover-letter-templates") or []
        _cl_names = [t["name"] for t in _cl_templates]
        if _loaded_tmpl != "Default (from Case Type)" and _loaded_tmpl not in _cl_names:
            st.session_state.inp_template_selection = "Default (from Case Type)"
        else:
            st.session_state.inp_template_selection = _loaded_tmpl
        st.session_state.inp_letter_subject = draft.get("letter_subject", "")
        st.session_state.inp_letter_body = draft.get("letter_body", "")
    else:
        # Old-format draft — convert custom_purpose/custom_closing to merged body
        _loaded_purpose = draft.get("custom_purpose", "")
        _loaded_closing = draft.get("custom_closing", "")
        _tpl = TEMPLATES.get(_loaded_case_type, {})
        _purpose = _loaded_purpose or _tpl.get("purpose_paragraph", "")
        _closing = _loaded_closing or _tpl.get("closing_paragraph", "")
        _merged = (_purpose + "\n\n" + _closing).strip() if _purpose or _closing else ""
        st.session_state.inp_template_selection = "Default (from Case Type)"
        st.session_state.inp_letter_subject = ""
        st.session_state.inp_letter_body = _merged

    st.session_state["_prev_case_type"] = _loaded_case_type
    st.session_state["_prev_template_selection"] = st.session_state.get("inp_template_selection", "Default (from Case Type)")


def _do_new() -> None:
    """Start a fresh cover letter."""
    st.session_state.draft_id = new_draft_id()
    st.session_state.last_saved_msg = ""
    st.session_state.custom_docs = []
    st.session_state.sf_task_docs = []
    st.session_state.doc_descriptions = {}
    for k in (
        "inp_client_name", "inp_a_number", "inp_receipt_number",
        "inp_attorney_staff",
        "inp_filing_office", "inp_case_type",
        "inp_recipient_type", "inp_recipient_category", "inp_recipient_name",
        "inp_client_address", "inp_client_salutation",
        "inp_purpose_paragraph", "inp_closing_paragraph",
        "inp_template_selection", "inp_letter_subject", "inp_letter_body",
        "inp_export_file_name",
        "_sf_autofill_cid", "_prev_case_type", "_prev_template_selection",
        "_gdoc_result_gdoc_export_cl", "google_doc_url",
        # EOIR tab keys
        "inp_eoir_submission_type", "inp_eoir_submission_line_1",
        "inp_eoir_submission_sub_line",
        "inp_eoir_dhs_address", "inp_eoir_served_by",
        "inp_eoir_service_method", "inp_eoir_legal_case_idx",
        "_eoir_beneficiaries", "_eoir_beneficiaries_case_id",
        "_eoir_autofill_case_id",
        "_case_bens", "_case_bens_case_id", "_case_bens_selected",
        "_case_bens_force_refresh", "_case_bens_loading", "_case_bens_error",
        # EOIR case detail keys
        "inp_eoir_cd_city", "inp_eoir_cd_applicant", "inp_eoir_cd_next_date_type",
        "inp_eoir_cd_next_govt_date",
        "inp_eoir_cd_judge", "inp_eoir_cd_attorney",
        "inp_eoir_cd_submission_type", "inp_eoir_cd_submission_line_1",
        "inp_eoir_cd_filing_method", "inp_eoir_cd_a_number",
        "_eoir_cd_lc_meta", "_eoir_cd_edit_mode",
        "_eoir_cd_autofill_case_id",
        "inp_eoir_draft", "_eoir_draft_fingerprint",
        "_eoir_canvas_order", "_eoir_canvas_edits",
        "_eoir_show_expand", "inp_eoir_expand_draft",
        # Brief Builder keys
        "inp_eoir_mode", "inp_eoir_bb_type", "_eoir_bb_case_id",
        "_eoir_bb_states", "_eoir_brief_preview_html",
    ):
        if k in st.session_state:
            del st.session_state[k]
    # Also clear brief builder content/checkbox keys
    _bb_keys = [k for k in st.session_state if k.startswith("_bb_content_") or k.startswith("_bb_chk_")]
    for k in _bb_keys:
        del st.session_state[k]


def _eoir_expand_on_save(edited_text: str) -> None:
    """Callback for the shared preview modal — splits text back into canvas blocks."""
    _new_blocks = split_eoir_into_blocks(edited_text)
    st.session_state["_eoir_canvas_order"] = [b["id"] for b in _new_blocks]
    st.session_state["_eoir_canvas_edits"] = {b["id"]: b["content"] for b in _new_blocks}
    st.session_state["_eoir_letter_text"] = edited_text
    st.session_state["_eoir_show_expand"] = False


def _open_eoir_expand():
    """Open the shared preview modal for EOIR cover page editing."""
    _expand_src = st.session_state.get("_eoir_letter_text", "")
    if show_preview_modal:
        show_preview_modal(
            title="EOIR Cover Page \u2014 Full View",
            preview_html="",
            plain_text=_expand_src,
            tool_name="cover-letters-eoir",
            edit_key="_eoir_edited_text",
            on_save=_eoir_expand_on_save,
        )
    else:
        # Fallback: inline dialog if shared module unavailable
        @st.dialog("EOIR Cover Page \u2014 Full View", width="large")
        def _fallback():
            if "inp_eoir_expand_draft" not in st.session_state:
                st.session_state["inp_eoir_expand_draft"] = _expand_src
            _edited = st.text_area(
                "Edit the full document below, then click Apply to update the canvas.",
                key="inp_eoir_expand_draft",
                height=650,
            )
            _c1, _c2, _ = st.columns([1, 1, 3])
            with _c1:
                if st.button("Apply Changes", type="primary", key="_expand_apply"):
                    _eoir_expand_on_save(_edited)
                    if "inp_eoir_expand_draft" in st.session_state:
                        del st.session_state["inp_eoir_expand_draft"]
                    st.rerun()
            with _c2:
                if st.button("Cancel", key="_expand_cancel"):
                    if "inp_eoir_expand_draft" in st.session_state:
                        del st.session_state["inp_eoir_expand_draft"]
                    st.session_state["_eoir_show_expand"] = False
                    st.rerun()
        _fallback()


@st.dialog("Add New Address")
def _add_address_dialog():
    """Dialog for adding a new recipient address."""
    _new_name = st.text_input("Name", placeholder="e.g. USCIS Field Office - Boston")
    _new_category = st.selectbox("Category", options=RECIPIENT_CATEGORIES)
    _new_address = st.text_area(
        "Address",
        height=100,
        placeholder="Organization Name\nStreet Address\nCity, State ZIP",
    )
    _new_salutation = st.text_input("Salutation", value="Dear Sir or Madam:")

    if st.button("Save Address", type="primary", use_container_width=True):
        if not _new_name or not _new_address:
            st.error("Name and address are required.")
        else:
            _addresses = get_recipient_addresses()
            _new_id = _new_name.lower().replace(" ", "_").replace("(", "").replace(")", "")[:40]
            _addresses.append({
                "id": _new_id,
                "name": _new_name,
                "category": _new_category,
                "address": _new_address,
                "salutation": _new_salutation,
            })
            save_recipient_addresses(_addresses)
            st.session_state["_show_add_address"] = False
            st.rerun()


def _build_preview_html(
    letter_text: str,
    case_type: str,
) -> str:
    """Convert the plain-text cover letter into styled HTML for the preview."""
    esc = html_mod.escape
    tpl = TEMPLATES.get(case_type, {})

    lines = letter_text.split("\n")
    parts: list[str] = []

    # If VAWA, show a confidentiality badge
    if "confidentiality_notice" in tpl:
        notice = tpl["confidentiality_notice"]
        parts.append(f'<div class="conf-notice">{esc(notice)}</div>')

    for line in lines:
        escaped = esc(line)
        if line.startswith("RE: "):
            parts.append(f"<strong>{escaped}</strong>")
        elif line.startswith("    A# ") or line.startswith("    Receipt#") or (line.startswith("    ") and not line.startswith("    " + "  ")):
            parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;{esc(line.strip())}")
        elif line == "CERTIFICATE OF SERVICE":
            parts.append(f'<div class="letter-cert"><strong>{escaped}</strong></div>')
        elif line.startswith("____"):
            parts.append(f"<br>{escaped}")
        else:
            parts.append(escaped)

    return "<br>".join(parts)


def _build_eoir_preview_html(text: str) -> str:
    """Convert EOIR submission plain text into styled HTML for the preview panel."""
    esc = html_mod.escape
    lines = text.split("\n")
    parts: list[str] = []

    # Centered headings
    _centered = {
        "UNITED STATES DEPARTMENT OF JUSTICE",
        "EXECUTIVE OFFICE FOR IMMIGRATION REVIEW",
        "IMMIGRATION COURT",
        "IN THE MATTERS OF:",
        "RESPONDENT'S SUBMISSION",
        "CERTIFICATE OF SERVICE",
    }

    for line in lines:
        stripped = line.strip()
        if stripped in _centered:
            if stripped == "CERTIFICATE OF SERVICE":
                parts.append(
                    f'<div class="letter-cert" style="text-align:center;">'
                    f"<strong>{esc(stripped)}</strong></div>"
                )
            else:
                parts.append(f'<div style="text-align:center;"><strong>{esc(stripped)}</strong></div>')
        elif line.startswith("    ") and len(line) > 4 and line.strip()[0:1].isdigit():
            # Numbered document list item
            parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;{esc(stripped)}")
        elif line.startswith("____"):
            parts.append(f"<br>{esc(line)}")
        elif stripped.startswith("Bar No."):
            parts.append(esc(stripped))
        elif stripped.startswith("Tel:") or stripped.startswith("Fax:") or stripped.startswith("Email:"):
            parts.append(esc(stripped))
        else:
            parts.append(esc(line))

    return "<br>".join(parts)


def _build_docx(letter_text: str, attorney_name: str, firm_name: str) -> bytes:
    """Build a Word document from the rendered cover letter text."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, txt, size=12, bold=False):
        r = para.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        return r

    lines = letter_text.split("\n")
    for line in lines:
        p = doc.add_paragraph()
        is_bold = line.startswith("RE:") or line == "CERTIFICATE OF SERVICE"
        _run(p, line, bold=is_bold)
        fmt = p.paragraph_format
        fmt.space_after = Pt(0)
        fmt.space_before = Pt(0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_cover_letter_pdf(letter_text: str) -> bytes:
    """Build a PDF from the rendered cover letter text using fpdf2.

    Handles RE: lines (bold), indented lines (10mm indent),
    CERTIFICATE OF SERVICE (bold + spacing), underscores (spacing),
    and normal text (multi_cell wrapping). 1-inch margins, Times 12pt.
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_margins(25.4, 25.4, 25.4)  # 1-inch margins — must precede add_page
    pdf.set_auto_page_break(auto=True, margin=25.4)
    pdf.add_page()
    pdf.set_font("Times", "", 12)

    # Explicit width avoids fpdf2 w=0 edge cases on page breaks
    _w = pdf.w - pdf.l_margin - pdf.r_margin
    _w_indent = _w - 10

    for line in letter_text.split("\n"):
        pdf.set_x(pdf.l_margin)
        if not line.strip():
            pdf.ln(6)
        elif line.startswith("RE: "):
            pdf.set_font("Times", "B", 12)
            pdf.multi_cell(_w, 6, line, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Times", "", 12)
        elif line == "CERTIFICATE OF SERVICE":
            pdf.ln(6)
            pdf.set_font("Times", "B", 12)
            pdf.multi_cell(_w, 6, line, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Times", "", 12)
            pdf.ln(3)
        elif line.startswith("____"):
            pdf.ln(6)
            pdf.multi_cell(_w, 6, line, new_x="LMARGIN", new_y="NEXT")
        elif line.startswith("    "):
            pdf.set_x(pdf.l_margin + 10)
            pdf.multi_cell(_w_indent, 6, line.strip(), new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.multi_cell(_w, 6, line, new_x="LMARGIN", new_y="NEXT")

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# -- Brief Builder helpers (for EOIR tab) -------------------------------------

_ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
          "XI", "XII", "XIII", "XIV", "XV"]

_BB_TITLE_MAP = {
    "Asylum Merits Brief": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR ASYLUM",
    "Motion to Reopen": "MOTION TO REOPEN PROCEEDINGS",
    "Appeal Brief": "RESPONDENT'S BRIEF ON APPEAL",
    "Bond Brief": "RESPONDENT'S BRIEF IN SUPPORT OF BOND REDETERMINATION",
    "Cancellation of Removal": "RESPONDENT'S BRIEF IN SUPPORT OF APPLICATION FOR CANCELLATION OF REMOVAL",
    "Motion to Terminate": "MOTION TO TERMINATE PROCEEDINGS",
    "Motion for Continuance": "MOTION FOR CONTINUANCE",
    "Motion to Change Venue": "MOTION TO CHANGE VENUE",
    "Motion to Suppress": "MOTION TO SUPPRESS EVIDENCE",
}


def _bb_content_key(section_key: str, sub_key: str | None = None) -> str:
    """Build a session-state key for a brief builder content text_area."""
    if sub_key:
        return f"_bb_content_{section_key}_{sub_key}"
    return f"_bb_content_{section_key}"


def _load_brief_sections(brief_type: str) -> list[dict]:
    """Load sections + boilerplate for a brief type from config store."""
    all_types = get_config_value("brief-builder", "brief_types", {})
    if brief_type not in all_types:
        return []
    all_bp = get_config_value("brief-builder", "boilerplate", {})
    bp = all_bp.get(brief_type, {})

    sections = []
    for section_def in all_types[brief_type]:
        section = {
            "key": section_def["key"],
            "heading": section_def["heading"],
            "subsections": [],
        }
        for sub in section_def.get("subsections", []):
            sub_entry = {"key": sub["key"], "heading": sub["heading"]}
            if sub["key"] in bp:
                sub_entry["boilerplate"] = bp[sub["key"]]
            section["subsections"].append(sub_entry)
        if section_def["key"] in bp:
            section["boilerplate"] = bp[section_def["key"]]
        sections.append(section)
    return sections


def _resolve_variables(text: str, var_map: dict[str, str]) -> str:
    """Replace {{variable}} placeholders in text with values from var_map."""
    result = text
    for key, val in var_map.items():
        result = result.replace("{{" + key + "}}", val)
    return result


def _build_brief_preview_html(
    brief_type: str,
    case_info: dict[str, str],
    sections: list[dict],
    sections_content: dict[str, str],
    submission_desc: str,
    checked_sections: set[str],
    attorney_name_val: str,
) -> str:
    """Render the brief as styled HTML for the live preview panel."""
    esc = html_mod.escape
    parts: list[str] = []

    # Title
    parts.append(
        f'<div class="brief-title">'
        f'{esc(_BB_TITLE_MAP.get(brief_type, brief_type.upper()))}'
        f'</div>'
    )

    # Caption
    caption_lines = []
    if case_info.get("client_name"):
        caption_lines.append(f"IN THE MATTER OF: {esc(case_info['client_name'].upper())}")
    if case_info.get("a_number"):
        caption_lines.append(f"A-Number: {esc(case_info['a_number'])}")
    if case_info.get("court"):
        caption_lines.append(f"Before the {esc(case_info['court'])}")
    if case_info.get("judge"):
        caption_lines.append(f"Immigration Judge {esc(case_info['judge'])}")
    if case_info.get("hearing_date"):
        caption_lines.append(f"Hearing Date: {esc(case_info['hearing_date'])}")
    if caption_lines:
        parts.append(f'<div class="brief-caption">{"<br>".join(caption_lines)}</div>')

    # Submission description as intro paragraph
    if submission_desc.strip():
        parts.append(f'<div class="brief-body">{esc(submission_desc.strip())}</div>')

    # Sections
    for idx, section in enumerate(sections):
        section_key = section["key"]
        if section_key not in checked_sections:
            continue
        roman = _ROMAN[idx] if idx < len(_ROMAN) else str(idx + 1)
        heading = section["heading"]
        subs = section.get("subsections", [])

        parts.append(f'<div class="brief-heading">{roman}. {esc(heading)}</div>')

        if subs:
            for sub_idx, sub in enumerate(subs):
                sub_letter = chr(ord("A") + sub_idx)
                parts.append(
                    f'<div class="brief-subheading">{sub_letter}. {esc(sub["heading"])}</div>'
                )
                ck = _bb_content_key(section_key, sub["key"])
                body = sections_content.get(ck, "").strip()
                if body:
                    for para in body.split("\n\n"):
                        para = para.strip()
                        if para:
                            parts.append(f'<div class="brief-body">{esc(para)}</div>')
        else:
            ck = _bb_content_key(section_key)
            body = sections_content.get(ck, "").strip()
            if body:
                for para in body.split("\n\n"):
                    para = para.strip()
                    if para:
                        parts.append(f'<div class="brief-body">{esc(para)}</div>')

    # Signature block
    _sig = (
        '<div class="brief-sig">'
        "Respectfully submitted,<br><br>"
        "____________________________<br>"
    )
    if attorney_name_val:
        _sig += f"{esc(attorney_name_val)}<br>"
    _sig += (
        "Attorney for Respondent<br>"
        f"Date: {date.today().strftime('%m/%d/%Y')}"
        "</div>"
    )
    parts.append(_sig)

    return "\n".join(parts)


def _bb_save_state(case_id: str) -> None:
    """Save current brief builder state for a case."""
    if not case_id:
        return
    states = st.session_state.get("_eoir_bb_states", {})
    state = {
        "brief_type": st.session_state.get("inp_eoir_bb_type", ""),
    }
    # Save all _bb_content_ and _bb_chk_ keys
    for k, v in st.session_state.items():
        if k.startswith("_bb_content_") or k.startswith("_bb_chk_"):
            state[k] = v
    states[case_id] = state
    st.session_state["_eoir_bb_states"] = states


def _bb_restore_state(case_id: str) -> None:
    """Restore brief builder state for a case."""
    if not case_id:
        return
    states = st.session_state.get("_eoir_bb_states", {})
    state = states.get(case_id, {})
    if not state:
        return
    if "brief_type" in state:
        st.session_state["inp_eoir_bb_type"] = state["brief_type"]
    for k, v in state.items():
        if k.startswith("_bb_content_") or k.startswith("_bb_chk_"):
            st.session_state[k] = v


def _bb_clear_content_keys() -> None:
    """Remove all _bb_content_ and _bb_chk_ keys from session state."""
    to_del = [k for k in st.session_state if k.startswith("_bb_content_") or k.startswith("_bb_chk_")]
    for k in to_del:
        del st.session_state[k]


# -- Sidebar (drafts only) ---------------------------------------------------

with st.sidebar:
    # Draft management
    st.markdown("#### Drafts")
    btn_cols = st.columns(2)
    with btn_cols[0]:
        if st.button("New", use_container_width=True):
            _do_new()
            st.rerun()
    with btn_cols[1]:
        save_clicked = st.button("Save", use_container_width=True, type="primary")

    saved_drafts = list_drafts()
    if saved_drafts:
        labels_map = {
            d["id"]: f"{d['client_name']} -- {d['case_type']}"
            for d in saved_drafts
        }
        draft_ids = list(labels_map.keys())
        selected_draft = st.selectbox(
            "Load a saved draft",
            options=[""] + draft_ids,
            format_func=lambda x: labels_map.get(x, "Select..."),
            label_visibility="collapsed",
        )
        load_cols = st.columns(2)
        with load_cols[0]:
            if selected_draft and st.button("Load", use_container_width=True):
                _do_load(selected_draft)
                st.rerun()
        with load_cols[1]:
            if selected_draft and st.button("Delete", use_container_width=True):
                delete_draft(selected_draft)
                st.rerun()

    if st.session_state.last_saved_msg:
        st.markdown(
            f'<div class="saved-toast">{html_mod.escape(st.session_state.last_saved_msg)}</div>',
            unsafe_allow_html=True,
        )

    render_tool_notes("cover-letters")


# -- Add Address dialog trigger -----------------------------------------------
if st.session_state.get("_show_add_address"):
    _add_address_dialog()


# -- Main area ----------------------------------------------------------------

# Pre-compute attorney/firm info from session state so the preview (rendered
# inside the tab BEFORE the footer selectbox) uses the current values.
_staff_dir = load_config("staff-directory") or []
_staff_dir = [m for m in _staff_dir if m.get("visible", True)]
_gs_global = load_config("global-settings") or {}
_comp_config_global = load_config("components") or {}
_eoir_cover_template = _comp_config_global.get("eoir_cover_template") or None
_doc_tpl_config = load_config("document-templates") or {}
_atty_sel_pre = st.session_state.get("inp_attorney_staff", "")
_matched_pre = next(
    (m for m in _staff_dir
     if f"{m.get('first_name', '')} {m.get('last_name', '')}".strip() == _atty_sel_pre),
    None,
)
attorney_name = _atty_sel_pre
bar_number = _matched_pre.get("bar_number", "") if _matched_pre else ""
firm_name = _gs_global.get("firm_name", "O'Brien Immigration Law")
firm_address = _gs_global.get("firm_address", "")

letter_text = ""

tab_cover, tab_eoir = st.tabs(["Cover Letter", "EOIR Submission"])

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 1: Cover Letter
# ═══════════════════════════════════════════════════════════════════════════════
with tab_cover:
    _cl_left, _cl_right = st.columns([3, 2], gap="large")

    with _cl_left:

        # ── Step 1: Case Type ────────────────────────────────────────────────
        st.markdown('<div class="section-label">Case Type</div>', unsafe_allow_html=True)
        case_type = st.selectbox(
            "Case Type",
            options=CASE_TYPES,
            key="inp_case_type",
            label_visibility="collapsed",
        )

        st.divider()

        # ── Step 2: Recipient ────────────────────────────────────────────────
        st.markdown('<div class="section-label">Recipient</div>', unsafe_allow_html=True)
        recipient_type = st.radio(
            "Recipient type",
            options=["Government Agency", "Client"],
            key="inp_recipient_type",
            horizontal=True,
            label_visibility="collapsed",
        )

        if recipient_type == "Government Agency":
            _all_addresses = get_recipient_addresses()

            _selected_category = st.selectbox(
                "Category",
                options=RECIPIENT_CATEGORIES,
                key="inp_recipient_category",
            )

            _filtered = [a for a in _all_addresses if a.get("category") == _selected_category]
            _addr_names = [a["name"] for a in _filtered]

            if _addr_names:
                _selected_name = st.selectbox(
                    "Address",
                    options=_addr_names,
                    key="inp_recipient_name",
                )
                _selected_entry = next((a for a in _filtered if a["name"] == _selected_name), None)
                if _selected_entry:
                    st.caption(_selected_entry["address"].replace("\n", "  \n"))
                    recipient_address = _selected_entry["address"]
                    salutation = _selected_entry.get("salutation", "Dear Sir or Madam:")
                else:
                    recipient_address = ""
                    salutation = "Dear Sir or Madam:"
            else:
                st.caption("No addresses in this category.")
                recipient_address = ""
                salutation = "Dear Sir or Madam:"

            if st.button("+ Add New Address", use_container_width=True):
                st.session_state["_show_add_address"] = True

            # Legacy fallback for filing_office (used by render_cover_letter if recipient_address empty)
            filing_office = ""

        else:
            # Client recipient
            _sf_client = st.session_state.get("sf_client")
            _default_addr = ""
            _default_sal = "Dear Sir or Madam:"
            if _sf_client:
                _parts = []
                _first = _sf_client.get("FirstName", "")
                _last = _sf_client.get("LastName", "")
                _full_name = f"{_first} {_last}".strip()
                if _full_name:
                    _parts.append(_full_name)
                _street = _sf_client.get("MailingStreet", "")
                if _street:
                    _parts.append(_street)
                _city = _sf_client.get("MailingCity", "")
                _state = _sf_client.get("MailingState", "")
                _zip = _sf_client.get("MailingPostalCode", "")
                _csz = ", ".join(filter(None, [_city, _state]))
                if _zip:
                    _csz = f"{_csz} {_zip}" if _csz else _zip
                if _csz:
                    _parts.append(_csz)
                _default_addr = "\n".join(_parts)
                if _first or _last:
                    _default_sal = f"Dear {_full_name}:"

            recipient_address = st.text_area(
                "Recipient Address",
                value=st.session_state.get("inp_client_address", _default_addr),
                key="inp_client_address",
                height=100,
                placeholder="Client name\nStreet address\nCity, State ZIP",
            )
            salutation = st.text_input(
                "Salutation",
                value=st.session_state.get("inp_client_salutation", _default_sal),
                key="inp_client_salutation",
            )
            filing_office = ""

        st.divider()

        # ── Step 3: Client / RE Block ────────────────────────────────────────
        st.markdown('<div class="section-label">Client</div>', unsafe_allow_html=True)
        client_name = st.text_input(
            "Client Name",
            key="inp_client_name",
            placeholder="e.g. Maria Garcia Lopez",
        )
        a_number = st.text_input(
            "A-Number",
            key="inp_a_number",
            placeholder="e.g. 123-456-789",
        )
        receipt_number = st.text_input(
            "Receipt Number",
            key="inp_receipt_number",
            placeholder="e.g. SRC-21-123-45678",
        )

        st.divider()

        # ── Step 4: Letter Body ──────────────────────────────────────────────
        st.markdown('<div class="section-label">Letter Body</div>', unsafe_allow_html=True)

        # Load client cover letter templates
        _cl_templates = load_config("client-cover-letter-templates") or []
        _cl_template_names = [t["name"] for t in _cl_templates]
        _template_options = ["Default (from Case Type)"] + _cl_template_names

        template_selection = st.selectbox(
            "Template",
            options=_template_options,
            key="inp_template_selection",
        )

        # ── Change detection for template + case type ──
        _prev_tmpl = st.session_state.get("_prev_template_selection")
        _prev_ct = st.session_state.get("_prev_case_type")
        _is_default = template_selection == "Default (from Case Type)"

        if _prev_tmpl is not None and _prev_tmpl != template_selection:
            # Template selection changed
            if _is_default:
                # Switched back to Default — populate from case type
                _ct_tpl = TEMPLATES.get(case_type, {})
                _purpose = _ct_tpl.get("purpose_paragraph", "")
                _closing = _ct_tpl.get("closing_paragraph", "")
                _merged = (_purpose + "\n\n" + _closing).strip()
                st.session_state["inp_letter_body"] = _merged
                st.session_state["inp_letter_subject"] = ""
            else:
                # Switched to a named template — populate from it
                _sel_tpl = next((t for t in _cl_templates if t["name"] == template_selection), None)
                if _sel_tpl:
                    st.session_state["inp_letter_subject"] = _sel_tpl.get("subject", "")
                    st.session_state["inp_letter_body"] = _sel_tpl.get("body", "")
        elif _is_default and _prev_ct is not None and _prev_ct != case_type:
            # Case type changed while in Default mode — repopulate body
            _ct_tpl = TEMPLATES.get(case_type, {})
            _purpose = _ct_tpl.get("purpose_paragraph", "")
            _closing = _ct_tpl.get("closing_paragraph", "")
            _merged = (_purpose + "\n\n" + _closing).strip()
            st.session_state["inp_letter_body"] = _merged
        elif _prev_tmpl is None:
            # First load — only set if not already in session state
            if "inp_letter_body" not in st.session_state:
                _ct_tpl = TEMPLATES.get(case_type, {})
                _purpose = _ct_tpl.get("purpose_paragraph", "")
                _closing = _ct_tpl.get("closing_paragraph", "")
                _merged = (_purpose + "\n\n" + _closing).strip()
                st.session_state["inp_letter_body"] = _merged

        st.session_state["_prev_template_selection"] = template_selection
        st.session_state["_prev_case_type"] = case_type

        # Subject field — only visible for non-Default templates
        if not _is_default:
            letter_subject = st.text_input(
                "Subject",
                key="inp_letter_subject",
                placeholder="e.g. Documents Needed for Your Case",
            )
        else:
            letter_subject = ""

        # Body field
        letter_body = st.text_area(
            "Body",
            key="inp_letter_body",
            height=250,
            help="Use {client_name}, {first_name}, {last_name}, {case_type}, {a_number} as placeholders.",
        )

        # ── Case Beneficiaries (Cover Letter) ─────────────────────────────
        # Show beneficiaries from the active legal case if available
        _cl_sf_client = st.session_state.get("sf_client")
        _cl_case_bens: list[dict] = st.session_state.get("_case_bens", [])
        if _cl_sf_client and _cl_case_bens:
            st.divider()
            st.markdown('<div class="section-label">Case Beneficiaries</div>', unsafe_allow_html=True)
            st.caption("Select beneficiaries to include in this cover letter.")
            _cl_sel_ids: list[str] = st.session_state.get("_case_bens_selected", [])

            _cl_all_ids = [b.get("Id") for b in _cl_case_bens]
            _cl_all_sel = all(bid in _cl_sel_ids for bid in _cl_all_ids)
            if st.checkbox("Select All", value=_cl_all_sel, key="_cl_bens_sel_all"):
                if not _cl_all_sel:
                    st.session_state["_case_bens_selected"] = list(_cl_all_ids)
                    st.rerun()
            elif _cl_all_sel:
                st.session_state["_case_bens_selected"] = []
                st.rerun()

            _cl_new_sel: list[str] = []
            for _ci, _cb in enumerate(_cl_case_bens):
                _cid = _cb.get("Id", "")
                _cname = _cb.get("Contact_Name", "") or "Unnamed"
                _crel = _cb.get("Role__c", "")
                _canum = _cb.get("Alien_Number_Dashed__c", "")
                _clabel = _cname
                if _crel:
                    _clabel += f" ({_crel})"
                if _canum:
                    _clabel += f" — A# {_canum}"
                if st.checkbox(_clabel, value=_cid in _cl_sel_ids, key=f"_cl_cben_chk_{_ci}"):
                    _cl_new_sel.append(_cid)
            st.session_state["_case_bens_selected"] = _cl_new_sel
        elif _cl_sf_client and not _cl_case_bens:
            st.divider()
            st.caption("No derivatives found for this case.")

        st.divider()

        # ── Step 5: Enclosed Documents ───────────────────────────────────────
        st.markdown('<div class="section-label">Enclosed Documents</div>', unsafe_allow_html=True)
        st.caption("Check the documents to include in this cover letter. Add descriptions for specificity.")

        # -- Salesforce Tasks (LC_Task__c) ------------------------------------
        _active_client = st.session_state.get("sf_client")
        _sf_contact_id = _active_client.get("Id", "") if _active_client else ""
        _sf_tasks: list[dict] = []

        if _sf_contact_id:
            try:
                from shared.salesforce_client import get_lc_tasks
                _sf_tasks = get_lc_tasks(_sf_contact_id)
            except Exception:
                _sf_tasks = []

        if _sf_tasks or _sf_contact_id:
            st.markdown('<div class="section-label" style="margin-top:4px;">Salesforce Tasks</div>', unsafe_allow_html=True)

        _pending_deletes: list[str] = list(st.session_state.get("sf_tasks_pending_delete", []))

        if _sf_tasks:
            _current_sf_checked = list(st.session_state.get("sf_task_docs", []))
            _new_sf_checked: list[str] = []
            _has_edits = False

            for _task in _sf_tasks:
                _task_id = _task.get("Id", "")
                _task_label = _task.get("For__c") or _task.get("Name") or "Untitled task"
                _is_pending_delete = _task_id in _pending_deletes
                _edit_key = f"_sf_task_edit_{_task_id}"

                if _is_pending_delete:
                    st.markdown(
                        f'<div class="sf-task-deleted">'
                        f'<div class="doc-item">{html_mod.escape(_task_label)}</div>'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
                    _undo_cols = st.columns([8, 1])
                    with _undo_cols[1]:
                        if st.button("Undo", key=f"_btn_undo_sft_{_task_id}", use_container_width=True):
                            _pending_deletes.remove(_task_id)
                            st.session_state.sf_tasks_pending_delete = _pending_deletes
                            st.rerun()
                else:
                    _edited_label = st.session_state.get(_edit_key, _task_label)
                    _was_checked = _edited_label in _current_sf_checked or _task_label in _current_sf_checked
                    _tc = st.columns([0.5, 8, 0.5])
                    with _tc[0]:
                        _tc_checked = st.checkbox(
                            _task_label,
                            value=_was_checked,
                            key=f"chk_sft_{_task_id}",
                            label_visibility="collapsed",
                        )
                    with _tc[1]:
                        _edited_val = st.text_input(
                            "Edit document name",
                            value=_edited_label,
                            key=_edit_key,
                            label_visibility="collapsed",
                        )
                    with _tc[2]:
                        if st.button("X", key=f"_btn_del_sft_{_task_id}", help="Mark for deletion"):
                            _pending_deletes.append(_task_id)
                            st.session_state.sf_tasks_pending_delete = _pending_deletes
                            st.rerun()

                    if _edited_val.strip() != _task_label:
                        _has_edits = True
                    if _tc_checked and _edited_val.strip():
                        _new_sf_checked.append(_edited_val.strip())

            st.session_state.sf_task_docs = _new_sf_checked

            # ── Save Changes button ──
            _has_changes = _has_edits or len(_pending_deletes) > 0
            if _has_changes:
                _badge_parts: list[str] = []
                _edit_count = sum(
                    1 for _t in _sf_tasks
                    if _t.get("Id", "") not in _pending_deletes
                    and st.session_state.get(f"_sf_task_edit_{_t.get('Id', '')}", "").strip()
                       != (_t.get("For__c") or _t.get("Name") or "Untitled task")
                    and st.session_state.get(f"_sf_task_edit_{_t.get('Id', '')}", "") != ""
                )
                if _edit_count:
                    _badge_parts.append(f'<span class="pending-badge edits">{_edit_count} edit{"s" if _edit_count != 1 else ""}</span>')
                if _pending_deletes:
                    _del_count = len(_pending_deletes)
                    _badge_parts.append(f'<span class="pending-badge deletes">{_del_count} deletion{"s" if _del_count != 1 else ""}</span>')
                st.markdown(f'{"".join(_badge_parts)}', unsafe_allow_html=True)

                if st.button("Save Changes to Salesforce", type="primary", use_container_width=True, key="_btn_save_all_sf"):
                    _save_errors: list[str] = []
                    try:
                        from shared.salesforce_client import update_lc_task, delete_lc_task

                        for _t in _sf_tasks:
                            _tid = _t.get("Id", "")
                            if _tid in _pending_deletes:
                                continue
                            _orig = _t.get("For__c") or _t.get("Name") or "Untitled task"
                            _new_val = st.session_state.get(f"_sf_task_edit_{_tid}", "").strip()
                            if _new_val and _new_val != _orig:
                                try:
                                    update_lc_task(_tid, _new_val)
                                except Exception as e:
                                    _save_errors.append(f"Edit '{_new_val}': {e}")

                        for _tid in _pending_deletes:
                            try:
                                delete_lc_task(_tid)
                            except Exception as e:
                                _save_errors.append(f"Delete {_tid}: {e}")

                        st.session_state.sf_tasks_pending_delete = []
                        for _t in _sf_tasks:
                            _ek = f"_sf_task_edit_{_t.get('Id', '')}"
                            if _ek in st.session_state:
                                del st.session_state[_ek]

                        if _save_errors:
                            st.warning(f"Saved with {len(_save_errors)} error(s): {'; '.join(_save_errors)}")
                        else:
                            st.success("Changes saved to Salesforce.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Save failed: {e}")

        elif _sf_contact_id:
            st.caption("No tasks found for this client.")

        from shared.document_adder import render_document_adder
        render_document_adder("cover-letters")

        descriptions = st.session_state.get("doc_descriptions", {})
        new_descriptions: dict[str, str] = dict(descriptions)

        # Custom documents
        custom_docs = list(st.session_state.get("custom_docs", []))
        new_custom: list[str] = []
        for doc_name in custom_docs:
            cols = st.columns([0.5, 5, 3, 1])
            with cols[0]:
                st.checkbox(
                    doc_name,
                    value=True,
                    key=f"chk_custom_{hash(doc_name)}",
                    label_visibility="collapsed",
                    disabled=True,
                )
            with cols[1]:
                st.markdown(
                    f'<div class="doc-item">{html_mod.escape(doc_name)}</div>',
                    unsafe_allow_html=True,
                )
            with cols[2]:
                desc_val = descriptions.get(doc_name, "")
                desc = st.text_input(
                    "Description",
                    value=desc_val,
                    key=f"desc_custom_{hash(doc_name)}",
                    label_visibility="collapsed",
                    placeholder="Optional description...",
                )
                if desc:
                    new_descriptions[doc_name] = desc
                elif doc_name in new_descriptions:
                    new_descriptions.pop(doc_name, None)
            with cols[3]:
                if st.button("X", key=f"rm_{hash(doc_name)}", help="Remove custom document"):
                    custom_docs.remove(doc_name)
                    st.session_state.custom_docs = custom_docs
                    new_descriptions.pop(doc_name, None)
                    st.rerun()
            new_custom.append(doc_name)

        st.session_state.custom_docs = new_custom
        st.session_state.doc_descriptions = new_descriptions

    # ── Cover Letter Preview (right column) ──────────────────────────────────
    with _cl_right:
        st.markdown('<div class="section-label">Cover Letter Preview</div>', unsafe_allow_html=True)
        if not client_name:
            st.info("Enter the client's name to see the live preview.")
        else:
            all_enclosed = _build_enclosed_docs_list()

            _is_default_mode = template_selection == "Default (from Case Type)"

            if _is_default_mode:
                # Default: split body on first \n\n → purpose / closing
                _body_text = letter_body or ""
                _split = _body_text.split("\n\n", 1)
                _raw_purpose = _split[0] if _split else ""
                _raw_closing = _split[1] if len(_split) > 1 else ""

                _purpose_rendered = _substitute_placeholders(_raw_purpose)
                _purpose_safe = _purpose_rendered.replace("{", "{{").replace("}", "}}") if _purpose_rendered else ""
                _closing_rendered = _substitute_placeholders(_raw_closing)
                _closing_safe = _closing_rendered.replace("{", "{{").replace("}", "}}") if _closing_rendered else ""

                letter_text = render_cover_letter(
                    case_type=case_type,
                    client_name=client_name,
                    a_number=a_number,
                    receipt_number=receipt_number,
                    filing_office=filing_office,
                    enclosed_docs=all_enclosed,
                    attorney_name=attorney_name,
                    bar_number=bar_number,
                    firm_name=firm_name,
                    firm_address=firm_address,
                    custom_purpose=_purpose_safe,
                    custom_closing=_closing_safe,
                    recipient_address=recipient_address,
                    salutation=salutation,
                )
            else:
                # Template mode: subject + body → enclosed docs → signature
                _subj_rendered = _substitute_placeholders(letter_subject or "")
                _body_rendered = _substitute_placeholders(letter_body or "")
                # Escape remaining braces
                _body_safe = _body_rendered.replace("{", "{{").replace("}", "}}") if _body_rendered else ""

                letter_text = render_cover_letter(
                    case_type=case_type,
                    client_name=client_name,
                    a_number=a_number,
                    receipt_number=receipt_number,
                    filing_office=filing_office,
                    enclosed_docs=all_enclosed,
                    attorney_name=attorney_name,
                    bar_number=bar_number,
                    firm_name=firm_name,
                    firm_address=firm_address,
                    recipient_address=recipient_address,
                    salutation=salutation,
                    custom_subject=_subj_rendered,
                    custom_body=_body_safe,
                )

            # Append selected beneficiaries to letter text
            _cl_sel_ben_ids = st.session_state.get("_case_bens_selected", [])
            _cl_bens_for_letter = [
                b for b in (st.session_state.get("_case_bens") or [])
                if b.get("Id") in _cl_sel_ben_ids
            ]
            if _cl_bens_for_letter:
                _ben_parts: list[str] = []
                for _b in _cl_bens_for_letter:
                    _bn = _b.get("Contact_Name", "") or "Unnamed"
                    _br = (_b.get("Role__c") or "").lower()
                    _ben_parts.append(f"{_br}: {_bn}" if _br else _bn)
                _ben_sentence = "Also filing for " + ", ".join(_ben_parts) + "."
                # Insert before the signature block (find "Sincerely" or "Respectfully")
                for _sig in ("Sincerely,", "Respectfully,", "Respectfully submitted,"):
                    if _sig in letter_text:
                        letter_text = letter_text.replace(_sig, f"{_ben_sentence}\n\n{_sig}", 1)
                        break
                else:
                    letter_text += f"\n\n{_ben_sentence}"

            preview_html = _build_preview_html(letter_text, case_type)
            st.markdown(
                f'<div class="preview-panel">{preview_html}</div>',
                unsafe_allow_html=True,
            )

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 2: EOIR Submission
# ═══════════════════════════════════════════════════════════════════════════════
with tab_eoir:
    _eoir_left, _eoir_right = st.columns([3, 2], gap="large")

    # Defaults for canvas and brief preview (set before columns so right column always has them)
    _eoir_blocks: list[dict] = []
    _reset_canvas = False
    _canvas_order: list[str] = []
    _canvas_edits: dict = {}
    _block_map: dict = {}
    _eoir_brief_preview_html = ""

    # Resolve legal cases from active client
    _eoir_sf_client = st.session_state.get("sf_client")
    _eoir_legal_cases: list[dict] = []
    if _eoir_sf_client:
        _eoir_legal_cases = _eoir_sf_client.get("legal_cases", [])

    with _eoir_left:
        if not _eoir_sf_client:
            st.info("Pull a client to use the EOIR Submission tab.")
        elif not _eoir_legal_cases:
            st.warning("No legal cases found for this client.")
        else:
            # ── Legal Case Selector ──────────────────────────────────────
            st.markdown('<div class="section-label">Legal Case</div>', unsafe_allow_html=True)
            _lc_labels = []
            for _lc in _eoir_legal_cases:
                _lc_name = _lc.get("Name", "")
                _lc_type = _lc.get("Legal_Case_Type__c", "")
                _lc_labels.append(f"{_lc_name} — {_lc_type}" if _lc_type else _lc_name)

            # Default to the selected_legal_case from Client Info if available
            _preselected_lc = _eoir_sf_client.get("selected_legal_case")
            _default_lc_idx = 0
            if _preselected_lc:
                for _i, _lc in enumerate(_eoir_legal_cases):
                    if _lc.get("Id") == _preselected_lc.get("Id"):
                        _default_lc_idx = _i
                        break

            _eoir_lc_idx = st.selectbox(
                "Select Legal Case",
                options=range(len(_eoir_legal_cases)),
                index=st.session_state.get("inp_eoir_legal_case_idx", _default_lc_idx),
                format_func=lambda i: _lc_labels[i],
                key="inp_eoir_legal_case_idx",
                label_visibility="collapsed",
            )
            _eoir_case = _eoir_legal_cases[_eoir_lc_idx]
            _eoir_case_id = _eoir_case.get("Id", "")

            # ── Auto-fill from case fields on case change ────────────────
            _prev_autofill = st.session_state.get("_eoir_autofill_case_id", "")
            if _eoir_case_id and _eoir_case_id != _prev_autofill:
                st.session_state["inp_eoir_dhs_address"] = _eoir_case.get("DHS_Address__c", "") or ""
                st.session_state["_eoir_autofill_case_id"] = _eoir_case_id

            # ── Fetch beneficiaries (cached per case) ────────────────────
            _cached_ben_case = st.session_state.get("_eoir_beneficiaries_case_id", "")
            if _eoir_case_id != _cached_ben_case:
                if get_beneficiaries and _eoir_case_id:
                    try:
                        st.session_state["_eoir_beneficiaries"] = get_beneficiaries(_eoir_case_id)
                    except Exception:
                        st.session_state["_eoir_beneficiaries"] = []
                else:
                    st.session_state["_eoir_beneficiaries"] = []
                st.session_state["_eoir_beneficiaries_case_id"] = _eoir_case_id
            _eoir_beneficiaries: list[dict] = st.session_state.get("_eoir_beneficiaries", [])

            # ── Fetch Case_Contact__c beneficiaries (reactive per case) ─────
            _cached_cb_case = st.session_state.get("_case_bens_case_id", "")
            _force_ben_refresh = st.session_state.pop("_case_bens_force_refresh", False)
            _need_ben_fetch = (_eoir_case_id != _cached_cb_case) or _force_ben_refresh
            if _need_ben_fetch:
                # Clear stale data immediately so UI never shows old case's data
                st.session_state["_case_bens"] = []
                st.session_state["_case_bens_selected"] = []
                st.session_state["_case_bens_case_id"] = _eoir_case_id
                st.session_state["_case_bens_error"] = ""
                if get_case_beneficiaries and _eoir_case_id:
                    # Log which ID we're querying (Record ID vs Case Number)
                    print(f"[EOIR] Fetching Case_Contact__c for Legal_Case__c ID: {_eoir_case_id}")
                    print(f"[EOIR]   (case Name/Number: {_eoir_case.get('Name', 'N/A')})")
                    try:
                        _fetched = get_case_beneficiaries(_eoir_case_id)
                        st.session_state["_case_bens"] = _fetched
                        print(f"[EOIR] Case Beneficiaries result: {len(_fetched)} records")
                        if _fetched:
                            for _fb in _fetched:
                                print(f"[EOIR]   -> {_fb.get('Contact_Name','')} ({_fb.get('Role__c','')})")
                    except Exception as _ben_err:
                        st.session_state["_case_bens"] = []
                        _err_msg = str(_ben_err)
                        st.session_state["_case_bens_error"] = _err_msg
                        print(f"[EOIR] Case Beneficiaries FAILED: {_err_msg}")
                        st.toast(f"Beneficiary fetch error: {_err_msg}", icon="⚠️")
                elif not get_case_beneficiaries:
                    print("[EOIR] get_case_beneficiaries not available (import failed)")
                    st.session_state["_case_bens_error"] = "get_case_beneficiaries not available"
                else:
                    print(f"[EOIR] No case selected, skipping beneficiary fetch")
                # Default: all selected
                st.session_state["_case_bens_selected"] = [
                    b.get("Id") for b in (st.session_state.get("_case_bens") or [])
                ]
            _case_bens: list[dict] = st.session_state.get("_case_bens", [])

            # ── Verify Case Details (collapsed) ──────────────────────────
            # Auto-fill case detail fields on case change
            _prev_cd_fill = st.session_state.get("_eoir_cd_autofill_case_id", "")
            if _eoir_case_id and _eoir_case_id != _prev_cd_fill:
                st.session_state["inp_eoir_cd_city"] = _eoir_case.get("Location_City__c", "") or ""
                st.session_state["inp_eoir_cd_applicant"] = (
                    _eoir_case.get("Primary_Applicant__r_Name", "") or _eoir_sf_client.get("Name", "")
                )
                st.session_state["inp_eoir_cd_next_date_type"] = _eoir_case.get("Type_of_next_date__c", "") or ""
                # Store date as Python date object for st.date_input
                _raw_date = _eoir_case.get("Next_Government_Date__c", "") or ""
                if _raw_date:
                    try:
                        st.session_state["inp_eoir_cd_next_govt_date"] = date.fromisoformat(_raw_date[:10])
                    except Exception:
                        st.session_state["inp_eoir_cd_next_govt_date"] = None
                else:
                    st.session_state["inp_eoir_cd_next_govt_date"] = None
                st.session_state["inp_eoir_cd_judge"] = _eoir_case.get("Immigration_Judge__c", "") or ""
                st.session_state["inp_eoir_cd_attorney"] = _eoir_case.get("Primary_Attorney__r_Name", "") or ""
                st.session_state["inp_eoir_cd_submission_type"] = _eoir_case.get("EOIR_Submission_Type__c", "") or ""
                st.session_state["inp_eoir_cd_submission_line_1"] = _eoir_case.get("EOIR_Submission_Line_1__c", "") or ""
                st.session_state["inp_eoir_cd_filing_method"] = _eoir_case.get("Paper_or_eRop__c", "") or ""
                # A-Number dashed (formula / read-only)
                st.session_state["inp_eoir_cd_a_number"] = (
                    _eoir_case.get("A_number_dashed__c", "") or _eoir_sf_client.get("A_Number__c", "")
                )
                st.session_state["_eoir_cd_autofill_case_id"] = _eoir_case_id
                st.session_state["_eoir_cd_edit_mode"] = False

            # Fetch Legal Case field metadata for picklists + updateable flags (cached)
            _lc_meta: dict = st.session_state.get("_eoir_cd_lc_meta") or {}
            if not _lc_meta and get_legal_case_field_metadata:
                try:
                    _lc_meta = get_legal_case_field_metadata()
                    st.session_state["_eoir_cd_lc_meta"] = _lc_meta
                except Exception:
                    _lc_meta = {}

            # Build picklist options from SF metadata
            def _pl_opts(api_name: str) -> list[str]:
                meta = _lc_meta.get(api_name, {})
                return [""] + [pv["value"] for pv in meta.get("picklistValues", [])]

            _city_opts = _pl_opts("Location_City__c")
            _next_type_opts = _pl_opts("Type_of_next_date__c")
            _judge_opts = _pl_opts("Immigration_Judge__c")
            _submission_type_opts = _pl_opts("EOIR_Submission_Type__c")
            _submission_line_1_opts = _pl_opts("EOIR_Submission_Line_1__c")
            _filing_method_opts = _pl_opts("Paper_or_eRop__c")

            # Load admin-configured visible fields (or defaults)
            _comp_config = load_config("components") or {}
            _CD_DEFAULTS = [
                "Location_City__c", "Type_of_next_date__c", "Next_Government_Date__c",
                "Immigration_Judge__c",
                "EOIR_Submission_Type__c", "EOIR_Submission_Line_1__c", "Paper_or_eRop__c",
                "Primary_Applicant__r_Name", "Primary_Attorney__r_Name", "A_number_dashed__c",
            ]
            _visible_cd_fields: list[str] = _comp_config.get("case_detail_fields", _CD_DEFAULTS)

            # Master field registry: api_key → (label, ss_key, widget, options, sf_api, readonly)
            _ALL_CD_FIELDS = {
                "Location_City__c":          ("City",              "inp_eoir_cd_city",           "picklist", _city_opts,      "Location_City__c",          False),
                "Type_of_next_date__c":      ("Next Date Type",    "inp_eoir_cd_next_date_type", "picklist", _next_type_opts, "Type_of_next_date__c",      False),
                "Next_Government_Date__c":   ("Next Govt Date",    "inp_eoir_cd_next_govt_date", "date",     None,            "Next_Government_Date__c",   False),
                "Immigration_Judge__c":      ("Immigration Judge", "inp_eoir_cd_judge",          "picklist", _judge_opts,     "Immigration_Judge__c",      False),
                "EOIR_Submission_Type__c":   ("EOIR Submission Type", "inp_eoir_cd_submission_type",  "picklist", _submission_type_opts, "EOIR_Submission_Type__c", False),
                "EOIR_Submission_Line_1__c": ("Submission Line 1",    "inp_eoir_cd_submission_line_1", "picklist", _submission_line_1_opts, "EOIR_Submission_Line_1__c", False),
                "Paper_or_eRop__c":          ("Filing Method",        "inp_eoir_cd_filing_method",    "picklist", _filing_method_opts,   "Paper_or_eRop__c",        False),
                "Primary_Applicant__r_Name": ("Primary Applicant", "inp_eoir_cd_applicant",     "text",     None,            None,                        True),
                "Primary_Attorney__r_Name":  ("Primary Attorney",  "inp_eoir_cd_attorney",      "text",     None,            None,                        True),
                "A_number_dashed__c":        ("A-Number (Dashed)", "inp_eoir_cd_a_number",      "text",     None,            None,                        True),
            }

            with st.expander("Verify Case Details", expanded=False):
                _edit_mode = st.session_state.get("_eoir_cd_edit_mode", False)

                # Bulk edit toggle buttons
                if not _edit_mode:
                    if st.button("Edit Case Information", key="_eoir_cd_enter_edit"):
                        st.session_state["_eoir_cd_edit_mode"] = True
                        st.rerun()

                # Build visible field lists
                _vis_editable = [(k, _ALL_CD_FIELDS[k]) for k in _visible_cd_fields
                                 if k in _ALL_CD_FIELDS and not _ALL_CD_FIELDS[k][5]]
                _vis_readonly = [(k, _ALL_CD_FIELDS[k]) for k in _visible_cd_fields
                                 if k in _ALL_CD_FIELDS and _ALL_CD_FIELDS[k][5]]

                _cd_c1, _cd_c2 = st.columns(2)
                _col_toggle = False

                if _edit_mode:
                    # ── Bulk edit mode: all editable fields as inputs ─────
                    for _fkey, (_label, _ss_key, _wtype, _extra, _sf_api, _ro) in _vis_editable:
                        _col = _cd_c1 if not _col_toggle else _cd_c2
                        _col_toggle = not _col_toggle
                        _val = st.session_state.get(_ss_key)
                        with _col:
                            if _wtype == "picklist" and _extra:
                                _cur = _val or ""
                                _idx = _extra.index(_cur) if _cur in _extra else 0
                                st.selectbox(f"✏️ {_label}", options=_extra, index=_idx, key=_ss_key)
                            elif _wtype == "date":
                                if _val is None:
                                    st.session_state[_ss_key] = date.today()
                                st.date_input(f"✏️ {_label}", format="MM/DD/YYYY", key=_ss_key)
                            elif _wtype == "toggle":
                                st.toggle(f"✏️ {_label}", key=_ss_key)
                            else:
                                st.text_input(f"✏️ {_label}", key=_ss_key)

                    # Read-only fields shown as disabled
                    for _fkey, (_label, _ss_key, _wtype, _extra, _sf_api, _ro) in _vis_readonly:
                        _col = _cd_c1 if not _col_toggle else _cd_c2
                        _col_toggle = not _col_toggle
                        with _col:
                            st.text_input(f"🔒 {_label}", key=_ss_key, disabled=True)

                    # Save / Cancel buttons
                    _btn_c1, _btn_c2, _btn_c3 = st.columns([2, 2, 4])
                    with _btn_c1:
                        _do_save = st.button("Save", type="primary", key="_eoir_cd_save")
                    with _btn_c2:
                        _do_cancel = st.button("Cancel", key="_eoir_cd_cancel")

                    if _do_cancel:
                        # Revert to SF values by re-triggering autofill
                        st.session_state["_eoir_cd_autofill_case_id"] = ""
                        st.session_state["_eoir_cd_edit_mode"] = False
                        st.rerun()

                    if _do_save and update_legal_case and _eoir_case_id:
                        # Build updates from visible editable fields only
                        # Use metadata to verify each field is updateable in SF
                        updates: dict = {}
                        for _fkey, (_label, _ss_key, _wtype, _extra, _sf_api, _ro) in _vis_editable:
                            if not _sf_api:
                                continue  # derived field, not in SF
                            # Check metadata: only push if field exists and is updateable
                            _fmeta = _lc_meta.get(_sf_api, {})
                            if _lc_meta and not _fmeta:
                                continue  # field not in SF org — store locally only
                            if _fmeta and not _fmeta.get("updateable", False):
                                continue  # skip formula / non-updateable
                            if _fmeta and _fmeta.get("formula", False):
                                continue
                            val = st.session_state.get(_ss_key)
                            if _wtype == "date":
                                val = val.strftime("%Y-%m-%d") if val else None
                            updates[_sf_api] = val
                        updates = {k: v for k, v in updates.items() if v is not None}

                        with st.spinner("Syncing to Salesforce..."):
                            try:
                                update_legal_case(_eoir_case_id, updates)
                                st.session_state["_eoir_cd_edit_mode"] = False
                                st.toast("Case details synced to Salesforce!")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Sync failed: {e}")

                else:
                    # ── View mode: compact display ───────────────────────
                    for _fkey, (_label, _ss_key, _wtype, _extra, _sf_api, _ro) in _vis_editable:
                        _col = _cd_c1 if not _col_toggle else _cd_c2
                        _col_toggle = not _col_toggle
                        _val = st.session_state.get(_ss_key)
                        if _wtype == "date":
                            _disp = _val.strftime("%m/%d/%Y") if _val else "—"
                        elif _wtype == "toggle":
                            _disp = "Yes" if _val else "No"
                        else:
                            _disp = str(_val) if _val else "—"
                        with _col:
                            st.markdown(f"✏️ **{_label}:** {_disp}")

                    for _fkey, (_label, _ss_key, _wtype, _extra, _sf_api, _ro) in _vis_readonly:
                        _col = _cd_c1 if not _col_toggle else _cd_c2
                        _col_toggle = not _col_toggle
                        _val = st.session_state.get(_ss_key, "") or "—"
                        with _col:
                            st.markdown(f"🔒 **{_label}:** {_val}")

            st.divider()

            # ── Case Beneficiaries (Case_Contact__c) ─────────────────
            _ben_hdr_cols = st.columns([5, 1, 1])
            with _ben_hdr_cols[0]:
                st.markdown("**Case Beneficiaries**")
            with _ben_hdr_cols[1]:
                if st.button("🔄", key="_case_bens_refresh_btn", help="Refresh beneficiaries from Salesforce"):
                    # Explicitly clear all local beneficiary state
                    st.session_state["_case_bens"] = []
                    st.session_state["_case_bens_selected"] = []
                    st.session_state["_case_bens_error"] = ""
                    st.session_state["_case_bens_force_refresh"] = True
                    st.rerun()
            with _ben_hdr_cols[2]:
                if st.button("🔍", key="_case_bens_describe_btn", help="Describe Case_Contact__c schema"):
                    if describe_case_contact_relationships:
                        try:
                            _refs = describe_case_contact_relationships()
                            print(f"[EOIR] Case_Contact__c lookup fields to Legal_Case__c: {_refs}")
                            if _refs:
                                _ref_info = "; ".join(f"{r['name']} ({r['label']})" for r in _refs)
                                st.toast(f"Lookup fields: {_ref_info}", icon="🔍")
                            else:
                                st.toast("No lookup fields to Legal_Case__c found on Case_Contact__c", icon="⚠️")
                        except Exception as _desc_err:
                            st.toast(f"Schema describe error: {_desc_err}", icon="⚠️")
                            print(f"[EOIR] Describe error: {_desc_err}")
                    else:
                        st.toast("describe_case_contact_relationships not available", icon="⚠️")
            # Show error from last fetch
            _ben_error = st.session_state.get("_case_bens_error", "")
            if _ben_error:
                st.error(f"Last fetch error: {_ben_error}")
            elif not _case_bens and _eoir_case_id:
                st.caption("No beneficiaries found for this case.")
            if _case_bens:
                _sel_ids: list[str] = st.session_state.get("_case_bens_selected", [])

                # Select All toggle
                _all_ids = [b.get("Id") for b in _case_bens]
                _all_selected = all(bid in _sel_ids for bid in _all_ids)
                if st.checkbox("Select All", value=_all_selected, key="_case_bens_sel_all"):
                    if not _all_selected:
                        st.session_state["_case_bens_selected"] = list(_all_ids)
                        st.rerun()
                elif _all_selected:
                    st.session_state["_case_bens_selected"] = []
                    st.rerun()

                # Individual checkboxes
                _new_sel: list[str] = []
                for _bi, _ben in enumerate(_case_bens):
                    _bid = _ben.get("Id", "")
                    _bname = _ben.get("Contact_Name", "") or "Unnamed"
                    _brel = _ben.get("Role__c", "")
                    _banum = _ben.get("Alien_Number_Dashed__c", "")
                    _bdob = _ben.get("DOB__c", "")
                    _bline = _bname
                    if _brel:
                        _bline += f" ({_brel})"
                    if _banum:
                        _bline += f" — A# {_banum}"
                    _checked = st.checkbox(_bline, value=_bid in _sel_ids, key=f"_cben_chk_{_bi}")
                    if _checked:
                        _new_sel.append(_bid)
                st.session_state["_case_bens_selected"] = _new_sel

                # ── Derivatives detail (expandable with pencil edits) ──
                _selected_bens = [b for b in _case_bens if b.get("Id") in _new_sel]
                if _selected_bens:
                    with st.expander(f"Derivatives ({len(_selected_bens)})", expanded=False):
                        for _di, _db in enumerate(_selected_bens):
                            _db_id = _db.get("Id", "")
                            _db_name = _db.get("Contact_Name", "") or "Unnamed"
                            _db_rel = _db.get("Role__c", "")
                            _db_anum = _db.get("Alien_Number_Dashed__c", "")
                            _db_dob = _db.get("DOB__c", "")
                            if _db_dob:
                                try:
                                    _db_dob = date.fromisoformat(_db_dob[:10]).strftime("%m/%d/%Y")
                                except Exception:
                                    pass

                            st.markdown(f"**{_db_name}**")

                            # Relationship — pencil editable
                            _rel_edit_key = f"_cben_edit_rel_{_di}"
                            if st.session_state.get(_rel_edit_key, False):
                                _new_rel = st.text_input(
                                    "✏️ Relationship", value=_db_rel,
                                    key=f"_cben_rel_val_{_di}",
                                )
                                if st.button("Done", key=f"_cben_rel_done_{_di}"):
                                    # Update in cached data
                                    _db["Role__c"] = _new_rel
                                    st.session_state[_rel_edit_key] = False
                                    st.rerun()
                            else:
                                _r_text, _r_btn = st.columns([5, 1])
                                with _r_text:
                                    st.caption(f"Relationship: {_db_rel or '—'}")
                                with _r_btn:
                                    if st.button("✏️", key=f"_cben_rel_pen_{_di}", help="Edit Relationship"):
                                        st.session_state[_rel_edit_key] = True
                                        st.rerun()

                            # A-Number — pencil editable
                            _anum_edit_key = f"_cben_edit_anum_{_di}"
                            if st.session_state.get(_anum_edit_key, False):
                                _new_anum = st.text_input(
                                    "✏️ A-Number", value=_db_anum,
                                    key=f"_cben_anum_val_{_di}",
                                )
                                if st.button("Done", key=f"_cben_anum_done_{_di}"):
                                    _db["Alien_Number_Dashed__c"] = _new_anum
                                    st.session_state[_anum_edit_key] = False
                                    st.rerun()
                            else:
                                _a_text, _a_btn = st.columns([5, 1])
                                with _a_text:
                                    st.caption(f"A-Number: {_db_anum or '—'}")
                                with _a_btn:
                                    if st.button("✏️", key=f"_cben_anum_pen_{_di}", help="Edit A-Number"):
                                        st.session_state[_anum_edit_key] = True
                                        st.rerun()

                            st.caption(f"DOB: {_db_dob or '—'}")

                            if _di < len(_selected_bens) - 1:
                                st.markdown("---")

                        # Sync edits back to Salesforce
                        if update_case_beneficiary:
                            if st.button("Sync Beneficiary Edits to Salesforce", key="_cben_sync"):
                                _sync_count = 0
                                with st.spinner("Syncing beneficiary edits..."):
                                    for _sb in _selected_bens:
                                        _sb_id = _sb.get("Id", "")
                                        if not _sb_id:
                                            continue
                                        _sb_updates = {
                                            "Alien_Number_Dashed__c": _sb.get("Alien_Number_Dashed__c", ""),
                                        }
                                        try:
                                            update_case_beneficiary(_sb_id, _sb_updates)
                                            _sync_count += 1
                                        except Exception as e:
                                            st.error(f"Failed to sync {_sb.get('Contact_Name', '')}: {e}")
                                if _sync_count:
                                    st.toast(f"Synced {_sync_count} beneficiar{'ies' if _sync_count != 1 else 'y'} to Salesforce!")
            else:
                st.caption("No derivatives found for this case.")

            st.divider()

            # ── Submission Mode Toggle ──────────────────────────────────
            _eoir_mode = st.radio(
                "Submission Mode",
                ["Document Submission", "Brief Builder"],
                key="inp_eoir_mode",
                horizontal=True,
                label_visibility="collapsed",
            )

            # Track brief preview HTML for right column
            _eoir_brief_preview_html = ""

            if _eoir_mode == "Document Submission":
                # ── Submission Details ───────────────────────────────────
                st.markdown('<div class="section-label">Submission Details</div>', unsafe_allow_html=True)

                # Caption fields — stacked vertically for the right-hand caption box
                _eoir_sub_line_1_val = st.session_state.get("inp_eoir_cd_submission_line_1", "") or ""
                _sub_line_1_display_opts = _submission_line_1_opts if len(_submission_line_1_opts) > 1 else [""]
                _sub_line_1_idx = _sub_line_1_display_opts.index(_eoir_sub_line_1_val) if _eoir_sub_line_1_val in _sub_line_1_display_opts else 0
                _eoir_submission_line_1 = st.selectbox(
                    "Submission Line 1 (SF Picklist)",
                    options=_sub_line_1_display_opts,
                    index=_sub_line_1_idx,
                    key="inp_eoir_submission_line_1",
                    help="Pulled from EOIR_Submission_Line_1__c in Salesforce",
                )
                _eoir_submission_sub_line = st.text_input(
                    "Submission Sub-line / Motion Title",
                    key="inp_eoir_submission_sub_line",
                    placeholder="e.g. Motion to Change Venue",
                )

                _eoir_submission_type = st.text_area(
                    "Submission Description",
                    key="inp_eoir_submission_type",
                    height=80,
                    placeholder="e.g. Respondent's Pre-Hearing Brief and Supporting Evidence",
                )

                # Show enclosed documents summary (shared with Cover Letter tab)
                _eoir_enclosed = _build_enclosed_docs_list()
                if _eoir_enclosed:
                    st.markdown(f"**Enclosed Documents** ({len(_eoir_enclosed)}):")
                    for _idx, _doc in enumerate(_eoir_enclosed, start=1):
                        _doc_label = _doc.get("name", "")
                        _doc_desc = _doc.get("description", "")
                        if _doc_desc:
                            st.caption(f"{_idx}. {_doc_label} — {_doc_desc}")
                        else:
                            st.caption(f"{_idx}. {_doc_label}")
                else:
                    st.caption("No enclosed documents. Add documents in the Cover Letter tab.")

                st.divider()

                # ── Certificate of Service ───────────────────────────────
                st.markdown('<div class="section-label">Certificate of Service</div>', unsafe_allow_html=True)
                _eoir_dhs_addr = st.text_area(
                    "DHS / OCC Address",
                    key="inp_eoir_dhs_address",
                    height=80,
                    placeholder="Office of the Chief Counsel address",
                )
                _eoir_service_method = st.selectbox(
                    "Service Method",
                    options=["first-class mail", "hand delivery", "electronic filing (ECAS)"],
                    key="inp_eoir_service_method",
                )

                # Served By — pull from Staff Directory
                _served_by_names = [""] + [
                    f"{m.get('first_name', '')} {m.get('last_name', '')}".strip()
                    for m in _staff_dir
                ]
                _eoir_served_by = st.selectbox(
                    "Served By",
                    options=_served_by_names,
                    key="inp_eoir_served_by",
                )
                _served_by_staff = next(
                    (m for m in _staff_dir
                     if f"{m.get('first_name', '')} {m.get('last_name', '')}".strip() == _eoir_served_by),
                    None,
                )
                _served_by_name = _eoir_served_by
                _served_by_bar = _served_by_staff.get("bar_number", "") if _served_by_staff else ""

                # ── Render EOIR text via template engine ─────────────────
                # Map selected Case_Contact__c beneficiaries to render format
                _sel_ben_ids = st.session_state.get("_case_bens_selected", [])
                _render_bens = []
                for _cb in _case_bens:
                    if _cb.get("Id") in _sel_ben_ids:
                        _render_bens.append({
                            "Name": _cb.get("Contact_Name", "") or "Unnamed",
                            "A_Number": _cb.get("Alien_Number_Dashed__c", ""),
                            "Type": _cb.get("Role__c", ""),
                        })
                # Fall back to LC_Contact__c if no Case_Contact__c data
                _final_bens = _render_bens if _case_bens else _eoir_beneficiaries

                # Defaults for canvas (set before conditional so right column always has them)
                _eoir_blocks: list[dict] = []
                _reset_canvas = False
                _canvas_order: list[str] = []
                _canvas_edits: dict = {}
                _block_map: dict = {}

                if _eoir_submission_type.strip():
                    # Source court/party data from Verify Case Details session state
                    _render_court_loc = st.session_state.get("inp_eoir_cd_city", "") or ""
                    _render_court_addr = _eoir_case.get("Address_of_next_hearing__c", "") or ""
                    _render_applicant = st.session_state.get("inp_eoir_cd_applicant", "") or (
                        _eoir_sf_client.get("Name", "")
                    )
                    _render_a_number = st.session_state.get("inp_eoir_cd_a_number", "") or ""
                    _render_case_type = _eoir_case.get("Legal_Case_Type__c", "")

                    # Compose effective template — merge document template
                    # blocks (firm header, court caption) into the EOIR
                    # cover template if admin has customized them.
                    _eff_template = _eoir_cover_template  # admin EOIR template or None
                    _fh_block = _doc_tpl_config.get("firm_header_template")
                    _cc_block = _doc_tpl_config.get("court_caption_template")
                    if _fh_block or _cc_block:
                        from app.templates import DEFAULT_EOIR_COVER_TEMPLATE as _DEF_TPL
                        _base = _eff_template or _DEF_TPL
                        if _fh_block:
                            # Replace default firm header section
                            _def_fh = (
                                "{attorney_name}\n{bar_number}\n{firm_name}\n"
                                "{firm_address}\n{contact_line}\n{email_line}"
                            )
                            if _def_fh in _base:
                                _base = _base.replace(_def_fh, _fh_block.rstrip())
                        if _cc_block:
                            _def_cc = (
                                "UNITED STATES DEPARTMENT OF JUSTICE\n"
                                "EXECUTIVE OFFICE FOR IMMIGRATION REVIEW\n"
                                "IMMIGRATION COURT\n{court_location}\n{court_address}"
                            )
                            if _def_cc in _base:
                                _base = _base.replace(_def_cc, _cc_block.rstrip())
                        _eff_template = _base

                    _eoir_rendered = render_eoir_from_template(
                        attorney_name=attorney_name,
                        bar_number=bar_number,
                        firm_name=firm_name,
                        firm_address=firm_address,
                        firm_phone=_gs_global.get("firm_phone", ""),
                        firm_fax=_gs_global.get("firm_fax", ""),
                        firm_email=_gs_global.get("firm_email", ""),
                        court_location=_render_court_loc,
                        court_address=_render_court_addr,
                        applicant_name=_render_applicant,
                        a_number=_render_a_number,
                        case_type=_render_case_type,
                        beneficiaries=_final_bens,
                        submission_type=_eoir_submission_type.strip(),
                        submission_line_1=_eoir_submission_line_1.strip() if _eoir_submission_line_1 else "",
                        submission_sub_line=_eoir_submission_sub_line.strip() if _eoir_submission_sub_line else "",
                        document_list=_eoir_enclosed,
                        dhs_address=_eoir_dhs_addr,
                        service_method=_eoir_service_method,
                        served_by_name=_served_by_name,
                        served_by_bar=_served_by_bar,
                        template_override=_eff_template,
                    )

                    # Split rendered text into blocks for the interactive canvas
                    _eoir_blocks = split_eoir_into_blocks(_eoir_rendered)

                    # Fingerprint for reset detection — when data changes,
                    # reset the canvas blocks; otherwise preserve user edits.
                    _eoir_case_id = _eoir_case.get("Id", "")
                    _eoir_fp = (
                        _eoir_case_id,
                        _eoir_submission_type.strip(),
                        _eoir_submission_line_1,
                        _eoir_submission_sub_line,
                        attorney_name,
                        _eoir_service_method,
                        _served_by_name,
                        len(_eoir_enclosed),
                        len(_final_bens),
                    )
                    _prev_fp = st.session_state.get("_eoir_draft_fingerprint")
                    _force_regen = st.session_state.pop("_eoir_force_regen", False)
                    _reset_canvas = (_prev_fp != _eoir_fp) or _force_regen
                    if _reset_canvas:
                        st.session_state["_eoir_draft_fingerprint"] = _eoir_fp
                        st.session_state["_eoir_canvas_order"] = [b["id"] for b in _eoir_blocks]
                        st.session_state["_eoir_canvas_edits"] = {}

                    # Assemble letter_text from blocks + user edits (Python-side
                    # so it's always up-to-date even before canvas emits)
                    _canvas_order = st.session_state.get(
                        "_eoir_canvas_order", [b["id"] for b in _eoir_blocks],
                    )
                    _canvas_edits = st.session_state.get("_eoir_canvas_edits", {})
                    _block_map = {b["id"]: b["content"] for b in _eoir_blocks}
                    letter_text = "\n\n".join(
                        _canvas_edits.get(bid, _block_map.get(bid, ""))
                        for bid in _canvas_order
                        if bid in _block_map or bid in _canvas_edits
                    )
                    st.session_state["_eoir_letter_text"] = letter_text

            else:
                # ── Brief Builder Mode ──────────────────────────────────
                st.markdown('<div class="section-label">Brief Builder</div>', unsafe_allow_html=True)

                # Case-ID-based state persistence
                _prev_bb_case = st.session_state.get("_eoir_bb_case_id", "")
                if _eoir_case_id and _eoir_case_id != _prev_bb_case:
                    if _prev_bb_case:
                        _bb_save_state(_prev_bb_case)
                    _bb_restore_state(_eoir_case_id)
                    st.session_state["_eoir_bb_case_id"] = _eoir_case_id

                # Brief Type selector
                _bb_types = list(get_config_value("brief-builder", "brief_types", {}).keys())
                if not _bb_types:
                    st.warning("No brief types configured. Configure them in the Brief Builder tool.")
                else:
                    # Detect brief type change to clear section content
                    _prev_bb_type = st.session_state.get("_eoir_bb_prev_type", "")
                    _bb_type = st.selectbox("Brief Type", _bb_types, key="inp_eoir_bb_type")
                    if _prev_bb_type and _prev_bb_type != _bb_type:
                        _bb_clear_content_keys()
                    st.session_state["_eoir_bb_prev_type"] = _bb_type

                    # Submission Description (shared key with Document Submission mode)
                    _eoir_submission_type = st.text_area(
                        "Submission Description",
                        key="inp_eoir_submission_type",
                        height=80,
                        placeholder="e.g. Respondent's Pre-Hearing Brief and Supporting Evidence",
                    )

                    # Load sections for current brief type
                    _bb_sections = _load_brief_sections(_bb_type)

                    # Dynamic section list
                    for _si, _section in enumerate(_bb_sections):
                        _s_key = _section["key"]
                        _s_heading = _section["heading"]
                        _s_bp = _section.get("boilerplate", "")
                        _s_subs = _section.get("subsections", [])

                        with st.expander(_s_heading, expanded=False):
                            # Checkbox to toggle inclusion
                            st.checkbox(
                                f"Include {_s_heading}",
                                value=True,
                                key=f"_bb_chk_{_s_key}",
                            )

                            # Section-level boilerplate
                            if _s_bp:
                                st.markdown(
                                    f'<div class="boilerplate-block">'
                                    f"<strong>Standard language:</strong><br>"
                                    f"{html_mod.escape(_s_bp)}"
                                    f"</div>",
                                    unsafe_allow_html=True,
                                )
                                _ck = _bb_content_key(_s_key)
                                if st.button("Insert Boilerplate", key=f"_bb_ins_{_s_key}"):
                                    st.session_state[_ck] = _s_bp
                                    st.rerun()

                            if _s_subs:
                                for _sub in _s_subs:
                                    _sub_key = _sub["key"]
                                    _sub_heading = _sub["heading"]
                                    _sub_bp = _sub.get("boilerplate", "")

                                    st.markdown(f"**{_sub_heading}**")

                                    if _sub_bp:
                                        st.markdown(
                                            f'<div class="boilerplate-block">'
                                            f"<strong>Standard language:</strong><br>"
                                            f"{html_mod.escape(_sub_bp)}"
                                            f"</div>",
                                            unsafe_allow_html=True,
                                        )
                                        _ck = _bb_content_key(_s_key, _sub_key)
                                        if st.button("Insert Boilerplate", key=f"_bb_ins_{_s_key}_{_sub_key}"):
                                            st.session_state[_ck] = _sub_bp
                                            st.rerun()

                                    _ck = _bb_content_key(_s_key, _sub_key)
                                    st.text_area(
                                        f"Content for {_sub_heading}",
                                        height=150,
                                        key=_ck,
                                        placeholder=f"Draft your {_sub_heading.lower()} argument here...",
                                        label_visibility="collapsed",
                                    )
                            else:
                                _ck = _bb_content_key(_s_key)
                                _height = 200 if _s_key in ("statement_of_facts", "country_conditions") else 150
                                st.text_area(
                                    f"Content for {_s_heading}",
                                    height=_height,
                                    key=_ck,
                                    placeholder=f"Draft your {_s_heading.lower()} here...",
                                    label_visibility="collapsed",
                                )

                    # Build variable map for resolving {{placeholders}}
                    _hearing_date_raw = _eoir_case.get("Next_Government_Date__c", "") or ""
                    if _hearing_date_raw:
                        try:
                            _hearing_date_fmt = date.fromisoformat(_hearing_date_raw[:10]).strftime("%m/%d/%Y")
                        except Exception:
                            _hearing_date_fmt = _hearing_date_raw
                    else:
                        _hearing_date_fmt = ""

                    _bb_var_map = {
                        "client_name": _eoir_sf_client.get("Name", ""),
                        "A_number": st.session_state.get("inp_eoir_cd_a_number", ""),
                        "court": st.session_state.get("inp_eoir_cd_city", ""),
                        "judge": st.session_state.get("inp_eoir_cd_judge", ""),
                        "hearing_date": _hearing_date_fmt,
                        "country": _eoir_sf_client.get("Country__c", ""),
                        "attorney_name": attorney_name,
                        "case_type": _eoir_case.get("Legal_Case_Type__c", ""),
                    }

                    # Collect section content with variables resolved
                    _bb_sections_content: dict[str, str] = {}
                    _bb_checked: set[str] = set()
                    for _section in _bb_sections:
                        _s_key = _section["key"]
                        if st.session_state.get(f"_bb_chk_{_s_key}", True):
                            _bb_checked.add(_s_key)
                        _s_subs = _section.get("subsections", [])
                        if _s_subs:
                            for _sub in _s_subs:
                                _ck = _bb_content_key(_s_key, _sub["key"])
                                _raw = st.session_state.get(_ck, "")
                                _bb_sections_content[_ck] = _resolve_variables(_raw, _bb_var_map)
                        else:
                            _ck = _bb_content_key(_s_key)
                            _raw = st.session_state.get(_ck, "")
                            _bb_sections_content[_ck] = _resolve_variables(_raw, _bb_var_map)

                    # Build case info dict for preview
                    _bb_case_info = {
                        "client_name": _bb_var_map["client_name"],
                        "a_number": _bb_var_map["A_number"],
                        "court": _bb_var_map["court"],
                        "judge": _bb_var_map["judge"],
                        "hearing_date": _bb_var_map["hearing_date"],
                    }

                    # Build brief preview HTML
                    _eoir_brief_preview_html = _build_brief_preview_html(
                        brief_type=_bb_type,
                        case_info=_bb_case_info,
                        sections=_bb_sections,
                        sections_content=_bb_sections_content,
                        submission_desc=_resolve_variables(
                            st.session_state.get("inp_eoir_submission_type", ""), _bb_var_map
                        ),
                        checked_sections=_bb_checked,
                        attorney_name_val=attorney_name,
                    )

                    # Build plain-text letter_text for export
                    _bb_text_parts: list[str] = []
                    _bb_text_parts.append(_BB_TITLE_MAP.get(_bb_type, _bb_type.upper()))
                    _bb_text_parts.append("")
                    if _bb_case_info["client_name"]:
                        _bb_text_parts.append(f"IN THE MATTER OF: {_bb_case_info['client_name'].upper()}")
                    if _bb_case_info["a_number"]:
                        _bb_text_parts.append(f"A-Number: {_bb_case_info['a_number']}")
                    if _bb_case_info["court"]:
                        _bb_text_parts.append(f"Before the {_bb_case_info['court']}")
                    if _bb_case_info["judge"]:
                        _bb_text_parts.append(f"Immigration Judge {_bb_case_info['judge']}")
                    if _bb_case_info["hearing_date"]:
                        _bb_text_parts.append(f"Hearing Date: {_bb_case_info['hearing_date']}")
                    _bb_text_parts.append("")

                    _sub_desc = _resolve_variables(
                        st.session_state.get("inp_eoir_submission_type", ""), _bb_var_map
                    ).strip()
                    if _sub_desc:
                        _bb_text_parts.append(_sub_desc)
                        _bb_text_parts.append("")

                    for _si, _section in enumerate(_bb_sections):
                        _s_key = _section["key"]
                        if _s_key not in _bb_checked:
                            continue
                        _roman = _ROMAN[_si] if _si < len(_ROMAN) else str(_si + 1)
                        _bb_text_parts.append(f"{_roman}. {_section['heading']}")
                        _bb_text_parts.append("")
                        _s_subs = _section.get("subsections", [])
                        if _s_subs:
                            for _sub_idx, _sub in enumerate(_s_subs):
                                _sub_letter = chr(ord("A") + _sub_idx)
                                _bb_text_parts.append(f"    {_sub_letter}. {_sub['heading']}")
                                _ck = _bb_content_key(_s_key, _sub["key"])
                                _body = _bb_sections_content.get(_ck, "").strip()
                                if _body:
                                    _bb_text_parts.append("")
                                    _bb_text_parts.append(_body)
                                _bb_text_parts.append("")
                        else:
                            _ck = _bb_content_key(_s_key)
                            _body = _bb_sections_content.get(_ck, "").strip()
                            if _body:
                                _bb_text_parts.append(_body)
                            _bb_text_parts.append("")

                    _bb_text_parts.append("")
                    _bb_text_parts.append("Respectfully submitted,")
                    _bb_text_parts.append("")
                    _bb_text_parts.append("____________________________")
                    if attorney_name:
                        _bb_text_parts.append(attorney_name)
                    _bb_text_parts.append("Attorney for Respondent")
                    _bb_text_parts.append(f"Date: {date.today().strftime('%m/%d/%Y')}")

                    letter_text = "\n".join(_bb_text_parts)
                    st.session_state["_eoir_letter_text"] = letter_text

    # ── EOIR Right Column (Canvas or Brief Preview) ──────────────────────
    with _eoir_right:
        _eoir_mode_val = st.session_state.get("inp_eoir_mode", "Document Submission")

        if _eoir_mode_val == "Brief Builder":
            # Brief Builder preview
            if _eoir_brief_preview_html:
                st.markdown(
                    f'<div class="preview-panel" style="white-space:normal;">'
                    f'{_eoir_brief_preview_html}</div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    '<div class="preview-panel" style="min-height:200px;display:flex;align-items:center;'
                    'justify-content:center;color:#94a3b8;font-style:italic;">'
                    'Select a brief type and add content to see the preview.'
                    '</div>',
                    unsafe_allow_html=True,
                )
        else:
            # Document Submission — existing canvas logic
            # Expand dialog trigger
            if st.session_state.get("_eoir_show_expand"):
                st.session_state["_eoir_show_expand"] = False
                _open_eoir_expand()

            if _eoir_blocks:
                _canvas_out = _eoir_canvas_component(
                    blocks=_eoir_blocks,
                    block_order=_canvas_order,
                    block_edits=_canvas_edits,
                    reset=_reset_canvas,
                    key="eoir_canvas",
                    default=None,
                )
                # Process canvas output — update session state with user edits/order
                if _canvas_out is not None:
                    _action = _canvas_out.get("action")
                    if _action == "expand":
                        st.session_state["_eoir_show_expand"] = True
                        st.rerun()
                    elif _action == "regenerate":
                        st.session_state["_eoir_force_regen"] = True
                        st.rerun()
                    else:
                        _new_order = _canvas_out.get("block_order")
                        _new_edits = _canvas_out.get("block_edits")
                        if _new_order is not None:
                            st.session_state["_eoir_canvas_order"] = _new_order
                        if _new_edits is not None:
                            st.session_state["_eoir_canvas_edits"] = _new_edits
                        # Reassemble letter_text with latest edits
                        _co = _new_order or _canvas_order
                        _ce = _new_edits if _new_edits is not None else _canvas_edits
                        letter_text = "\n\n".join(
                            _ce.get(bid, _block_map.get(bid, ""))
                            for bid in _co
                            if bid in _block_map or bid in _ce
                        )
                        st.session_state["_eoir_letter_text"] = letter_text
            else:
                st.markdown(
                    '<div class="preview-panel" style="min-height:200px;display:flex;align-items:center;'
                    'justify-content:center;color:#94a3b8;font-style:italic;">'
                    'Fill in the submission description to see the document canvas.'
                    '</div>',
                    unsafe_allow_html=True,
                )

# ═══════════════════════════════════════════════════════════════════════════════
# PERSISTENT FOOTER — Attorney / Signature + Export
# ═══════════════════════════════════════════════════════════════════════════════
st.divider()
_footer_left, _footer_right = st.columns([3, 2], gap="large")

with _footer_left:
    # ── Attorney / Signature Block ───────────────────────────────────────────
    st.markdown('<div class="section-label">Attorney / Signature</div>', unsafe_allow_html=True)
    _staff_names = [
        f"{m.get('first_name', '')} {m.get('last_name', '')}".strip()
        for m in _staff_dir
    ]
    _staff_options = [""] + _staff_names
    _cur_sel = st.session_state.get("inp_attorney_staff", "")
    _sel_idx = _staff_options.index(_cur_sel) if _cur_sel in _staff_options else 0
    _selected_staff = st.selectbox(
        "Attorney",
        options=_staff_options,
        index=_sel_idx,
        format_func=lambda x: x if x else "Select attorney...",
        key="inp_attorney_staff",
    )
    # Update attorney_name and bar_number to match widget value
    _matched = next(
        (m for m in _staff_dir
         if f"{m.get('first_name', '')} {m.get('last_name', '')}".strip() == _selected_staff),
        None,
    )
    attorney_name = _selected_staff
    bar_number = _matched.get("bar_number", "") if _matched else ""
    if bar_number:
        st.caption(f"Bar #: {bar_number}")
    firm_name = _gs_global.get("firm_name", "O'Brien Immigration Law")
    firm_address = _gs_global.get("firm_address", "")
    st.caption(f"{firm_name}")
    if firm_address:
        st.caption(firm_address.replace("\n", "  \n"))

with _footer_right:
    if letter_text:
        st.markdown("---")
        st.markdown('<div class="section-label">Export</div>', unsafe_allow_html=True)

        # File name input — smart default, no explicit value (Streamlit pitfall)
        _is_eoir_content = "EXECUTIVE OFFICE FOR IMMIGRATION REVIEW" in letter_text
        _is_brief_content = st.session_state.get("inp_eoir_mode") == "Brief Builder" and letter_text.strip()
        _doc_type_label = "Brief" if _is_brief_content else ("EOIR Submission" if _is_eoir_content else "Cover Letter")
        _eoir_applicant = st.session_state.get("inp_eoir_court_location", "")
        _file_name_client = client_name
        # For EOIR/Brief, use the applicant name from the case if available
        if (_is_eoir_content or _is_brief_content) and _eoir_sf_client:
            _eoir_case_for_name = _eoir_legal_cases[st.session_state.get("inp_eoir_legal_case_idx", 0)] if _eoir_legal_cases else {}
            _eoir_app_name = _eoir_case_for_name.get("Primary_Applicant__r_Name", "") or _eoir_sf_client.get("Name", "")
            if _eoir_app_name:
                _file_name_client = _eoir_app_name
        _default_file_name = f"{date.today().isoformat()} - {_file_name_client} - {_doc_type_label}"
        if "inp_export_file_name" not in st.session_state:
            st.session_state["inp_export_file_name"] = ""
        _export_input = st.text_input(
            "File name",
            key="inp_export_file_name",
            placeholder=_default_file_name,
            label_visibility="collapsed",
        )
        export_file_name = _export_input.strip() if _export_input.strip() else _default_file_name

        # Three-column download row
        exp_cols = st.columns(3)
        with exp_cols[0]:
            st.download_button(
                "Download .txt",
                data=letter_text,
                file_name=f"{export_file_name}.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with exp_cols[1]:
            docx_bytes = _build_docx(letter_text, attorney_name, firm_name)
            st.download_button(
                "Download .docx",
                data=docx_bytes,
                file_name=f"{export_file_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with exp_cols[2]:
            pdf_bytes = _build_cover_letter_pdf(letter_text)
            st.download_button(
                "Download .pdf",
                data=pdf_bytes,
                file_name=f"{export_file_name}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

        # Export to Google Doc
        # Pick template based on context (Cover Letter vs EOIR)
        _gd_folder = _gs_global.get("google_drive_folder_id", "")
        if _is_eoir_content or _is_brief_content:
            # EOIR/Brief content is fully rendered — prefer direct DOCX upload
            # Only use template if explicitly configured; don't fall back to cover letter template
            _gd_template = _gs_global.get("eoir_template_id", "")
        else:
            _gd_template = _gs_global.get("cover_letter_template_id", "")
        _sf_contact_id_export = (_sf_client.get("Id", "") if _sf_client else "")

        if _sf_contact_id_export and _gd_template and render_google_doc_button and not _is_eoir_content and not _is_brief_content:
            # Full export: Google Doc (from template) + Salesforce Google_Doc__c record
            _selected_lc = _sf_client.get("selected_legal_case") if _sf_client else None
            _sf_legal_case_id = _selected_lc.get("Id") if _selected_lc else (
                _sf_client.get("Legal_Case__c") if _sf_client else None
            )
            render_google_doc_button(
                template_id=_gd_template,
                file_name=export_file_name,
                contact_id=_sf_contact_id_export,
                legal_case_id=_sf_legal_case_id,
                replacements={"{{CONTENT}}": letter_text},
                folder_id=_gd_folder,
                key="gdoc_export_cl",
            )
        else:
            # Direct DOCX upload (no template copy needed — used for EOIR
            # where content is fully rendered, or when template ID is missing)
            if st.button("Upload to Google Docs", use_container_width=True):
                with st.spinner("Uploading to Google Docs..."):
                    try:
                        url = upload_to_google_docs(
                            docx_bytes, export_file_name, folder_id=_gd_folder
                        )
                        st.session_state.google_doc_url = url
                        st.toast("Google Doc created!", icon="✅")
                    except Exception as e:
                        st.error(f"Upload failed: {e}")
            if st.session_state.get("google_doc_url"):
                st.markdown(f"[Open Google Doc]({st.session_state.google_doc_url})")

        # Save to Salesforce (DOCX file upload to Contact Files)
        if _sf_contact_id_export and upload_file_to_contact:
            if st.button("Save to Salesforce", use_container_width=True):
                with st.spinner("Saving to Salesforce..."):
                    try:
                        upload_file_to_contact(
                            contact_sf_id=_sf_contact_id_export,
                            file_bytes=docx_bytes,
                            file_name=export_file_name,
                            file_extension="docx",
                            title=export_file_name,
                        )
                        st.success("Saved to Salesforce.")
                    except Exception as e:
                        st.error(f"Salesforce upload failed: {e}")

        if not _sf_contact_id_export:
            st.caption("Pull a client to enable Salesforce features.")


# -- Handle save (after all widgets render) -----------------------------------

if save_clicked:
    _do_save(case_type)
    st.rerun()

# -- Draft Box (below everything) --------------------------------------------
if render_draft_box is not None:
    _all_enc = _build_enclosed_docs_list()
    _draft_content = ""
    if client_name:
        _db_tmpl_sel = st.session_state.get("inp_template_selection", "Default (from Case Type)")
        _db_is_default = _db_tmpl_sel == "Default (from Case Type)"
        _db_body = st.session_state.get("inp_letter_body", "")
        _db_subject = st.session_state.get("inp_letter_subject", "")

        if _db_is_default:
            _db_split = _db_body.split("\n\n", 1)
            _db_purpose = _substitute_placeholders(_db_split[0] if _db_split else "")
            _db_closing = _substitute_placeholders(_db_split[1] if len(_db_split) > 1 else "")
            _db_purpose_safe = _db_purpose.replace("{", "{{").replace("}", "}}") if _db_purpose else ""
            _db_closing_safe = _db_closing.replace("{", "{{").replace("}", "}}") if _db_closing else ""
            _draft_content = render_cover_letter(
                case_type=case_type, client_name=client_name,
                a_number=a_number, receipt_number=receipt_number,
                filing_office=filing_office, enclosed_docs=_all_enc,
                attorney_name=attorney_name, bar_number=bar_number,
                firm_name=firm_name, firm_address=firm_address,
                custom_purpose=_db_purpose_safe,
                custom_closing=_db_closing_safe,
                recipient_address=recipient_address, salutation=salutation,
            )
        else:
            _db_subj_rendered = _substitute_placeholders(_db_subject)
            _db_body_rendered = _substitute_placeholders(_db_body)
            _db_body_safe = _db_body_rendered.replace("{", "{{").replace("}", "}}") if _db_body_rendered else ""
            _draft_content = render_cover_letter(
                case_type=case_type, client_name=client_name,
                a_number=a_number, receipt_number=receipt_number,
                filing_office=filing_office, enclosed_docs=_all_enc,
                attorney_name=attorney_name, bar_number=bar_number,
                firm_name=firm_name, firm_address=firm_address,
                recipient_address=recipient_address, salutation=salutation,
                custom_subject=_db_subj_rendered,
                custom_body=_db_body_safe,
            )
    render_draft_box("cover-letters", {
        "document_type": "cover letter",
        "client_name": client_name,
        "case_id": st.session_state.get("draft_id", ""),
        "content": _draft_content,
        "enclosed_docs": [doc["name"] for doc in _all_enc],
        "recipient_type": recipient_type,
        "recipient_address": recipient_address,
    })
