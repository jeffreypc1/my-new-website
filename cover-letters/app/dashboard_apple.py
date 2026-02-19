"""Cover Letter Generator — Apple Design System sample.

Duplicate of dashboard.py with Apple-inspired styling:
- Background color shifts instead of borders
- Standardized 12px radii
- Preview-as-hero layout (2:3 ratio, inputs = editing tools)
- System colors (SF Blue #007AFF, System Grey #8E8E93)
- Elegant empty state with ghosted letter skeleton
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from datetime import date
from pathlib import Path

import streamlit as st
from docx import Document
from docx.shared import Inches, Pt

from app.drafts import delete_draft, list_drafts, load_draft, new_draft_id, save_draft
from app.templates import (
    CASE_TYPES,
    RECIPIENT_CATEGORIES,
    TEMPLATES,
    get_filing_office_address,
    get_filing_offices,
    get_recipient_addresses,
    render_cover_letter,
    save_recipient_addresses,
)

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.client_banner import render_client_banner
from shared.tool_notes import render_tool_notes
try:
    from shared.draft_box import render_draft_box
except ImportError:
    render_draft_box = None
try:
    from shared.tool_help import render_tool_help
except ImportError:
    render_tool_help = None
try:
    from shared.feedback_button import render_feedback_button
except ImportError:
    render_feedback_button = None

# -- Page config --------------------------------------------------------------

st.set_page_config(
    page_title="Cover Letter Generator — O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -- Apple Design System CSS --------------------------------------------------

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* ─── Reset & Foundation ─── */
#MainMenu, footer,
div[data-testid="stToolbar"] { display: none !important; }

.stApp {
    font-family: -apple-system, 'SF Pro Display', 'SF Pro Text', 'Inter',
                 BlinkMacSystemFont, 'Helvetica Neue', sans-serif;
    background: #FFFFFF;
    -webkit-font-smoothing: antialiased;
}

/* ─── Sidebar: tool palette feel ─── */
section[data-testid="stSidebar"] {
    background: #F5F5F7 !important;
    border-right: none !important;
}
section[data-testid="stSidebar"] > div {
    background: #F5F5F7 !important;
}

/* ─── Standardize all radii to 12px ─── */
section[data-testid="stSidebar"] input,
section[data-testid="stSidebar"] textarea,
section[data-testid="stSidebar"] [data-baseweb="select"] > div,
.stTextInput input,
.stTextArea textarea,
[data-baseweb="select"] > div,
.stButton > button,
.stDownloadButton > button {
    border-radius: 12px !important;
}

/* ─── Inputs: borderless with background shift ─── */
section[data-testid="stSidebar"] .stTextInput input,
section[data-testid="stSidebar"] .stTextArea textarea {
    border: none !important;
    background: #FFFFFF !important;
    box-shadow: none !important;
    padding: 10px 14px !important;
    font-size: 0.88rem !important;
    color: #1D1D1F !important;
    transition: background 0.2s ease;
}
section[data-testid="stSidebar"] .stTextInput input:focus,
section[data-testid="stSidebar"] .stTextArea textarea:focus {
    background: #FFFFFF !important;
    box-shadow: 0 0 0 3px rgba(0, 122, 255, 0.25) !important;
}
section[data-testid="stSidebar"] .stTextInput input::placeholder,
section[data-testid="stSidebar"] .stTextArea textarea::placeholder {
    color: #8E8E93 !important;
}

/* ─── Selectbox: borderless ─── */
section[data-testid="stSidebar"] [data-baseweb="select"] > div {
    border: none !important;
    background: #FFFFFF !important;
    box-shadow: none !important;
}
section[data-testid="stSidebar"] [data-baseweb="select"] > div:focus-within {
    box-shadow: 0 0 0 3px rgba(0, 122, 255, 0.25) !important;
}

/* ─── Labels: System Grey ─── */
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] .stRadio label {
    color: #8E8E93 !important;
    font-size: 0.78rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.02em !important;
    text-transform: uppercase !important;
}

/* ─── Section headers in sidebar ─── */
section[data-testid="stSidebar"] h4 {
    color: #1D1D1F !important;
    font-size: 0.82rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.01em !important;
    text-transform: uppercase !important;
    margin-top: 16px !important;
    margin-bottom: 8px !important;
}

/* ─── Buttons: SF Blue primary, ghost secondary ─── */
.stButton > button[kind="primary"],
.stButton > button[data-testid="stBaseButton-primary"] {
    background: #007AFF !important;
    color: #FFFFFF !important;
    border: none !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    padding: 8px 20px !important;
    transition: all 0.2s ease !important;
}
.stButton > button[kind="primary"]:hover,
.stButton > button[data-testid="stBaseButton-primary"]:hover {
    background: #0066D6 !important;
    transform: scale(1.01);
}

.stButton > button:not([kind="primary"]):not([data-testid="stBaseButton-primary"]),
.stDownloadButton > button {
    background: transparent !important;
    color: #007AFF !important;
    border: 1.5px solid #D2D2D7 !important;
    font-weight: 500 !important;
    font-size: 0.85rem !important;
    transition: all 0.2s ease !important;
}
.stButton > button:not([kind="primary"]):not([data-testid="stBaseButton-primary"]):hover,
.stDownloadButton > button:hover {
    background: #F5F5F7 !important;
    border-color: #007AFF !important;
}

/* ─── Dividers: subtle ─── */
section[data-testid="stSidebar"] hr {
    border: none !important;
    height: 1px !important;
    background: rgba(0,0,0,0.06) !important;
    margin: 16px 0 !important;
}

/* ─── Radio buttons: pill style ─── */
section[data-testid="stSidebar"] .stRadio > div[role="radiogroup"] {
    gap: 0 !important;
    background: #E8E8ED !important;
    border-radius: 12px !important;
    padding: 3px !important;
}
section[data-testid="stSidebar"] .stRadio > div[role="radiogroup"] label {
    text-transform: none !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    color: #1D1D1F !important;
    padding: 6px 14px !important;
    border-radius: 10px !important;
    transition: all 0.15s ease !important;
}
section[data-testid="stSidebar"] .stRadio > div[role="radiogroup"] label[data-checked="true"],
section[data-testid="stSidebar"] .stRadio > div[role="radiogroup"] label:has(input:checked) {
    background: #FFFFFF !important;
    box-shadow: 0 1px 3px rgba(0,0,0,0.08) !important;
    font-weight: 600 !important;
}

/* ─── Navigation bar ─── */
.nav-bar {
    display: flex;
    align-items: center;
    padding: 12px 4px;
    margin: -1rem 0 1.2rem 0;
    /* no border — color shift instead */
    background: linear-gradient(to bottom, #FFFFFF, #FAFAFA);
}
.nav-back {
    display: flex;
    align-items: center;
    gap: 6px;
    font-size: 0.88rem;
    font-weight: 500;
    color: #007AFF;
    text-decoration: none;
    min-width: 150px;
}
.nav-back:hover { color: #0066D6; text-decoration: none; opacity: 0.8; }
.nav-title {
    flex: 1;
    text-align: center;
    font-size: 1.1rem;
    font-weight: 700;
    color: #1D1D1F;
    letter-spacing: -0.02em;
}
.nav-firm {
    font-weight: 400;
    color: #8E8E93;
    font-size: 0.85rem;
    margin-left: 8px;
}
.nav-spacer { min-width: 150px; }

/* ─── Section labels (main area) ─── */
.section-label {
    font-size: 0.72rem;
    font-weight: 700;
    color: #8E8E93;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    margin-bottom: 6px;
    margin-top: 12px;
}

/* ─── Document items ─── */
.doc-item {
    font-size: 0.88rem;
    padding: 4px 0;
    color: #1D1D1F;
}

/* ─── SF task rows ─── */
.sf-task-deleted {
    background: #FFF2F2;
    border: none;
    border-radius: 12px;
    padding: 8px 12px;
    margin: 3px 0;
}
.sf-task-deleted .doc-item {
    text-decoration: line-through;
    color: #FF3B30;
}
.sf-task-row {
    padding: 4px 10px;
    margin: 2px 0;
    border-radius: 12px;
    border: none;
}
.sf-task-row:hover { background: #F5F5F7; }

/* ─── Pending badges ─── */
.pending-badge {
    display: inline-block;
    font-size: 0.7rem;
    font-weight: 600;
    padding: 3px 10px;
    border-radius: 20px;
    margin-left: 6px;
}
.pending-badge.edits {
    background: #FFF9E6;
    color: #A06800;
}
.pending-badge.deletes {
    background: #FFF2F2;
    color: #FF3B30;
}

/* ─── Preview panel: THE HERO ─── */
.preview-panel {
    font-family: 'Times New Roman', 'Georgia', Times, serif;
    font-size: 0.88rem;
    line-height: 1.75;
    background: #FFFFFF;
    border: none;
    border-radius: 16px;
    padding: 40px 44px;
    box-shadow: 0 2px 12px rgba(0,0,0,0.06), 0 0 0 1px rgba(0,0,0,0.03);
    min-height: 60vh;
    max-height: 80vh;
    overflow-y: auto;
    white-space: pre-wrap;
    word-wrap: break-word;
}
.preview-panel .letter-cert {
    margin-top: 32px;
    padding-top: 20px;
    border-top: 1px solid #E8E8ED;
}
.preview-panel .conf-notice {
    background: #FFF9E6;
    border: none;
    border-radius: 12px;
    padding: 12px 16px;
    margin-bottom: 16px;
    font-size: 0.82rem;
    color: #A06800;
    font-style: italic;
}

/* ─── Empty state: ghosted skeleton ─── */
.skeleton-letter {
    border: none;
    border-radius: 16px;
    padding: 40px 44px;
    min-height: 60vh;
    background: #FFFFFF;
    box-shadow: 0 2px 12px rgba(0,0,0,0.06), 0 0 0 1px rgba(0,0,0,0.03);
    position: relative;
    overflow: hidden;
}
.skeleton-line {
    height: 12px;
    border-radius: 6px;
    background: linear-gradient(90deg, #F0F0F5 25%, #E8E8ED 50%, #F0F0F5 75%);
    background-size: 200% 100%;
    animation: shimmer 2.5s ease-in-out infinite;
    margin-bottom: 14px;
}
.skeleton-line.short { width: 35%; }
.skeleton-line.medium { width: 55%; }
.skeleton-line.long { width: 80%; }
.skeleton-line.full { width: 100%; }
.skeleton-line.indent { margin-left: 24px; width: 50%; }
.skeleton-line.spacer { height: 0; margin-bottom: 28px; }
.skeleton-line.bold { height: 14px; width: 45%; }
.skeleton-hint {
    text-align: center;
    color: #C7C7CC;
    font-size: 0.82rem;
    font-weight: 500;
    margin-top: 24px;
    letter-spacing: 0.02em;
}
@keyframes shimmer {
    0% { background-position: 200% 0; }
    100% { background-position: -200% 0; }
}

/* ─── Saved toast ─── */
.saved-toast {
    font-size: 0.8rem;
    color: #34C759;
    font-weight: 600;
}

/* ─── Draft badge ─── */
.draft-badge {
    display: inline-block;
    padding: 3px 10px;
    font-size: 0.7rem;
    font-weight: 600;
    background: #E8F0FE;
    color: #007AFF;
    border-radius: 12px;
    margin-bottom: 8px;
}

/* ─── Export area ─── */
.export-area {
    background: #F5F5F7;
    border-radius: 16px;
    padding: 20px 24px;
    margin-top: 16px;
}

/* ─── Enclosed docs panel ─── */
.docs-panel {
    background: #F5F5F7;
    border-radius: 16px;
    padding: 20px 24px;
}

/* ─── Caption text ─── */
.stCaption, .stCaption p {
    color: #8E8E93 !important;
}
</style>
""",
    unsafe_allow_html=True,
)

# -- Navigation bar -----------------------------------------------------------

st.markdown(
    """
<div class="nav-bar">
    <a href="http://localhost:8502" class="nav-back">&#8592; Staff Dashboard</a>
    <div class="nav-title">Cover Letter Generator<span class="nav-firm">&mdash; O'Brien Immigration Law</span></div>
    <div class="nav-spacer"></div>
</div>
""",
    unsafe_allow_html=True,
)

render_client_banner()
if render_tool_help:
    render_tool_help("cover-letters")
if render_feedback_button:
    render_feedback_button("cover-letters")

# -- Session state defaults ---------------------------------------------------

_DEFAULTS: dict = {
    "draft_id": None,
    "last_saved_msg": "",
    "custom_docs": [],
    "sf_task_docs": [],
    "doc_descriptions": {},
    "sf_tasks_pending_delete": [],
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state.draft_id is None:
    st.session_state.draft_id = new_draft_id()


# -- Helpers ------------------------------------------------------------------

def _build_enclosed_docs_list() -> list[dict[str, str]]:
    """Collect the currently selected enclosed documents with descriptions."""
    docs: list[dict[str, str]] = []
    for doc_name in st.session_state.get("sf_task_docs", []):
        docs.append({"name": doc_name, "description": ""})
    for doc_name in st.session_state.get("custom_docs", []):
        desc = st.session_state.get("doc_descriptions", {}).get(doc_name, "")
        docs.append({"name": doc_name, "description": desc})
    return docs


def _do_save(case_type: str) -> None:
    """Save the current state as a draft."""
    client = {
        "name": st.session_state.get("inp_client_name", ""),
        "a_number": st.session_state.get("inp_a_number", ""),
        "receipt_number": st.session_state.get("inp_receipt_number", ""),
    }
    attorney = {
        "name": st.session_state.get("inp_attorney_name", ""),
        "bar_number": st.session_state.get("inp_bar_number", ""),
        "firm_name": st.session_state.get("inp_firm_name", "O'Brien Immigration Law"),
        "firm_address": st.session_state.get("inp_firm_address", ""),
    }
    filing_office = st.session_state.get("inp_filing_office", "")
    enclosed = _build_enclosed_docs_list()

    _rt = st.session_state.get("inp_recipient_type", "Government Agency")
    _rtype = "client" if _rt == "Client" else "agency"
    if _rtype == "agency":
        _all_addr = get_recipient_addresses()
        _sel_name = st.session_state.get("inp_recipient_name", "")
        _entry = next((a for a in _all_addr if a["name"] == _sel_name), None)
        _raddr = _entry["address"] if _entry else ""
        _rsal = _entry.get("salutation", "Dear Sir or Madam:") if _entry else "Dear Sir or Madam:"
    else:
        _raddr = st.session_state.get("inp_client_address", "")
        _rsal = st.session_state.get("inp_client_salutation", "Dear Sir or Madam:")

    save_draft(
        st.session_state.draft_id,
        case_type,
        client,
        attorney,
        filing_office,
        enclosed,
        recipient_type=_rtype,
        recipient_address=_raddr,
        salutation=_rsal,
    )
    name = client["name"] or "draft"
    st.session_state.last_saved_msg = f"Saved -- {name}"


def _do_load(draft_id: str) -> None:
    """Load a draft into session state."""
    draft = load_draft(draft_id)
    if not draft:
        return
    st.session_state.draft_id = draft["id"]
    st.session_state.inp_case_type = draft.get("case_type", CASE_TYPES[0])

    c = draft.get("client", {})
    st.session_state.inp_client_name = c.get("name", "")
    st.session_state.inp_a_number = c.get("a_number", "")
    st.session_state.inp_receipt_number = c.get("receipt_number", "")

    a = draft.get("attorney", {})
    st.session_state.inp_attorney_name = a.get("name", "")
    st.session_state.inp_bar_number = a.get("bar_number", "")
    st.session_state.inp_firm_name = a.get("firm_name", "O'Brien Immigration Law")
    st.session_state.inp_firm_address = a.get("firm_address", "")

    st.session_state.inp_filing_office = draft.get("filing_office", "")

    _rtype = draft.get("recipient_type", "")
    _raddr = draft.get("recipient_address", "")
    _rsal = draft.get("salutation", "Dear Sir or Madam:")

    if _rtype == "client":
        st.session_state.inp_recipient_type = "Client"
        st.session_state.inp_client_address = _raddr
        st.session_state.inp_client_salutation = _rsal
    elif _rtype == "agency" and _raddr:
        st.session_state.inp_recipient_type = "Government Agency"
        _all_addr = get_recipient_addresses()
        _match = next((a for a in _all_addr if a["address"] == _raddr), None)
        if _match:
            st.session_state.inp_recipient_category = _match["category"]
            st.session_state.inp_recipient_name = _match["name"]
    elif not _rtype and draft.get("filing_office"):
        st.session_state.inp_recipient_type = "Government Agency"
        _all_addr = get_recipient_addresses()
        _match = next((a for a in _all_addr if a["name"] == draft["filing_office"]), None)
        if _match:
            st.session_state.inp_recipient_category = _match["category"]
            st.session_state.inp_recipient_name = _match["name"]

    enclosed = draft.get("enclosed_docs", [])
    custom = []
    descs: dict[str, str] = {}
    for doc in enclosed:
        name = doc.get("name", "")
        desc = doc.get("description", "")
        custom.append(name)
        if desc:
            descs[name] = desc
    st.session_state.custom_docs = custom
    st.session_state.doc_descriptions = descs


def _do_new() -> None:
    """Start a fresh cover letter."""
    st.session_state.draft_id = new_draft_id()
    st.session_state.last_saved_msg = ""
    st.session_state.custom_docs = []
    st.session_state.sf_task_docs = []
    st.session_state.doc_descriptions = {}
    for k in (
        "inp_client_name", "inp_a_number", "inp_receipt_number",
        "inp_attorney_name", "inp_bar_number", "inp_firm_address",
        "inp_filing_office", "inp_case_type",
        "inp_recipient_type", "inp_recipient_category", "inp_recipient_name",
        "inp_client_address", "inp_client_salutation",
    ):
        if k in st.session_state:
            del st.session_state[k]
    if "inp_firm_name" in st.session_state:
        del st.session_state["inp_firm_name"]


@st.dialog("Add New Address")
def _add_address_dialog():
    """Dialog for adding a new recipient address."""
    _new_name = st.text_input("Name", placeholder="e.g. USCIS Field Office - Boston")
    _new_category = st.selectbox("Category", options=RECIPIENT_CATEGORIES)
    _new_address = st.text_area(
        "Address",
        height=100,
        placeholder="Organization Name\nStreet Address\nCity, State ZIP",
    )
    _new_salutation = st.text_input("Salutation", value="Dear Sir or Madam:")

    if st.button("Save Address", type="primary", use_container_width=True):
        if not _new_name or not _new_address:
            st.error("Name and address are required.")
        else:
            _addresses = get_recipient_addresses()
            _new_id = _new_name.lower().replace(" ", "_").replace("(", "").replace(")", "")[:40]
            _addresses.append({
                "id": _new_id,
                "name": _new_name,
                "category": _new_category,
                "address": _new_address,
                "salutation": _new_salutation,
            })
            save_recipient_addresses(_addresses)
            st.session_state["_show_add_address"] = False
            st.rerun()


def _build_preview_html(letter_text: str, case_type: str) -> str:
    """Convert the plain-text cover letter into styled HTML for the preview."""
    esc = html_mod.escape
    tpl = TEMPLATES.get(case_type, {})

    lines = letter_text.split("\n")
    parts: list[str] = []

    if "confidentiality_notice" in tpl:
        notice = tpl["confidentiality_notice"]
        parts.append(f'<div class="conf-notice">{esc(notice)}</div>')

    for line in lines:
        escaped = esc(line)
        if line.startswith("RE: "):
            parts.append(f"<strong>{escaped}</strong>")
        elif line.startswith("    A# ") or line.startswith("    Receipt#") or (line.startswith("    ") and not line.startswith("    ") + "  "):
            parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;{esc(line.strip())}")
        elif line == "CERTIFICATE OF SERVICE":
            parts.append(f'<div class="letter-cert"><strong>{escaped}</strong></div>')
        elif line.startswith("____"):
            parts.append(f"<br>{escaped}")
        else:
            parts.append(escaped)

    return "<br>".join(parts)


def _build_skeleton_html() -> str:
    """Build a ghosted letter skeleton for the empty state."""
    return """
    <div class="skeleton-letter">
        <!-- Date -->
        <div class="skeleton-line short"></div>
        <div class="skeleton-line spacer"></div>

        <!-- Address block -->
        <div class="skeleton-line medium"></div>
        <div class="skeleton-line medium" style="width:45%"></div>
        <div class="skeleton-line short"></div>
        <div class="skeleton-line spacer"></div>

        <!-- RE block -->
        <div class="skeleton-line bold"></div>
        <div class="skeleton-line indent" style="width:35%"></div>
        <div class="skeleton-line indent" style="width:28%"></div>
        <div class="skeleton-line spacer"></div>

        <!-- Salutation -->
        <div class="skeleton-line" style="width:25%"></div>
        <div class="skeleton-line spacer"></div>

        <!-- Body paragraphs -->
        <div class="skeleton-line full"></div>
        <div class="skeleton-line full"></div>
        <div class="skeleton-line long"></div>
        <div class="skeleton-line spacer"></div>

        <!-- Enclosed docs -->
        <div class="skeleton-line long"></div>
        <div class="skeleton-line spacer" style="margin-bottom:14px"></div>
        <div class="skeleton-line indent" style="width:60%"></div>
        <div class="skeleton-line indent" style="width:55%"></div>
        <div class="skeleton-line indent" style="width:65%"></div>
        <div class="skeleton-line indent" style="width:50%"></div>
        <div class="skeleton-line spacer"></div>

        <!-- Closing -->
        <div class="skeleton-line full"></div>
        <div class="skeleton-line long"></div>
        <div class="skeleton-line spacer"></div>

        <!-- Signature -->
        <div class="skeleton-line" style="width:22%"></div>
        <div class="skeleton-line spacer" style="margin-bottom:20px"></div>
        <div class="skeleton-line" style="width:30%; height:2px; background:#D2D2D7"></div>
        <div class="skeleton-line" style="width:28%; margin-top:8px"></div>
        <div class="skeleton-line short"></div>

        <div class="skeleton-hint">Fill in the sidebar to see your letter come to life</div>
    </div>
    """


def _build_docx(letter_text: str, attorney_name: str, firm_name: str) -> bytes:
    """Build a Word document from the rendered cover letter text."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, txt, size=12, bold=False):
        r = para.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        return r

    lines = letter_text.split("\n")
    for line in lines:
        p = doc.add_paragraph()
        is_bold = line.startswith("RE:") or line == "CERTIFICATE OF SERVICE"
        _run(p, line, bold=is_bold)
        fmt = p.paragraph_format
        fmt.space_after = Pt(0)
        fmt.space_before = Pt(0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# -- Sidebar ------------------------------------------------------------------

with st.sidebar:
    st.markdown("#### Drafts")
    btn_cols = st.columns(2)
    with btn_cols[0]:
        if st.button("New", use_container_width=True):
            _do_new()
            st.rerun()
    with btn_cols[1]:
        save_clicked = st.button("Save", use_container_width=True, type="primary")

    saved_drafts = list_drafts()
    if saved_drafts:
        labels_map = {
            d["id"]: f"{d['client_name']} -- {d['case_type']}"
            for d in saved_drafts
        }
        draft_ids = list(labels_map.keys())
        selected_draft = st.selectbox(
            "Load a saved draft",
            options=[""] + draft_ids,
            format_func=lambda x: labels_map.get(x, "Select..."),
            label_visibility="collapsed",
        )
        load_cols = st.columns(2)
        with load_cols[0]:
            if selected_draft and st.button("Load", use_container_width=True):
                _do_load(selected_draft)
                st.rerun()
        with load_cols[1]:
            if selected_draft and st.button("Delete", use_container_width=True):
                delete_draft(selected_draft)
                st.rerun()

    if st.session_state.last_saved_msg:
        st.markdown(
            f'<div class="saved-toast">{html_mod.escape(st.session_state.last_saved_msg)}</div>',
            unsafe_allow_html=True,
        )

    st.divider()

    case_type = st.selectbox(
        "Case Type",
        options=CASE_TYPES,
        key="inp_case_type",
    )

    st.divider()

    st.markdown("#### Attorney / Firm")
    attorney_name = st.text_input("Attorney Name", key="inp_attorney_name")
    bar_number = st.text_input("Bar Number", key="inp_bar_number")
    firm_name = st.text_input(
        "Firm Name",
        value="O'Brien Immigration Law",
        key="inp_firm_name",
    )
    firm_address = st.text_area(
        "Firm Address",
        key="inp_firm_address",
        height=80,
        placeholder="123 Main Street\nSuite 400\nCity, State ZIP",
    )

    st.divider()

    st.markdown("#### Client")
    client_name = st.text_input(
        "Client Name",
        key="inp_client_name",
        placeholder="e.g. Maria Garcia Lopez",
    )
    a_number = st.text_input(
        "A-Number",
        key="inp_a_number",
        placeholder="e.g. 123-456-789",
    )
    receipt_number = st.text_input(
        "Receipt Number",
        key="inp_receipt_number",
        placeholder="e.g. SRC-21-123-45678",
    )

    st.divider()

    st.markdown("#### Recipient")
    recipient_type = st.radio(
        "Recipient type",
        options=["Government Agency", "Client"],
        key="inp_recipient_type",
        horizontal=True,
        label_visibility="collapsed",
    )

    if recipient_type == "Government Agency":
        _all_addresses = get_recipient_addresses()

        _selected_category = st.selectbox(
            "Category",
            options=RECIPIENT_CATEGORIES,
            key="inp_recipient_category",
        )

        _filtered = [a for a in _all_addresses if a.get("category") == _selected_category]
        _addr_names = [a["name"] for a in _filtered]

        if _addr_names:
            _selected_name = st.selectbox(
                "Address",
                options=_addr_names,
                key="inp_recipient_name",
            )
            _selected_entry = next((a for a in _filtered if a["name"] == _selected_name), None)
            if _selected_entry:
                st.caption(_selected_entry["address"].replace("\n", "  \n"))
                recipient_address = _selected_entry["address"]
                salutation = _selected_entry.get("salutation", "Dear Sir or Madam:")
            else:
                recipient_address = ""
                salutation = "Dear Sir or Madam:"
        else:
            st.caption("No addresses in this category.")
            recipient_address = ""
            salutation = "Dear Sir or Madam:"

        if st.button("+ Add New Address", use_container_width=True):
            st.session_state["_show_add_address"] = True

        filing_office = ""

    else:
        _sf_client = st.session_state.get("sf_client")
        _default_addr = ""
        _default_sal = "Dear Sir or Madam:"
        if _sf_client:
            _parts = []
            _first = _sf_client.get("FirstName", "")
            _last = _sf_client.get("LastName", "")
            _full_name = f"{_first} {_last}".strip()
            if _full_name:
                _parts.append(_full_name)
            _street = _sf_client.get("MailingStreet", "")
            if _street:
                _parts.append(_street)
            _city = _sf_client.get("MailingCity", "")
            _state = _sf_client.get("MailingState", "")
            _zip = _sf_client.get("MailingPostalCode", "")
            _csz = ", ".join(filter(None, [_city, _state]))
            if _zip:
                _csz = f"{_csz} {_zip}" if _csz else _zip
            if _csz:
                _parts.append(_csz)
            _default_addr = "\n".join(_parts)
            if _first or _last:
                _default_sal = f"Dear {_full_name}:"

        recipient_address = st.text_area(
            "Recipient Address",
            value=st.session_state.get("inp_client_address", _default_addr),
            key="inp_client_address",
            height=100,
            placeholder="Client name\nStreet address\nCity, State ZIP",
        )
        salutation = st.text_input(
            "Salutation",
            value=st.session_state.get("inp_client_salutation", _default_sal),
            key="inp_client_salutation",
        )
        filing_office = ""

    render_tool_notes("cover-letters")


# -- Add Address dialog trigger -----------------------------------------------
if st.session_state.get("_show_add_address"):
    _add_address_dialog()

# -- Handle save (after sidebar renders) -------------------------------------

if save_clicked:
    _do_save(case_type)
    st.rerun()


# -- Main area — Preview-first layout (2:3 = tools:preview) ------------------

tools_col, preview_col = st.columns([2, 3], gap="large")

# -- Left column: Enclosed documents (the "editing tools") --------------------

with tools_col:
    st.markdown('<div class="section-label">Enclosed Documents</div>', unsafe_allow_html=True)
    st.caption("Check documents to include. Add descriptions for specificity.")

    _active_client = st.session_state.get("sf_client")
    _sf_contact_id = _active_client.get("Id", "") if _active_client else ""
    _sf_tasks: list[dict] = []

    if _sf_contact_id:
        try:
            from shared.salesforce_client import get_lc_tasks, create_lc_task
            _sf_tasks = get_lc_tasks(_sf_contact_id)
        except Exception:
            _sf_tasks = []

    if _sf_tasks or _sf_contact_id:
        st.markdown('<div class="section-label" style="margin-top:4px;">Salesforce Tasks</div>', unsafe_allow_html=True)

    _pending_deletes: list[str] = list(st.session_state.get("sf_tasks_pending_delete", []))

    if _sf_tasks:
        _current_sf_checked = list(st.session_state.get("sf_task_docs", []))
        _new_sf_checked: list[str] = []
        _has_edits = False

        for _task in _sf_tasks:
            _task_id = _task.get("Id", "")
            _task_label = _task.get("For__c") or _task.get("Name") or "Untitled task"
            _is_pending_delete = _task_id in _pending_deletes
            _edit_key = f"_sf_task_edit_{_task_id}"

            if _is_pending_delete:
                st.markdown(
                    f'<div class="sf-task-deleted">'
                    f'<div class="doc-item">{html_mod.escape(_task_label)}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
                _undo_cols = st.columns([8, 1])
                with _undo_cols[1]:
                    if st.button("Undo", key=f"_btn_undo_sft_{_task_id}", use_container_width=True):
                        _pending_deletes.remove(_task_id)
                        st.session_state.sf_tasks_pending_delete = _pending_deletes
                        st.rerun()
            else:
                _edited_label = st.session_state.get(_edit_key, _task_label)
                _was_checked = _edited_label in _current_sf_checked or _task_label in _current_sf_checked
                _tc = st.columns([0.5, 8, 0.5])
                with _tc[0]:
                    _tc_checked = st.checkbox(
                        _task_label,
                        value=_was_checked,
                        key=f"chk_sft_{_task_id}",
                        label_visibility="collapsed",
                    )
                with _tc[1]:
                    _edited_val = st.text_input(
                        "Edit document name",
                        value=_edited_label,
                        key=_edit_key,
                        label_visibility="collapsed",
                    )
                with _tc[2]:
                    if st.button("X", key=f"_btn_del_sft_{_task_id}", help="Mark for deletion"):
                        _pending_deletes.append(_task_id)
                        st.session_state.sf_tasks_pending_delete = _pending_deletes
                        st.rerun()

                if _edited_val.strip() != _task_label:
                    _has_edits = True
                if _tc_checked and _edited_val.strip():
                    _new_sf_checked.append(_edited_val.strip())

        st.session_state.sf_task_docs = _new_sf_checked

        _has_changes = _has_edits or len(_pending_deletes) > 0
        if _has_changes:
            _badge_parts: list[str] = []
            _edit_count = sum(
                1 for _t in _sf_tasks
                if _t.get("Id", "") not in _pending_deletes
                and st.session_state.get(f"_sf_task_edit_{_t.get('Id', '')}", "").strip()
                   != (_t.get("For__c") or _t.get("Name") or "Untitled task")
                and st.session_state.get(f"_sf_task_edit_{_t.get('Id', '')}", "") != ""
            )
            if _edit_count:
                _badge_parts.append(f'<span class="pending-badge edits">{_edit_count} edit{"s" if _edit_count != 1 else ""}</span>')
            if _pending_deletes:
                _del_count = len(_pending_deletes)
                _badge_parts.append(f'<span class="pending-badge deletes">{_del_count} deletion{"s" if _del_count != 1 else ""}</span>')
            st.markdown(f'{"".join(_badge_parts)}', unsafe_allow_html=True)

            if st.button("Save Changes to Salesforce", type="primary", use_container_width=True, key="_btn_save_all_sf"):
                _save_errors: list[str] = []
                try:
                    from shared.salesforce_client import update_lc_task, delete_lc_task

                    for _t in _sf_tasks:
                        _tid = _t.get("Id", "")
                        if _tid in _pending_deletes:
                            continue
                        _orig = _t.get("For__c") or _t.get("Name") or "Untitled task"
                        _new_val = st.session_state.get(f"_sf_task_edit_{_tid}", "").strip()
                        if _new_val and _new_val != _orig:
                            try:
                                update_lc_task(_tid, _new_val)
                            except Exception as e:
                                _save_errors.append(f"Edit '{_new_val}': {e}")

                    for _tid in _pending_deletes:
                        try:
                            delete_lc_task(_tid)
                        except Exception as e:
                            _save_errors.append(f"Delete {_tid}: {e}")

                    st.session_state.sf_tasks_pending_delete = []
                    for _t in _sf_tasks:
                        _ek = f"_sf_task_edit_{_t.get('Id', '')}"
                        if _ek in st.session_state:
                            del st.session_state[_ek]

                    if _save_errors:
                        st.warning(f"Saved with {len(_save_errors)} error(s): {'; '.join(_save_errors)}")
                    else:
                        st.success("Changes saved to Salesforce.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Save failed: {e}")

    elif _sf_contact_id:
        st.caption("No tasks found for this client.")

    if _sf_contact_id:
        st.markdown("---")

        _add_task_cols = st.columns([5, 1])
        with _add_task_cols[0]:
            _new_task_desc = st.text_input(
                "New SF task",
                key="_inp_new_sf_task",
                placeholder="Add a document...",
                label_visibility="collapsed",
            )
        with _add_task_cols[1]:
            if st.button("Add", use_container_width=True, key="_btn_add_sf_task") and _new_task_desc:
                try:
                    from shared.salesforce_client import create_lc_task
                    create_lc_task(_sf_contact_id, _new_task_desc.strip())
                    st.rerun()
                except Exception as e:
                    st.error(f"Failed to create task: {e}")

        with st.expander("Bulk Add Documents"):
            _mass_input = st.text_area(
                "Paste comma-separated document names",
                key="_inp_mass_upload",
                height=80,
                placeholder="e.g.: I-589 Application, Passport Copy, Birth Certificate",
                label_visibility="collapsed",
            )
            if st.button("Upload All to Salesforce", use_container_width=True, key="_btn_mass_upload") and _mass_input.strip():
                _items = [item.strip() for item in _mass_input.split(",") if item.strip()]
                if not _items:
                    st.warning("No documents found. Separate items with commas.")
                else:
                    _created = 0
                    _errors = 0
                    try:
                        from shared.salesforce_client import create_lc_task
                        for _item in _items:
                            try:
                                create_lc_task(_sf_contact_id, _item)
                                _created += 1
                            except Exception:
                                _errors += 1
                        if _errors:
                            st.warning(f"Created {_created} record(s), {_errors} failed.")
                        else:
                            st.success(f"Created {_created} record(s) in Salesforce.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Failed to upload: {e}")

    descriptions = st.session_state.get("doc_descriptions", {})
    new_descriptions: dict[str, str] = dict(descriptions)

    custom_docs = list(st.session_state.get("custom_docs", []))
    new_custom: list[str] = []
    for doc_name in custom_docs:
        cols = st.columns([0.5, 5, 3, 1])
        with cols[0]:
            st.checkbox(
                doc_name,
                value=True,
                key=f"chk_custom_{hash(doc_name)}",
                label_visibility="collapsed",
                disabled=True,
            )
        with cols[1]:
            st.markdown(
                f'<div class="doc-item">{html_mod.escape(doc_name)}</div>',
                unsafe_allow_html=True,
            )
        with cols[2]:
            desc_val = descriptions.get(doc_name, "")
            desc = st.text_input(
                "Description",
                value=desc_val,
                key=f"desc_custom_{hash(doc_name)}",
                label_visibility="collapsed",
                placeholder="Optional description...",
            )
            if desc:
                new_descriptions[doc_name] = desc
            elif doc_name in new_descriptions:
                new_descriptions.pop(doc_name, None)
        with cols[3]:
            if st.button("X", key=f"rm_{hash(doc_name)}", help="Remove custom document"):
                custom_docs.remove(doc_name)
                st.session_state.custom_docs = custom_docs
                new_descriptions.pop(doc_name, None)
                st.rerun()
        new_custom.append(doc_name)

    st.session_state.custom_docs = new_custom
    st.session_state.doc_descriptions = new_descriptions

# -- Right column: Cover letter preview (THE HERO) ---------------------------

with preview_col:
    st.markdown('<div class="section-label">Cover Letter Preview</div>', unsafe_allow_html=True)

    if not client_name:
        # Elegant empty state: ghosted skeleton letter
        st.markdown(_build_skeleton_html(), unsafe_allow_html=True)
    else:
        all_enclosed = _build_enclosed_docs_list()

        letter_text = render_cover_letter(
            case_type=case_type,
            client_name=client_name,
            a_number=a_number,
            receipt_number=receipt_number,
            filing_office=filing_office,
            enclosed_docs=all_enclosed,
            attorney_name=attorney_name,
            bar_number=bar_number,
            firm_name=firm_name,
            firm_address=firm_address,
            recipient_address=recipient_address,
            salutation=salutation,
        )

        preview_html = _build_preview_html(letter_text, case_type)
        st.markdown(
            f'<div class="preview-panel">{preview_html}</div>',
            unsafe_allow_html=True,
        )

        # Export area with background shift
        st.markdown('<div class="section-label" style="margin-top:20px">Export</div>', unsafe_allow_html=True)

        exp_cols = st.columns(3)
        with exp_cols[0]:
            st.download_button(
                "Download .txt",
                data=letter_text,
                file_name=f"Cover_Letter_{client_name.replace(' ', '_')}.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with exp_cols[1]:
            docx_bytes = _build_docx(letter_text, attorney_name, firm_name)
            st.download_button(
                "Download .docx",
                data=docx_bytes,
                file_name=f"Cover_Letter_{client_name.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with exp_cols[2]:
            if st.button("Google Docs", use_container_width=True):
                with st.spinner("Uploading..."):
                    try:
                        from shared.google_upload import upload_to_google_docs
                        url = upload_to_google_docs(docx_bytes, f"Cover Letter - {client_name}")
                        st.session_state.google_doc_url = url
                    except Exception as e:
                        st.error(f"Upload failed: {e}")
            if st.session_state.get("google_doc_url"):
                st.markdown(f"[Open Google Doc]({st.session_state.google_doc_url})")

# -- Draft Box (below both columns) ------------------------------------------
if render_draft_box is not None:
    _all_enc = _build_enclosed_docs_list()
    _draft_content = ""
    if client_name:
        _draft_content = render_cover_letter(
            case_type=case_type, client_name=client_name,
            a_number=a_number, receipt_number=receipt_number,
            filing_office=filing_office, enclosed_docs=_all_enc,
            attorney_name=attorney_name, bar_number=bar_number,
            firm_name=firm_name, firm_address=firm_address,
            recipient_address=recipient_address, salutation=salutation,
        )
    render_draft_box("cover-letters", {
        "document_type": "cover letter",
        "client_name": client_name,
        "case_id": st.session_state.get("draft_id", ""),
        "content": _draft_content,
        "enclosed_docs": [doc["name"] for doc in _all_enc],
        "recipient_type": recipient_type,
        "recipient_address": recipient_address,
    })
