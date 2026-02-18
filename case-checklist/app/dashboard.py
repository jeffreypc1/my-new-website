"""Case Checklist -- Streamlit dashboard.

Full case-management UI for tracking immigration case checklists with
deadline monitoring, progress tracking, and Word/text export.
Works entirely with local persistence (no API server required).

Part of the O'Brien Immigration Law tool suite.
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.checklists import (
    CASE_TYPES,
    add_custom_item,
    create_case,
    delete_case,
    get_case_progress,
    get_deadline_status,
    list_cases,
    load_case,
    save_case,
    update_item,
)

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.google_upload import upload_to_google_docs
from shared.client_banner import render_client_banner
from shared.tool_notes import render_tool_notes
try:
    from shared.tool_help import render_tool_help
except ImportError:
    render_tool_help = None
try:
    from shared.feedback_button import render_feedback_button
except ImportError:
    render_feedback_button = None

# ── Page config ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Case Checklist -- O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ──────────────────────────────────────────────────────────────────────

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* Hide Streamlit chrome */
#MainMenu, footer,
div[data-testid="stToolbar"] { display: none !important; }

.stApp {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

/* Navigation bar */
.nav-bar {
    display: flex;
    align-items: center;
    padding: 10px 4px;
    margin: -1rem 0 1.2rem 0;
    border-bottom: 1px solid rgba(0,0,0,0.07);
}
.nav-back {
    display: flex;
    align-items: center;
    gap: 6px;
    font-family: 'Inter', sans-serif;
    font-size: 0.85rem;
    font-weight: 500;
    color: #0066CC;
    text-decoration: none;
    min-width: 150px;
}
.nav-back:hover { color: #004499; text-decoration: underline; }
.nav-title {
    flex: 1;
    text-align: center;
    font-family: 'Inter', sans-serif;
    font-size: 1.15rem;
    font-weight: 700;
    color: #1a2744;
    letter-spacing: -0.02em;
}
.nav-firm {
    font-weight: 400;
    color: #86868b;
    font-size: 0.85rem;
    margin-left: 8px;
}
.nav-spacer { min-width: 150px; }

/* Case list cards */
.case-card {
    background: #fff;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 14px 16px;
    margin-bottom: 10px;
    cursor: pointer;
    transition: border-color 0.15s, box-shadow 0.15s;
}
.case-card:hover {
    border-color: #4a7ddb;
    box-shadow: 0 2px 8px rgba(74, 125, 219, 0.10);
}
.case-card-selected {
    border-color: #1a2744;
    box-shadow: 0 2px 8px rgba(26, 39, 68, 0.12);
}
.case-card-name {
    font-weight: 700;
    color: #1a2744;
    font-size: 0.95rem;
    margin-bottom: 2px;
}
.case-card-meta {
    font-size: 0.78rem;
    color: #5a6a85;
    margin-bottom: 6px;
}

/* Case type badge */
.case-type-badge {
    display: inline-block;
    padding: 2px 9px;
    font-size: 0.7rem;
    font-weight: 600;
    background: #e8f0fe;
    color: #1a73e8;
    border-radius: 12px;
    margin-right: 6px;
}

/* Progress bar (custom) */
.progress-track {
    background: #e9ecef;
    border-radius: 6px;
    height: 8px;
    overflow: hidden;
    margin: 6px 0 4px 0;
}
.progress-fill {
    height: 100%;
    border-radius: 6px;
    background: #1a2744;
    transition: width 0.3s ease;
}
.progress-text {
    font-size: 0.75rem;
    font-weight: 600;
    color: #5a6a85;
}

/* Category headers */
.cat-header {
    display: flex;
    align-items: center;
    gap: 8px;
    margin: 20px 0 10px 0;
    padding-bottom: 6px;
    border-bottom: 2px solid #e2e8f0;
}
.cat-header-label {
    font-weight: 700;
    font-size: 0.88rem;
    color: #1a2744;
    text-transform: uppercase;
    letter-spacing: 0.04em;
}
.cat-header-count {
    font-size: 0.75rem;
    font-weight: 600;
    color: #5a6a85;
    background: #f0f4fa;
    padding: 2px 8px;
    border-radius: 10px;
}

/* Checklist item rows */
.item-row {
    display: flex;
    align-items: center;
    padding: 8px 12px;
    border-radius: 8px;
    margin-bottom: 4px;
    gap: 10px;
}
.item-row-overdue {
    background: #fef2f2;
    border-left: 3px solid #dc3545;
}
.item-row-due-soon {
    background: #fffbeb;
    border-left: 3px solid #ffc107;
}
.item-row-completed {
    background: #f0fdf4;
}

/* Deadline labels */
.dl-overdue {
    font-size: 0.75rem;
    font-weight: 600;
    color: #dc3545;
    background: #fef2f2;
    padding: 2px 8px;
    border-radius: 4px;
}
.dl-due-soon {
    font-size: 0.75rem;
    font-weight: 600;
    color: #92400e;
    background: #fef3c7;
    padding: 2px 8px;
    border-radius: 4px;
}
.dl-on-track {
    font-size: 0.75rem;
    color: #5a6a85;
}
.dl-completed {
    font-size: 0.75rem;
    font-weight: 600;
    color: #198754;
}

/* Case detail header */
.detail-header {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 18px 22px;
    margin-bottom: 16px;
}
.detail-name {
    font-size: 1.3rem;
    font-weight: 800;
    color: #1a2744;
    letter-spacing: -0.02em;
    margin-bottom: 4px;
}
.detail-meta {
    font-size: 0.85rem;
    color: #5a6a85;
    line-height: 1.6;
}
.detail-meta strong {
    color: #334155;
}

/* Empty states */
.empty-state {
    text-align: center;
    padding: 60px 20px;
    color: #5a6a85;
}
.empty-state-icon {
    font-size: 2.5rem;
    margin-bottom: 12px;
    opacity: 0.4;
}
.empty-state-title {
    font-size: 1.1rem;
    font-weight: 700;
    color: #1a2744;
    margin-bottom: 6px;
}
.empty-state-desc {
    font-size: 0.88rem;
}

/* Status badge */
.status-active {
    display: inline-block;
    padding: 2px 10px;
    font-size: 0.72rem;
    font-weight: 600;
    background: #dcfce7;
    color: #166534;
    border-radius: 12px;
}
.status-completed {
    display: inline-block;
    padding: 2px 10px;
    font-size: 0.72rem;
    font-weight: 600;
    background: #e0e7ff;
    color: #3730a3;
    border-radius: 12px;
}
</style>
""",
    unsafe_allow_html=True,
)

from shared.auth import require_auth, render_logout
require_auth()

# ── Navigation bar ───────────────────────────────────────────────────────────

st.markdown(
    """
<div class="nav-bar">
    <a href="http://localhost:8502" class="nav-back">&#8592; Staff Dashboard</a>
    <div class="nav-title">Case Checklist<span class="nav-firm">&mdash; O'Brien Immigration Law</span></div>
    <div class="nav-spacer"></div>
</div>
""",
    unsafe_allow_html=True,
)
render_logout()

render_client_banner()
if render_tool_help:
    render_tool_help("case-checklist")
if render_feedback_button:
    render_feedback_button("case-checklist")

# ── Session state defaults ───────────────────────────────────────────────────

_DEFAULTS: dict[str, Any] = {
    "selected_case_id": None,
    "filter_case_type": "All",
    "filter_status": "All",
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ── Helpers ──────────────────────────────────────────────────────────────────


def _esc(text: str) -> str:
    """HTML-escape a string."""
    return html_mod.escape(str(text))


def _progress_bar_html(completed: int, total: int) -> str:
    """Render a custom progress bar as HTML."""
    pct = round((completed / total) * 100) if total > 0 else 0
    return (
        f'<div class="progress-track">'
        f'<div class="progress-fill" style="width:{pct}%"></div>'
        f'</div>'
        f'<div class="progress-text">{completed} of {total} items completed ({pct}%)</div>'
    )


def _deadline_html(deadline: str | None, is_completed: bool) -> str:
    """Render deadline status as colored HTML."""
    if is_completed:
        return '<span class="dl-completed">Done</span>'
    status = get_deadline_status(deadline)
    urgency = status["urgency"]
    label = status["label"]
    if urgency == "overdue":
        return f'<span class="dl-overdue">{_esc(label)}</span>'
    elif urgency == "due_soon":
        return f'<span class="dl-due-soon">{_esc(label)}</span>'
    elif urgency == "on_track":
        return f'<span class="dl-on-track">{_esc(label)}</span>'
    return ""


def _build_docx(case_data: dict[str, Any]) -> bytes:
    """Build a Word document checklist export and return its bytes."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, text: str, size: int = 11, bold: bool = False, italic: bool = False):
        r = para.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return r

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "CASE CHECKLIST", size=14, bold=True)

    # Case info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_parts = [case_data.get("client_name", "")]
    if case_data.get("a_number"):
        info_parts.append(f"A# {case_data['a_number']}")
    info_parts.append(case_data.get("case_type", ""))
    _run(p, " | ".join(info_parts), size=10)

    if case_data.get("attorney"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"Attorney: {case_data['attorney']}", size=10, italic=True)

    doc.add_paragraph()  # spacer

    # Progress summary
    progress = get_case_progress(case_data)
    p = doc.add_paragraph()
    _run(
        p,
        f"Progress: {progress['completed']} of {progress['total']} items completed "
        f"({progress['pct']}%)",
        size=10,
        bold=True,
    )

    doc.add_paragraph()

    # Items grouped by category
    items = case_data.get("items", [])
    categories_order = ["Filing", "Evidence", "Preparation", "Administrative"]
    grouped: dict[str, list[dict]] = {}
    for item in items:
        cat = item.get("category", "General")
        grouped.setdefault(cat, []).append(item)

    for cat in categories_order:
        cat_items = grouped.get(cat, [])
        if not cat_items:
            continue

        p = doc.add_paragraph()
        _run(p, cat.upper(), size=11, bold=True)

        for item in cat_items:
            check = "[X]" if item.get("is_completed") else "[ ]"
            title = item.get("title", "")
            deadline = item.get("deadline", "")
            notes = item.get("notes", "")

            line = f"  {check}  {title}"
            if deadline:
                dl_status = get_deadline_status(deadline)
                line += f"  --  {dl_status['label']}" if dl_status["label"] else ""
            p = doc.add_paragraph()
            _run(p, line, size=10)
            if notes:
                p = doc.add_paragraph()
                _run(p, f"        Notes: {notes}", size=9, italic=True)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, f"Generated: {date.today().strftime('%m/%d/%Y')}", size=9, italic=True)
    p = doc.add_paragraph()
    _run(p, "O'Brien Immigration Law", size=9, italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_plain_text(case_data: dict[str, Any]) -> str:
    """Build a plain-text checklist export."""
    lines: list[str] = []
    lines.append("CASE CHECKLIST")
    lines.append("=" * 60)
    lines.append(f"Client: {case_data.get('client_name', '')}")
    if case_data.get("a_number"):
        lines.append(f"A-Number: {case_data['a_number']}")
    lines.append(f"Case Type: {case_data.get('case_type', '')}")
    if case_data.get("attorney"):
        lines.append(f"Attorney: {case_data['attorney']}")

    progress = get_case_progress(case_data)
    lines.append(
        f"Progress: {progress['completed']}/{progress['total']} ({progress['pct']}%)"
    )
    lines.append("=" * 60)
    lines.append("")

    items = case_data.get("items", [])
    categories_order = ["Filing", "Evidence", "Preparation", "Administrative"]
    grouped: dict[str, list[dict]] = {}
    for item in items:
        cat = item.get("category", "General")
        grouped.setdefault(cat, []).append(item)

    for cat in categories_order:
        cat_items = grouped.get(cat, [])
        if not cat_items:
            continue
        lines.append(f"--- {cat.upper()} ---")
        for item in cat_items:
            check = "[X]" if item.get("is_completed") else "[ ]"
            title = item.get("title", "")
            deadline = item.get("deadline", "")
            dl_label = ""
            if deadline:
                dl_status = get_deadline_status(deadline)
                dl_label = f"  ({dl_status['label']})" if dl_status["label"] else ""
            lines.append(f"  {check} {title}{dl_label}")
            notes = item.get("notes", "")
            if notes:
                lines.append(f"       Notes: {notes}")
        lines.append("")

    lines.append(f"Generated: {date.today().strftime('%m/%d/%Y')}")
    lines.append("O'Brien Immigration Law")
    return "\n".join(lines)


# ── New Case Dialog ──────────────────────────────────────────────────────────


@st.dialog("New Case")
def _new_case_dialog() -> None:
    """Modal dialog for creating a new case."""
    case_type = st.selectbox("Case Type", options=CASE_TYPES)
    client_name = st.text_input("Client Name", placeholder="e.g. Maria Garcia Lopez")
    a_number = st.text_input("A-Number", placeholder="e.g. 123-456-789")
    attorney = st.text_input("Attorney", placeholder="e.g. John O'Brien")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Create Case", type="primary", use_container_width=True):
            if not client_name.strip():
                st.error("Client name is required.")
                return
            new_case = create_case(
                client_name=client_name.strip(),
                case_type=case_type,
                a_number=a_number.strip(),
                attorney=attorney.strip(),
            )
            st.session_state.selected_case_id = new_case["id"]
            st.rerun()
    with col2:
        if st.button("Cancel", use_container_width=True):
            st.rerun()


# ── Sidebar ──────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("#### Filters")

    filter_type = st.selectbox(
        "Case Type",
        options=["All"] + CASE_TYPES,
        key="filter_case_type",
    )
    filter_status = st.selectbox(
        "Status",
        options=["All", "Active", "Completed"],
        key="filter_status",
    )

    st.divider()

    # When a case is selected, show deadline-setting controls
    if st.session_state.selected_case_id:
        case_for_sidebar = load_case(st.session_state.selected_case_id)
        if case_for_sidebar:
            st.markdown("#### Set Deadlines")
            items = case_for_sidebar.get("items", [])
            incomplete_items = [
                item for item in items if not item.get("is_completed")
            ]
            if incomplete_items:
                item_labels = {
                    item["id"]: item["title"] for item in incomplete_items
                }
                selected_item_id = st.selectbox(
                    "Select item",
                    options=list(item_labels.keys()),
                    format_func=lambda x: item_labels.get(x, x),
                    label_visibility="collapsed",
                )
                if selected_item_id:
                    # Find current deadline
                    current_item = next(
                        (i for i in items if i["id"] == selected_item_id), None
                    )
                    current_dl = None
                    if current_item and current_item.get("deadline"):
                        try:
                            current_dl = datetime.strptime(
                                current_item["deadline"], "%Y-%m-%d"
                            ).date()
                        except ValueError:
                            pass

                    new_deadline = st.date_input(
                        "Deadline",
                        value=current_dl,
                        key=f"dl_{selected_item_id}",
                    )
                    if st.button("Set Deadline", use_container_width=True):
                        dl_str = new_deadline.isoformat() if new_deadline else ""
                        update_item(
                            st.session_state.selected_case_id,
                            selected_item_id,
                            {"deadline": dl_str},
                        )
                        st.rerun()
            else:
                st.caption("All items are completed.")

            st.divider()

            # Delete case
            st.markdown("#### Manage Case")
            case_status = case_for_sidebar.get("status", "Active")
            if case_status == "Active":
                if st.button("Mark Completed", use_container_width=True):
                    case_for_sidebar["status"] = "Completed"
                    save_case(case_for_sidebar)
                    st.rerun()
            else:
                if st.button("Reopen Case", use_container_width=True):
                    case_for_sidebar["status"] = "Active"
                    save_case(case_for_sidebar)
                    st.rerun()

            if st.button("Delete Case", use_container_width=True):
                delete_case(st.session_state.selected_case_id)
                st.session_state.selected_case_id = None
                st.rerun()

    render_tool_notes("case-checklist")


# ── Load cases ───────────────────────────────────────────────────────────────

all_cases = list_cases()

# Apply filters
filtered_cases = all_cases
if st.session_state.filter_case_type != "All":
    filtered_cases = [
        c for c in filtered_cases
        if c.get("case_type") == st.session_state.filter_case_type
    ]
if st.session_state.filter_status != "All":
    filtered_cases = [
        c for c in filtered_cases
        if c.get("status") == st.session_state.filter_status
    ]


# ── Two-panel layout ────────────────────────────────────────────────────────

list_col, detail_col = st.columns([2, 3], gap="large")


# ── Left panel: case list ────────────────────────────────────────────────────

with list_col:
    # Header row with New Case button
    hdr_left, hdr_right = st.columns([3, 2])
    with hdr_left:
        st.markdown(f"**Cases** ({len(filtered_cases)})")
    with hdr_right:
        if st.button("New Case", use_container_width=True, type="primary"):
            _new_case_dialog()

    if not filtered_cases:
        st.markdown(
            '<div class="empty-state">'
            '<div class="empty-state-icon">&#128203;</div>'
            '<div class="empty-state-title">No cases found</div>'
            '<div class="empty-state-desc">'
            "Create a new case to get started, or adjust your filters."
            "</div></div>",
            unsafe_allow_html=True,
        )
    else:
        for case in filtered_cases:
            case_id = case.get("id", "")
            client_name = case.get("client_name", "Unknown")
            case_type = case.get("case_type", "")
            status = case.get("status", "Active")
            progress = get_case_progress(case)
            updated = case.get("updated_at", "")[:10]

            is_selected = st.session_state.selected_case_id == case_id
            card_cls = "case-card case-card-selected" if is_selected else "case-card"
            status_cls = "status-completed" if status == "Completed" else "status-active"

            st.markdown(
                f'<div class="{card_cls}">'
                f'<div class="case-card-name">{_esc(client_name)}</div>'
                f'<div class="case-card-meta">'
                f'<span class="case-type-badge">{_esc(case_type)}</span>'
                f'<span class="{status_cls}">{_esc(status)}</span>'
                f"</div>"
                f"{_progress_bar_html(progress['completed'], progress['total'])}"
                f'<div class="case-card-meta">Updated: {_esc(updated)}</div>'
                f"</div>",
                unsafe_allow_html=True,
            )

            if st.button(
                "Select" if not is_selected else "Selected",
                key=f"sel_{case_id}",
                use_container_width=True,
                disabled=is_selected,
            ):
                st.session_state.selected_case_id = case_id
                st.rerun()


# ── Right panel: case detail ─────────────────────────────────────────────────

with detail_col:
    if not st.session_state.selected_case_id:
        st.markdown(
            '<div class="empty-state">'
            '<div class="empty-state-icon">&#9745;</div>'
            '<div class="empty-state-title">Select a case</div>'
            '<div class="empty-state-desc">'
            "Choose a case from the list to view and manage its checklist."
            "</div></div>",
            unsafe_allow_html=True,
        )
    else:
        case_data = load_case(st.session_state.selected_case_id)

        if case_data is None:
            st.warning("Could not load the selected case. It may have been deleted.")
            st.session_state.selected_case_id = None
        else:
            # ── Case header ──────────────────────────────────────────────
            progress = get_case_progress(case_data)
            status = case_data.get("status", "Active")
            status_cls = "status-completed" if status == "Completed" else "status-active"

            created_display = case_data.get("created_at", "")[:10]
            header_parts = []
            if case_data.get("a_number"):
                header_parts.append(
                    f'<strong>A-Number:</strong> {_esc(case_data["a_number"])}'
                )
            header_parts.append(
                f'<strong>Case Type:</strong> {_esc(case_data.get("case_type", ""))}'
            )
            if case_data.get("attorney"):
                header_parts.append(
                    f'<strong>Attorney:</strong> {_esc(case_data["attorney"])}'
                )
            header_parts.append(f"<strong>Created:</strong> {_esc(created_display)}")

            st.markdown(
                f'<div class="detail-header">'
                f'<div class="detail-name">{_esc(case_data.get("client_name", ""))}'
                f' <span class="{status_cls}">{_esc(status)}</span></div>'
                f'<div class="detail-meta">{" &nbsp;&bull;&nbsp; ".join(header_parts)}</div>'
                f"{_progress_bar_html(progress['completed'], progress['total'])}"
                f"</div>",
                unsafe_allow_html=True,
            )

            # ── Checklist items grouped by category ──────────────────────
            items = case_data.get("items", [])
            categories_order = ["Filing", "Evidence", "Preparation", "Administrative"]
            grouped: dict[str, list[dict[str, Any]]] = {}
            for item in items:
                cat = item.get("category", "General")
                grouped.setdefault(cat, []).append(item)

            # Also include any "General" or unlisted categories
            for cat in grouped:
                if cat not in categories_order:
                    categories_order.append(cat)

            for cat in categories_order:
                cat_items = grouped.get(cat, [])
                if not cat_items:
                    continue

                cat_done = sum(1 for i in cat_items if i.get("is_completed"))
                cat_total = len(cat_items)

                st.markdown(
                    f'<div class="cat-header">'
                    f'<span class="cat-header-label">{_esc(cat)}</span>'
                    f'<span class="cat-header-count">{cat_done}/{cat_total}</span>'
                    f"</div>",
                    unsafe_allow_html=True,
                )

                for item in cat_items:
                    item_id = item.get("id", "")
                    title = item.get("title", "")
                    is_completed = item.get("is_completed", False)
                    deadline = item.get("deadline")
                    notes = item.get("notes", "")

                    # Determine row styling based on deadline urgency
                    row_cls = "item-row"
                    if is_completed:
                        row_cls += " item-row-completed"
                    elif deadline:
                        dl_status = get_deadline_status(deadline)
                        if dl_status["urgency"] == "overdue":
                            row_cls += " item-row-overdue"
                        elif dl_status["urgency"] == "due_soon":
                            row_cls += " item-row-due-soon"

                    # Checkbox for toggling completion
                    chk_col, label_col, dl_col = st.columns([0.5, 4, 2])

                    with chk_col:
                        new_val = st.checkbox(
                            title,
                            value=is_completed,
                            key=f"chk_{case_data['id']}_{item_id}",
                            label_visibility="collapsed",
                        )
                        if new_val != is_completed:
                            update_item(
                                case_data["id"],
                                item_id,
                                {"is_completed": new_val},
                            )
                            st.rerun()

                    with label_col:
                        if is_completed:
                            st.markdown(
                                f"~~{_esc(title)}~~",
                                unsafe_allow_html=True,
                            )
                        else:
                            st.markdown(title)

                    with dl_col:
                        dl_html = _deadline_html(deadline, is_completed)
                        if dl_html:
                            st.markdown(dl_html, unsafe_allow_html=True)

                    # Notes field (collapsible)
                    if notes or not is_completed:
                        with st.expander(
                            "Notes" if not notes else f"Notes: {notes[:40]}...",
                            expanded=False,
                        ):
                            new_notes = st.text_area(
                                "Notes",
                                value=notes,
                                key=f"notes_{case_data['id']}_{item_id}",
                                label_visibility="collapsed",
                                height=68,
                            )
                            if new_notes != notes:
                                if st.button(
                                    "Save Note",
                                    key=f"savenote_{case_data['id']}_{item_id}",
                                ):
                                    update_item(
                                        case_data["id"],
                                        item_id,
                                        {"notes": new_notes},
                                    )
                                    st.rerun()

                # Add Custom Item button at bottom of each category
                with st.expander(f"Add item to {cat}", expanded=False):
                    new_title = st.text_input(
                        "Item title",
                        key=f"new_item_{case_data['id']}_{cat}",
                        placeholder=f"New {cat.lower()} item...",
                        label_visibility="collapsed",
                    )
                    new_dl = st.date_input(
                        "Deadline (optional)",
                        value=None,
                        key=f"new_item_dl_{case_data['id']}_{cat}",
                    )
                    if st.button("Add", key=f"add_item_{case_data['id']}_{cat}"):
                        if new_title.strip():
                            dl_str = new_dl.isoformat() if new_dl else None
                            add_custom_item(
                                case_data["id"],
                                new_title.strip(),
                                cat,
                                dl_str,
                            )
                            st.rerun()

            # ── Export controls ───────────────────────────────────────────
            st.divider()
            st.markdown("**Export Checklist**")
            exp_col1, exp_col2 = st.columns(2)

            with exp_col1:
                plain_text = _build_plain_text(case_data)
                safe_name = case_data.get("client_name", "case").replace(" ", "_")
                st.download_button(
                    "Download .txt",
                    data=plain_text,
                    file_name=f"Checklist_{safe_name}.txt",
                    mime="text/plain",
                    use_container_width=True,
                )

            with exp_col2:
                docx_bytes = _build_docx(case_data)
                st.download_button(
                    "Download .docx",
                    data=docx_bytes,
                    file_name=f"Checklist_{safe_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
                if st.button("Upload to Google Docs", use_container_width=True):
                    with st.spinner("Uploading to Google Docs..."):
                        try:
                            url = upload_to_google_docs(docx_bytes, f"Case Checklist - {case_data.get('client_name', 'Case')}")
                            st.session_state.google_doc_url = url
                        except Exception as e:
                            st.error(f"Upload failed: {e}")
                if st.session_state.get("google_doc_url"):
                    st.markdown(f"[Open Google Doc]({st.session_state.google_doc_url})")
