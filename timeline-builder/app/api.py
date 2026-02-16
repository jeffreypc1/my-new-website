"""FastAPI backend for the Timeline Builder tool.

Provides REST endpoints for timeline CRUD, event management within timelines,
and Word document export. The dashboard works without this server (using
local persistence), but the API enables integration with other tools.
"""

from __future__ import annotations

import io
from datetime import date

from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.events import (
    EVENT_CATEGORIES,
    TimelineEvent,
    add_event,
    delete_event,
    delete_timeline,
    list_timelines,
    load_timeline,
    new_timeline,
    save_timeline,
    update_event,
)

app = FastAPI(title="Timeline Builder API", version="1.0.0")


# ── Request / response schemas ───────────────────────────────────────────────


class TimelineCreate(BaseModel):
    case_name: str = ""
    client_name: str = ""


class EventIn(BaseModel):
    title: str
    date_text: str
    category: str
    description: str = ""
    end_date_text: str = ""


class EventUpdate(BaseModel):
    title: str | None = None
    date_text: str | None = None
    category: str | None = None
    description: str | None = None
    end_date_text: str | None = None


# ── Timeline CRUD ────────────────────────────────────────────────────────────


@app.get("/api/timelines")
def api_list_timelines() -> list[dict]:
    """Return summary info for all saved timelines."""
    return list_timelines()


@app.get("/api/timelines/{timeline_id}")
def api_get_timeline(timeline_id: str) -> dict:
    """Return the full timeline with events."""
    timeline = load_timeline(timeline_id)
    if timeline is None:
        raise HTTPException(status_code=404, detail="Timeline not found")
    return timeline


@app.post("/api/timelines", status_code=201)
def api_create_timeline(body: TimelineCreate) -> dict:
    """Create a new empty timeline and return it."""
    tl = new_timeline(case_name=body.case_name, client_name=body.client_name)
    save_timeline(tl)
    return tl


@app.delete("/api/timelines/{timeline_id}")
def api_delete_timeline(timeline_id: str) -> dict:
    """Delete a timeline by id."""
    deleted = delete_timeline(timeline_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Timeline not found")
    return {"deleted": True, "id": timeline_id}


# ── Event CRUD within a timeline ─────────────────────────────────────────────


@app.post("/api/timelines/{timeline_id}/events", status_code=201)
def api_add_event(timeline_id: str, body: EventIn) -> dict:
    """Add a new event to a timeline."""
    timeline = load_timeline(timeline_id)
    if timeline is None:
        raise HTTPException(status_code=404, detail="Timeline not found")

    if body.category not in EVENT_CATEGORIES:
        raise HTTPException(
            status_code=422,
            detail=f"Invalid category '{body.category}'. "
                   f"Must be one of: {', '.join(EVENT_CATEGORIES)}",
        )

    ev = TimelineEvent.create(
        title=body.title,
        date_text=body.date_text,
        category=body.category,
        description=body.description,
        end_date_text=body.end_date_text,
    )
    updated = add_event(timeline, ev)
    return updated


@app.put("/api/timelines/{timeline_id}/events/{event_id}")
def api_update_event(timeline_id: str, event_id: str, body: EventUpdate) -> dict:
    """Update an existing event within a timeline."""
    timeline = load_timeline(timeline_id)
    if timeline is None:
        raise HTTPException(status_code=404, detail="Timeline not found")

    # Verify event exists
    event_ids = [e.get("id") for e in timeline.get("events", [])]
    if event_id not in event_ids:
        raise HTTPException(status_code=404, detail="Event not found")

    if body.category is not None and body.category not in EVENT_CATEGORIES:
        raise HTTPException(
            status_code=422,
            detail=f"Invalid category '{body.category}'. "
                   f"Must be one of: {', '.join(EVENT_CATEGORIES)}",
        )

    updates = body.model_dump(exclude_none=True)
    updated = update_event(timeline, event_id, updates)
    return updated


@app.delete("/api/timelines/{timeline_id}/events/{event_id}")
def api_delete_event(timeline_id: str, event_id: str) -> dict:
    """Delete an event from a timeline."""
    timeline = load_timeline(timeline_id)
    if timeline is None:
        raise HTTPException(status_code=404, detail="Timeline not found")

    event_ids = [e.get("id") for e in timeline.get("events", [])]
    if event_id not in event_ids:
        raise HTTPException(status_code=404, detail="Event not found")

    updated = delete_event(timeline, event_id)
    return updated


# ── Export endpoints ─────────────────────────────────────────────────────────


@app.get("/api/timelines/{timeline_id}/export/docx")
def api_export_docx(timeline_id: str):
    """Export a timeline as a Word document with a chronological table.

    Returns the .docx file as a downloadable attachment.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    timeline = load_timeline(timeline_id)
    if timeline is None:
        raise HTTPException(status_code=404, detail="Timeline not found")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, text: str, size: int = 11, bold: bool = False, italic: bool = False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return r

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "CASE TIMELINE", size=14, bold=True)

    # Client / case info
    client = timeline.get("client_name", "")
    case = timeline.get("case_name", "")
    if client or case:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_parts = []
        if client:
            info_parts.append(f"Client: {client}")
        if case:
            info_parts.append(f"Case: {case}")
        _run(p, " | ".join(info_parts), size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"Prepared {date.today().strftime('%B %d, %Y')}", size=9, italic=True)

    doc.add_paragraph()

    # Table
    events = timeline.get("events", [])
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Date", "Event", "Category", "Description"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(header_text)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

    for ev in events:
        row = table.add_row()
        date_display = ev.get("date_text", "")
        end_date = ev.get("end_date_text", "")
        if end_date:
            date_display += f" to {end_date}"

        values = [
            date_display,
            ev.get("title", ""),
            ev.get("category", ""),
            ev.get("description", ""),
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)

    for row in table.rows:
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(1.0)
        row.cells[3].width = Inches(3.0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"Timeline_{client or 'export'}.docx".replace(" ", "_")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
