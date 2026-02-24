"""Timeline Builder — Streamlit dashboard.

Visual timeline creation for immigration case preparation. Attorneys can
build chronological timelines with approximate dates, color-coded categories,
and export to Word documents for hearing prep.
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from collections import Counter
from datetime import date
from pathlib import Path

import streamlit as st
import streamlit.components.v1
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.events import (
    CATEGORY_DESCRIPTIONS,
    EVENT_CATEGORIES,
    TimelineEvent,
    add_event,
    delete_event,
    delete_timeline,
    list_timelines,
    load_timeline,
    new_timeline,
    new_timeline_id,
    parse_approximate_date,
    parsed_date_to_display,
    save_timeline,
)
from app.doc_extractor import compile_exhibit_pdf, extract_events_from_documents, extract_pages_from_pdf

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.google_upload import upload_to_google_docs
from shared.client_banner import render_client_banner
from shared.tool_notes import render_tool_notes
try:
    from shared.draft_box import render_draft_box
except ImportError:
    render_draft_box = None
try:
    from shared.tool_help import render_tool_help
except ImportError:
    render_tool_help = None
try:
    from shared.feedback_button import render_feedback_button
except ImportError:
    render_feedback_button = None
try:
    from shared.box_folder_browser import render_box_folder_browser
    from shared.box_client import parse_folder_id as _parse_folder_id
except ImportError:
    render_box_folder_browser = None
    _parse_folder_id = None

# ── Page config ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Timeline Builder — O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ──────────────────────────────────────────────────────────────────────

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* Hide Streamlit chrome */
#MainMenu, footer,
div[data-testid="stToolbar"] { display: none !important; }

.stApp {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

/* Navigation bar */
.nav-bar {
    display: flex;
    align-items: center;
    padding: 10px 4px;
    margin: -1rem 0 1.2rem 0;
    border-bottom: 1px solid rgba(0,0,0,0.07);
}
.nav-back {
    display: flex;
    align-items: center;
    gap: 6px;
    font-family: 'Inter', sans-serif;
    font-size: 0.85rem;
    font-weight: 500;
    color: #0066CC;
    text-decoration: none;
    min-width: 150px;
}
.nav-back:hover { color: #004499; text-decoration: underline; }
.nav-title {
    flex: 1;
    text-align: center;
    font-family: 'Inter', sans-serif;
    font-size: 1.15rem;
    font-weight: 700;
    color: #1a2744;
    letter-spacing: -0.02em;
}
.nav-firm {
    font-weight: 400;
    color: #86868b;
    font-size: 0.85rem;
    margin-left: 8px;
}
.nav-spacer { min-width: 150px; }

/* Summary stats */
.stats-row {
    display: flex;
    gap: 16px;
    margin-bottom: 20px;
    flex-wrap: wrap;
}
.stat-card {
    background: #fff;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    padding: 14px 20px;
    min-width: 140px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
}
.stat-value {
    font-size: 1.5rem;
    font-weight: 700;
    color: #1a2744;
    line-height: 1.2;
}
.stat-label {
    font-size: 0.78rem;
    color: #5a6a85;
    font-weight: 500;
    margin-top: 2px;
}

/* Category legend */
.cat-legend {
    display: flex;
    gap: 16px;
    flex-wrap: wrap;
    margin-bottom: 20px;
}
.cat-item {
    display: flex;
    align-items: center;
    gap: 6px;
    font-size: 0.8rem;
    color: #5a6a85;
    font-weight: 500;
}
.cat-dot {
    width: 10px;
    height: 10px;
    border-radius: 50%;
    display: inline-block;
    flex-shrink: 0;
}
.cat-count {
    background: #f0f4fa;
    border-radius: 10px;
    padding: 1px 7px;
    font-size: 0.72rem;
    font-weight: 600;
    color: #4a7ddb;
    margin-left: 2px;
}

/* Saved toast */
.saved-toast {
    font-size: 0.8rem;
    color: #2e7d32;
    font-weight: 600;
}

/* Section header */
.section-header {
    color: #1a2744;
    font-size: 1.05rem;
    font-weight: 700;
    margin-bottom: 12px;
    padding-bottom: 8px;
    border-bottom: 2px solid #e2e8f0;
}

/* Print-friendly styles */
@media print {
    .nav-bar, .stSidebar { display: none !important; }
    .stApp { background: #fff !important; }
}
</style>
""",
    unsafe_allow_html=True,
)

from shared.auth import require_auth, render_logout
require_auth()

# ── Navigation bar ───────────────────────────────────────────────────────────

st.markdown(
    """
<div class="nav-bar">
    <a href="http://localhost:8502" class="nav-back">&#8592; Staff Dashboard</a>
    <div class="nav-title">Timeline Builder<span class="nav-firm">&mdash; O'Brien Immigration Law</span></div>
    <div class="nav-spacer"></div>
</div>
""",
    unsafe_allow_html=True,
)
render_logout()

render_client_banner()
if render_tool_help:
    render_tool_help("timeline-builder")
if render_feedback_button:
    render_feedback_button("timeline-builder")

# ── Session state defaults ───────────────────────────────────────────────────

_DEFAULTS: dict = {
    "timeline_id": None,
    "timeline": None,
    "last_saved_msg": "",
    "extracted_events": [],
    "extraction_selections": {},
    "box_docs": [],
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ── Helpers ──────────────────────────────────────────────────────────────────


def _ensure_timeline() -> dict:
    """Make sure we have a current timeline in session state."""
    if st.session_state.timeline is None:
        tl = new_timeline()
        st.session_state.timeline_id = tl["id"]
        st.session_state.timeline = tl
    return st.session_state.timeline


def _get_client_name() -> str:
    """Get the client name from SF data or the current timeline."""
    sf = st.session_state.get("sf_client")
    if sf:
        return sf.get("Name", "")
    tl = st.session_state.get("timeline")
    if tl:
        return tl.get("client_name", "")
    return ""


def _do_save() -> None:
    """Save the current timeline to disk."""
    tl = _ensure_timeline()
    tl["client_name"] = _get_client_name()
    save_timeline(tl)
    label = tl["client_name"] or "timeline"
    st.session_state.last_saved_msg = f"Saved — {label}"


def _do_new() -> None:
    """Start a fresh timeline."""
    tl = new_timeline()
    st.session_state.timeline_id = tl["id"]
    st.session_state.timeline = tl
    st.session_state.last_saved_msg = ""


def _do_load(timeline_id: str) -> None:
    """Load a timeline from disk into session state."""
    tl = load_timeline(timeline_id)
    if tl is None:
        return
    st.session_state.timeline_id = tl["id"]
    st.session_state.timeline = tl
    st.session_state.last_saved_msg = ""


def _do_delete_event(event_id: str) -> None:
    """Delete an event from the current timeline."""
    tl = _ensure_timeline()
    delete_event(tl, event_id)


_STATS_CSS = """\
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
body { margin: 0; padding: 0; font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif; background: transparent; }
.stats-row { display: flex; gap: 16px; margin-bottom: 12px; flex-wrap: wrap; }
.stat-card { background: #fff; border: 1px solid #e2e8f0; border-radius: 8px; padding: 14px 20px; min-width: 140px; box-shadow: 0 1px 3px rgba(0,0,0,0.04); }
.stat-value { font-size: 1.5rem; font-weight: 700; color: #1a2744; line-height: 1.2; }
.stat-label { font-size: 0.78rem; color: #5a6a85; font-weight: 500; margin-top: 2px; }
.cat-legend { display: flex; gap: 16px; flex-wrap: wrap; }
.cat-item { display: flex; align-items: center; gap: 6px; font-size: 0.8rem; color: #5a6a85; font-weight: 500; }
.cat-dot { width: 10px; height: 10px; border-radius: 50%; display: inline-block; flex-shrink: 0; }
.cat-count { background: #f0f4fa; border-radius: 10px; padding: 1px 7px; font-size: 0.72rem; font-weight: 600; color: #4a7ddb; margin-left: 2px; }
"""


def _build_stats_html(events: list[dict]) -> str:
    """Build summary statistics HTML."""
    if not events:
        return ""

    total = len(events)

    # Date range
    dates = [e.get("parsed_date", "9999-99-99") for e in events]
    valid_dates = [d for d in dates if d != "9999-99-99"]
    if valid_dates:
        earliest = parsed_date_to_display(min(valid_dates))
        latest = parsed_date_to_display(max(valid_dates))
        date_range = f"{earliest} &mdash; {latest}"
    else:
        date_range = "No dates"

    # Category counts
    cat_counts = Counter(e.get("category", "Personal") for e in events)

    parts: list[str] = ['<div class="stats-row">']
    parts.append(
        f'<div class="stat-card">'
        f'<div class="stat-value">{total}</div>'
        f'<div class="stat-label">Total Events</div>'
        f"</div>"
    )
    parts.append(
        f'<div class="stat-card">'
        f'<div class="stat-value" style="font-size:1rem;">{date_range}</div>'
        f'<div class="stat-label">Date Range</div>'
        f"</div>"
    )
    parts.append("</div>")

    # Category legend with counts
    parts.append('<div class="cat-legend">')
    for cat, color in EVENT_CATEGORIES.items():
        count = cat_counts.get(cat, 0)
        if count > 0:
            parts.append(
                f'<div class="cat-item">'
                f'<span class="cat-dot" style="background:{color};"></span>'
                f"{html_mod.escape(cat)}"
                f'<span class="cat-count">{count}</span>'
                f"</div>"
            )
    parts.append("</div>")

    return f"<style>{_STATS_CSS}</style>" + "".join(parts)


def _build_docx(timeline: dict) -> bytes:
    """Build a Word document with a formatted event table."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, text: str, size: int = 11, bold: bool = False, italic: bool = False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return r

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "CASE TIMELINE", size=14, bold=True)

    # Client / case info
    client = timeline.get("client_name", "")
    case = timeline.get("case_name", "")
    if client or case:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_parts = []
        if client:
            info_parts.append(f"Client: {client}")
        if case:
            info_parts.append(f"Case: {case}")
        _run(p, " | ".join(info_parts), size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"Prepared {date.today().strftime('%m/%d/%Y')}", size=9, italic=True)

    doc.add_paragraph()  # spacer

    # Table
    events = timeline.get("events", [])
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    headers = ["Date", "Event", "Category", "Description"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(header_text)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

    # Data rows
    for ev in events:
        row = table.add_row()
        date_display = ev.get("date_text", "")
        end_date = ev.get("end_date_text", "")
        if end_date:
            date_display += f" to {end_date}"

        values = [
            date_display,
            ev.get("title", ""),
            ev.get("category", ""),
            ev.get("description", ""),
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(1.0)
        row.cells[3].width = Inches(3.0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_plain_text(timeline: dict) -> str:
    """Build a plain-text export of the timeline."""
    lines: list[str] = []
    client = timeline.get("client_name", "")
    case = timeline.get("case_name", "")

    lines.append("CASE TIMELINE")
    lines.append("=" * 60)
    if client:
        lines.append(f"Client: {client}")
    if case:
        lines.append(f"Case:   {case}")
    lines.append(f"Date:   {date.today().strftime('%m/%d/%Y')}")
    lines.append("=" * 60)
    lines.append("")

    events = timeline.get("events", [])
    if not events:
        lines.append("(No events)")
    else:
        for ev in events:
            date_text = ev.get("date_text", "Unknown")
            end_date = ev.get("end_date_text", "")
            if end_date:
                date_text += f" to {end_date}"
            title = ev.get("title", "")
            cat = ev.get("category", "")
            desc = ev.get("description", "")

            lines.append(f"[{date_text}] [{cat}]")
            lines.append(f"  {title}")
            if desc:
                lines.append(f"  {desc}")
            lines.append("")

    lines.append("=" * 60)
    lines.append(f"Total events: {len(events)}")
    return "\n".join(lines)


# ── Sidebar ──────────────────────────────────────────────────────────────────

tl = _ensure_timeline()
client_name = _get_client_name()

with st.sidebar:
    # Timeline management
    st.markdown("#### Timelines")
    btn_cols = st.columns(2)
    with btn_cols[0]:
        if st.button("New", use_container_width=True):
            _do_new()
            st.rerun()
    with btn_cols[1]:
        save_clicked = st.button("Save", use_container_width=True, type="primary")

    saved_timelines = list_timelines()
    if saved_timelines:
        labels_map = {
            t["id"]: f"{t['client_name'] or 'Unnamed'} — {t['case_name'] or 'No case'} ({t['event_count']} events)"
            for t in saved_timelines
        }
        tl_ids = list(labels_map.keys())
        selected_tl = st.selectbox(
            "Load a saved timeline",
            options=[""] + tl_ids,
            format_func=lambda x: labels_map.get(x, "Select..."),
            label_visibility="collapsed",
        )
        load_cols = st.columns(2)
        with load_cols[0]:
            if selected_tl and st.button("Load", use_container_width=True):
                _do_load(selected_tl)
                st.rerun()
        with load_cols[1]:
            if selected_tl and st.button("Delete", use_container_width=True):
                delete_timeline(selected_tl)
                if st.session_state.timeline_id == selected_tl:
                    _do_new()
                st.rerun()

    if st.session_state.last_saved_msg:
        st.markdown(
            f'<div class="saved-toast">{html_mod.escape(st.session_state.last_saved_msg)}</div>',
            unsafe_allow_html=True,
        )

    st.divider()

    # Add Event form
    st.markdown("#### Add Event")

    event_title = st.text_input(
        "Event Title",
        key="inp_event_title",
        placeholder="e.g. Fled home country",
    )
    event_date = st.text_input(
        "Date",
        key="inp_event_date",
        placeholder='e.g. "March 2019", "Summer 2020"',
    )
    event_end_date = st.text_input(
        "End Date (optional, for ranges)",
        key="inp_event_end_date",
        placeholder='e.g. "June 2019"',
    )
    event_category = st.selectbox(
        "Category",
        options=list(EVENT_CATEGORIES.keys()),
        format_func=lambda x: f"{x} — {CATEGORY_DESCRIPTIONS.get(x, '')}",
        key="inp_event_category",
    )
    event_description = st.text_area(
        "Description",
        key="inp_event_desc",
        height=100,
        placeholder="Details about this event...",
    )

    add_disabled = not (event_title and event_date)
    if st.button("Add Event", type="primary", use_container_width=True, disabled=add_disabled):
        new_ev = TimelineEvent.create(
            title=event_title,
            date_text=event_date,
            category=event_category,
            description=event_description,
            end_date_text=event_end_date,
        )
        add_event(tl, new_ev)
        # Clear form fields
        for k in ("inp_event_title", "inp_event_date", "inp_event_end_date", "inp_event_desc"):
            if k in st.session_state:
                del st.session_state[k]
        st.rerun()

    if event_date:
        parsed = parse_approximate_date(event_date)
        display = parsed_date_to_display(parsed)
        st.caption(f"Parsed as: {display}")

    st.divider()

    # Box folder browser — picker mode for extracting PDFs
    if render_box_folder_browser and _parse_folder_id:
        _sf = st.session_state.get("sf_client")
        _box_raw = (_sf.get("Box_Folder_Id__c", "") or "") if _sf else ""
        if _box_raw:

            def _on_box_select(files: list[dict]) -> None:
                """Download selected PDFs from Box for extraction."""
                from shared.box_client import get_file_content

                box_docs: list[dict] = st.session_state.get("box_docs", [])
                existing_ids = {d.get("box_id") for d in box_docs}

                with st.spinner("Downloading documents from Box..."):
                    for f in files:
                        if f["id"] in existing_ids:
                            continue
                        if (f.get("extension", "") or "").lower() != "pdf":
                            continue
                        pdf_bytes = get_file_content(f["id"])
                        pages = extract_pages_from_pdf(pdf_bytes)
                        box_docs.append({
                            "name": f["name"],
                            "pages": pages,
                            "box_id": f["id"],
                            "pdf_bytes": pdf_bytes,
                        })

                st.session_state["box_docs"] = box_docs

            already_ids = {d.get("box_id") for d in st.session_state.get("box_docs", [])}

            render_box_folder_browser(
                _parse_folder_id(_box_raw),
                mode="picker",
                on_select=_on_box_select,
                already_selected_ids=already_ids,
                key_prefix="_tl_box",
                header_label="Client Documents",
                add_button_label="Add to Extraction",
            )

    render_tool_notes("timeline-builder")


# ── Handle save (after sidebar widgets render) ───────────────────────────────

if save_clicked:
    _do_save()
    st.rerun()


# ── Main area ────────────────────────────────────────────────────────────────

tab_timeline, tab_extract = st.tabs(["Timeline", "Extract from Documents"])

# ── Timeline tab ─────────────────────────────────────────────────────────────

with tab_timeline:
    events = tl.get("events", [])

    # Summary stats
    stats_html = _build_stats_html(events)
    if stats_html:
        st.components.v1.html(stats_html, height=100)

    # Timeline events as native Streamlit cards with inline delete
    if not events:
        st.info("No events yet. Use the sidebar form or Extract tab to add events.")
    else:
        import re as _re

        for ev in events:
            event_id = ev.get("id", "")
            title = ev.get("title", "Untitled")
            date_text = ev.get("date_text", "")
            end_date = ev.get("end_date_text", "")
            desc = ev.get("description", "")
            cat = ev.get("category", "Personal")
            color = EVENT_CATEGORIES.get(cat, "#6c757d")

            with st.container(border=True):
                cols = st.columns([12, 1])
                with cols[0]:
                    # Date + optional range + category badge
                    date_html = f'<span style="color:{color}; font-weight:600; font-size:0.85rem;">{html_mod.escape(date_text)}</span>'
                    if end_date:
                        date_html += f'<span style="color:#86868b; font-size:0.78rem; margin-left:6px;">to {html_mod.escape(end_date)}</span>'
                    date_html += (
                        f'<span style="display:inline-block; background:{color}; color:#fff; '
                        f'padding:1px 8px; font-size:0.7rem; font-weight:600; border-radius:12px; '
                        f'margin-left:10px;">{html_mod.escape(cat)}</span>'
                    )
                    st.markdown(date_html, unsafe_allow_html=True)

                    # Title
                    st.markdown(f"**{html_mod.escape(title)}**")

                    # Description + source citation
                    if desc:
                        # Show source citation separately if present
                        source_match = _re.search(r"\[([^\]]+\.pdf[^\]]*)\]", desc, _re.IGNORECASE)
                        if source_match:
                            source_label = source_match.group(0)
                            desc_clean = desc.replace(source_label, "").strip()
                            display_parts = []
                            if desc_clean:
                                display_parts.append(desc_clean)
                            display_parts.append(f"*{html_mod.escape(source_label)}*")
                            st.caption(" ".join(display_parts))
                        else:
                            st.caption(desc)
                with cols[1]:
                    if st.button("✕", key=f"del_{event_id}"):
                        tl["events"] = [e for e in tl.get("events", []) if e.get("id") != event_id]
                        save_timeline(tl)
                        st.rerun()

    # Draft Box
    if render_draft_box is not None:
        plain_text_for_draft = _build_plain_text(tl)
        render_draft_box("timeline-builder", {
            "document_type": "timeline narrative",
            "client_name": client_name,
            "case_id": st.session_state.get("timeline_id", ""),
            "content": plain_text_for_draft,
        })

    # Export controls
    st.markdown("---")
    exp_cols = st.columns(2)

    with exp_cols[0]:
        plain_text = _build_plain_text(tl)
        file_label = (client_name or "Timeline").replace(" ", "_")
        st.download_button(
            "Download .txt",
            data=plain_text,
            file_name=f"Timeline_{file_label}.txt",
            mime="text/plain",
            use_container_width=True,
            disabled=not events,
        )

    with exp_cols[1]:
        if events:
            docx_bytes = _build_docx(tl)
            st.download_button(
                "Download .docx",
                data=docx_bytes,
                file_name=f"Timeline_{file_label}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            if st.button("Upload to Google Docs", use_container_width=True):
                with st.spinner("Uploading to Google Docs..."):
                    try:
                        url = upload_to_google_docs(docx_bytes, f"Timeline - {client_name or 'Timeline'}")
                        st.session_state.google_doc_url = url
                    except Exception as e:
                        st.error(f"Upload failed: {e}")
            if st.session_state.get("google_doc_url"):
                st.markdown(f"[Open Google Doc]({st.session_state.google_doc_url})")
        else:
            st.button("Download .docx", disabled=True, use_container_width=True)

# ── Extract from Documents tab ───────────────────────────────────────────────

with tab_extract:
    # Box-selected docs (from sidebar picker)
    box_docs: list[dict] = st.session_state.get("box_docs", [])
    if box_docs:
        st.markdown("**Documents from Box:**")
        for i, bd in enumerate(box_docs):
            bd_cols = st.columns([6, 1])
            with bd_cols[0]:
                st.caption(f"{bd['name']} — {len(bd['pages'])} page{'s' if len(bd['pages']) != 1 else ''}")
            with bd_cols[1]:
                if st.button("Remove", key=f"rm_box_{i}", use_container_width=True):
                    st.session_state["box_docs"].pop(i)
                    st.rerun()

    # Manual file uploads
    uploaded_files = st.file_uploader(
        "Upload documents (PDF)",
        type=["pdf"],
        accept_multiple_files=True,
        key="extract_uploader",
    )

    # Build combined document list
    all_doc_data: list[dict] = []

    # Add Box docs (already parsed, with pdf_bytes)
    for bd in box_docs:
        entry: dict = {"name": bd["name"], "pages": bd["pages"]}
        if "pdf_bytes" in bd:
            entry["pdf_bytes"] = bd["pdf_bytes"]
        all_doc_data.append(entry)

    # Add uploaded docs
    if uploaded_files:
        st.markdown("**Uploaded documents:**")
        for uf in uploaded_files:
            pdf_bytes = uf.read()
            uf.seek(0)
            pages = extract_pages_from_pdf(pdf_bytes)
            all_doc_data.append({"name": uf.name, "pages": pages, "pdf_bytes": pdf_bytes})
            st.caption(f"{uf.name} — {len(pages)} page{'s' if len(pages) != 1 else ''}")

    if all_doc_data:
        # Cache combined list for extraction
        st.session_state["_extract_doc_data"] = all_doc_data

        # Extract button
        if st.button(
            "Extract Timeline Events",
            type="primary",
            use_container_width=True,
        ):
            st.session_state["_run_extraction"] = True
            st.rerun()

    # Run extraction if flagged (runs on the rerun after button click)
    if st.session_state.get("_run_extraction"):
        st.session_state["_run_extraction"] = False
        doc_data = st.session_state.get("_extract_doc_data", [])
        if doc_data:
            categories = list(EVENT_CATEGORIES.keys())
            with st.spinner("Extracting timeline events from documents..."):
                try:
                    extracted = extract_events_from_documents(
                        doc_data, categories,
                    )
                except Exception as exc:
                    st.error(f"Extraction failed: {exc}")
                    extracted = []

            if extracted:
                st.success(f"Extracted {len(extracted)} events.")
                st.session_state.extracted_events = extracted
                st.session_state.extraction_selections = {
                    i: True for i in range(len(extracted))
                }
            else:
                st.warning("No events were extracted from the documents.")

    # Show extracted events for review
    extracted_events = st.session_state.get("extracted_events", [])
    if extracted_events:
        st.markdown(f"#### Extracted Events ({len(extracted_events)})")

        # Select all / deselect all
        sel_cols = st.columns(3)
        with sel_cols[0]:
            if st.button("Select All"):
                st.session_state.extraction_selections = {
                    i: True for i in range(len(extracted_events))
                }
                st.rerun()
        with sel_cols[1]:
            if st.button("Deselect All"):
                st.session_state.extraction_selections = {
                    i: False for i in range(len(extracted_events))
                }
                st.rerun()

        selections = st.session_state.get("extraction_selections", {})

        for idx, ev in enumerate(extracted_events):
            cat = ev.get("category", "Personal")
            color = EVENT_CATEGORIES.get(cat, "#6c757d")
            checked = selections.get(idx, True)

            cols = st.columns([0.5, 1.5, 2, 3, 2])
            with cols[0]:
                new_val = st.checkbox(
                    "sel",
                    value=checked,
                    key=f"ext_sel_{idx}",
                    label_visibility="collapsed",
                )
                if new_val != checked:
                    st.session_state.extraction_selections[idx] = new_val
            with cols[1]:
                st.markdown(f"**{html_mod.escape(ev.get('date', ''))}**")
            with cols[2]:
                st.markdown(
                    f'<span style="color:{color}; font-weight:600;">{html_mod.escape(cat)}</span> '
                    f"&mdash; {html_mod.escape(ev.get('title', ''))}",
                    unsafe_allow_html=True,
                )
            with cols[3]:
                desc = ev.get("description", "")
                st.caption(desc[:120] + ("..." if len(desc) > 120 else ""))
            with cols[4]:
                st.caption(ev.get("source", ""))

        # Add selected events to timeline
        selected_count = sum(1 for v in selections.values() if v)
        st.markdown("---")
        if st.button(
            f"Add {selected_count} Selected Events to Timeline",
            type="primary",
            disabled=selected_count == 0,
            use_container_width=True,
        ):
            added = 0
            for idx, ev in enumerate(extracted_events):
                if not selections.get(idx, False):
                    continue
                new_ev = TimelineEvent.create(
                    title=ev.get("title", ""),
                    date_text=ev.get("date", ""),
                    category=ev.get("category", "Personal"),
                    description=ev.get("description", ""),
                )
                add_event(tl, new_ev)
                added += 1
            st.session_state.extracted_events = []
            st.session_state.extraction_selections = {}
            st.success(f"Added {added} events to timeline.")
            st.rerun()

    # ── Compile Exhibit PDF ──────────────────────────────────────────────
    # Available whenever source docs with pdf_bytes exist
    docs_with_bytes = [
        d for d in st.session_state.get("_extract_doc_data", [])
        if d.get("pdf_bytes")
    ]
    if docs_with_bytes:
        st.markdown("---")
        st.markdown("#### Compile Exhibit PDF")
        st.caption(
            "Merge all source PDFs into a single exhibit package with a "
            "Table of Contents and separator pages so pin citations match."
        )
        if st.button("Compile Exhibit PDF", type="primary", use_container_width=True):
            with st.spinner("Compiling exhibit PDF..."):
                try:
                    pdf_bytes_out, page_map = compile_exhibit_pdf(docs_with_bytes)
                    st.session_state["_exhibit_pdf"] = pdf_bytes_out
                    st.session_state["_exhibit_page_map"] = page_map
                except Exception as exc:
                    st.error(f"Compilation failed: {exc}")

    # Show results if compiled
    exhibit_pdf = st.session_state.get("_exhibit_pdf")
    exhibit_map = st.session_state.get("_exhibit_page_map")
    if exhibit_pdf and exhibit_map:
        st.markdown("**Page Map:**")
        map_rows = []
        for pm in exhibit_map:
            map_rows.append({
                "Document": pm["name"],
                "Pages in Exhibit": f"{pm['start_page']}–{pm['end_page']}",
                "Document Pages": str(pm["num_pages"]),
            })
        st.dataframe(map_rows, use_container_width=True, hide_index=True)

        file_label = (client_name or "Exhibit").replace(" ", "_")
        st.download_button(
            "Download Exhibit PDF",
            data=exhibit_pdf,
            file_name=f"Exhibit_{file_label}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
