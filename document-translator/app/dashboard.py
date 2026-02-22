"""Document Translator — Streamlit dashboard.

Upload documents (PDF, DOCX, images), auto-detect language, translate
to English or another target language via Google Translate API, review
and edit translations, and export to .txt, .docx, or Google Docs.
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from pathlib import Path

from dotenv import load_dotenv

load_dotenv(Path(__file__).resolve().parent.parent / ".env")

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.drafts import delete_draft, list_drafts, load_draft, new_draft_id, save_draft
from app.translator import (
    CERTIFICATE_TYPES,
    LANGUAGES,
    TARGET_LANGUAGES,
    build_certificate,
    certification_header,
    detect_language,
    extract_text,
    language_name,
    translate_paragraphs,
)

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.google_upload import upload_to_google_docs
from shared.client_banner import render_client_banner
from shared.tool_notes import render_tool_notes
try:
    from shared.tool_help import render_tool_help
except ImportError:
    render_tool_help = None
try:
    from shared.feedback_button import render_feedback_button
except ImportError:
    render_feedback_button = None
try:
    from shared.box_folder_browser import render_box_folder_browser
    from shared.box_client import parse_folder_id as _parse_folder_id
except ImportError:
    render_box_folder_browser = None
    _parse_folder_id = None

# -- Page config --------------------------------------------------------------

st.set_page_config(
    page_title="Document Translator — O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -- CSS ----------------------------------------------------------------------

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* Hide Streamlit chrome */
#MainMenu, footer,
div[data-testid="stToolbar"] { display: none !important; }

.stApp {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

/* Navigation bar */
.nav-bar {
    display: flex;
    align-items: center;
    padding: 10px 4px;
    margin: -1rem 0 1.2rem 0;
    border-bottom: 1px solid rgba(0,0,0,0.07);
}
.nav-back {
    display: flex;
    align-items: center;
    gap: 6px;
    font-family: 'Inter', sans-serif;
    font-size: 0.85rem;
    font-weight: 500;
    color: #0066CC;
    text-decoration: none;
    min-width: 150px;
}
.nav-back:hover { color: #004499; text-decoration: underline; }
.nav-title {
    flex: 1;
    text-align: center;
    font-family: 'Inter', sans-serif;
    font-size: 1.15rem;
    font-weight: 700;
    color: #1a2744;
    letter-spacing: -0.02em;
}
.nav-firm {
    font-weight: 400;
    color: #86868b;
    font-size: 0.85rem;
    margin-left: 8px;
}
.nav-spacer { min-width: 150px; }

/* Header */
.dt-header {
    color: #1a2744;
    font-size: 1.8rem;
    font-weight: 800;
    letter-spacing: -0.03em;
    margin-bottom: 0;
}
.dt-sub {
    color: #5a6a85;
    font-size: 0.95rem;
    margin-bottom: 0.5rem;
}

/* Section labels */
.section-label {
    font-size: 0.78rem;
    font-weight: 600;
    color: #5a6a85;
    text-transform: uppercase;
    letter-spacing: 0.04em;
    margin-bottom: 4px;
    margin-top: 12px;
}

/* Language badge */
.lang-badge {
    display: inline-block;
    padding: 4px 14px;
    font-size: 0.8rem;
    font-weight: 600;
    background: #e8f0fe;
    color: #1a73e8;
    border-radius: 12px;
    margin: 4px 0 8px;
}

/* Certification notice */
.cert-notice {
    background: #fff8f0;
    border: 1px solid #f0d0a0;
    border-radius: 6px;
    padding: 10px 14px;
    font-size: 0.82rem;
    color: #8b4513;
    font-style: italic;
    margin-bottom: 12px;
}

/* Preview panel */
.preview-panel {
    font-family: 'Times New Roman', Times, serif;
    font-size: 0.85rem;
    line-height: 1.7;
    background: #fff;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    padding: 28px 32px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
    max-height: 75vh;
    overflow-y: auto;
    white-space: pre-wrap;
    word-wrap: break-word;
}

/* Side-by-side original */
.original-text {
    font-size: 0.82rem;
    color: #666;
    font-style: italic;
    border-left: 3px solid #e2e8f0;
    padding-left: 10px;
    margin-bottom: 4px;
}

/* Draft badge */
.draft-badge {
    display: inline-block;
    padding: 3px 10px;
    font-size: 0.7rem;
    font-weight: 600;
    background: #e8f0fe;
    color: #1a73e8;
    border-radius: 12px;
    margin-bottom: 8px;
}

/* Saved confirmation */
.saved-toast {
    font-size: 0.8rem;
    color: #2e7d32;
    font-weight: 600;
}

/* Paragraph number */
.para-num {
    font-size: 0.7rem;
    color: #999;
    font-weight: 600;
    margin-bottom: 2px;
}
</style>
""",
    unsafe_allow_html=True,
)

from shared.auth import require_auth, render_logout
require_auth()

# -- Navigation bar -----------------------------------------------------------

st.markdown(
    """
<div class="nav-bar">
    <a href="http://localhost:8502" class="nav-back">&#8592; Staff Dashboard</a>
    <div class="nav-title">Document Translator<span class="nav-firm">&mdash; O'Brien Immigration Law</span></div>
    <div class="nav-spacer"></div>
</div>
""",
    unsafe_allow_html=True,
)
render_logout()

sf_record = render_client_banner()
if render_tool_help:
    render_tool_help("document-translator")
if render_feedback_button:
    render_feedback_button("document-translator")
if sf_record and not st.session_state.get("inp_client_name"):
    st.session_state.inp_client_name = sf_record.get("Name", "")

# -- Session state defaults ---------------------------------------------------

_DEFAULTS: dict = {
    "draft_id": None,
    "last_saved_msg": "",
    "paragraphs": [],
    "translated": [],
    "source_filename": "",
    "detected_lang": "",
    "detected_confidence": 0.0,
    "translation_done": False,
    "google_doc_url": "",
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state.draft_id is None:
    st.session_state.draft_id = new_draft_id()


# -- Helpers ------------------------------------------------------------------


def _do_save(
    client_name: str,
    target_lang: str,
    notes: str,
    include_original: bool,
    export_format: str,
    certificate_type: str,
    translator_info: dict,
    client_pronoun: str,
    show_disclaimer: bool,
) -> None:
    """Save the current state as a draft."""
    # Build paragraphs list with any staff edits
    paras = []
    for i, t in enumerate(st.session_state.get("translated", [])):
        edited = st.session_state.get(f"edit_{i}", t.get("translated", ""))
        paras.append(
            {
                "original": t.get("original", ""),
                "translated": edited,
                "detected_lang": t.get("detected_lang", ""),
            }
        )

    save_draft(
        draft_id=st.session_state.draft_id,
        client_name=client_name,
        source_filename=st.session_state.get("source_filename", ""),
        source_lang=st.session_state.get("detected_lang", ""),
        target_lang=target_lang,
        paragraphs=paras,
        notes=notes,
        include_original=include_original,
        export_format=export_format,
        certificate_type=certificate_type,
        translator_info=translator_info,
        client_pronoun=client_pronoun,
        show_disclaimer=show_disclaimer,
    )
    name = client_name or "draft"
    st.session_state.last_saved_msg = f"Saved -- {name}"


def _do_load(draft_id: str) -> None:
    """Load a draft into session state."""
    draft = load_draft(draft_id)
    if not draft:
        return
    st.session_state.draft_id = draft["id"]
    st.session_state.source_filename = draft.get("source_filename", "")
    st.session_state.detected_lang = draft.get("source_lang", "")
    st.session_state.inp_client_name = draft.get("client_name", "")
    st.session_state.inp_notes = draft.get("notes", "")
    st.session_state.inp_include_original = draft.get("include_original", True)

    paras = draft.get("paragraphs", [])
    st.session_state.translated = paras
    st.session_state.translation_done = bool(paras)
    st.session_state.paragraphs = [p.get("original", "") for p in paras]

    # Set target language
    tl = draft.get("target_lang", "English")
    st.session_state.inp_target_lang = tl

    # Restore new fields
    st.session_state.inp_export_format = draft.get("export_format", "Translation only")
    st.session_state.inp_cert_type = draft.get("certificate_type", "None")
    st.session_state.inp_client_pronoun = draft.get("client_pronoun", "they")
    st.session_state.inp_show_disclaimer = draft.get("show_disclaimer", True)

    ti = draft.get("translator_info", {})
    st.session_state.inp_translator_name = ti.get("name", "")
    st.session_state.inp_translator_address = ti.get("address", "")
    st.session_state.inp_translator_phone = ti.get("phone", "")

    # Restore edited text
    for i, p in enumerate(paras):
        st.session_state[f"edit_{i}"] = p.get("translated", "")


def _do_new() -> None:
    """Start a fresh translation."""
    st.session_state.draft_id = new_draft_id()
    st.session_state.last_saved_msg = ""
    st.session_state.paragraphs = []
    st.session_state.translated = []
    st.session_state.source_filename = ""
    st.session_state.detected_lang = ""
    st.session_state.detected_confidence = 0.0
    st.session_state.translation_done = False
    st.session_state.google_doc_url = ""
    # Clean up edit keys
    for key in list(st.session_state.keys()):
        if key.startswith("edit_"):
            del st.session_state[key]
    for k in (
        "inp_client_name", "inp_notes", "inp_export_format", "inp_cert_type",
        "inp_translator_name", "inp_translator_address", "inp_translator_phone",
        "inp_client_pronoun", "inp_show_disclaimer", "sf_client",
    ):
        if k in st.session_state:
            del st.session_state[k]


def _get_edited_paragraphs() -> list[dict]:
    """Return translated paragraphs with any staff edits applied."""
    result = []
    for i, t in enumerate(st.session_state.get("translated", [])):
        edited = st.session_state.get(f"edit_{i}", t.get("translated", ""))
        result.append(
            {
                "original": t.get("original", ""),
                "translated": edited,
                "detected_lang": t.get("detected_lang", ""),
            }
        )
    return result


def _build_plain_text(
    paras: list[dict], include_original: bool, cert_text: str = "", show_disclaimer: bool = True
) -> str:
    """Build a plain-text export of the translation."""
    lines = []
    if show_disclaimer:
        lines.extend([certification_header(), ""])

    if st.session_state.get("source_filename"):
        lines.append(f"Source: {st.session_state.source_filename}")
    detected = st.session_state.get("detected_lang", "")
    if detected:
        lines.append(f"Detected language: {language_name(detected)}")
    lines.append("")
    lines.append("=" * 60)
    lines.append("")

    for i, p in enumerate(paras):
        if include_original:
            lines.append(f"[Original] {p['original']}")
            lines.append("")
        lines.append(p["translated"])
        lines.append("")

    if cert_text:
        lines.append("=" * 60)
        lines.append("")
        lines.append(cert_text)
        lines.append("")

    return "\n".join(lines)


def _build_docx(
    paras: list[dict],
    include_original: bool,
    client_name: str,
    cert_text: str = "",
    show_disclaimer: bool = True,
) -> bytes:
    """Build a Word document from the translated paragraphs."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, txt, size=12, bold=False, italic=False, color=None, align=None):
        r = para.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            from docx.shared import RGBColor

            r.font.color.rgb = RGBColor(*color)
        if align is not None:
            para.alignment = align
        return r

    # Certification header
    if show_disclaimer:
        cert_para = doc.add_paragraph()
        _run(cert_para, certification_header(), size=10, italic=True, color=(139, 69, 19))
        cert_para.paragraph_format.space_after = Pt(12)

    # File info
    if st.session_state.get("source_filename"):
        info = doc.add_paragraph()
        _run(info, f"Source: {st.session_state.source_filename}", size=10, color=(100, 100, 100))
        info.paragraph_format.space_after = Pt(2)

    detected = st.session_state.get("detected_lang", "")
    if detected:
        info2 = doc.add_paragraph()
        _run(info2, f"Detected language: {language_name(detected)}", size=10, color=(100, 100, 100))
        info2.paragraph_format.space_after = Pt(12)

    # Separator
    sep = doc.add_paragraph()
    _run(sep, "_" * 60, size=10, color=(200, 200, 200))
    sep.paragraph_format.space_after = Pt(12)

    if include_original:
        # Side-by-side table
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Header row
        hdr = table.rows[0]
        h1 = hdr.cells[0].paragraphs[0]
        _run(h1, "Original", size=10, bold=True, color=(100, 100, 100))
        h2 = hdr.cells[1].paragraphs[0]
        _run(h2, "Translation", size=10, bold=True, color=(100, 100, 100))

        for p in paras:
            row = table.add_row()
            c1 = row.cells[0].paragraphs[0]
            _run(c1, p["original"], size=11, italic=True, color=(100, 100, 100))
            c2 = row.cells[1].paragraphs[0]
            _run(c2, p["translated"], size=12)
    else:
        # Translation only
        for p in paras:
            para = doc.add_paragraph()
            _run(para, p["translated"], size=12)
            para.paragraph_format.space_after = Pt(6)

    # EOIR certificate on new page
    if cert_text:
        doc.add_page_break()
        cert_lines = cert_text.split("\n")
        # Title line — centered, bold, 14pt
        if cert_lines:
            title_p = doc.add_paragraph()
            _run(title_p, cert_lines[0], size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            title_p.paragraph_format.space_after = Pt(18)

        # Body lines
        for line in cert_lines[1:]:
            p = doc.add_paragraph()
            _run(p, line, size=12)
            p.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# -- Sidebar ------------------------------------------------------------------

with st.sidebar:
    # Draft management
    st.markdown("#### Drafts")
    btn_cols = st.columns(2)
    with btn_cols[0]:
        if st.button("New", use_container_width=True):
            _do_new()
            st.rerun()
    with btn_cols[1]:
        save_clicked = st.button("Save", use_container_width=True, type="primary")

    saved_drafts = list_drafts()
    if saved_drafts:
        labels_map = {
            d["id"]: f"{d['client_name']} -- {d.get('source_filename', '')}"
            for d in saved_drafts
        }
        draft_ids = list(labels_map.keys())
        selected_draft = st.selectbox(
            "Load a saved draft",
            options=[""] + draft_ids,
            format_func=lambda x: labels_map.get(x, "Select..."),
            label_visibility="collapsed",
        )
        load_cols = st.columns(2)
        with load_cols[0]:
            if selected_draft and st.button("Load", use_container_width=True):
                _do_load(selected_draft)
                st.rerun()
        with load_cols[1]:
            if selected_draft and st.button("Delete", use_container_width=True):
                delete_draft(selected_draft)
                st.rerun()

    if st.session_state.last_saved_msg:
        st.markdown(
            f'<div class="saved-toast">{html_mod.escape(st.session_state.last_saved_msg)}</div>',
            unsafe_allow_html=True,
        )

    st.divider()

    # Client info
    st.markdown("#### Client")
    client_name = st.text_input(
        "Client Name",
        key="inp_client_name",
        placeholder="e.g. Maria Garcia Lopez",
    )

    st.divider()

    # Translation settings
    st.markdown("#### Translation Settings")
    target_lang = st.selectbox(
        "Target Language",
        options=TARGET_LANGUAGES,
        key="inp_target_lang",
    )

    _FORMAT_OPTIONS = ["Translation only", "Side-by-side (Original | Translated)"]
    export_format = st.selectbox(
        "Export Format",
        options=_FORMAT_OPTIONS,
        key="inp_export_format",
    )
    include_original = export_format == _FORMAT_OPTIONS[1]

    show_disclaimer = st.checkbox(
        "Include machine translation disclaimer",
        value=True,
        key="inp_show_disclaimer",
    )

    cert_type = st.selectbox(
        "Certificate",
        options=CERTIFICATE_TYPES,
        key="inp_cert_type",
    )

    if cert_type == "Certificate of Interpretation":
        client_pronoun = st.selectbox(
            "Client Pronoun",
            options=["they", "he", "she"],
            key="inp_client_pronoun",
        )
    else:
        client_pronoun = st.session_state.get("inp_client_pronoun", "they")

    # Translator / Interpreter info (shown when a cert is selected)
    if cert_type != "None":
        st.divider()
        st.markdown("#### Translator / Interpreter")
        translator_name = st.text_input(
            "Full Name",
            key="inp_translator_name",
            placeholder="e.g. John Smith",
        )
        translator_address = st.text_area(
            "Address",
            key="inp_translator_address",
            height=68,
            placeholder="e.g. 123 Main St, City, State 00000",
        )
        translator_phone = st.text_input(
            "Phone",
            key="inp_translator_phone",
            placeholder="e.g. (555) 123-4567",
        )
    else:
        translator_name = st.session_state.get("inp_translator_name", "")
        translator_address = st.session_state.get("inp_translator_address", "")
        translator_phone = st.session_state.get("inp_translator_phone", "")

    st.divider()

    # Notes
    st.markdown("#### Notes")
    notes = st.text_area(
        "Notes",
        key="inp_notes",
        height=80,
        placeholder="Internal notes about this translation...",
        label_visibility="collapsed",
    )

    # Box folder browser
    if render_box_folder_browser and _parse_folder_id:
        _sf = st.session_state.get("sf_client")
        _box_raw = (_sf.get("Box_Folder_Id__c", "") or "") if _sf else ""
        if _box_raw:
            st.divider()
            render_box_folder_browser(
                _parse_folder_id(_box_raw),
                mode="viewer",
                key_prefix="_dt_box",
                header_label="Client Documents",
            )

    render_tool_notes("document-translator")


# -- Handle save (after sidebar renders) -------------------------------------

if save_clicked:
    _do_save(
        client_name,
        target_lang,
        notes,
        include_original,
        export_format,
        cert_type,
        {"name": translator_name, "address": translator_address, "phone": translator_phone},
        client_pronoun,
        show_disclaimer,
    )
    st.rerun()


# -- Main area ----------------------------------------------------------------

work_col, preview_col = st.columns([3, 2], gap="large")

# -- Left column: Upload, detect, translate, edit ----------------------------

with work_col:
    st.markdown('<div class="section-label">Upload Document</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "Upload a document to translate",
        type=["pdf", "docx", "jpg", "jpeg", "png", "tiff", "bmp", "webp"],
        label_visibility="collapsed",
    )

    # Process uploaded file
    if uploaded_file is not None and uploaded_file.name != st.session_state.get("source_filename"):
        file_bytes = uploaded_file.getvalue()
        st.session_state.source_filename = uploaded_file.name
        st.session_state.translation_done = False
        st.session_state.translated = []
        st.session_state.google_doc_url = ""

        # Clean up old edit keys
        for key in list(st.session_state.keys()):
            if key.startswith("edit_"):
                del st.session_state[key]

        with st.spinner("Extracting text..."):
            try:
                paragraphs = extract_text(file_bytes, uploaded_file.name)
                st.session_state.paragraphs = paragraphs
            except Exception as e:
                st.error(f"Failed to extract text: {e}")
                st.session_state.paragraphs = []

        # Auto-detect language
        if st.session_state.paragraphs:
            combined = " ".join(st.session_state.paragraphs[:5])
            try:
                lang_code, confidence = detect_language(combined)
                st.session_state.detected_lang = lang_code
                st.session_state.detected_confidence = confidence
            except Exception as e:
                st.warning(f"Language detection failed: {e}")
                st.session_state.detected_lang = ""
                st.session_state.detected_confidence = 0.0

    # Show file info
    if st.session_state.get("source_filename"):
        st.markdown(f"**File:** {html_mod.escape(st.session_state.source_filename)}")

        para_count = len(st.session_state.get("paragraphs", []))
        st.markdown(f"**Paragraphs extracted:** {para_count}")

        if st.session_state.get("detected_lang"):
            lang = language_name(st.session_state.detected_lang)
            conf = st.session_state.get("detected_confidence", 0)
            st.markdown(
                f'<span class="lang-badge">Detected: {html_mod.escape(lang)} ({conf:.0%})</span>',
                unsafe_allow_html=True,
            )

    # Translate button
    if st.session_state.get("paragraphs") and not st.session_state.get("translation_done"):
        st.markdown("---")
        if st.button("Translate", type="primary", use_container_width=True):
            progress_bar = st.progress(0, text="Translating...")

            def update_progress(completed: int, total: int) -> None:
                progress_bar.progress(
                    completed / total,
                    text=f"Translating... {completed}/{total} paragraphs",
                )

            try:
                results = translate_paragraphs(
                    st.session_state.paragraphs,
                    target_lang=target_lang,
                    source_lang=None,  # let API auto-detect
                    on_progress=update_progress,
                )
                st.session_state.translated = results
                st.session_state.translation_done = True
                progress_bar.progress(1.0, text="Translation complete!")

                # Initialize edit keys
                for i, r in enumerate(results):
                    st.session_state[f"edit_{i}"] = r["translated"]

                st.rerun()
            except Exception as e:
                st.error(f"Translation failed: {e}")

    # Editable translation results
    if st.session_state.get("translation_done") and st.session_state.get("translated"):
        st.markdown("---")

        edit_hdr_cols = st.columns([3, 1])
        with edit_hdr_cols[0]:
            st.markdown(
                '<div class="section-label">Review &amp; Edit Translation</div>',
                unsafe_allow_html=True,
            )
        with edit_hdr_cols[1]:
            @st.dialog("Original Document", width="large")
            def _show_original_dialog():
                detected = st.session_state.get("detected_lang", "")
                if detected:
                    st.markdown(
                        f'<span class="lang-badge">{html_mod.escape(language_name(detected))}</span>',
                        unsafe_allow_html=True,
                    )
                for idx, t in enumerate(st.session_state.translated):
                    st.markdown(
                        f'<div class="para-num">Paragraph {idx + 1}</div>',
                        unsafe_allow_html=True,
                    )
                    st.markdown(t.get("original", ""))
                    if idx < len(st.session_state.translated) - 1:
                        st.divider()

            if st.button("View Original", use_container_width=True):
                _show_original_dialog()

        st.caption("Edit any paragraph below. Changes are reflected in the preview and export.")

        for i, t in enumerate(st.session_state.translated):
            if include_original:
                st.markdown(
                    f'<div class="original-text">{html_mod.escape(t["original"])}</div>',
                    unsafe_allow_html=True,
                )
            st.text_area(
                f"Paragraph {i + 1}",
                value=st.session_state.get(f"edit_{i}", t["translated"]),
                key=f"edit_{i}",
                height=80,
                label_visibility="collapsed",
            )

    # Show placeholder when no file uploaded
    if not st.session_state.get("source_filename"):
        st.info("Upload a PDF, Word document, or image to get started.")


# -- Right column: Preview and export ----------------------------------------

with preview_col:
    st.markdown('<div class="section-label">Translation Preview</div>', unsafe_allow_html=True)

    if st.session_state.get("translation_done") and st.session_state.get("translated"):
        # Build preview
        paras = _get_edited_paragraphs()

        # Build certificate text if applicable
        cert_text = ""
        if cert_type != "None" and translator_name:
            detected = st.session_state.get("detected_lang", "")
            cert_text = build_certificate(
                cert_type=cert_type,
                translator_name=translator_name,
                translator_address=translator_address,
                translator_phone=translator_phone,
                source_lang=language_name(detected) if detected else "Unknown",
                target_lang=target_lang,
                source_filename=st.session_state.get("source_filename", ""),
                client_name=client_name,
                client_pronoun=client_pronoun,
            )

        preview_parts = []
        if show_disclaimer:
            preview_parts.append(
                f'<div class="cert-notice">{html_mod.escape(certification_header())}</div>'
            )

        if st.session_state.get("source_filename"):
            preview_parts.append(
                f"<em>Source: {html_mod.escape(st.session_state.source_filename)}</em><br>"
            )
        if st.session_state.get("detected_lang"):
            preview_parts.append(
                f"<em>Detected: {html_mod.escape(language_name(st.session_state.detected_lang))}</em><br>"
            )
        preview_parts.append("<hr>")

        for p in paras:
            if include_original:
                preview_parts.append(
                    f'<div class="original-text">{html_mod.escape(p["original"])}</div>'
                )
            preview_parts.append(f"<p>{html_mod.escape(p['translated'])}</p>")

        # Certificate preview
        if cert_text:
            preview_parts.append("<hr>")
            preview_parts.append(
                f"<pre style='white-space:pre-wrap;font-family:Times New Roman,serif;'>"
                f"{html_mod.escape(cert_text)}</pre>"
            )

        preview_html = "\n".join(preview_parts)
        st.markdown(
            f'<div class="preview-panel">{preview_html}</div>',
            unsafe_allow_html=True,
        )

        # Export controls
        st.markdown("---")
        st.markdown('<div class="section-label">Export</div>', unsafe_allow_html=True)

        safe_name = (client_name or "translation").replace(" ", "_")

        exp_cols = st.columns(2)
        with exp_cols[0]:
            plain_text = _build_plain_text(paras, include_original, cert_text, show_disclaimer)
            st.download_button(
                "Download .txt",
                data=plain_text,
                file_name=f"Translation_{safe_name}.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with exp_cols[1]:
            docx_bytes = _build_docx(paras, include_original, client_name, cert_text, show_disclaimer)
            st.download_button(
                "Download .docx",
                data=docx_bytes,
                file_name=f"Translation_{safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        if st.button("Upload to Google Docs", use_container_width=True):
            with st.spinner("Uploading to Google Docs..."):
                try:
                    docx_data = _build_docx(paras, include_original, client_name, cert_text, show_disclaimer)
                    url = upload_to_google_docs(
                        docx_data, f"Translation - {client_name or 'Document'}"
                    )
                    st.session_state.google_doc_url = url
                except Exception as e:
                    st.error(f"Upload failed: {e}")

        if st.session_state.get("google_doc_url"):
            st.markdown(f"[Open Google Doc]({st.session_state.google_doc_url})")
    else:
        st.info("Translation will appear here after you upload and translate a document.")
