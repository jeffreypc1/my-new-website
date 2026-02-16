"""Legal Research -- Streamlit dashboard.

Full legal research UI for searching landmark immigration decisions,
saving relevant case law to collections, and exporting research bundles.
Works entirely with local persistence (no API server required).

Part of the O'Brien Immigration Law tool suite.
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from datetime import date
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.case_law import (
    KEY_DECISIONS,
    LEGAL_TOPICS,
    CaseLaw,
    delete_collection,
    get_by_citation,
    list_collections,
    load_collection,
    new_collection_id,
    save_collection,
    search_decisions,
)

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.google_upload import upload_to_google_docs
from shared.client_banner import render_client_banner

# -- Page config --------------------------------------------------------------

st.set_page_config(
    page_title="Legal Research -- O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -- CSS ----------------------------------------------------------------------

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* Hide Streamlit chrome */
#MainMenu, header[data-testid="stHeader"], footer,
div[data-testid="stToolbar"] { display: none !important; }

.stApp {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

/* Navigation bar */
.nav-bar {
    display: flex;
    align-items: center;
    padding: 10px 4px;
    margin: -1rem 0 1.2rem 0;
    border-bottom: 1px solid rgba(0,0,0,0.07);
}
.nav-back {
    display: flex;
    align-items: center;
    gap: 6px;
    font-family: 'Inter', sans-serif;
    font-size: 0.85rem;
    font-weight: 500;
    color: #0066CC;
    text-decoration: none;
    min-width: 150px;
}
.nav-back:hover { color: #004499; text-decoration: underline; }
.nav-title {
    flex: 1;
    text-align: center;
    font-family: 'Inter', sans-serif;
    font-size: 1.15rem;
    font-weight: 700;
    color: #1a2744;
    letter-spacing: -0.02em;
}
.nav-firm {
    font-weight: 400;
    color: #86868b;
    font-size: 0.85rem;
    margin-left: 8px;
}
.nav-spacer { min-width: 150px; }

/* Section labels */
.section-label {
    font-size: 0.78rem;
    font-weight: 600;
    color: #5a6a85;
    text-transform: uppercase;
    letter-spacing: 0.04em;
    margin-bottom: 4px;
    margin-top: 12px;
}

/* Court badge */
.court-badge {
    display: inline-block;
    background: #1a2744;
    color: #ffffff;
    padding: 2px 10px;
    border-radius: 4px;
    font-size: 0.78rem;
    font-weight: 600;
    margin-right: 8px;
}

/* Topic tags */
.topic-tag {
    display: inline-block;
    background: #e2e8f0;
    color: #334155;
    padding: 2px 8px;
    border-radius: 12px;
    font-size: 0.73rem;
    margin-right: 4px;
    margin-top: 4px;
}

/* Citation */
.citation {
    font-size: 0.88rem;
    color: #2b5797;
    font-style: italic;
}

/* Holding text */
.holding-text {
    font-size: 0.88rem;
    line-height: 1.65;
    color: #334155;
    margin-top: 6px;
}

/* Case card */
.case-card {
    background: #fff;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 16px 18px;
    margin-bottom: 12px;
    transition: border-color 0.15s, box-shadow 0.15s;
}
.case-card:hover {
    border-color: #4a7ddb;
    box-shadow: 0 2px 8px rgba(74, 125, 219, 0.10);
}

/* Preview panel */
.preview-panel {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 18px 22px;
    margin-bottom: 12px;
}

/* Saved toast */
.saved-toast {
    font-size: 0.8rem;
    color: #2e7d32;
    font-weight: 600;
}

/* Detail header */
.detail-header {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 18px 22px;
    margin-bottom: 16px;
}
.detail-name {
    font-size: 1.25rem;
    font-weight: 800;
    color: #1a2744;
    letter-spacing: -0.02em;
    margin-bottom: 4px;
}
.detail-meta {
    font-size: 0.85rem;
    color: #5a6a85;
    line-height: 1.6;
}
.detail-meta strong {
    color: #334155;
}

/* Saved decision row */
.saved-row {
    display: flex;
    align-items: center;
    padding: 8px 12px;
    border-radius: 8px;
    margin-bottom: 4px;
    background: #f0f4fa;
    gap: 10px;
}
.saved-row-name {
    font-weight: 600;
    font-size: 0.88rem;
    color: #1a2744;
}
.saved-row-cite {
    font-size: 0.8rem;
    color: #2b5797;
    font-style: italic;
}

/* Empty state */
.empty-state {
    text-align: center;
    padding: 60px 20px;
    color: #5a6a85;
}
.empty-state-title {
    font-size: 1.1rem;
    font-weight: 700;
    color: #1a2744;
    margin-bottom: 6px;
}
.empty-state-desc {
    font-size: 0.88rem;
}
</style>
""",
    unsafe_allow_html=True,
)

# -- Navigation bar -----------------------------------------------------------

st.markdown(
    """
<div class="nav-bar">
    <a href="http://localhost:8502" class="nav-back">&#8592; Staff Dashboard</a>
    <div class="nav-title">Legal Research<span class="nav-firm">&mdash; O'Brien Immigration Law</span></div>
    <div class="nav-spacer"></div>
</div>
""",
    unsafe_allow_html=True,
)

render_client_banner()

# -- Session state defaults ---------------------------------------------------

_DEFAULTS: dict = {
    "collection_id": None,
    "last_saved_msg": "",
    "saved_decisions": [],
    "search_query": "",
    "selected_decision": None,
    "case_name": "",
    "a_number": "",
    "notes": "",
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

if st.session_state.collection_id is None:
    st.session_state.collection_id = new_collection_id()


# -- Helpers ------------------------------------------------------------------


def _esc(text: str) -> str:
    """HTML-escape a string."""
    return html_mod.escape(str(text))


def _do_save() -> None:
    """Save the current collection to disk."""
    decisions_to_save = [
        {
            "name": d.get("name", ""),
            "citation": d.get("citation", ""),
            "court": d.get("court", ""),
            "date": d.get("date", ""),
            "holding": d.get("holding", ""),
            "topics": d.get("topics", []),
        }
        for d in st.session_state.saved_decisions
    ]
    save_collection(
        collection_id=st.session_state.collection_id,
        case_name=st.session_state.get("case_name", ""),
        a_number=st.session_state.get("a_number", ""),
        decisions=decisions_to_save,
        notes=st.session_state.get("notes", ""),
    )
    name = st.session_state.get("case_name", "") or "collection"
    st.session_state.last_saved_msg = f"Saved -- {name}"


def _do_load(collection_id: str) -> None:
    """Load a saved collection into session state."""
    data = load_collection(collection_id)
    if not data:
        return
    st.session_state.collection_id = data["id"]
    st.session_state.case_name = data.get("case_name", "")
    st.session_state.a_number = data.get("a_number", "")
    st.session_state.notes = data.get("notes", "")
    st.session_state.saved_decisions = list(data.get("decisions", []))
    st.session_state.selected_decision = None
    st.session_state.last_saved_msg = ""


def _do_new() -> None:
    """Start a fresh collection."""
    st.session_state.collection_id = new_collection_id()
    st.session_state.last_saved_msg = ""
    st.session_state.saved_decisions = []
    st.session_state.selected_decision = None
    st.session_state.notes = ""
    for k in ("case_name", "a_number"):
        if k in st.session_state:
            st.session_state[k] = ""


def _build_docx(case_name: str, decisions: list[dict]) -> bytes:
    """Build a Word document summarizing saved decisions."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, text: str, size: int = 11, bold: bool = False, italic: bool = False):
        r = para.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return r

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "LEGAL RESEARCH -- CASE LAW SUMMARY", size=14, bold=True)

    # Case info
    if case_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, case_name, size=11)

    doc.add_paragraph()  # spacer

    # Table header
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, label in enumerate(["Case Name", "Citation", "Holding"]):
        p = hdr_cells[i].paragraphs[0]
        _run(p, label, size=10, bold=True)

    # Decision rows
    for d in decisions:
        row_cells = table.add_row().cells
        p = row_cells[0].paragraphs[0]
        _run(p, d.get("name", ""), size=10)
        p = row_cells[1].paragraphs[0]
        _run(p, d.get("citation", ""), size=10, italic=True)
        p = row_cells[2].paragraphs[0]
        _run(p, d.get("holding", ""), size=9)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, f"Generated: {date.today().strftime('%B %d, %Y')}", size=9, italic=True)
    p = doc.add_paragraph()
    _run(p, "O'Brien Immigration Law", size=9, italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_plain_text(case_name: str, decisions: list[dict]) -> str:
    """Build a plain-text summary of saved decisions."""
    lines: list[str] = []
    lines.append("LEGAL RESEARCH -- CASE LAW SUMMARY")
    lines.append("=" * 60)
    if case_name:
        lines.append(f"Case: {case_name}")
    lines.append(f"Decisions: {len(decisions)}")
    lines.append("=" * 60)
    lines.append("")

    for i, d in enumerate(decisions, 1):
        lines.append(f"{i}. {d.get('name', 'Unknown')}")
        lines.append(f"   {d.get('citation', '')}")
        lines.append(f"   Court: {d.get('court', '')}")
        lines.append(f"   Date: {d.get('date', '')}")
        lines.append(f"   Holding: {d.get('holding', '')}")
        topics = d.get("topics", [])
        if topics:
            lines.append(f"   Topics: {', '.join(topics)}")
        lines.append("")

    lines.append(f"Generated: {date.today().strftime('%B %d, %Y')}")
    lines.append("O'Brien Immigration Law")
    return "\n".join(lines)


# -- Sidebar ------------------------------------------------------------------

with st.sidebar:
    # Collection management
    st.markdown("#### Collections")
    btn_cols = st.columns(2)
    with btn_cols[0]:
        if st.button("New", use_container_width=True):
            _do_new()
            st.rerun()
    with btn_cols[1]:
        save_clicked = st.button("Save", use_container_width=True, type="primary")

    saved_collections = list_collections()
    if saved_collections:
        labels_map = {
            c["id"]: f"{c['case_name']} ({c['decision_count']} decisions)"
            for c in saved_collections
        }
        coll_ids = list(labels_map.keys())
        selected_coll = st.selectbox(
            "Load a saved collection",
            options=[""] + coll_ids,
            format_func=lambda x: labels_map.get(x, "Select..."),
            label_visibility="collapsed",
        )
        load_cols = st.columns(2)
        with load_cols[0]:
            if selected_coll and st.button("Load", use_container_width=True):
                _do_load(selected_coll)
                st.rerun()
        with load_cols[1]:
            if selected_coll and st.button("Delete", use_container_width=True):
                delete_collection(selected_coll)
                if st.session_state.collection_id == selected_coll:
                    _do_new()
                st.rerun()

    if st.session_state.last_saved_msg:
        st.markdown(
            f'<div class="saved-toast">{_esc(st.session_state.last_saved_msg)}</div>',
            unsafe_allow_html=True,
        )

    st.divider()

    # Case info
    st.markdown("#### Case Info")
    case_name = st.text_input(
        "Case Name",
        key="case_name",
        placeholder="e.g. Garcia-Lopez",
    )
    a_number = st.text_input(
        "A-Number",
        key="a_number",
        placeholder="e.g. A-123-456-789",
    )

    st.divider()

    # Topic filters
    st.markdown("#### Topic Filters")
    selected_topics = st.multiselect(
        "Filter by topic",
        options=LEGAL_TOPICS,
        default=[],
        label_visibility="collapsed",
    )

    st.divider()

    # Saved count
    saved_count = len(st.session_state.saved_decisions)
    st.caption(f"{saved_count} decision(s) in collection")


# -- Handle save (after sidebar renders) --------------------------------------

if save_clicked:
    _do_save()
    st.rerun()


# -- Main area: two columns ---------------------------------------------------

search_col, detail_col = st.columns([3, 2], gap="large")


# -- Left column: search + results --------------------------------------------

with search_col:
    st.markdown('<div class="section-label">Search Case Law</div>', unsafe_allow_html=True)

    query = st.text_input(
        "Search decisions",
        placeholder="e.g. particular social group, credibility, cancellation hardship",
        label_visibility="collapsed",
    )

    # Perform search
    if query:
        results = search_decisions(query, topics=selected_topics if selected_topics else None)
    elif selected_topics:
        results = search_decisions("", topics=selected_topics)
    else:
        results = []

    if query or selected_topics:
        if results:
            st.markdown(f"**{len(results)} result(s)**")

            for idx, decision in enumerate(results):
                # Check if already saved
                is_saved = any(
                    d.get("citation") == decision.citation
                    for d in st.session_state.saved_decisions
                )

                # Build card HTML
                tags_html = " ".join(
                    f'<span class="topic-tag">{_esc(t)}</span>'
                    for t in decision.topics
                )

                st.markdown(
                    f'<div class="case-card">'
                    f'<span class="court-badge">{_esc(decision.court)}</span>'
                    f'<strong>{_esc(decision.name)}</strong><br>'
                    f'<span class="citation">{_esc(decision.citation)}</span>'
                    f' &nbsp; {_esc(decision.date)}<br>'
                    f'{tags_html}'
                    f'<div class="holding-text">{_esc(decision.holding)}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

                btn_cols = st.columns([1, 1, 3])
                with btn_cols[0]:
                    if is_saved:
                        st.button(
                            "Saved",
                            key=f"saved_{idx}",
                            disabled=True,
                            use_container_width=True,
                        )
                    else:
                        if st.button("Save", key=f"save_{idx}", use_container_width=True):
                            st.session_state.saved_decisions.append({
                                "name": decision.name,
                                "citation": decision.citation,
                                "court": decision.court,
                                "date": decision.date,
                                "holding": decision.holding,
                                "topics": decision.topics,
                            })
                            st.rerun()
                with btn_cols[1]:
                    if st.button("View", key=f"view_{idx}", use_container_width=True):
                        st.session_state.selected_decision = {
                            "name": decision.name,
                            "citation": decision.citation,
                            "court": decision.court,
                            "date": decision.date,
                            "holding": decision.holding,
                            "topics": decision.topics,
                        }
                        st.rerun()
        else:
            st.info("No results found. Try different search terms or broaden your topic filters.")
    else:
        st.info(
            "Enter a search term above to find relevant case law, BIA decisions, "
            "and federal court holdings. You can also filter by topic in the sidebar."
        )


# -- Right column: detail view + saved collection -----------------------------

with detail_col:
    # Decision detail view
    selected = st.session_state.selected_decision
    if selected:
        st.markdown('<div class="section-label">Decision Detail</div>', unsafe_allow_html=True)

        tags_html = " ".join(
            f'<span class="topic-tag">{_esc(t)}</span>'
            for t in selected.get("topics", [])
        )

        st.markdown(
            f'<div class="detail-header">'
            f'<div class="detail-name">{_esc(selected.get("name", ""))}</div>'
            f'<div class="detail-meta">'
            f'<span class="court-badge">{_esc(selected.get("court", ""))}</span>'
            f'<strong>Citation:</strong> <span class="citation">{_esc(selected.get("citation", ""))}</span>'
            f' &nbsp;&bull;&nbsp; <strong>Date:</strong> {_esc(selected.get("date", ""))}'
            f'</div>'
            f'<div style="margin-top: 6px;">{tags_html}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

        st.markdown("**Holding**")
        st.markdown(
            f'<div class="holding-text">{_esc(selected.get("holding", ""))}</div>',
            unsafe_allow_html=True,
        )

        # Quick-save from detail view
        is_detail_saved = any(
            d.get("citation") == selected.get("citation")
            for d in st.session_state.saved_decisions
        )
        detail_btn_cols = st.columns([1, 1, 2])
        with detail_btn_cols[0]:
            if is_detail_saved:
                st.button("Saved", key="detail_saved", disabled=True, use_container_width=True)
            else:
                if st.button("Save", key="detail_save", type="primary", use_container_width=True):
                    st.session_state.saved_decisions.append(dict(selected))
                    st.rerun()
        with detail_btn_cols[1]:
            if st.button("Close", key="detail_close", use_container_width=True):
                st.session_state.selected_decision = None
                st.rerun()

        st.divider()

    # Saved collection
    st.markdown('<div class="section-label">Saved Decisions</div>', unsafe_allow_html=True)

    saved = st.session_state.saved_decisions
    if not saved:
        st.markdown(
            '<div class="empty-state">'
            '<div class="empty-state-title">No decisions saved</div>'
            '<div class="empty-state-desc">'
            'Search for case law and click "Save" to add decisions to your collection.'
            '</div></div>',
            unsafe_allow_html=True,
        )
    else:
        for s_idx, s_dec in enumerate(saved):
            st.markdown(
                f'<div class="saved-row">'
                f'<div>'
                f'<span class="saved-row-name">{_esc(s_dec.get("name", ""))}</span><br>'
                f'<span class="saved-row-cite">{_esc(s_dec.get("citation", ""))}</span>'
                f'</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            rm_cols = st.columns([1, 1, 3])
            with rm_cols[0]:
                if st.button("Remove", key=f"rm_{s_idx}", use_container_width=True):
                    st.session_state.saved_decisions.pop(s_idx)
                    st.rerun()
            with rm_cols[1]:
                if st.button("View", key=f"sv_{s_idx}", use_container_width=True):
                    st.session_state.selected_decision = dict(s_dec)
                    st.rerun()

        # Export controls
        st.divider()
        st.markdown('<div class="section-label">Export</div>', unsafe_allow_html=True)

        exp_cols = st.columns(2)
        safe_name = (st.session_state.get("case_name", "") or "research").replace(" ", "_")

        with exp_cols[0]:
            plain_text = _build_plain_text(
                st.session_state.get("case_name", ""),
                saved,
            )
            st.download_button(
                "Download .txt",
                data=plain_text,
                file_name=f"Legal_Research_{safe_name}.txt",
                mime="text/plain",
                use_container_width=True,
            )

        with exp_cols[1]:
            docx_bytes = _build_docx(
                st.session_state.get("case_name", ""),
                saved,
            )
            st.download_button(
                "Download .docx",
                data=docx_bytes,
                file_name=f"Legal_Research_{safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            if st.button("Upload to Google Docs", use_container_width=True):
                with st.spinner("Uploading to Google Docs..."):
                    try:
                        url = upload_to_google_docs(docx_bytes, f"Legal Research - {st.session_state.get('case_name', 'Research')}")
                        st.session_state.google_doc_url = url
                    except Exception as e:
                        st.error(f"Upload failed: {e}")
            if st.session_state.get("google_doc_url"):
                st.markdown(f"[Open Google Doc]({st.session_state.google_doc_url})")
