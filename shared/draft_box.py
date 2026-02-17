"""Reusable AI Draft Box component for tool dashboards.

Renders a text area for instructions + Generate Draft button. Sends the
page's current content + instructions to Claude, which drafts a document.
Export as .txt, .docx, .pdf, or upload to Google Docs — with optional
page numbers (bottom-right).

Draft history is kept per case/draft ID so users don't see unrelated
drafts from other cases.
"""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from shared.config_store import load_config, save_config

_HISTORY_CONFIG = "draft-history"

# ── Default system prompts ──────────────────────────────────────────────────

_DEFAULT_GLOBAL_PROMPT = (
    "You are a legal drafting assistant for O'Brien Immigration Law. "
    "Draft professional, well-structured legal documents. Use formal legal "
    "language appropriate for USCIS filings and immigration court proceedings. "
    "Be thorough but concise. Do not include placeholder text — use the "
    "provided case details to produce a complete, ready-to-review draft."
)

_DEFAULT_TOOL_PROMPTS = {
    "cover-letters": (
        "Draft a professional cover letter addressed to USCIS or the "
        "immigration court. Include attorney information, client identification, "
        "case type, enclosed documents, and a formal closing."
    ),
    "brief-builder": (
        "Draft a legal brief section for immigration proceedings. Use proper "
        "legal citation format, reference relevant INA sections and case law, "
        "and present arguments persuasively with clear headings."
    ),
    "declaration-drafter": (
        "Draft a declaration in the first person based on the provided answers. "
        "Use clear, specific language. Organize chronologically. Include standard "
        "declaration preamble and signature block."
    ),
    "timeline-builder": (
        "Draft a narrative summary connecting the timeline events. Highlight key "
        "dates, patterns, and cumulative impact. Use transitional language that "
        "tells the client's story cohesively."
    ),
}


def _get_prompts(tool_name: str) -> tuple[str, str]:
    """Load global + tool-specific prompts from components.json or defaults."""
    config = load_config("components") or {}
    draft_box = config.get("draft-box", {})
    global_prompt = draft_box.get("global_prompt", _DEFAULT_GLOBAL_PROMPT)
    tool_prompts = draft_box.get("tool_prompts", {})
    tool_prompt = tool_prompts.get(tool_name, _DEFAULT_TOOL_PROMPTS.get(tool_name, ""))
    return global_prompt, tool_prompt


def _history_key(tool_name: str, case_id: str) -> str:
    return f"{tool_name}:{case_id or '_unsaved'}"


def _load_history(tool_name: str, case_id: str) -> list[dict]:
    all_history = load_config(_HISTORY_CONFIG) or {}
    return all_history.get(_history_key(tool_name, case_id), [])


def _save_history_entry(tool_name: str, case_id: str, entry: dict) -> None:
    all_history = load_config(_HISTORY_CONFIG) or {}
    key = _history_key(tool_name, case_id)
    entries = all_history.get(key, [])
    entries.insert(0, entry)
    entries = entries[:10]  # cap at 10
    all_history[key] = entries
    save_config(_HISTORY_CONFIG, all_history)


# ── Document builders ───────────────────────────────────────────────────────


def _add_docx_page_number(section) -> None:
    """Add a right-aligned page number field to the section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # PAGE field: begin
    run1 = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run1._r.append(fld_begin)

    # PAGE field: instruction
    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    # PAGE field: end
    run3 = para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_end)

    for run in (run1, run2, run3):
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)


def _build_draft_docx(text: str, title: str, page_numbers: bool = False) -> bytes:
    """Build a .docx from draft text with standard formatting."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        if page_numbers:
            _add_docx_page_number(section)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Times New Roman"
    doc.add_paragraph()  # blank line

    # Body — split on double newlines for paragraphs
    for para_text in text.split("\n\n"):
        stripped = para_text.strip()
        if stripped:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_draft_pdf(text: str, title: str, page_numbers: bool = False) -> bytes:
    """Build a PDF from draft text using fpdf2."""
    from fpdf import FPDF

    class _PDF(FPDF):
        _show_page_numbers = page_numbers

        def footer(self):
            if self._show_page_numbers:
                self.set_y(-15)
                self.set_font("Times", "", 10)
                self.cell(0, 10, str(self.page_no()), 0, 0, "R")

    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=25 if page_numbers else 15)
    pdf.add_page()
    pdf.set_margins(25.4, 25.4, 25.4)  # 1-inch margins

    # Title
    pdf.set_font("Times", "B", 14)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Body
    pdf.set_font("Times", "", 12)
    for para_text in text.split("\n\n"):
        stripped = para_text.strip()
        if stripped:
            pdf.multi_cell(0, 6, stripped)
            pdf.ln(3)

    return bytes(pdf.output())


# ── Main component ──────────────────────────────────────────────────────────


def render_draft_box(tool_name: str, context: dict) -> None:
    """Render the Draft Box component.

    Parameters
    ----------
    tool_name : str
        The tool's slug (e.g. "cover-letters").
    context : dict
        Must contain:
        - document_type: str — what kind of document to draft
        - client_name: str — for file naming
        - case_id: str — scopes draft history to this case
        - content: str — the page's current content as text
    """
    document_type = context.get("document_type", "document")
    client_name = context.get("client_name", "")
    case_id = context.get("case_id", "")
    content = context.get("content", "")

    st.divider()
    st.markdown('<div class="section-label" style="font-weight:600; margin-bottom:4px;">Draft Box</div>', unsafe_allow_html=True)
    st.caption(f"Use AI to draft a {document_type} based on the current content.")

    instructions = st.text_area(
        "Instructions",
        key=f"_draft_box_instructions_{tool_name}",
        height=100,
        placeholder=f"Describe what you'd like in the {document_type}...",
        label_visibility="collapsed",
    )

    if st.button("Generate Draft", type="primary", key=f"_draft_box_generate_{tool_name}", use_container_width=True):
        if not instructions.strip():
            st.warning("Enter instructions describing what you'd like drafted.")
        else:
            with st.spinner("Drafting with Claude..."):
                try:
                    from shared.claude_client import draft_with_claude

                    global_prompt, tool_prompt = _get_prompts(tool_name)
                    system_prompt = global_prompt
                    if tool_prompt:
                        system_prompt += "\n\n" + tool_prompt

                    if content.strip():
                        user_message = (
                            f"## Current {document_type} content\n\n{content}\n\n"
                            f"## Instructions\n\n{instructions}"
                        )
                    else:
                        user_message = (
                            f"## Instructions\n\n{instructions}\n\n"
                            f"(No content on the page yet — draft from scratch based on the instructions.)"
                        )

                    draft_text = draft_with_claude(system_prompt, user_message, tool_name=tool_name)

                    # Save to session state
                    st.session_state[f"_draft_box_latest_{tool_name}"] = draft_text
                    st.session_state[f"_draft_box_latest_instructions_{tool_name}"] = instructions

                    # Save to history
                    entry = {
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "instructions": instructions,
                        "draft": draft_text,
                    }
                    _save_history_entry(tool_name, case_id, entry)
                    st.rerun()

                except Exception as e:
                    st.error(f"Draft generation failed: {e}")

    # Show latest draft
    latest = st.session_state.get(f"_draft_box_latest_{tool_name}", "")
    if latest:
        with st.expander("Latest Draft", expanded=True):
            st.markdown(latest)

            safe_name = (client_name or "Draft").replace(" ", "_")
            title = f"{document_type.title()} Draft — {client_name}" if client_name else f"{document_type.title()} Draft"

            # Export options
            st.markdown("---")
            opt_cols = st.columns([3, 2])
            with opt_cols[0]:
                fmt = st.radio(
                    "Format",
                    options=["Google Docs", "Word (.docx)", "PDF (.pdf)", "Text (.txt)"],
                    key=f"_draft_box_fmt_{tool_name}",
                    horizontal=True,
                    label_visibility="collapsed",
                )
            with opt_cols[1]:
                page_nums = st.checkbox(
                    "Page numbers",
                    key=f"_draft_box_pn_{tool_name}",
                )

            if fmt == "Google Docs":
                if st.button("Upload to Google Docs", use_container_width=True, key=f"_draft_box_action_{tool_name}"):
                    with st.spinner("Uploading to Google Docs..."):
                        try:
                            from shared.google_upload import upload_to_google_docs

                            docx_bytes = _build_draft_docx(latest, title, page_numbers=page_nums)
                            url = upload_to_google_docs(docx_bytes, title)
                            st.session_state[f"_draft_box_gdoc_url_{tool_name}"] = url
                        except Exception as e:
                            st.error(f"Upload failed: {e}")
                if st.session_state.get(f"_draft_box_gdoc_url_{tool_name}"):
                    st.markdown(f"[Open Google Doc]({st.session_state[f'_draft_box_gdoc_url_{tool_name}']})")

            elif fmt == "Word (.docx)":
                docx_bytes = _build_draft_docx(latest, title, page_numbers=page_nums)
                st.download_button(
                    "Download .docx",
                    data=docx_bytes,
                    file_name=f"Draft_{safe_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"_draft_box_action_{tool_name}",
                )

            elif fmt == "PDF (.pdf)":
                pdf_bytes = _build_draft_pdf(latest, title, page_numbers=page_nums)
                st.download_button(
                    "Download .pdf",
                    data=pdf_bytes,
                    file_name=f"Draft_{safe_name}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"_draft_box_action_{tool_name}",
                )

            else:  # Text (.txt)
                st.download_button(
                    "Download .txt",
                    data=latest,
                    file_name=f"Draft_{safe_name}.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key=f"_draft_box_action_{tool_name}",
                )

    # Draft history
    history = _load_history(tool_name, case_id)
    # Exclude the current latest from history display (it's already shown above)
    display_history = history[1:] if latest and history else history

    if display_history:
        with st.expander(f"Previous Drafts ({len(display_history)})"):
            for i, entry in enumerate(display_history):
                ts = entry.get("timestamp", "")
                instr = entry.get("instructions", "")
                draft = entry.get("draft", "")
                first_line = instr.split("\n")[0][:80] if instr else "No instructions"

                st.markdown(f"**{ts}** — {first_line}")
                with st.expander("Show full draft", expanded=False):
                    st.markdown(draft)
                if st.button("Use This Draft", key=f"_draft_box_use_{tool_name}_{i}"):
                    st.session_state[f"_draft_box_latest_{tool_name}"] = draft
                    st.session_state[f"_draft_box_latest_instructions_{tool_name}"] = instr
                    st.rerun()
                st.markdown("---")
