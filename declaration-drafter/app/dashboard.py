"""Declaration Drafter — Streamlit dashboard.

Guided declaration drafting for immigration cases with live preview,
draft persistence, coaching tips, and Word export.
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from datetime import date
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.drafts import delete_draft, list_drafts, load_draft, new_draft_id, save_draft
from app.prompts import (
    DECLARATION_PROMPTS,
    DECLARATION_TYPES,
    INTERPRETER_CERT,
    PERJURY_CLAUSE,
    format_numbered_paragraphs,
    get_declaration_prompts,
    get_declaration_types,
)

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.google_upload import upload_to_google_docs
from shared.client_banner import render_client_banner
from shared.tool_notes import render_tool_notes
try:
    from shared.draft_box import render_draft_box
except ImportError:
    render_draft_box = None
try:
    from shared.tool_help import render_tool_help
except ImportError:
    render_tool_help = None
try:
    from shared.feedback_button import render_feedback_button
except ImportError:
    render_feedback_button = None
try:
    from shared.box_folder_browser import render_box_folder_browser
    from shared.box_client import parse_folder_id as _parse_folder_id
except ImportError:
    render_box_folder_browser = None
    _parse_folder_id = None

# ── Page config ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Declaration Drafter — O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ──────────────────────────────────────────────────────────────────────

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* Hide Streamlit chrome */
#MainMenu, footer,
div[data-testid="stToolbar"] { display: none !important; }

.stApp {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

/* Navigation bar */
.nav-bar {
    display: flex;
    align-items: center;
    padding: 10px 4px;
    margin: -1rem 0 1.2rem 0;
    border-bottom: 1px solid rgba(0,0,0,0.07);
}
.nav-back {
    display: flex;
    align-items: center;
    gap: 6px;
    font-family: 'Inter', sans-serif;
    font-size: 0.85rem;
    font-weight: 500;
    color: #0066CC;
    text-decoration: none;
    min-width: 150px;
}
.nav-back:hover { color: #004499; text-decoration: underline; }
.nav-title {
    flex: 1;
    text-align: center;
    font-family: 'Inter', sans-serif;
    font-size: 1.15rem;
    font-weight: 700;
    color: #1a2744;
    letter-spacing: -0.02em;
}
.nav-firm {
    font-weight: 400;
    color: #86868b;
    font-size: 0.85rem;
    margin-left: 8px;
}
.nav-spacer { min-width: 150px; }

/* Header */
.dd-header {
    color: #1a2744;
    font-size: 1.8rem;
    font-weight: 800;
    letter-spacing: -0.03em;
    margin-bottom: 0;
}
.dd-sub {
    color: #5a6a85;
    font-size: 0.95rem;
    margin-bottom: 0.5rem;
}

/* Progress bar label */
.progress-label {
    font-size: 0.8rem;
    font-weight: 600;
    color: #5a6a85;
    margin-bottom: 4px;
}

/* Section instructions */
.sec-instructions {
    color: #5a6a85;
    font-size: 0.88rem;
    line-height: 1.5;
    margin-bottom: 1rem;
    padding: 10px 14px;
    background: #f0f4fa;
    border-radius: 8px;
    border-left: 3px solid #4a7ddb;
}

/* Coaching tip */
.coaching-tip {
    font-size: 0.82rem;
    color: #4a6741;
    background: #f0f7ee;
    border: 1px solid #d4e8cf;
    border-radius: 6px;
    padding: 10px 14px;
    margin-top: -8px;
    margin-bottom: 12px;
    line-height: 1.5;
}

/* Preview panel */
.preview-panel {
    font-family: 'Times New Roman', Times, serif;
    font-size: 0.88rem;
    line-height: 1.8;
    background: #fff;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    padding: 24px 28px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
    max-height: 70vh;
    overflow-y: auto;
}
.preview-panel .doc-title {
    text-align: center;
    font-weight: bold;
    font-size: 1.05rem;
    margin-bottom: 4px;
}
.preview-panel .doc-caption {
    text-align: center;
    font-size: 0.82rem;
    color: #666;
    margin-bottom: 8px;
}
.preview-panel .doc-type {
    text-align: center;
    font-style: italic;
    font-size: 0.88rem;
    margin-bottom: 16px;
}
.preview-panel .doc-para {
    text-indent: 0;
    margin-bottom: 8px;
}
.preview-panel .doc-sig {
    margin-top: 24px;
    line-height: 2;
}
.preview-panel .doc-interp {
    margin-top: 28px;
    padding-top: 16px;
    border-top: 1px solid #ddd;
}

/* Draft badge */
.draft-badge {
    display: inline-block;
    padding: 3px 10px;
    font-size: 0.7rem;
    font-weight: 600;
    background: #e8f0fe;
    color: #1a73e8;
    border-radius: 12px;
    margin-bottom: 8px;
}

/* Saved confirmation */
.saved-toast {
    font-size: 0.8rem;
    color: #2e7d32;
    font-weight: 600;
}
</style>
""",
    unsafe_allow_html=True,
)

from shared.auth import require_auth, render_logout
require_auth()

# ── Navigation bar ───────────────────────────────────────────────────────────

st.markdown(
    """
<div class="nav-bar">
    <a href="http://localhost:8502" class="nav-back">&#8592; Staff Dashboard</a>
    <div class="nav-title">Declaration Drafter<span class="nav-firm">&mdash; O'Brien Immigration Law</span></div>
    <div class="nav-spacer"></div>
</div>
""",
    unsafe_allow_html=True,
)
render_logout()

render_client_banner()
if render_tool_help:
    render_tool_help("declaration-drafter")
if render_feedback_button:
    render_feedback_button("declaration-drafter")

# ── Session state defaults ───────────────────────────────────────────────────

_DEFAULTS: dict = {
    "draft_id": None,
    "last_saved_msg": "",
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state.draft_id is None:
    st.session_state.draft_id = new_draft_id()

# -- SF auto-fill for client fields -------------------------------------------
_sf_client = st.session_state.get("sf_client")
if _sf_client:
    _sf_cid = _sf_client.get("Id", "")
    _prev_cid = st.session_state.get("_sf_autofill_cid", "")
    if _sf_cid and _sf_cid != _prev_cid:
        _sf_name = _sf_client.get("Name", "")
        if _sf_name:
            st.session_state["inp_name"] = _sf_name
        _sf_country = _sf_client.get("Country__c", "")
        if _sf_country:
            st.session_state["inp_country"] = _sf_country
        _sf_anum = _sf_client.get("A_Number__c", "")
        if _sf_anum:
            st.session_state["inp_anumber"] = _sf_anum
        st.session_state["_sf_autofill_cid"] = _sf_cid


# ── Helpers ──────────────────────────────────────────────────────────────────

def _collect_answers(declaration_type: str) -> dict[str, str]:
    """Gather current answers from widget keys."""
    sections = get_declaration_prompts().get(declaration_type, [])
    answers: dict[str, str] = {}
    for section in sections:
        for q in section["questions"]:
            answers[q["id"]] = st.session_state.get(f"ta_{q['id']}", "")
    return answers


def _clear_all_answer_keys() -> None:
    """Clear all text_area widget keys across all declaration types."""
    for sections in get_declaration_prompts().values():
        for section in sections:
            for q in section["questions"]:
                key = f"ta_{q['id']}"
                if key in st.session_state:
                    del st.session_state[key]


def _do_save(declaration_type: str) -> None:
    """Save the current state as a draft."""
    answers = _collect_answers(declaration_type)
    name = st.session_state.get("inp_name", "")
    declarant = {
        "name": name,
        "country_of_origin": st.session_state.get("inp_country", ""),
        "a_number": st.session_state.get("inp_anumber", ""),
        "language": st.session_state.get("inp_language", "English"),
        "interpreter_name": st.session_state.get("inp_interpreter", ""),
    }
    save_draft(st.session_state.draft_id, declaration_type, declarant, answers)
    st.session_state.last_saved_msg = f"Saved — {name or 'draft'}"


def _do_load(draft_id: str) -> None:
    """Load a draft into session state and rerun."""
    draft = load_draft(draft_id)
    if not draft:
        return
    _clear_all_answer_keys()
    st.session_state.draft_id = draft["id"]
    st.session_state.inp_decl_type = draft["declaration_type"]
    d = draft.get("declarant", {})
    st.session_state.inp_name = d.get("name", "")
    st.session_state.inp_country = d.get("country_of_origin", "")
    st.session_state.inp_anumber = d.get("a_number", "")
    st.session_state.inp_language = d.get("language", "English")
    st.session_state.inp_interpreter = d.get("interpreter_name", "")
    for qid, answer in draft.get("answers", {}).items():
        st.session_state[f"ta_{qid}"] = answer


def _do_new() -> None:
    """Start a fresh declaration."""
    _clear_all_answer_keys()
    st.session_state.draft_id = new_draft_id()
    st.session_state.last_saved_msg = ""
    for k in ("inp_name", "inp_country", "inp_anumber", "inp_interpreter"):
        if k in st.session_state:
            del st.session_state[k]
    if "inp_language" in st.session_state:
        del st.session_state["inp_language"]
    if "inp_decl_type" in st.session_state:
        del st.session_state["inp_decl_type"]


def _build_preview_html(
    name: str,
    a_number: str,
    country: str,
    dec_type: str,
    paragraphs: list[str],
    language: str,
    interpreter_name: str,
) -> str:
    """Build document-like HTML preview."""
    esc = html_mod.escape
    parts: list[str] = []

    parts.append(f'<div class="doc-title">DECLARATION OF {esc(name.upper())}</div>')

    caption_bits = []
    if a_number:
        caption_bits.append(f"A# {esc(a_number)}")
    if country:
        caption_bits.append(f"Country of Origin: {esc(country)}")
    if caption_bits:
        parts.append(f'<div class="doc-caption">{" &middot; ".join(caption_bits)}</div>')

    parts.append(f'<div class="doc-type">{esc(dec_type)}</div>')

    parts.append(
        f'<div class="doc-para">I, {esc(name)}, hereby declare under penalty '
        "of perjury that the following statements are true and correct:</div>"
    )

    for idx, para in enumerate(paragraphs, start=1):
        parts.append(f'<div class="doc-para">{idx}. {esc(para)}</div>')

    parts.append(
        f'<div class="doc-para" style="margin-top:16px;">'
        f"{esc(PERJURY_CLAUSE.format(name=name))}</div>"
    )

    parts.append(
        f'<div class="doc-sig">'
        f"____________________________<br>{esc(name)}<br>"
        f"Date: {date.today().strftime('%m/%d/%Y')}</div>"
    )

    if language != "English" and interpreter_name:
        cert = INTERPRETER_CERT.format(
            interpreter_name=interpreter_name,
            language=language,
            declarant_name=name,
        )
        parts.append(
            f'<div class="doc-interp">'
            f'<div class="doc-title" style="font-size:0.92rem;">INTERPRETER CERTIFICATION</div>'
            f'<div class="doc-para">{esc(cert)}</div>'
            f'<div class="doc-sig">'
            f"____________________________<br>{esc(interpreter_name)}<br>"
            f"Date: {date.today().strftime('%m/%d/%Y')}</div>"
            f"</div>"
        )

    return "\n".join(parts)


def _build_docx(
    name: str,
    a_number: str,
    country: str,
    dec_type: str,
    paragraphs: list[str],
    language: str,
    interpreter_name: str,
) -> bytes:
    """Build a Word document and return its bytes."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, text, size=12, bold=False, italic=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return r

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"DECLARATION OF {name.upper()}", size=14, bold=True)

    # Caption
    if a_number or country:
        bits = []
        if a_number:
            bits.append(f"A# {a_number}")
        if country:
            bits.append(f"Country of Origin: {country}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, " | ".join(bits), size=10)

    # Type subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, dec_type, size=11, italic=True)

    # Opening statement
    p = doc.add_paragraph()
    _run(
        p,
        f"I, {name}, hereby declare under penalty of perjury "
        "that the following statements are true and correct:",
    )

    # Numbered paragraphs
    for idx, para_text in enumerate(paragraphs, start=1):
        p = doc.add_paragraph()
        _run(p, f"{idx}. {para_text}")
        fmt = p.paragraph_format
        fmt.space_after = Pt(6)
        fmt.line_spacing = 2.0

    # Perjury clause
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, PERJURY_CLAUSE.format(name=name))

    # Signature block
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, "____________________________")
    p = doc.add_paragraph()
    _run(p, name)
    p = doc.add_paragraph()
    _run(p, f"Date: {date.today().strftime('%m/%d/%Y')}")

    # Interpreter certification
    if language != "English" and interpreter_name:
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "INTERPRETER CERTIFICATION", size=12, bold=True)
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(
            p,
            INTERPRETER_CERT.format(
                interpreter_name=interpreter_name,
                language=language,
                declarant_name=name,
            ),
        )
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, "____________________________")
        p = doc.add_paragraph()
        _run(p, interpreter_name)
        p = doc.add_paragraph()
        _run(p, f"Date: {date.today().strftime('%m/%d/%Y')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Sidebar ──────────────────────────────────────────────────────────────────

with st.sidebar:
    # Draft management
    st.markdown("#### Drafts")
    btn_cols = st.columns(2)
    with btn_cols[0]:
        if st.button("New", use_container_width=True):
            _do_new()
            st.rerun()
    with btn_cols[1]:
        save_clicked = st.button("Save", use_container_width=True, type="primary")

    saved_drafts = list_drafts()
    if saved_drafts:
        labels_map = {d["id"]: f"{d['declarant_name']} — {d['declaration_type']}" for d in saved_drafts}
        draft_ids = list(labels_map.keys())
        selected_draft = st.selectbox(
            "Load a saved draft",
            options=[""] + draft_ids,
            format_func=lambda x: labels_map.get(x, "Select..."),
            label_visibility="collapsed",
        )
        load_cols = st.columns(2)
        with load_cols[0]:
            if selected_draft and st.button("Load", use_container_width=True):
                _do_load(selected_draft)
                st.rerun()
        with load_cols[1]:
            if selected_draft and st.button("Delete", use_container_width=True):
                delete_draft(selected_draft)
                st.rerun()

    if st.session_state.last_saved_msg:
        st.markdown(
            f'<div class="saved-toast">{html_mod.escape(st.session_state.last_saved_msg)}</div>',
            unsafe_allow_html=True,
        )

    st.divider()

    # Declaration type
    declaration_type = st.selectbox(
        "Declaration Type",
        options=get_declaration_types(),
        key="inp_decl_type",
    )

    st.divider()

    # Box folder browser
    if render_box_folder_browser and _parse_folder_id:
        _sf = st.session_state.get("sf_client")
        _box_raw = (_sf.get("Box_Folder_Id__c", "") or "") if _sf else ""
        if _box_raw:
            render_box_folder_browser(
                _parse_folder_id(_box_raw),
                mode="viewer",
                key_prefix="_dd_box",
                header_label="Client Documents",
            )
            st.divider()

    # Language / interpreter
    st.markdown("#### Language")
    LANGUAGES = [
        "English", "Spanish", "French", "Portuguese", "Arabic",
        "Mandarin", "Cantonese", "Hindi", "Urdu", "Bengali",
        "Tagalog", "Korean", "Russian", "Haitian Creole",
        "Amharic", "Tigrinya", "Somali", "Swahili", "Nepali",
        "Burmese", "Vietnamese", "Other",
    ]
    language = st.selectbox("Language", options=LANGUAGES, key="inp_language")

    interpreter_name = ""
    if language != "English":
        interpreter_name = st.text_input(
            "Interpreter Name",
            key="inp_interpreter",
            help="Required for interpreter certification block",
        )

    render_tool_notes("declaration-drafter")



# ── Handle save (after sidebar widgets render so we can read values) ─────

if save_clicked:
    _do_save(declaration_type)
    st.rerun()


# ── Main area ────────────────────────────────────────────────────────────────

sections = get_declaration_prompts().get(declaration_type, [])

# Progress bar
total_qs = sum(len(s["questions"]) for s in sections)
filled_qs = sum(
    1
    for s in sections
    for q in s["questions"]
    if st.session_state.get(f"ta_{q['id']}", "").strip()
)
if total_qs > 0:
    pct = filled_qs / total_qs
    st.markdown(
        f'<div class="progress-label">{filled_qs} of {total_qs} questions answered</div>',
        unsafe_allow_html=True,
    )
    st.progress(pct)

# Two-column layout: questions left, preview right
q_col, preview_col = st.columns([3, 2], gap="large")

with q_col:
    if not sections:
        st.info("No prompts defined for this declaration type yet.")
    else:
        # Build tab labels with completion indicators
        tab_labels: list[str] = []
        for section in sections:
            n_total = len(section["questions"])
            n_filled = sum(
                1
                for q in section["questions"]
                if st.session_state.get(f"ta_{q['id']}", "").strip()
            )
            check = "+" if n_filled == n_total else f"{n_filled}/{n_total}"
            tab_labels.append(f"{section['title']} [{check}]")

        tabs = st.tabs(tab_labels)
        for tab, section in zip(tabs, sections):
            with tab:
                st.markdown(
                    f'<div class="sec-instructions">{html_mod.escape(section["instructions"])}</div>',
                    unsafe_allow_html=True,
                )
                for q in section["questions"]:
                    st.text_area(
                        q["label"],
                        key=f"ta_{q['id']}",
                        height=130,
                    )
                    tip = q.get("tip", "")
                    if tip:
                        with st.expander("Attorney tip", icon=":material/lightbulb:"):
                            st.markdown(tip)

with preview_col:
    st.markdown("**Declaration Preview**")

    if not declarant_name:
        st.info("Enter the declarant's name in the sidebar to see the preview.")
    else:
        answers = _collect_answers(declaration_type)
        paragraphs = format_numbered_paragraphs(answers, declaration_type)

        if not paragraphs:
            st.caption("Start answering questions to see the declaration take shape.")
        else:
            preview_html = _build_preview_html(
                name=declarant_name,
                a_number=a_number,
                country=country,
                dec_type=declaration_type,
                paragraphs=paragraphs,
                language=language,
                interpreter_name=interpreter_name,
            )
            st.markdown(
                f'<div class="preview-panel">{preview_html}</div>',
                unsafe_allow_html=True,
            )

        # Export controls
        st.markdown("---")
        st.markdown("**Export**")

        # Plain text
        plain_text = ""
        if paragraphs:
            from app.prompts import build_declaration_text

            plain_text = build_declaration_text(
                answers=answers,
                declaration_type=declaration_type,
                declarant_name=declarant_name,
                language=language,
                interpreter_name=interpreter_name,
            )

        exp_cols = st.columns(2)
        with exp_cols[0]:
            st.download_button(
                "Download .txt",
                data=plain_text or "No content yet.",
                file_name=f"Declaration_{declarant_name.replace(' ', '_')}.txt",
                mime="text/plain",
                use_container_width=True,
                disabled=not paragraphs,
            )
        with exp_cols[1]:
            if paragraphs:
                docx_bytes = _build_docx(
                    name=declarant_name,
                    a_number=a_number,
                    country=country,
                    dec_type=declaration_type,
                    paragraphs=paragraphs,
                    language=language,
                    interpreter_name=interpreter_name,
                )
                st.download_button(
                    "Download .docx",
                    data=docx_bytes,
                    file_name=f"Declaration_{declarant_name.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
                if st.button("Upload to Google Docs", use_container_width=True):
                    with st.spinner("Uploading to Google Docs..."):
                        try:
                            url = upload_to_google_docs(docx_bytes, f"Declaration - {declarant_name}")
                            st.session_state.google_doc_url = url
                        except Exception as e:
                            st.error(f"Upload failed: {e}")
                if st.session_state.get("google_doc_url"):
                    st.markdown(f"[Open Google Doc]({st.session_state.google_doc_url})")
            else:
                st.button("Download .docx", disabled=True, use_container_width=True)

    # Draft Box (always visible — uses whatever data is on the page)
    if render_draft_box is not None:
        _answers = _collect_answers(declaration_type)
        _paras = format_numbered_paragraphs(_answers, declaration_type)
        _draft_content = ""
        if _paras:
            from app.prompts import build_declaration_text as _build_decl_text
            _draft_content = _build_decl_text(
                answers=_answers,
                declaration_type=declaration_type,
                declarant_name=declarant_name or "Declarant",
                language=language,
                interpreter_name=interpreter_name,
            )
        render_draft_box("declaration-drafter", {
            "document_type": "declaration",
            "client_name": declarant_name,
            "case_id": st.session_state.get("draft_id", ""),
            "content": _draft_content,
        })
