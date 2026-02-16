"""FastAPI backend for the Declaration Drafter tool."""

from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.drafts import delete_draft, list_drafts, load_draft, new_draft_id, save_draft
from app.prompts import (
    DECLARATION_PROMPTS,
    DECLARATION_TYPES,
    INTERPRETER_CERT,
    PERJURY_CLAUSE,
    build_declaration_text,
    format_numbered_paragraphs,
)

app = FastAPI(title="Declaration Drafter API")


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------


class DeclarantInfo(BaseModel):
    name: str
    country_of_origin: str = ""
    language: str = "English"
    a_number: str = ""
    interpreter_name: str = ""


class GenerateRequest(BaseModel):
    declaration_type: str
    declarant: DeclarantInfo
    answers: dict[str, str]


class ExportDocxRequest(BaseModel):
    declaration_type: str
    declarant: DeclarantInfo
    answers: dict[str, str]


class SaveDraftRequest(BaseModel):
    draft_id: str | None = None
    declaration_type: str
    declarant: DeclarantInfo
    answers: dict[str, str]


# ---------------------------------------------------------------------------
# Declaration endpoints
# ---------------------------------------------------------------------------


@app.get("/api/declaration-types")
def list_declaration_types() -> list[str]:
    """Return the available declaration types."""
    return DECLARATION_TYPES


@app.get("/api/prompts/{declaration_type}")
def get_prompts(declaration_type: str) -> list[dict]:
    """Return the guided prompt sections for a declaration type."""
    sections = DECLARATION_PROMPTS.get(declaration_type, [])
    if not sections:
        raise HTTPException(404, f"Unknown declaration type: {declaration_type}")
    return sections


@app.post("/api/generate")
def generate_declaration(req: GenerateRequest) -> dict:
    """Assemble a declaration from answers."""
    text = build_declaration_text(
        answers=req.answers,
        declaration_type=req.declaration_type,
        declarant_name=req.declarant.name,
        language=req.declarant.language,
        interpreter_name=req.declarant.interpreter_name,
    )
    paragraphs = format_numbered_paragraphs(req.answers, req.declaration_type)
    return {
        "declaration_text": text,
        "paragraph_count": len(paragraphs),
        "declaration_type": req.declaration_type,
        "declarant_name": req.declarant.name,
    }


@app.post("/api/export/docx")
def export_docx(req: ExportDocxRequest):
    """Export the declaration to a Word document."""
    paragraphs = format_numbered_paragraphs(req.answers, req.declaration_type)
    doc = _build_docx(
        declarant_name=req.declarant.name,
        country_of_origin=req.declarant.country_of_origin,
        a_number=req.declarant.a_number,
        declaration_type=req.declaration_type,
        paragraphs=paragraphs,
        language=req.declarant.language,
        interpreter_name=req.declarant.interpreter_name,
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = req.declarant.name.replace(" ", "_")
    filename = f"Declaration_{safe_name}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Draft endpoints
# ---------------------------------------------------------------------------


@app.get("/api/drafts")
def api_list_drafts() -> list[dict]:
    """List all saved drafts (newest first)."""
    return list_drafts()


@app.get("/api/drafts/{draft_id}")
def api_get_draft(draft_id: str) -> dict:
    """Load a specific draft."""
    draft = load_draft(draft_id)
    if not draft:
        raise HTTPException(404, f"Draft not found: {draft_id}")
    return draft


@app.post("/api/drafts")
def api_save_draft(req: SaveDraftRequest) -> dict:
    """Create or update a draft."""
    draft_id = req.draft_id or new_draft_id()
    declarant = {
        "name": req.declarant.name,
        "country_of_origin": req.declarant.country_of_origin,
        "a_number": req.declarant.a_number,
        "language": req.declarant.language,
        "interpreter_name": req.declarant.interpreter_name,
    }
    saved = save_draft(draft_id, req.declaration_type, declarant, req.answers)
    return saved


@app.delete("/api/drafts/{draft_id}")
def api_delete_draft(draft_id: str) -> dict:
    """Delete a draft."""
    if not delete_draft(draft_id):
        raise HTTPException(404, f"Draft not found: {draft_id}")
    return {"status": "deleted", "id": draft_id}


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------


def _build_docx(
    declarant_name: str,
    country_of_origin: str,
    a_number: str,
    declaration_type: str,
    paragraphs: list[str],
    language: str = "English",
    interpreter_name: str = "",
) -> Document:
    """Build a python-docx Document with proper declaration formatting."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, text, size=12, bold=False, italic=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return r

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"DECLARATION OF {declarant_name.upper()}", size=14, bold=True)

    # Caption
    if a_number or country_of_origin:
        bits = []
        if a_number:
            bits.append(f"A# {a_number}")
        if country_of_origin:
            bits.append(f"Country of Origin: {country_of_origin}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, " | ".join(bits), size=10)

    # Type subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, declaration_type, size=11, italic=True)

    # Opening statement
    p = doc.add_paragraph()
    _run(
        p,
        f"I, {declarant_name}, hereby declare under penalty of perjury "
        "that the following statements are true and correct:",
    )

    # Numbered paragraphs
    for idx, para_text in enumerate(paragraphs, start=1):
        p = doc.add_paragraph()
        _run(p, f"{idx}. {para_text}")
        fmt = p.paragraph_format
        fmt.space_after = Pt(6)
        fmt.line_spacing = 2.0

    # Perjury clause
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, PERJURY_CLAUSE.format(name=declarant_name))

    # Signature block
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, "____________________________")
    p = doc.add_paragraph()
    _run(p, declarant_name)
    p = doc.add_paragraph()
    _run(p, f"Date: {date.today().strftime('%B %d, %Y')}")

    # Interpreter certification
    if language != "English" and interpreter_name:
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "INTERPRETER CERTIFICATION", size=12, bold=True)
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(
            p,
            INTERPRETER_CERT.format(
                interpreter_name=interpreter_name,
                language=language,
                declarant_name=declarant_name,
            ),
        )
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, "____________________________")
        p = doc.add_paragraph()
        _run(p, interpreter_name)
        p = doc.add_paragraph()
        _run(p, f"Date: {date.today().strftime('%B %d, %Y')}")

    return doc
