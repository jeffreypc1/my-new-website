"""Evidence management for the Evidence Indexer tool.

Provides structured handling of evidence packages for immigration cases,
including document categorization, exhibit lettering, index generation,
case persistence via JSON files, and Word document export.

Part of the O'Brien Immigration Law tool suite.
"""

from __future__ import annotations

import io
import json
import uuid
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Any

import sys as _sys
_sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.config_store import get_config_value


# ---------------------------------------------------------------------------
# Storage
# ---------------------------------------------------------------------------

DATA_DIR = Path(__file__).resolve().parent.parent / "data" / "cases"
DATA_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------


@dataclass
class EvidenceItem:
    """A single document in an evidence package."""

    exhibit_letter: str
    title: str
    category: str
    page_count: int = 0
    date_added: str = ""
    box_url: str = ""
    description: str = ""
    doc_id: str = ""


# ---------------------------------------------------------------------------
# Standard immigration evidence categories
# ---------------------------------------------------------------------------

_DEFAULT_DOCUMENT_CATEGORIES: list[str] = [
    "Identity Documents",
    "Country Conditions",
    "Medical/Psychological",
    "Expert Reports",
    "Declarations",
    "Photographs",
    "Government Documents",
    "Correspondence",
    "Employment Records",
    "Financial Records",
    "Legal Documents",
    "Other",
]

# ── Config-aware loading (JSON override with hardcoded fallback) ─────────────
DOCUMENT_CATEGORIES: list[str] = get_config_value("evidence-indexer", "document_categories", _DEFAULT_DOCUMENT_CATEGORIES)


# ---------------------------------------------------------------------------
# Exhibit lettering utilities
# ---------------------------------------------------------------------------


def _exhibit_letter(index: int) -> str:
    """Convert 0-based index to exhibit letter: 0->A, 1->B, ..., 25->Z, 26->AA.

    This mirrors the implementation in country-reports-tool/app/exhibit_compiler.py.
    """
    if index < 26:
        return chr(ord("A") + index)
    return chr(ord("A") + index // 26 - 1) + chr(ord("A") + index % 26)


def auto_assign_letters(items: list[EvidenceItem]) -> list[EvidenceItem]:
    """Assign sequential exhibit letters (Tab A, Tab B, ...) to evidence items.

    Preserves any manually-assigned letters and fills gaps with auto-assigned
    sequential letters for items with empty exhibit_letter fields.

    Args:
        items: List of EvidenceItem objects in desired exhibit order.

    Returns:
        The same list with exhibit_letter fields populated.
    """
    next_index = 0
    for item in items:
        if not item.exhibit_letter:
            item.exhibit_letter = _exhibit_letter(next_index)
        next_index += 1
    return items


def reorder_exhibits(
    items: list[EvidenceItem],
    new_order: list[int],
) -> list[EvidenceItem]:
    """Reorder exhibits according to a new index ordering and reassign letters.

    Args:
        items: Current list of EvidenceItem objects.
        new_order: List of current indices in the desired new order.
            e.g. [2, 0, 1] means: current item 2 becomes first, then 0, then 1.

    Returns:
        Reordered list with reassigned exhibit letters.
    """
    reordered = [items[i] for i in new_order if 0 <= i < len(items)]

    # Reassign sequential letters after reordering
    for idx, item in enumerate(reordered):
        item.exhibit_letter = _exhibit_letter(idx)

    return reordered


# ---------------------------------------------------------------------------
# Index generation
# ---------------------------------------------------------------------------


def generate_index(items: list[EvidenceItem]) -> list[dict]:
    """Generate a structured exhibit index suitable for Word doc export.

    Args:
        items: List of EvidenceItem objects in exhibit order.

    Returns:
        List of dicts with keys: exhibit_letter, title, category, page_count,
        description. Suitable for rendering into a Word table.
    """
    index_rows: list[dict] = []
    for item in items:
        index_rows.append({
            "exhibit_letter": item.exhibit_letter,
            "title": item.title,
            "category": item.category,
            "page_count": item.page_count,
            "description": item.description,
            "date_added": item.date_added,
        })
    return index_rows


def generate_index_docx(items: list[EvidenceItem], case_name: str = "") -> bytes:
    """Generate a Word document exhibit index with proper legal formatting.

    Creates a formatted Word document with 1-inch margins, Times New Roman 12pt,
    centered title, case caption, and a table listing all exhibits. Suitable for
    filing with the immigration court or USCIS.

    Args:
        items: List of EvidenceItem objects in exhibit order.
        case_name: Client name or case caption for the document header.

    Returns:
        Bytes of the generated .docx file.
    """
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        doc = Document()

        # Set 1-inch margins on all sections
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        def _run(para, text: str, size: int = 12, bold: bool = False, italic: bool = False):
            r = para.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
            r.bold = bold
            r.italic = italic
            return r

        # Centered title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "EXHIBIT INDEX", size=14, bold=True)

        # Case caption
        if case_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, f"In the Matter of: {case_name}", size=12, italic=True)

        doc.add_paragraph()  # spacer

        # Create exhibit table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Bold header row
        header_cells = table.rows[0].cells
        headers = ["Exhibit", "Document Title", "Category", "Pages"]
        for i, header_text in enumerate(headers):
            header_cells[i].text = ""
            p = header_cells[i].paragraphs[0]
            _run(p, header_text, size=11, bold=True)

        # Data rows
        for item in items:
            row = table.add_row().cells
            row_data = [
                f"Tab {item.exhibit_letter}",
                item.title,
                item.category,
                str(item.page_count) if item.page_count else "",
            ]
            for i, cell_text in enumerate(row_data):
                row[i].text = ""
                p = row[i].paragraphs[0]
                _run(p, cell_text, size=11)

        # Footer
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p, f"Generated: {date.today().strftime('%B %d, %Y')}", size=9, italic=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p, "O'Brien Immigration Law", size=9, italic=True)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    except ImportError:
        # Graceful degradation if python-docx is not installed yet
        return b""


# ---------------------------------------------------------------------------
# Persistence helpers
# ---------------------------------------------------------------------------


def _case_path(case_id: str) -> Path:
    """Return the JSON file path for a given case ID."""
    # Sanitize to prevent path traversal
    safe_id = "".join(c for c in case_id if c.isalnum() or c in "-_")
    return DATA_DIR / f"{safe_id}.json"


def new_case_id() -> str:
    """Generate a unique case ID (8 hex chars)."""
    return uuid.uuid4().hex[:8]


def _make_doc_id() -> str:
    """Generate a unique document ID (12 hex chars)."""
    return uuid.uuid4().hex[:12]


# ---------------------------------------------------------------------------
# CRUD functions
# ---------------------------------------------------------------------------


def save_case(
    case_id: str,
    client_name: str,
    a_number: str,
    documents: list[dict[str, Any]],
) -> dict[str, Any]:
    """Save a case with its documents to disk.

    Args:
        case_id: Unique case identifier.
        client_name: Full legal name of the client.
        a_number: Alien registration number.
        documents: List of document dicts, each containing doc_id, title,
            category, description, page_count, box_url, exhibit_letter,
            date_added.

    Returns:
        The saved case dict.
    """
    now = datetime.now().isoformat(timespec="seconds")

    # Check if the case already exists (preserve created_at)
    existing = load_case(case_id)
    created_at = existing["created_at"] if existing else now

    case_data: dict[str, Any] = {
        "id": case_id,
        "client_name": client_name,
        "a_number": a_number,
        "documents": documents,
        "created_at": created_at,
        "updated_at": now,
    }

    path = _case_path(case_id)
    path.write_text(json.dumps(case_data, indent=2, default=str))
    return case_data


def load_case(case_id: str) -> dict[str, Any] | None:
    """Load a single case by its ID.

    Returns:
        The case dict, or None if not found or unreadable.
    """
    path = _case_path(case_id)
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text())
    except (json.JSONDecodeError, OSError):
        return None


def list_cases() -> list[dict[str, Any]]:
    """Return summary info for all saved cases, sorted by most-recently updated.

    Returns:
        List of dicts with case_id, client_name, document_count, updated_at.
    """
    cases: list[dict[str, Any]] = []
    if not DATA_DIR.exists():
        return cases
    for path in DATA_DIR.glob("*.json"):
        try:
            data = json.loads(path.read_text())
            cases.append({
                "case_id": data.get("id", ""),
                "client_name": data.get("client_name", ""),
                "document_count": len(data.get("documents", [])),
                "updated_at": data.get("updated_at", ""),
            })
        except (json.JSONDecodeError, OSError):
            continue
    cases.sort(key=lambda c: c.get("updated_at", ""), reverse=True)
    return cases


def delete_case(case_id: str) -> bool:
    """Delete a case from disk.

    Returns:
        True if the file was deleted, False if it did not exist.
    """
    path = _case_path(case_id)
    if path.exists():
        path.unlink()
        return True
    return False


def add_document(
    case_id: str,
    title: str,
    category: str,
    description: str = "",
    page_count: int = 0,
    box_url: str = "",
) -> dict[str, Any] | None:
    """Add a document to an existing case.

    Auto-assigns an exhibit letter based on position. Returns the new
    document dict, or None if the case was not found.

    Args:
        case_id: The case identifier.
        title: Document title.
        category: Document category (from DOCUMENT_CATEGORIES).
        description: Brief description of the document.
        page_count: Number of pages.
        box_url: Link to the document in Box.

    Returns:
        The newly added document dict with auto-assigned exhibit letter.
    """
    case_data = load_case(case_id)
    if case_data is None:
        return None

    documents = case_data.get("documents", [])

    doc: dict[str, Any] = {
        "doc_id": _make_doc_id(),
        "title": title,
        "category": category,
        "description": description,
        "page_count": page_count,
        "box_url": box_url,
        "exhibit_letter": "",
        "date_added": date.today().isoformat(),
    }
    documents.append(doc)

    # Auto-assign exhibit letters to all documents
    items = _docs_to_items(documents)
    items = auto_assign_letters(items)
    for i, item in enumerate(items):
        documents[i]["exhibit_letter"] = item.exhibit_letter

    save_case(
        case_id=case_data["id"],
        client_name=case_data.get("client_name", ""),
        a_number=case_data.get("a_number", ""),
        documents=documents,
    )

    return doc


def remove_document(case_id: str, doc_id: str) -> bool:
    """Remove a document from a case and reassign exhibit letters.

    Args:
        case_id: The case identifier.
        doc_id: The document ID to remove.

    Returns:
        True if the document was removed, False otherwise.
    """
    case_data = load_case(case_id)
    if case_data is None:
        return False

    documents = case_data.get("documents", [])
    original_count = len(documents)
    documents = [d for d in documents if d.get("doc_id") != doc_id]

    if len(documents) == original_count:
        return False

    # Reassign exhibit letters after removal
    items = _docs_to_items(documents)
    for idx, item in enumerate(items):
        item.exhibit_letter = _exhibit_letter(idx)
        documents[idx]["exhibit_letter"] = item.exhibit_letter

    save_case(
        case_id=case_data["id"],
        client_name=case_data.get("client_name", ""),
        a_number=case_data.get("a_number", ""),
        documents=documents,
    )

    return True


def update_document(
    case_id: str,
    doc_id: str,
    **fields: Any,
) -> dict[str, Any] | None:
    """Update fields on an existing document.

    Args:
        case_id: The case identifier.
        doc_id: The document ID to update.
        **fields: Keyword arguments for fields to update (title, category,
            description, page_count, box_url).

    Returns:
        The updated document dict, or None if not found.
    """
    case_data = load_case(case_id)
    if case_data is None:
        return None

    documents = case_data.get("documents", [])
    target = None
    for doc in documents:
        if doc.get("doc_id") == doc_id:
            target = doc
            break

    if target is None:
        return None

    allowed_fields = {"title", "category", "description", "page_count", "box_url"}
    for key, value in fields.items():
        if key in allowed_fields:
            target[key] = value

    save_case(
        case_id=case_data["id"],
        client_name=case_data.get("client_name", ""),
        a_number=case_data.get("a_number", ""),
        documents=documents,
    )

    return target


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _docs_to_items(documents: list[dict[str, Any]]) -> list[EvidenceItem]:
    """Convert stored document dicts to EvidenceItem objects."""
    return [
        EvidenceItem(
            exhibit_letter=d.get("exhibit_letter", ""),
            title=d.get("title", ""),
            category=d.get("category", "Other"),
            page_count=d.get("page_count", 0),
            date_added=d.get("date_added", ""),
            box_url=d.get("box_url", ""),
            description=d.get("description", ""),
            doc_id=d.get("doc_id", ""),
        )
        for d in documents
    ]
