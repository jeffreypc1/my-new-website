"""Convert template canvas sections → Word (.docx) and PDF documents.

Parses the Quill-generated HTML from each section, mapping formatting to
python-docx paragraph/run styles and fpdf2 PDF output. Follows the same
conventions as shared/draft_box.py (1-inch margins, Times New Roman 12pt,
BytesIO → bytes).

Merge field resolution: call ``resolve_merge_fields_for_export()`` to build
a {alias: value} dict from the active SF client + legal case, falling back
to realistic sample data when no client is loaded.
"""

from __future__ import annotations

import io
import re
from datetime import date
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Twips


def _get_letterhead_path() -> Path | None:
    """Return the path to the letterhead image if one has been uploaded."""
    config_dir = Path(__file__).resolve().parent.parent.parent / "data" / "config"
    for ext in ("png", "jpg", "jpeg"):
        p = config_dir / f"letterhead.{ext}"
        if p.exists():
            return p
    return None


# Max letterhead dimensions — keeps header compact regardless of upload size
_LH_MAX_WIDTH_IN = 6.5   # full page width (8.5 - 1" margins)
_LH_MAX_HEIGHT_IN = 1.5  # prevent oversized headers


def _letterhead_dims_inches(path: Path) -> tuple[float, float]:
    """Return (width, height) in inches, fit within max bounds keeping aspect ratio."""
    from PIL import Image

    with Image.open(path) as img:
        w_px, h_px = img.size

    aspect = w_px / h_px
    # Start at max width, compute height
    w = _LH_MAX_WIDTH_IN
    h = w / aspect
    # If too tall, constrain by height instead
    if h > _LH_MAX_HEIGHT_IN:
        h = _LH_MAX_HEIGHT_IN
        w = h * aspect
    return w, h


# ── Sample data for preview when no client is loaded ─────────────────────────

_SAMPLE_VALUES: dict[str, str] = {
    # Contact
    "first_name": "Maria",
    "last_name": "Garcia Lopez",
    "name": "Maria Garcia Lopez",
    "a_number": "123-456-789",
    "country": "Guatemala",
    "language": "Spanish",
    "email": "maria.garcia@email.com",
    "phone": "(555) 234-5678",
    "mobile": "(555) 987-6543",
    "dob": "1990-03-15",
    "gender": "Female",
    "marital_status": "Married",
    "immigration_status": "Respondent in Removal Proceedings",
    "court": "San Francisco Immigration Court",
    "case_type": "Asylum",
    "case_number": "A123-456-789",
    "client_status": "Active",
    "city_of_birth": "Guatemala City",
    "spouse": "Carlos Garcia",
    "customer_id": "CLI-2024-0042",
    # Legal Case
    "hearing_date": "2026-04-15",
    "hearing_time": "1:00 PM",
    "court_location": "San Francisco, CA",
    "dhs_address": "100 Montgomery St, Suite 800, San Francisco, CA 94104",
    "judge_name": "Hon. Dana L. Marks",
    "next_date_type": "Individual Hearing",
    "court_address": "120 Montgomery St, Suite 800, San Francisco, CA 94104",
    "filing_method": "Paper",
    "bar_number": "123456",
    "receipt_number": "ZSF-26-0001234",
    # Attorney / User
    "attorney_name": "Jeffrey P. O'Brien",
    "attorney_email": "jeff@obrienimmigration.com",
    "attorney_phone": "(555) 123-4567",
    # Firm (not from SF but commonly used in templates)
    "firm_name": "O'Brien Immigration Law",
    "firm_address": "123 Main Street, Suite 400, San Francisco, CA 94105",
    "date": date.today().strftime("%B %d, %Y"),
}


def resolve_merge_fields_for_export(
    sf_client: dict | None = None,
    legal_case: dict | None = None,
) -> dict[str, str]:
    """Build {alias: value} for merge field resolution.

    If *sf_client* is provided, pull real values from the Contact and
    (optionally) its selected_legal_case. Any alias not found in real data
    falls back to ``_SAMPLE_VALUES``.
    """
    from app.merge_fields_store import get_enabled_merge_fields_for_resolution

    field_map = get_enabled_merge_fields_for_resolution()
    values: dict[str, str] = dict(_SAMPLE_VALUES)  # start with sample

    if sf_client:
        # Resolve Contact fields
        for alias, info in field_map.items():
            obj = info["sf_object"]
            api = info["sf_api_name"]
            if obj == "Contact":
                val = sf_client.get(api)
                if val:
                    values[alias] = str(val)
            elif obj == "Legal_Case__c" and legal_case:
                val = legal_case.get(api)
                if val:
                    values[alias] = str(val)
            elif obj == "User":
                # Attorney info sometimes stored on Legal Case owner
                if legal_case:
                    val = legal_case.get(api)
                    if val:
                        values[alias] = str(val)

    # Always update date to today
    values["date"] = date.today().strftime("%B %d, %Y")
    return values


def _apply_merge_fields(html: str, values: dict[str, str]) -> str:
    """Replace {alias} placeholders in HTML with resolved values."""
    if not values:
        return html

    def _replace(m):
        alias = m.group(1)
        return values.get(alias, m.group(0))  # leave unresolved as-is

    # Match {word} but not HTML attributes or URLs
    return re.sub(r"\{([a-zA-Z_][a-zA-Z0-9_]*)\}", _replace, html)


# ── HTML parser ──────────────────────────────────────────────────────────────


class _QuillHTMLParser(HTMLParser):
    """Parse Quill HTML into a list of paragraph dicts.

    Each dict: {"text_runs": [{"text": str, "bold": bool, "italic": bool,
    "underline": bool}], "align": str, "list_type": str|None, "heading": int}
    """

    def __init__(self):
        super().__init__()
        self.paragraphs: list[dict] = []
        self._current_para: dict | None = None
        self._fmt_stack: list[dict] = []
        self._in_li = False
        self._list_stack: list[str] = []  # "ol" or "ul"

    def _ensure_para(self):
        if self._current_para is None:
            self._current_para = {
                "text_runs": [],
                "align": "left",
                "list_type": None,
                "heading": 0,
            }

    def _current_fmt(self) -> dict:
        return self._fmt_stack[-1] if self._fmt_stack else {"bold": False, "italic": False, "underline": False}

    def _close_para(self):
        if self._current_para is not None:
            self.paragraphs.append(self._current_para)
            self._current_para = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]):
        attr_dict = dict(attrs)
        cls = attr_dict.get("class", "")

        if tag == "p":
            self._close_para()
            self._ensure_para()
            if "ql-align-center" in cls:
                self._current_para["align"] = "center"
            elif "ql-align-right" in cls:
                self._current_para["align"] = "right"
            elif "ql-align-justify" in cls:
                self._current_para["align"] = "justify"

        elif tag in ("h1", "h2", "h3"):
            self._close_para()
            self._ensure_para()
            self._current_para["heading"] = int(tag[1])

        elif tag == "ol":
            self._list_stack.append("ol")
        elif tag == "ul":
            self._list_stack.append("ul")
        elif tag == "li":
            self._close_para()
            self._ensure_para()
            self._in_li = True
            if self._list_stack:
                self._current_para["list_type"] = self._list_stack[-1]

        elif tag == "br":
            self._ensure_para()
            # Treat <br> as a newline within the paragraph
            self._current_para["text_runs"].append({**self._current_fmt(), "text": "\n"})

        elif tag == "strong" or tag == "b":
            fmt = {**self._current_fmt(), "bold": True}
            self._fmt_stack.append(fmt)
        elif tag == "em" or tag == "i":
            fmt = {**self._current_fmt(), "italic": True}
            self._fmt_stack.append(fmt)
        elif tag == "u":
            fmt = {**self._current_fmt(), "underline": True}
            self._fmt_stack.append(fmt)
        elif tag == "span":
            # Merge field spans — just pass through text
            pass

    def handle_endtag(self, tag: str):
        if tag == "p" or tag in ("h1", "h2", "h3"):
            self._close_para()
        elif tag == "li":
            self._close_para()
            self._in_li = False
        elif tag == "ol" or tag == "ul":
            if self._list_stack:
                self._list_stack.pop()
        elif tag in ("strong", "b", "em", "i", "u"):
            if self._fmt_stack:
                self._fmt_stack.pop()

    def handle_data(self, data: str):
        if not data:
            return
        self._ensure_para()
        fmt = self._current_fmt()
        self._current_para["text_runs"].append({**fmt, "text": data})

    def close(self):
        self._close_para()
        super().close()


def _parse_html(html: str) -> list[dict]:
    """Parse Quill HTML into structured paragraph dicts."""
    parser = _QuillHTMLParser()
    parser.feed(html or "")
    parser.close()
    return parser.paragraphs


# ── Docx builder ─────────────────────────────────────────────────────────────

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _run(para, text: str, *, bold=False, italic=False, underline=False, size=12):
    """Add a formatted run to a paragraph."""
    r = para.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    return r


def build_template_docx(
    foundation_id: str,
    sections: list[dict],
    block_order: list[str],
    block_edits: dict[str, str],
    disabled_ids: list[str],
    merge_values: dict[str, str] | None = None,
    include_letterhead: bool = True,
) -> bytes:
    """Build a Word document from template canvas state.

    If *merge_values* is provided, {alias} placeholders are resolved first.
    Returns .docx file as bytes.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # Insert letterhead into page header if available
    letterhead_path = _get_letterhead_path() if include_letterhead else None
    if letterhead_path:
        lh_w, lh_h = _letterhead_dims_inches(letterhead_path)
        header = doc.sections[0].header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run()
        run.add_picture(str(letterhead_path), width=Inches(lh_w), height=Inches(lh_h))

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    sec_map = {s["id"]: s for s in sections}
    ol_counter = 0

    for sec_id in block_order:
        if sec_id in disabled_ids:
            continue
        sec = sec_map.get(sec_id)
        if not sec:
            continue

        html = block_edits.get(sec_id) or sec.get("content") or sec.get("default_content") or ""
        if merge_values:
            html = _apply_merge_fields(html, merge_values)
        if not html.strip():
            continue

        paragraphs = _parse_html(html)
        ol_counter = 0

        for pdata in paragraphs:
            runs = pdata["text_runs"]
            text_content = "".join(r["text"] for r in runs).strip()
            if not text_content and not any(r["text"].strip() for r in runs):
                # Empty paragraph — single blank line
                bp = doc.add_paragraph()
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(0)
                continue

            heading = pdata.get("heading", 0)
            list_type = pdata.get("list_type")

            if list_type == "ol":
                ol_counter += 1
                prefix = f"{ol_counter}. "
                p = doc.add_paragraph()
                _run(p, prefix, size=12)
            elif list_type == "ul":
                prefix = "\u2022 "
                p = doc.add_paragraph()
                _run(p, prefix, size=12)
            else:
                ol_counter = 0
                p = doc.add_paragraph()

            p.alignment = _ALIGN_MAP.get(pdata.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)

            size = 12
            if heading == 1:
                size = 16
            elif heading == 2:
                size = 14

            for rd in runs:
                txt = rd["text"]
                if not txt:
                    continue
                _run(
                    p, txt,
                    bold=rd.get("bold", False) or heading > 0,
                    italic=rd.get("italic", False),
                    underline=rd.get("underline", False),
                    size=size,
                )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF builder ──────────────────────────────────────────────────────────────

def build_template_pdf(
    foundation_id: str,
    sections: list[dict],
    block_order: list[str],
    block_edits: dict[str, str],
    disabled_ids: list[str],
    merge_values: dict[str, str] | None = None,
    include_letterhead: bool = True,
) -> bytes:
    """Build a PDF from template canvas state.

    If *merge_values* is provided, {alias} placeholders are resolved first.
    Returns PDF file as bytes.
    """
    from fpdf import FPDF

    letterhead_path = _get_letterhead_path() if include_letterhead else None
    lh_w_mm, lh_h_mm = (0.0, 0.0)
    if letterhead_path:
        lh_w, lh_h = _letterhead_dims_inches(letterhead_path)
        lh_w_mm = lh_w * 25.4
        lh_h_mm = lh_h * 25.4

    class _TemplPDF(FPDF):
        def header(self):
            if letterhead_path:
                # Center horizontally within page
                page_w = self.w - 25.4 * 2  # usable width
                x = 25.4 + (page_w - lh_w_mm) / 2
                self.image(str(letterhead_path), x=x, y=10, w=lh_w_mm, h=lh_h_mm)
                # Reset cursor to left margin, below the image
                self.set_xy(25.4, 10 + lh_h_mm + 5)

    pdf = _TemplPDF()
    pdf.set_margins(25.4, 25.4, 25.4)  # 1-inch — must be set before add_page
    pdf.set_auto_page_break(auto=True, margin=25.4)
    pdf.add_page()

    sec_map = {s["id"]: s for s in sections}
    ol_counter = 0

    for sec_id in block_order:
        if sec_id in disabled_ids:
            continue
        sec = sec_map.get(sec_id)
        if not sec:
            continue

        html = block_edits.get(sec_id) or sec.get("content") or sec.get("default_content") or ""
        if merge_values:
            html = _apply_merge_fields(html, merge_values)
        if not html.strip():
            continue

        paragraphs = _parse_html(html)
        ol_counter = 0

        for pdata in paragraphs:
            runs = pdata["text_runs"]
            text_content = "".join(r["text"] for r in runs).strip()
            if not text_content:
                pdf.ln(5)  # single blank line height
                continue

            heading = pdata.get("heading", 0)
            list_type = pdata.get("list_type")
            align = pdata.get("align", "left")

            # Map alignment
            align_map = {"left": "L", "center": "C", "right": "R", "justify": "J"}
            pdf_align = align_map.get(align, "L")

            size = 12
            if heading == 1:
                size = 16
            elif heading == 2:
                size = 14

            prefix = ""
            if list_type == "ol":
                ol_counter += 1
                prefix = f"{ol_counter}. "
            elif list_type == "ul":
                prefix = "\u2022 "
            else:
                ol_counter = 0

            # Build full text with prefix
            full_text = prefix + text_content

            # Determine dominant style
            has_bold = heading > 0 or any(r.get("bold") for r in runs)
            has_italic = any(r.get("italic") for r in runs)

            style = ""
            if has_bold:
                style += "B"
            if has_italic:
                style += "I"

            pdf.set_font("Times", style, size)
            pdf.set_x(25.4)  # ensure cursor is at left margin
            pdf.multi_cell(0, size * 0.4, full_text, align=pdf_align)

    return bytes(pdf.output())
