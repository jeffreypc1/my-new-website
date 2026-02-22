"""PDF compiler — merge exhibits with tab pages, stamp pagination, generate TOC.

Adapted from country-reports-tool/app/exhibit_compiler.py for the Document Assembler.
"""

from __future__ import annotations

import io
import os
from datetime import date

import pymupdf

# US Letter dimensions in points
_PAGE_W = 612
_PAGE_H = 792


_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp"}


def _to_pdf_bytes(raw: bytes, filename: str) -> bytes | None:
    """Convert raw file bytes to PDF. Returns None if unsupported format."""
    if not raw:
        return None

    # Already a PDF — return as-is
    if raw[:5] == b"%PDF-":
        return raw

    ext = os.path.splitext(filename)[1].lower()

    # Images — pymupdf opens them natively and converts to PDF
    if ext in _IMAGE_EXTS:
        try:
            doc = pymupdf.open(stream=raw, filetype=ext.lstrip("."))
            pdf_bytes = doc.convert_to_pdf()
            doc.close()
            return pdf_bytes
        except Exception:
            return None

    # DOCX — extract text via python-docx, render onto pymupdf pages
    if ext == ".docx":
        try:
            from docx import Document as DocxDocument

            docx_doc = DocxDocument(io.BytesIO(raw))
            pdf_doc = pymupdf.open()
            page = pdf_doc.new_page(width=_PAGE_W, height=_PAGE_H)
            margin = 72
            y = margin
            line_h = 14

            for para in docx_doc.paragraphs:
                text = para.text.strip()
                if not text:
                    y += line_h
                    if y > _PAGE_H - margin:
                        page = pdf_doc.new_page(width=_PAGE_W, height=_PAGE_H)
                        y = margin
                    continue

                # Wrap long text into the available width
                rect = pymupdf.Rect(margin, y, _PAGE_W - margin, _PAGE_H - margin)
                rc = page.insert_textbox(
                    rect,
                    text,
                    fontsize=11,
                    fontname="helv",
                )
                # rc < 0 means overflow — the text didn't fit
                if rc < 0:
                    page = pdf_doc.new_page(width=_PAGE_W, height=_PAGE_H)
                    y = margin
                    rect = pymupdf.Rect(margin, y, _PAGE_W - margin, _PAGE_H - margin)
                    rc = page.insert_textbox(
                        rect,
                        text,
                        fontsize=11,
                        fontname="helv",
                    )
                    # Advance y by the space used (estimate from return code)
                    y = _PAGE_H - margin - max(rc, 0)
                else:
                    y = _PAGE_H - margin - rc
                y += 4  # small gap between paragraphs

                if y > _PAGE_H - margin:
                    page = pdf_doc.new_page(width=_PAGE_W, height=_PAGE_H)
                    y = margin

            pdf_bytes = pdf_doc.tobytes()
            pdf_doc.close()
            return pdf_bytes
        except Exception:
            return None

    # Unsupported format
    return None


def _exhibit_letter(index: int) -> str:
    """Convert 0-based index to exhibit letter: 0→A, 1→B, ..., 25→Z, 26→AA."""
    if index < 26:
        return chr(ord("A") + index)
    return chr(ord("A") + index // 26 - 1) + chr(ord("A") + index % 26)


def _insert_tab_page(merged: pymupdf.Document, letter: str) -> None:
    """Insert a blank US-Letter page with centered 'TAB {letter}' label."""
    page = merged.new_page(width=_PAGE_W, height=_PAGE_H)
    text_rect = pymupdf.Rect(50, 340, _PAGE_W - 50, 440)
    page.insert_textbox(
        text_rect,
        f"TAB {letter}",
        fontsize=36,
        fontname="helv",
        align=1,  # center
        color=(0, 0, 0),
    )


def _stamp_page_number(page: pymupdf.Page, page_num: int) -> None:
    """Stamp a page number 0.5in from bottom, 0.5in from right.

    Draws a white mask to cover any existing content, then overlays
    the page number in 12pt Helvetica.
    """
    margin_x = 36  # 0.5 inch in points
    margin_y = 36
    box_w = 50
    box_h = 20

    rect = page.rect
    # Bottom-right position
    x1 = rect.width - margin_x
    x0 = x1 - box_w
    y1 = rect.height - margin_y
    y0 = y1 - box_h

    # White mask
    pad = 4
    mask = pymupdf.Rect(x0 - pad, y0 - pad, x1 + pad, y1 + pad)
    page.draw_rect(mask, color=(1, 1, 1), fill=(1, 1, 1))

    # Page number
    num_rect = pymupdf.Rect(x0, y0, x1, y1)
    page.insert_textbox(
        num_rect,
        str(page_num),
        fontsize=12,
        fontname="helv",
        align=2,  # right-align
        color=(0, 0, 0),
    )


def compile_exhibit_package(
    exhibits: list[dict],
    on_progress=None,
) -> tuple[bytes, list[dict]]:
    """Merge exhibits into a single paginated PDF with tab dividers.

    Args:
        exhibits: List of dicts with keys: id, letter, title, pdf_bytes.
        on_progress: Optional callback(message: str) for status updates.

    Returns:
        (compiled_pdf_bytes, toc_entries) where toc_entries is:
        [{letter, title, start_page, end_page, page_count}]
        Page numbers are content-only (tabs excluded).
    """
    merged = pymupdf.open()
    toc_entries: list[dict] = []

    # Track which physical pages are tab dividers
    tab_pages: set[int] = set()
    # Track physical page ranges per exhibit
    content_ranges: list[tuple[str, str, int, int]] = []  # (letter, title, start, end)

    for exhibit in exhibits:
        letter = exhibit["letter"]
        title = exhibit["title"]
        pdf_bytes = exhibit["pdf_bytes"]

        if on_progress:
            on_progress(f"Adding Tab {letter}: {title}...")

        # Insert tab divider
        tab_phys = len(merged)
        _insert_tab_page(merged, letter)
        tab_pages.add(tab_phys)

        # Convert non-PDF files before merging
        pdf_bytes = _to_pdf_bytes(pdf_bytes, title)
        if pdf_bytes is None:
            if on_progress:
                on_progress(f"Skipping {title} (unsupported file format)")
            continue

        # Insert exhibit pages
        try:
            doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
            start_phys = len(merged)
            merged.insert_pdf(doc)
            end_phys = len(merged) - 1
            doc.close()
            content_ranges.append((letter, title, start_phys, end_phys))
        except Exception as e:
            if on_progress:
                on_progress(f"Error adding {title}: {e}")
            continue

    if len(merged) == 0:
        merged.close()
        raise ValueError("No exhibits could be compiled.")

    # Build content-only page numbering (skip tab pages)
    content_num = 0
    phys_to_content: dict[int, int] = {}
    for phys in range(len(merged)):
        if phys not in tab_pages:
            content_num += 1
            phys_to_content[phys] = content_num

    # Stamp only content pages
    if on_progress:
        on_progress("Stamping page numbers...")
    for phys, num in phys_to_content.items():
        _stamp_page_number(merged[phys], num)

    # Build TOC entries with content-only page numbers
    for letter, title, start_phys, end_phys in content_ranges:
        start_content = phys_to_content.get(start_phys, 0)
        end_content = phys_to_content.get(end_phys, 0)
        toc_entries.append({
            "letter": letter,
            "title": title,
            "start_page": start_content,
            "end_page": end_content,
            "page_count": end_content - start_content + 1,
        })

    if on_progress:
        on_progress("Compilation complete.")

    result_bytes = merged.tobytes()
    merged.close()
    return result_bytes, toc_entries


def generate_toc_pdf(toc_entries: list[dict], case_name: str = "") -> bytes:
    """Generate a single-page Table of Contents PDF.

    Lists each exhibit with tab letter, title, and page range.
    """
    doc = pymupdf.open()
    page = doc.new_page(width=_PAGE_W, height=_PAGE_H)

    margin = 72
    y = margin

    # Title
    title_rect = pymupdf.Rect(margin, y, _PAGE_W - margin, y + 30)
    page.insert_textbox(
        title_rect,
        "TABLE OF CONTENTS",
        fontsize=16,
        fontname="helv",
        align=1,
        color=(0, 0, 0),
    )
    y += 36

    # Case name
    if case_name:
        case_rect = pymupdf.Rect(margin, y, _PAGE_W - margin, y + 20)
        page.insert_textbox(
            case_rect,
            f"In the Matter of: {case_name}",
            fontsize=11,
            fontname="helv",
            align=1,
            color=(0.3, 0.3, 0.3),
        )
        y += 28

    # Divider
    page.draw_line(
        pymupdf.Point(margin, y),
        pymupdf.Point(_PAGE_W - margin, y),
        color=(0.5, 0.5, 0.5),
        width=0.5,
    )
    y += 20

    # Column headers
    col_tab = margin
    col_title = margin + 60
    col_pages = _PAGE_W - margin - 80

    header_h = 18
    page.insert_textbox(
        pymupdf.Rect(col_tab, y, col_tab + 55, y + header_h),
        "Exhibit",
        fontsize=11,
        fontname="helv",
        align=0,
        color=(0, 0, 0),
    )
    page.insert_textbox(
        pymupdf.Rect(col_title, y, col_pages - 5, y + header_h),
        "Document Title",
        fontsize=11,
        fontname="helv",
        align=0,
        color=(0, 0, 0),
    )
    page.insert_textbox(
        pymupdf.Rect(col_pages, y, _PAGE_W - margin, y + header_h),
        "Pages",
        fontsize=11,
        fontname="helv",
        align=2,
        color=(0, 0, 0),
    )
    y += 22

    # Divider under headers
    page.draw_line(
        pymupdf.Point(margin, y),
        pymupdf.Point(_PAGE_W - margin, y),
        color=(0.8, 0.8, 0.8),
        width=0.5,
    )
    y += 10

    # Entries
    for entry in toc_entries:
        if y + 20 > _PAGE_H - margin:
            page = doc.new_page(width=_PAGE_W, height=_PAGE_H)
            y = margin

        row_h = 18
        page.insert_textbox(
            pymupdf.Rect(col_tab, y, col_tab + 55, y + row_h),
            f"Tab {entry['letter']}",
            fontsize=11,
            fontname="tiro",
            align=0,
            color=(0, 0, 0),
        )
        page.insert_textbox(
            pymupdf.Rect(col_title, y, col_pages - 5, y + row_h),
            entry["title"],
            fontsize=11,
            fontname="tiro",
            align=0,
            color=(0, 0, 0),
        )
        page.insert_textbox(
            pymupdf.Rect(col_pages, y, _PAGE_W - margin, y + row_h),
            f"{entry['start_page']}-{entry['end_page']}",
            fontsize=11,
            fontname="tiro",
            align=2,
            color=(0, 0, 0),
        )
        y += 22

    # Footer
    y = _PAGE_H - margin - 10
    footer_rect = pymupdf.Rect(margin, y, _PAGE_W - margin, y + 14)
    page.insert_textbox(
        footer_rect,
        f"Generated: {date.today().strftime('%m/%d/%Y')}",
        fontsize=8,
        fontname="helv",
        align=2,
        color=(0.5, 0.5, 0.5),
    )

    result = doc.tobytes()
    doc.close()
    return result


def generate_toc_docx(toc_entries: list[dict], case_name: str = "") -> bytes:
    """Generate a Word document Table of Contents.

    Same format as evidence.py generate_index_docx() but with TOC data.
    """
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        def _run(para, text: str, size: int = 12, bold: bool = False, italic: bool = False):
            r = para.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
            r.bold = bold
            r.italic = italic
            return r

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "TABLE OF CONTENTS", size=14, bold=True)

        # Case caption
        if case_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, f"In the Matter of: {case_name}", size=12, italic=True)

        doc.add_paragraph()  # spacer

        # Table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = table.rows[0].cells
        headers = ["Exhibit", "Document Title", "Start Page", "End Page"]
        for i, header_text in enumerate(headers):
            header_cells[i].text = ""
            p = header_cells[i].paragraphs[0]
            _run(p, header_text, size=11, bold=True)

        for entry in toc_entries:
            row = table.add_row().cells
            row_data = [
                f"Tab {entry['letter']}",
                entry["title"],
                str(entry["start_page"]),
                str(entry["end_page"]),
            ]
            for i, cell_text in enumerate(row_data):
                row[i].text = ""
                p = row[i].paragraphs[0]
                _run(p, cell_text, size=11)

        # Footer
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p, f"Generated: {date.today().strftime('%m/%d/%Y')}", size=9, italic=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p, "O'Brien Immigration Law", size=9, italic=True)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    except ImportError:
        return b""
