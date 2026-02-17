"""Forms Assistant -- Streamlit dashboard.

Production-quality UI for completing USCIS immigration forms with
field-by-field guidance, validation, draft persistence, live preview,
and Word/text export.  Works entirely offline without the API server.
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from datetime import date
from pathlib import Path

import streamlit as st
from docx import Document
from docx.shared import Inches, Pt

from app.form_definitions import (
    SUPPORTED_FORMS,
    check_completeness,
    delete_form_draft,
    get_fields_for_form,
    list_form_drafts,
    load_form_draft,
    new_draft_id,
    save_form_draft,
    validate_field,
)
from app.pdf_form_store import (
    get_all_forms,
    get_all_fields,
    get_template_pdf_bytes,
    is_uploaded_form,
    get_field_roles,
    get_field_sf_mappings,
)

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.google_upload import upload_to_google_docs
from shared.client_banner import render_client_banner
from shared.tool_notes import render_tool_notes
try:
    from shared.tool_help import render_tool_help
except ImportError:
    render_tool_help = None

# -- Page config --------------------------------------------------------------

st.set_page_config(
    page_title="Forms Assistant -- O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -- CSS ----------------------------------------------------------------------

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* Hide Streamlit chrome */
#MainMenu, footer,
div[data-testid="stToolbar"] { display: none !important; }

.stApp {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

/* Navigation bar */
.nav-bar {
    display: flex;
    align-items: center;
    padding: 10px 4px;
    margin: -1rem 0 1.2rem 0;
    border-bottom: 1px solid rgba(0,0,0,0.07);
}
.nav-back {
    display: flex;
    align-items: center;
    gap: 6px;
    font-family: 'Inter', sans-serif;
    font-size: 0.85rem;
    font-weight: 500;
    color: #0066CC;
    text-decoration: none;
    min-width: 150px;
}
.nav-back:hover { color: #004499; text-decoration: underline; }
.nav-title {
    flex: 1;
    text-align: center;
    font-family: 'Inter', sans-serif;
    font-size: 1.15rem;
    font-weight: 700;
    color: #1a2744;
    letter-spacing: -0.02em;
}
.nav-firm {
    font-weight: 400;
    color: #86868b;
    font-size: 0.85rem;
    margin-left: 8px;
}
.nav-spacer { min-width: 150px; }

/* Section labels */
.section-label {
    font-size: 0.78rem;
    font-weight: 600;
    color: #5a6a85;
    text-transform: uppercase;
    letter-spacing: 0.04em;
    margin-bottom: 4px;
    margin-top: 12px;
}

/* Section header */
.section-header {
    font-size: 1.1rem;
    font-weight: 600;
    color: #1a2744;
    margin-top: 0.5rem;
    margin-bottom: 0.5rem;
}

/* Form metadata */
.form-meta {
    font-size: 0.85rem;
    color: #5a6a85;
    line-height: 1.6;
}

/* Field error */
.field-error {
    font-size: 0.82rem;
    color: #c62828;
    margin-top: -8px;
    margin-bottom: 8px;
}

/* Progress bar */
.progress-bar {
    background: #e8ecf0;
    border-radius: 6px;
    height: 10px;
    overflow: hidden;
    margin-bottom: 4px;
}
.progress-fill {
    height: 100%;
    background: #1a73e8;
    border-radius: 6px;
    transition: width 0.3s ease;
}

/* Preview panel */
.preview-panel {
    font-family: 'Inter', sans-serif;
    font-size: 0.82rem;
    line-height: 1.7;
    background: #fff;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    padding: 20px 24px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
    max-height: 75vh;
    overflow-y: auto;
}
.preview-panel .pv-section {
    font-weight: 700;
    font-size: 0.85rem;
    color: #1a2744;
    margin-top: 14px;
    margin-bottom: 6px;
    border-bottom: 1px solid #e8ecf0;
    padding-bottom: 3px;
}
.preview-panel .pv-field {
    display: flex;
    gap: 8px;
    padding: 2px 0;
}
.preview-panel .pv-label {
    font-weight: 600;
    color: #5a6a85;
    min-width: 160px;
    flex-shrink: 0;
}
.preview-panel .pv-value {
    color: #1a2744;
}
.preview-panel .pv-empty {
    color: #b0b8c4;
    font-style: italic;
}

/* Saved toast */
.saved-toast {
    font-size: 0.8rem;
    color: #2e7d32;
    font-weight: 600;
}

/* Help text below fields */
.help-text {
    font-size: 0.78rem;
    color: #86868b;
    margin-top: -6px;
    margin-bottom: 10px;
}
</style>
""",
    unsafe_allow_html=True,
)

# -- Navigation bar -----------------------------------------------------------

st.markdown(
    """
<div class="nav-bar">
    <a href="http://localhost:8502" class="nav-back">&#8592; Staff Dashboard</a>
    <div class="nav-title">Forms Assistant<span class="nav-firm">&mdash; O'Brien Immigration Law</span></div>
    <div class="nav-spacer"></div>
</div>
""",
    unsafe_allow_html=True,
)

render_client_banner()
if render_tool_help:
    render_tool_help("forms-assistant")

# -- Session state defaults ---------------------------------------------------

_DEFAULTS: dict = {
    "draft_id": None,
    "last_saved_msg": "",
    "form_data": {},
    "current_section": 0,
    "validation_errors": {},
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

if st.session_state.draft_id is None:
    st.session_state.draft_id = new_draft_id()


# -- Helpers ------------------------------------------------------------------

def _get_fields(form_id: str) -> dict:
    """Return the field definitions dict for a given form (hardcoded or uploaded)."""
    return get_all_fields(form_id)


def _get_display_label(form_id: str, field_name: str) -> str:
    """Get a human-readable label for a field.

    For uploaded forms, uses the display_label from config.
    For hardcoded forms, derives from the field name.
    """
    if is_uploaded_form(form_id):
        from shared.config_store import load_config
        cfg = load_config("forms-assistant") or {}
        uploaded = cfg.get("uploaded_forms", {})
        form_cfg = uploaded.get(form_id, {})
        for fd in form_cfg.get("fields", []):
            if fd.get("pdf_field_name") == field_name:
                return fd.get("display_label", field_name.replace("_", " ").title())
    return field_name.replace("_", " ").title()


def _do_save(form_id: str) -> None:
    """Save the current form data as a draft."""
    save_form_draft(
        st.session_state.draft_id,
        form_id,
        dict(st.session_state.form_data),
        st.session_state.current_section,
    )
    # Derive client name for toast
    name = ""
    for key in ("full_name", "applicant_name", "petitioner_name", "appellant_name"):
        if st.session_state.form_data.get(key, "").strip():
            name = st.session_state.form_data[key].strip()
            break
    st.session_state.last_saved_msg = f"Saved -- {name or 'draft'}"


def _do_load(draft_id: str) -> None:
    """Load a draft into session state."""
    draft = load_form_draft(draft_id)
    if not draft:
        return
    st.session_state.draft_id = draft["id"]
    st.session_state.form_data = dict(draft.get("form_data", {}))
    st.session_state.current_section = draft.get("current_section", 0)
    st.session_state.validation_errors = {}
    st.session_state.last_saved_msg = ""
    # Store loaded form_id so sidebar selectbox picks it up
    st.session_state._loaded_form_id = draft.get("form_id", "I-589")


def _do_new() -> None:
    """Start a fresh form."""
    st.session_state.draft_id = new_draft_id()
    st.session_state.last_saved_msg = ""
    st.session_state.form_data = {}
    st.session_state.current_section = 0
    st.session_state.validation_errors = {}
    if "_loaded_form_id" in st.session_state:
        del st.session_state["_loaded_form_id"]


def _build_preview_html(form_id: str, form_data: dict) -> str:
    """Build an HTML preview of all filled fields, organized by section."""
    esc = html_mod.escape
    fields_dict = _get_fields(form_id)
    form_meta = get_all_forms().get(form_id, {})
    parts: list[str] = []

    parts.append(f'<div style="font-weight:700; font-size:0.95rem; color:#1a2744; margin-bottom:8px;">'
                 f'{esc(form_id)} -- {esc(form_meta.get("title", ""))}</div>')

    for section_name, fields in fields_dict.items():
        if not fields:
            continue
        parts.append(f'<div class="pv-section">{esc(section_name)}</div>')
        for f in fields:
            label = _get_display_label(form_id, f.name)
            val = form_data.get(f.name, "")
            val_str = str(val).strip() if val else ""
            if val_str:
                parts.append(
                    f'<div class="pv-field">'
                    f'<span class="pv-label">{esc(label)}:</span>'
                    f'<span class="pv-value">{esc(val_str)}</span>'
                    f'</div>'
                )
            else:
                parts.append(
                    f'<div class="pv-field">'
                    f'<span class="pv-label">{esc(label)}:</span>'
                    f'<span class="pv-empty">--</span>'
                    f'</div>'
                )
    return "\n".join(parts)


def _build_docx(form_id: str, form_data: dict) -> bytes:
    """Build a Word document from the form data as a table."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    form_meta = get_all_forms().get(form_id, {})

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{form_id} -- {form_meta.get('title', '')}")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_para.paragraph_format.space_after = Pt(12)

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}")
    date_run.font.name = "Arial"
    date_run.font.size = Pt(10)
    date_para.paragraph_format.space_after = Pt(12)

    fields_dict = _get_fields(form_id)

    for section_name, fields in fields_dict.items():
        if not fields:
            continue

        # Section heading
        heading = doc.add_paragraph()
        h_run = heading.add_run(section_name)
        h_run.font.name = "Arial"
        h_run.font.size = Pt(11)
        h_run.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)

        # Table for fields
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        for i, f in enumerate(fields):
            label = _get_display_label(form_id, f.name)
            val = str(form_data.get(f.name, "")).strip()
            cell_label = table.cell(i, 0)
            cell_value = table.cell(i, 1)
            cell_label.text = label
            cell_value.text = val if val else "--"
            # Style cells
            for paragraph in cell_label.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.bold = True
            for paragraph in cell_value.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_plain_text(form_id: str, form_data: dict) -> str:
    """Build a plain text export of the form data."""
    form_meta = get_all_forms().get(form_id, {})
    lines: list[str] = []

    lines.append(f"{form_id} -- {form_meta.get('title', '')}")
    lines.append(f"Generated: {date.today().strftime('%B %d, %Y')}")
    lines.append("=" * 60)
    lines.append("")

    fields_dict = _get_fields(form_id)

    for section_name, fields in fields_dict.items():
        if not fields:
            continue
        lines.append(section_name)
        lines.append("-" * len(section_name))
        for f in fields:
            label = _get_display_label(form_id, f.name)
            val = str(form_data.get(f.name, "")).strip()
            lines.append(f"  {label}: {val if val else '--'}")
        lines.append("")

    return "\n".join(lines)


# -- Sidebar ------------------------------------------------------------------

with st.sidebar:
    # Draft management
    st.markdown("#### Drafts")
    btn_cols = st.columns(2)
    with btn_cols[0]:
        if st.button("New", use_container_width=True):
            _do_new()
            st.rerun()
    with btn_cols[1]:
        save_clicked = st.button("Save", use_container_width=True, type="primary")

    saved_drafts = list_form_drafts()
    if saved_drafts:
        labels_map = {
            d["id"]: f"{d['client_name']} -- {d['form_id']}"
            for d in saved_drafts
        }
        draft_ids = list(labels_map.keys())
        selected_draft = st.selectbox(
            "Load a saved draft",
            options=[""] + draft_ids,
            format_func=lambda x: labels_map.get(x, "Select..."),
            label_visibility="collapsed",
        )
        load_cols = st.columns(2)
        with load_cols[0]:
            if selected_draft and st.button("Load", use_container_width=True):
                _do_load(selected_draft)
                st.rerun()
        with load_cols[1]:
            if selected_draft and st.button("Delete", use_container_width=True):
                delete_form_draft(selected_draft)
                st.rerun()

    if st.session_state.last_saved_msg:
        st.markdown(
            f'<div class="saved-toast">{html_mod.escape(st.session_state.last_saved_msg)}</div>',
            unsafe_allow_html=True,
        )

    st.divider()

    # Form selector — merged hardcoded + uploaded
    all_forms = get_all_forms()
    form_ids = list(all_forms.keys())
    form_labels = {}
    for fid, meta in all_forms.items():
        badge = " (PDF)" if meta.get("_uploaded") else ""
        form_labels[fid] = f"{fid}{badge} -- {meta.get('title', '')}"

    # If a draft was just loaded, use its form_id as default
    default_idx = 0
    if "_loaded_form_id" in st.session_state:
        loaded_fid = st.session_state._loaded_form_id
        if loaded_fid in form_ids:
            default_idx = form_ids.index(loaded_fid)

    selected_form = st.selectbox(
        "Form",
        options=form_ids,
        index=default_idx,
        format_func=lambda x: form_labels.get(x, x),
        key="inp_form_selector",
    )

    # Form metadata
    form_meta = all_forms.get(selected_form, {})
    st.markdown(
        f'<div class="form-meta">'
        f'<strong>Filing Fee:</strong> {html_mod.escape(form_meta.get("filing_fee", "N/A"))}<br>'
        f'<strong>Processing:</strong> {html_mod.escape(form_meta.get("processing_time", "N/A"))}<br>'
        f'<strong>Agency:</strong> {html_mod.escape(form_meta.get("agency", "N/A"))}'
        f'</div>',
        unsafe_allow_html=True,
    )

    # Auto-fill sections for uploaded forms
    if is_uploaded_form(selected_form):
        roles = get_field_roles(selected_form)
        sf_mappings = get_field_sf_mappings(selected_form)

        has_attorney_roles = any(r.startswith("attorney_") for r in roles.values())
        has_preparer_roles = any(r.startswith("preparer_") for r in roles.values())

        # -- Client auto-fill from SF via direct sf_field mappings --
        if sf_mappings:
            sf_client = st.session_state.get("sf_client")
            if sf_client:
                filled = 0
                for field_name, sf_key in sf_mappings.items():
                    if sf_client.get(sf_key) and not st.session_state.form_data.get(field_name):
                        st.session_state.form_data[field_name] = str(sf_client[sf_key])
                        filled += 1
                if filled:
                    st.divider()
                    st.caption(f"Auto-filled {filled} field(s) from Salesforce")

            # -- Sync to Salesforce button --
            st.divider()
            st.markdown("#### Salesforce Sync")
            if sf_client:
                changed: dict[str, str] = {}
                for field_name, sf_key in sf_mappings.items():
                    form_val = str(st.session_state.form_data.get(field_name, "")).strip()
                    sf_val = str(sf_client.get(sf_key, "")).strip()
                    if form_val and form_val != sf_val:
                        changed[sf_key] = form_val
                if changed:
                    st.caption(f"{len(changed)} field(s) differ from Salesforce")
                    if st.button("Sync to Salesforce", type="primary", use_container_width=True, key="btn_sf_sync"):
                        try:
                            from shared.salesforce_client import update_client
                            sf_id = sf_client.get("Id")
                            if sf_id:
                                update_client(sf_id, changed)
                                # Refresh local cache
                                for sf_key, val in changed.items():
                                    sf_client[sf_key] = val
                                st.toast(f"Pushed {len(changed)} field(s) to Salesforce")
                                st.rerun()
                            else:
                                st.error("No Salesforce record ID found.")
                        except Exception as e:
                            st.error(f"Sync failed: {e}")
                else:
                    st.caption("All mapped fields match Salesforce")
            else:
                st.caption("Pull a client to enable Salesforce sync.")

        # -- Attorney picklist --
        if has_attorney_roles:
            st.divider()
            st.markdown("#### Attorney")
            try:
                from shared.attorney_store import load_attorneys
                attorneys = load_attorneys()
            except ImportError:
                attorneys = []
            if attorneys:
                atty_options = ["(None)"] + [a["name"] for a in attorneys]
                atty_sel = st.selectbox(
                    "Select attorney",
                    atty_options,
                    key="inp_attorney",
                    label_visibility="collapsed",
                )
                if atty_sel != "(None)":
                    atty = next((a for a in attorneys if a["name"] == atty_sel), None)
                    if atty:
                        role_to_atty_key = {
                            "attorney_name": "name",
                            "attorney_bar_number": "bar_number",
                            "attorney_firm": "firm",
                            "attorney_address": "address",
                            "attorney_phone": "phone",
                            "attorney_email": "email",
                        }
                        for field_name, role in roles.items():
                            atty_key = role_to_atty_key.get(role)
                            if atty_key and atty.get(atty_key):
                                st.session_state.form_data[field_name] = atty[atty_key]
            else:
                st.caption("No attorneys configured. Add them in Admin Panel.")

        # -- Preparer picklist --
        if has_preparer_roles:
            st.divider()
            st.markdown("#### Preparer")
            try:
                from shared.preparer_store import load_preparers
                preparers = load_preparers()
            except ImportError:
                preparers = []
            if preparers:
                prep_options = ["(None)"] + [p["name"] for p in preparers]
                prep_sel = st.selectbox(
                    "Select preparer",
                    prep_options,
                    key="inp_preparer",
                    label_visibility="collapsed",
                )
                if prep_sel != "(None)":
                    prep = next((p for p in preparers if p["name"] == prep_sel), None)
                    if prep:
                        role_to_prep_key = {
                            "preparer_name": "name",
                            "preparer_firm": "firm",
                            "preparer_address": "address",
                            "preparer_phone": "phone",
                            "preparer_email": "email",
                            "preparer_bar_number": "bar_number",
                        }
                        for field_name, role in roles.items():
                            prep_key = role_to_prep_key.get(role)
                            if prep_key and prep.get(prep_key):
                                st.session_state.form_data[field_name] = prep[prep_key]
            else:
                st.caption("No preparers configured. Add them in Admin Panel.")

    st.divider()

    # Progress bar
    st.markdown("#### Progress")
    fields_dict = _get_fields(selected_form)
    all_defined_fields = [f for section_fields in fields_dict.values() for f in section_fields]
    total_fields = len(all_defined_fields)
    filled_count = sum(
        1 for f in all_defined_fields
        if str(st.session_state.form_data.get(f.name, "")).strip()
    )
    pct = round((filled_count / total_fields) * 100) if total_fields > 0 else 0

    st.markdown(
        f'<div class="progress-bar"><div class="progress-fill" style="width:{pct}%"></div></div>',
        unsafe_allow_html=True,
    )
    st.caption(f"{pct}% complete ({filled_count}/{total_fields} fields)")

    st.divider()

    # Validate button
    validate_clicked = st.button("Validate Form", use_container_width=True)

    render_tool_notes("forms-assistant")


# -- Handle save (after sidebar renders) -------------------------------------

if save_clicked:
    _do_save(selected_form)
    st.rerun()

# -- Handle validate ----------------------------------------------------------

if validate_clicked:
    errors: dict[str, list[str]] = {}
    fields_dict = _get_fields(selected_form)
    for _section_name, fields in fields_dict.items():
        for field_def in fields:
            val = str(st.session_state.form_data.get(field_def.name, ""))
            field_errors = validate_field(field_def, val)
            if field_errors:
                errors[field_def.name] = field_errors
    st.session_state.validation_errors = errors

# -- Main area ----------------------------------------------------------------

form_col, preview_col = st.columns([3, 2], gap="large")

# -- Left column: Form fields ------------------------------------------------

with form_col:
    fields_dict = _get_fields(selected_form)
    section_names = list(fields_dict.keys())

    if not section_names:
        st.info("No field definitions are available for this form yet.")
    else:
        # Clamp current_section to valid range
        if st.session_state.current_section >= len(section_names):
            st.session_state.current_section = 0

        selected_section = st.radio(
            "Section",
            section_names,
            horizontal=True,
            index=st.session_state.current_section,
            key="section_radio",
        )
        st.session_state.current_section = section_names.index(selected_section)

        st.markdown(
            f'<div class="section-header">{html_mod.escape(selected_section)}</div>',
            unsafe_allow_html=True,
        )

        fields = fields_dict.get(selected_section, [])

        if fields:
            for field_def in fields:
                field_name = field_def.name
                field_type = field_def.field_type
                required = field_def.required
                help_text = field_def.help_text
                options = field_def.options

                label = _get_display_label(selected_form, field_name)
                if required:
                    label += " *"

                current_value = st.session_state.form_data.get(field_name, "")

                if field_type in ("select", "combo") and options:
                    all_options = [""] + options
                    idx = 0
                    if current_value and current_value in options:
                        idx = options.index(current_value) + 1
                    value = st.selectbox(
                        label,
                        options=all_options,
                        index=idx,
                        help=help_text,
                        key=f"field_{field_name}",
                    )
                elif field_type == "textarea":
                    value = st.text_area(
                        label,
                        value=current_value,
                        help=help_text,
                        key=f"field_{field_name}",
                        height=120,
                    )
                elif field_type == "checkbox":
                    checked = st.checkbox(
                        label,
                        value=bool(current_value),
                        help=help_text,
                        key=f"field_{field_name}",
                    )
                    value = "Yes" if checked else ""
                elif field_type == "date":
                    value = st.text_input(
                        label,
                        value=current_value,
                        help=help_text,
                        placeholder="mm/dd/yyyy",
                        key=f"field_{field_name}",
                    )
                else:
                    # text, phone, email, etc.
                    value = st.text_input(
                        label,
                        value=current_value,
                        help=help_text,
                        key=f"field_{field_name}",
                    )

                st.session_state.form_data[field_name] = value

                # Show validation errors
                field_errors = st.session_state.validation_errors.get(field_name, [])
                if field_errors:
                    for err in field_errors:
                        st.markdown(
                            f'<div class="field-error">{html_mod.escape(err)}</div>',
                            unsafe_allow_html=True,
                        )
        else:
            st.info(
                "Detailed field definitions are not yet available for this section. "
                "Additional sections will be added in future updates."
            )

        # Navigation buttons
        st.markdown("---")
        nav_left, nav_right = st.columns(2)
        with nav_left:
            if st.session_state.current_section > 0:
                if st.button("Previous Section", use_container_width=True):
                    st.session_state.current_section -= 1
                    st.rerun()
        with nav_right:
            if st.session_state.current_section < len(section_names) - 1:
                if st.button("Next Section", use_container_width=True):
                    st.session_state.current_section += 1
                    st.rerun()

# -- Right column: Preview + Export -------------------------------------------

with preview_col:
    st.markdown('<div class="section-label">Live Preview</div>', unsafe_allow_html=True)

    preview_html = _build_preview_html(selected_form, st.session_state.form_data)
    st.markdown(
        f'<div class="preview-panel">{preview_html}</div>',
        unsafe_allow_html=True,
    )

    # Validation summary
    if st.session_state.validation_errors:
        error_count = sum(len(v) for v in st.session_state.validation_errors.values())
        st.error(f"{error_count} validation error(s) found. See field-level messages on the left.")
    elif validate_clicked:
        if is_uploaded_form(selected_form):
            # Simple completeness check for uploaded forms
            all_flds = [f for flds in _get_fields(selected_form).values() for f in flds]
            req_missing = [f.name for f in all_flds if f.required and not str(st.session_state.form_data.get(f.name, "")).strip()]
            if req_missing:
                st.warning(f"{len(req_missing)} required field(s) still empty.")
            else:
                st.success("All required fields complete. No validation errors.")
        else:
            completeness = check_completeness(selected_form, st.session_state.form_data)
            missing = completeness.get("required_missing", [])
            if missing:
                st.warning(f"{len(missing)} required field(s) still empty.")
            else:
                st.success("All required fields complete. No validation errors.")

    # Export controls
    st.markdown("---")
    st.markdown('<div class="section-label">Export</div>', unsafe_allow_html=True)

    # Filled PDF download (uploaded forms only)
    if is_uploaded_form(selected_form):
        template_bytes = get_template_pdf_bytes(selected_form)
        if template_bytes:
            # Build field value mapping: pdf_field_name -> user value
            field_mapping: dict[str, str] = {}
            for _sn, flds in _get_fields(selected_form).items():
                for fld in flds:
                    val = str(st.session_state.form_data.get(fld.name, "")).strip()
                    if val:
                        field_mapping[fld.name] = val
            try:
                from shared.pdf_form_extractor import fill_pdf_form
                filled_pdf = fill_pdf_form(template_bytes, field_mapping)
                client_name_pdf = ""
                for key in ("full_name", "applicant_name", "petitioner_name", "appellant_name"):
                    if st.session_state.form_data.get(key, "").strip():
                        client_name_pdf = st.session_state.form_data[key].strip()
                        break
                safe_pdf_name = client_name_pdf.replace(" ", "_") if client_name_pdf else "form"
                st.download_button(
                    "Download Filled PDF",
                    data=filled_pdf,
                    file_name=f"{selected_form}_{safe_pdf_name}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary",
                )
            except Exception as e:
                st.error(f"PDF fill failed: {e}")

    exp_cols = st.columns(2)
    with exp_cols[0]:
        plain_text = _build_plain_text(selected_form, st.session_state.form_data)
        # Derive filename from client name
        client_name = ""
        for key in ("full_name", "applicant_name", "petitioner_name", "appellant_name"):
            if st.session_state.form_data.get(key, "").strip():
                client_name = st.session_state.form_data[key].strip()
                break
        safe_name = client_name.replace(" ", "_") if client_name else "form"
        st.download_button(
            "Download .txt",
            data=plain_text,
            file_name=f"{selected_form}_{safe_name}.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with exp_cols[1]:
        docx_bytes = _build_docx(selected_form, st.session_state.form_data)
        st.download_button(
            "Download .docx",
            data=docx_bytes,
            file_name=f"{selected_form}_{safe_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        if st.button("Upload to Google Docs", use_container_width=True):
            with st.spinner("Uploading to Google Docs..."):
                try:
                    url = upload_to_google_docs(docx_bytes, f"{selected_form} - {client_name or 'Form'}")
                    st.session_state.google_doc_url = url
                except Exception as e:
                    st.error(f"Upload failed: {e}")
        if st.session_state.get("google_doc_url"):
            st.markdown(f"[Open Google Doc]({st.session_state.google_doc_url})")
