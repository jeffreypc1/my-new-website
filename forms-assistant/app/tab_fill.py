"""Fill Forms tab -- primary workflow for data entry, preview, and export.

Renders form fields grouped by section with live HTML preview, validation
summaries, and export options (filled PDF, .txt, .docx, Google Docs).

Supports both single-form mode (section-by-section navigation) and
multi-form mode (shared fields first, then per-form unique fields).

Public API:
    render_fill_tab()  -- called by the dashboard orchestrator.

Session state consumed:
    form_data          -- dict of field_name -> value
    current_section    -- int index of the current section
    validation_errors  -- dict of field_name -> list[str]
    selected_forms     -- list of form_ids selected for filling
    sf_client          -- optional Salesforce Contact dict for diff indicators
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from datetime import date
from pathlib import Path

import streamlit as st
from docx import Document
from docx.shared import Inches, Pt

from app.form_definitions import (
    SUPPORTED_FORMS,
    check_completeness,
    validate_field,
    get_fields_for_form,
)
from app.pdf_form_store import (
    get_all_forms,
    get_all_fields,
    get_template_pdf_bytes,
    is_uploaded_form,
    get_field_roles,
    get_field_sf_mappings,
)
from app.mapping_store import load_mapping_set
from app.multi_form import merge_form_schemas, split_form_data, get_shared_field_key
from app.schema import FormSchema
from app.ingestion import load_form_schema

# -- Shared imports (monorepo root) ------------------------------------------

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.google_upload import upload_to_google_docs
from shared.pdf_form_extractor import fill_pdf_form


# ============================================================================
# Private helpers
# ============================================================================

def _get_fields(form_id: str) -> dict:
    """Return the field definitions dict for a given form (hardcoded or uploaded)."""
    return get_all_fields(form_id)


def _get_display_label(form_id: str, field_name: str) -> str:
    """Get a human-readable label for a field.

    For uploaded forms, uses the display_label from config.
    For hardcoded forms, derives from the field name.
    """
    if is_uploaded_form(form_id):
        from shared.config_store import load_config

        cfg = load_config("forms-assistant") or {}
        uploaded = cfg.get("uploaded_forms", {})
        form_cfg = uploaded.get(form_id, {})
        for fd in form_cfg.get("fields", []):
            if fd.get("pdf_field_name") == field_name:
                return fd.get("display_label", field_name.replace("_", " ").title())
    return field_name.replace("_", " ").title()


def _build_preview_html(form_id: str, form_data: dict) -> str:
    """Build an HTML preview of all filled fields, organized by section."""
    esc = html_mod.escape
    fields_dict = _get_fields(form_id)
    form_meta = get_all_forms().get(form_id, {})
    parts: list[str] = []

    title = form_meta.get("title", "")
    preview_header = f'{esc(form_id)} -- {esc(title)}' if title else esc(form_id)
    parts.append(
        f'<div style="font-weight:700; font-size:0.95rem; color:#1a2744; '
        f'margin-bottom:8px;">'
        f'{preview_header}</div>'
    )

    for section_name, fields in fields_dict.items():
        if not fields:
            continue
        parts.append(f'<div class="pv-section">{esc(section_name)}</div>')
        for f in fields:
            label = _get_display_label(form_id, f.name)
            val = form_data.get(f.name, "")
            val_str = str(val).strip() if val else ""
            if val_str:
                parts.append(
                    f'<div class="pv-field">'
                    f'<span class="pv-label">{esc(label)}:</span>'
                    f'<span class="pv-value">{esc(val_str)}</span>'
                    f'</div>'
                )
            else:
                parts.append(
                    f'<div class="pv-field">'
                    f'<span class="pv-label">{esc(label)}:</span>'
                    f'<span class="pv-empty">--</span>'
                    f'</div>'
                )
    return "\n".join(parts)


def _build_docx(form_id: str, form_data: dict) -> bytes:
    """Build a Word document from the form data as a table."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    form_meta = get_all_forms().get(form_id, {})

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{form_id} -- {form_meta.get('title', '')}")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_para.paragraph_format.space_after = Pt(12)

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated: {date.today().strftime('%m/%d/%Y')}")
    date_run.font.name = "Arial"
    date_run.font.size = Pt(10)
    date_para.paragraph_format.space_after = Pt(12)

    fields_dict = _get_fields(form_id)

    for section_name, fields in fields_dict.items():
        if not fields:
            continue

        # Section heading
        heading = doc.add_paragraph()
        h_run = heading.add_run(section_name)
        h_run.font.name = "Arial"
        h_run.font.size = Pt(11)
        h_run.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)

        # Table for fields
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        for i, f in enumerate(fields):
            label = _get_display_label(form_id, f.name)
            val = str(form_data.get(f.name, "")).strip()
            cell_label = table.cell(i, 0)
            cell_value = table.cell(i, 1)
            cell_label.text = label
            cell_value.text = val if val else "--"
            # Style cells
            for paragraph in cell_label.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.bold = True
            for paragraph in cell_value.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_plain_text(form_id: str, form_data: dict) -> str:
    """Build a plain text export of the form data."""
    form_meta = get_all_forms().get(form_id, {})
    lines: list[str] = []

    lines.append(f"{form_id} -- {form_meta.get('title', '')}")
    lines.append(f"Generated: {date.today().strftime('%m/%d/%Y')}")
    lines.append("=" * 60)
    lines.append("")

    fields_dict = _get_fields(form_id)

    for section_name, fields in fields_dict.items():
        if not fields:
            continue
        lines.append(section_name)
        lines.append("-" * len(section_name))
        for f in fields:
            label = _get_display_label(form_id, f.name)
            val = str(form_data.get(f.name, "")).strip()
            lines.append(f"  {label}: {val if val else '--'}")
        lines.append("")

    return "\n".join(lines)


def _derive_client_name(form_data: dict) -> str:
    """Extract a client name from form data for file-naming purposes."""
    for key in ("full_name", "applicant_name", "petitioner_name", "appellant_name"):
        name = form_data.get(key, "")
        if isinstance(name, str) and name.strip():
            return name.strip()
    return ""


def _render_diff_indicator(
    form_id: str,
    field_name: str,
    form_value: str,
    sf_client: dict | None,
) -> None:
    """Show a colored dot indicating SF mapping sync status.

    Green  -- form value matches the SF value.
    Yellow -- form value differs from the SF value.
    Gray   -- no approved SF mapping for this field.
    """
    if sf_client is None:
        return

    # Determine the SF field mapped to this form field
    sf_field_name = ""

    # Check uploaded-form direct mappings first
    if is_uploaded_form(form_id):
        sf_mappings = get_field_sf_mappings(form_id)
        sf_field_name = sf_mappings.get(field_name, "")

    # Fall back to MappingSet approved mappings
    if not sf_field_name:
        mapping_set = load_mapping_set(form_id)
        if mapping_set:
            mapping = mapping_set.get_mapping(field_name)
            if mapping and mapping.approved and mapping.sf_field:
                sf_field_name = mapping.sf_field

    if not sf_field_name:
        # No mapping -- gray dot
        st.markdown(
            '<span style="color:#b0b8c4; font-size:0.7rem;" '
            'title="No SF mapping">&#9679;</span>',
            unsafe_allow_html=True,
        )
        return

    sf_value = str(sf_client.get(sf_field_name, "")).strip()
    form_val_stripped = str(form_value).strip()

    if form_val_stripped and form_val_stripped == sf_value:
        # Match -- green dot
        st.markdown(
            '<span style="color:#2e7d32; font-size:0.7rem;" '
            'title="Matches Salesforce">&#9679;</span>',
            unsafe_allow_html=True,
        )
    elif form_val_stripped:
        # Differs -- yellow dot
        st.markdown(
            '<span style="color:#f9a825; font-size:0.7rem;" '
            'title="Differs from Salesforce">&#9679;</span>',
            unsafe_allow_html=True,
        )
    else:
        # Empty form value with mapping -- gray dot
        st.markdown(
            '<span style="color:#b0b8c4; font-size:0.7rem;" '
            'title="No SF mapping">&#9679;</span>',
            unsafe_allow_html=True,
        )


def _render_field_widget(
    form_id: str,
    field_def,
    form_data: dict,
    sf_client: dict | None,
    key_prefix: str = "field",
) -> None:
    """Render the appropriate Streamlit widget for a single field definition.

    Reads and writes values in *form_data* (a mutable dict, typically
    ``st.session_state.form_data``).  Shows validation errors and an
    SF diff indicator beneath the widget.
    """
    field_name = field_def.name
    field_type = field_def.field_type
    required = field_def.required
    help_text = field_def.help_text
    options = field_def.options

    label = _get_display_label(form_id, field_name)
    if required:
        label += " *"

    current_value = form_data.get(field_name, "")
    widget_key = f"{key_prefix}_{field_name}"

    if field_type in ("select", "combo") and options:
        all_options = [""] + options
        idx = 0
        if current_value and current_value in options:
            idx = options.index(current_value) + 1
        value = st.selectbox(
            label,
            options=all_options,
            index=idx,
            help=help_text,
            key=widget_key,
        )
    elif field_type == "textarea":
        value = st.text_area(
            label,
            value=current_value,
            help=help_text,
            key=widget_key,
            height=120,
        )
    elif field_type == "checkbox":
        checked = st.checkbox(
            label,
            value=bool(current_value),
            help=help_text,
            key=widget_key,
        )
        value = "Yes" if checked else ""
    elif field_type == "date":
        value = st.text_input(
            label,
            value=current_value,
            help=help_text,
            placeholder="mm/dd/yyyy",
            key=widget_key,
        )
    else:
        # text, phone, email, etc.
        value = st.text_input(
            label,
            value=current_value,
            help=help_text,
            key=widget_key,
        )

    form_data[field_name] = value

    # Diff indicator
    _render_diff_indicator(form_id, field_name, value, sf_client)

    # Validation errors
    field_errors = st.session_state.get("validation_errors", {}).get(field_name, [])
    if field_errors:
        for err in field_errors:
            st.markdown(
                f'<div class="field-error">{html_mod.escape(err)}</div>',
                unsafe_allow_html=True,
            )


# ============================================================================
# Single-form mode: left column
# ============================================================================

def _render_single_form_fields(form_id: str, form_data: dict) -> None:
    """Render section radio + field widgets for a single selected form."""
    fields_dict = _get_fields(form_id)
    section_names = list(fields_dict.keys())

    if not section_names:
        st.info("No field definitions are available for this form yet.")
        return

    # Form title header
    form_meta = get_all_forms().get(form_id, {})
    title = form_meta.get("title", "")
    header = f"{form_id} -- {title}" if title else form_id
    st.markdown(
        f'<div class="section-header">{html_mod.escape(header)}</div>',
        unsafe_allow_html=True,
    )

    # Clamp current_section to valid range
    if st.session_state.current_section >= len(section_names):
        st.session_state.current_section = 0

    selected_section = st.radio(
        "Section",
        section_names,
        horizontal=True,
        index=st.session_state.current_section,
        key="fill_section_radio",
    )
    st.session_state.current_section = section_names.index(selected_section)

    st.markdown(
        f'<div class="section-header">{html_mod.escape(selected_section)}</div>',
        unsafe_allow_html=True,
    )

    fields = fields_dict.get(selected_section, [])
    sf_client = st.session_state.get("sf_client")

    if fields:
        for field_def in fields:
            _render_field_widget(
                form_id, field_def, form_data, sf_client, key_prefix="field"
            )
    else:
        st.info(
            "Detailed field definitions are not yet available for this section. "
            "Additional sections will be added in future updates."
        )

    # Navigation buttons
    st.markdown("---")
    nav_left, nav_right = st.columns(2)
    with nav_left:
        if st.session_state.current_section > 0:
            if st.button("Previous Section", use_container_width=True, key="fill_prev"):
                st.session_state.current_section -= 1
                st.rerun()
    with nav_right:
        if st.session_state.current_section < len(section_names) - 1:
            if st.button("Next Section", use_container_width=True, key="fill_next"):
                st.session_state.current_section += 1
                st.rerun()


# ============================================================================
# Multi-form mode: left column
# ============================================================================

def _render_multi_form_fields(selected_forms: list[str], form_data: dict) -> None:
    """Render shared fields then per-form unique fields for multiple forms."""
    # Load FormSchemas for each selected form
    schemas: list[FormSchema] = []
    for fid in selected_forms:
        schema = load_form_schema(fid)
        if schema is not None:
            schemas.append(schema)

    if not schemas:
        # Fallback: if no saved schemas, render each form sequentially
        for fid in selected_forms:
            st.markdown(f"### {fid}")
            _render_single_form_fields(fid, form_data)
        return

    merged = merge_form_schemas(schemas)
    shared_fields = merged.get("shared_fields", [])
    form_specific = merged.get("form_specific", {})

    sf_client = st.session_state.get("sf_client")

    # -- Shared Fields section ------------------------------------------------
    if shared_fields:
        st.markdown(
            '<div class="section-header">Shared Fields</div>',
            unsafe_allow_html=True,
        )
        st.caption(
            f"These {len(shared_fields)} field(s) are shared across the "
            f"selected forms. Fill them once and the values apply everywhere."
        )
        for fld_schema in shared_fields:
            # Create a lightweight adapter so _render_field_widget works.
            # FormFieldSchema has field_id/display_label; FormField has name.
            # We build a simple namespace to bridge the two.
            _FieldProxy = type(
                "_FieldProxy",
                (),
                {
                    "name": fld_schema.field_id,
                    "field_type": fld_schema.field_type,
                    "required": fld_schema.required,
                    "help_text": fld_schema.help_text,
                    "options": fld_schema.options,
                },
            )
            proxy = _FieldProxy()
            # Use the shared key in form_data so split_form_data can find it
            shared_key = get_shared_field_key(fld_schema.sf_field)
            proxy.name = shared_key  # type: ignore[attr-defined]

            # Use the first form that contains this field for label/diff purposes
            first_form_id = selected_forms[0]
            _render_field_widget(
                first_form_id,
                proxy,
                form_data,
                sf_client,
                key_prefix="shared",
            )

    # -- Per-form unique fields -----------------------------------------------
    for fid, unique_fields in form_specific.items():
        if not unique_fields:
            continue
        st.markdown("---")
        form_meta = get_all_forms().get(fid, {})
        title = form_meta.get("title", "")
        st.markdown(
            f'<div class="section-header">{html_mod.escape(fid)}'
            f'{(" -- " + html_mod.escape(title)) if title else ""}'
            f" (unique fields)</div>",
            unsafe_allow_html=True,
        )
        for fld_schema in unique_fields:
            _FieldProxy = type(
                "_FieldProxy",
                (),
                {
                    "name": fld_schema.field_id,
                    "field_type": fld_schema.field_type,
                    "required": fld_schema.required,
                    "help_text": fld_schema.help_text,
                    "options": fld_schema.options,
                },
            )
            proxy = _FieldProxy()
            _render_field_widget(
                fid, proxy, form_data, sf_client, key_prefix=f"uniq_{fid}"
            )


# ============================================================================
# Right column: preview + export
# ============================================================================

def _render_preview_and_export(selected_forms: list[str], form_data: dict) -> None:
    """Render the live preview panel and export controls."""
    st.markdown(
        '<div class="section-label">Live Preview</div>',
        unsafe_allow_html=True,
    )

    # Build preview for each selected form
    preview_parts: list[str] = []
    for fid in selected_forms:
        preview_parts.append(_build_preview_html(fid, form_data))

    combined_preview = "\n<hr style='border:none;border-top:1px solid #e8ecf0;margin:12px 0;'>\n".join(
        preview_parts
    ) if len(preview_parts) > 1 else (preview_parts[0] if preview_parts else "")

    st.markdown(
        f'<div class="preview-panel">{combined_preview}</div>',
        unsafe_allow_html=True,
    )

    # -- Validation summary ---------------------------------------------------
    validation_errors = st.session_state.get("validation_errors", {})
    if validation_errors:
        error_count = sum(len(v) for v in validation_errors.values())
        st.error(
            f"{error_count} validation error(s) found. "
            "See field-level messages on the left."
        )

    # -- Export controls ------------------------------------------------------
    st.markdown("---")
    st.markdown(
        '<div class="section-label">Export</div>',
        unsafe_allow_html=True,
    )

    client_name = _derive_client_name(form_data)
    safe_name = client_name.replace(" ", "_") if client_name else "form"

    # Render export controls per form
    for fid in selected_forms:
        if len(selected_forms) > 1:
            st.markdown(f"**{fid}**")

        # Filled PDF download (uploaded forms with templates only)
        if is_uploaded_form(fid):
            template_bytes = get_template_pdf_bytes(fid)
            if template_bytes:
                field_mapping: dict[str, str] = {}
                for _sn, flds in _get_fields(fid).items():
                    for fld in flds:
                        val = str(form_data.get(fld.name, "")).strip()
                        if val:
                            field_mapping[fld.name] = val
                try:
                    filled_pdf = fill_pdf_form(template_bytes, field_mapping)
                    safe_pdf_name = (
                        client_name.replace(" ", "_") if client_name else "form"
                    )
                    st.download_button(
                        "Download Filled PDF",
                        data=filled_pdf,
                        file_name=f"{fid}_{safe_pdf_name}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        type="primary",
                        key=f"dl_pdf_{fid}",
                    )
                except Exception as e:
                    st.error(f"PDF fill failed: {e}")

        # Text and DOCX exports
        exp_cols = st.columns(2)
        with exp_cols[0]:
            plain_text = _build_plain_text(fid, form_data)
            st.download_button(
                "Download .txt",
                data=plain_text,
                file_name=f"{fid}_{safe_name}.txt",
                mime="text/plain",
                use_container_width=True,
                key=f"dl_txt_{fid}",
            )
        with exp_cols[1]:
            docx_bytes = _build_docx(fid, form_data)
            st.download_button(
                "Download .docx",
                data=docx_bytes,
                file_name=f"{fid}_{safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"dl_docx_{fid}",
            )

        # Google Docs upload
        gdoc_key = f"google_doc_url_{fid}"
        if st.button(
            "Upload to Google Docs",
            use_container_width=True,
            key=f"btn_gdoc_{fid}",
        ):
            docx_for_upload = _build_docx(fid, form_data)
            with st.spinner("Uploading to Google Docs..."):
                try:
                    url = upload_to_google_docs(
                        docx_for_upload,
                        f"{fid} - {client_name or 'Form'}",
                    )
                    st.session_state[gdoc_key] = url
                except Exception as e:
                    st.error(f"Upload failed: {e}")

        if st.session_state.get(gdoc_key):
            st.markdown(f"[Open Google Doc]({st.session_state[gdoc_key]})")

        if len(selected_forms) > 1:
            st.markdown("---")


# ============================================================================
# Public entry point
# ============================================================================

def render_fill_tab() -> None:
    """Render the Fill Forms tab -- the primary workflow tab.

    Reads ``st.session_state.form_data``, ``st.session_state.current_section``,
    ``st.session_state.validation_errors``, and ``st.session_state.selected_forms``
    to drive a two-column layout:

    - **Left column**: field entry widgets (single-form or multi-form mode).
    - **Right column**: live HTML preview, validation summary, and export buttons.
    """
    selected_forms: list[str] = st.session_state.get("selected_forms", [])

    if not selected_forms:
        st.info("Select one or more forms from the sidebar to begin filling.")
        return

    form_data: dict = st.session_state.setdefault("form_data", {})

    form_col, preview_col = st.columns([3, 2], gap="large")

    # -- Left column: field entry --------------------------------------------
    with form_col:
        if len(selected_forms) == 1:
            _render_single_form_fields(selected_forms[0], form_data)
        else:
            _render_multi_form_fields(selected_forms, form_data)

    # -- Right column: preview + export --------------------------------------
    with preview_col:
        _render_preview_and_export(selected_forms, form_data)
