"""Hearing Prep — Streamlit dashboard.

Oral Q&A simulation for ICE cross-examination practice. Supports
attorney-guided and solo modes with audio recording, transcription,
and Claude evaluation.
"""

from __future__ import annotations

import html as html_mod
import io
import sys
from datetime import date
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.database import (
    add_evaluation,
    add_turn,
    create_session,
    delete_session,
    get_session,
    get_session_transcript,
    get_turns,
    init_db,
    list_sessions,
    update_session,
)
from app.evaluator import build_conversation_history, evaluate_answer
from app.prompts import (
    ICE_EVALUATOR_SYSTEM_PROMPT,
    get_all_questions,
    get_case_types,
    get_question_banks,
)
from app.transcription import SUPPORTED_LANGUAGES, transcribe_audio

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.google_upload import upload_to_google_docs
from shared.client_banner import render_client_banner
from shared.tool_notes import render_tool_notes

try:
    from shared.tool_help import render_tool_help
except ImportError:
    render_tool_help = None
try:
    from shared.feedback_button import render_feedback_button
except ImportError:
    render_feedback_button = None

# ── Initialize database ────────────────────────────────────────────────────

init_db()

# ── Page config ─────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Hearing Prep — O'Brien Immigration Law",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ─────────────────────────────────────────────────────────────────────

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* Hide Streamlit chrome */
#MainMenu, footer,
div[data-testid="stToolbar"] { display: none !important; }

.stApp {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

/* Navigation bar */
.nav-bar {
    display: flex;
    align-items: center;
    padding: 10px 4px;
    margin: -1rem 0 1.2rem 0;
    border-bottom: 1px solid rgba(0,0,0,0.07);
}
.nav-back {
    display: flex;
    align-items: center;
    gap: 6px;
    font-family: 'Inter', sans-serif;
    font-size: 0.85rem;
    font-weight: 500;
    color: #0066CC;
    text-decoration: none;
    min-width: 150px;
}
.nav-back:hover { color: #004499; text-decoration: underline; }
.nav-title {
    flex: 1;
    text-align: center;
    font-family: 'Inter', sans-serif;
    font-size: 1.15rem;
    font-weight: 700;
    color: #1a2744;
    letter-spacing: -0.02em;
}
.nav-firm {
    font-weight: 400;
    color: #86868b;
    font-size: 0.85rem;
    margin-left: 8px;
}
.nav-spacer { min-width: 150px; }

/* Section instructions */
.sec-instructions {
    color: #5a6a85;
    font-size: 0.88rem;
    line-height: 1.5;
    margin-bottom: 1rem;
    padding: 10px 14px;
    background: #f0f4fa;
    border-radius: 8px;
    border-left: 3px solid #4a7ddb;
}

/* Coaching tip */
.coaching-tip {
    font-size: 0.82rem;
    color: #4a6741;
    background: #f0f7ee;
    border: 1px solid #d4e8cf;
    border-radius: 6px;
    padding: 10px 14px;
    margin-top: -8px;
    margin-bottom: 12px;
    line-height: 1.5;
}

/* Score display */
.score-stars {
    font-size: 1.4rem;
    letter-spacing: 2px;
    margin: 4px 0;
}
.score-label {
    font-size: 0.78rem;
    color: #86868b;
    font-weight: 500;
}

/* Evaluation card */
.eval-card {
    background: #f8f9fb;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 16px 20px;
    margin-bottom: 12px;
}
.eval-card h4 {
    margin: 0 0 8px 0;
    font-size: 0.9rem;
    color: #1a2744;
}
.eval-card p {
    font-size: 0.85rem;
    color: #444;
    line-height: 1.6;
    margin: 0;
}

/* Strength/weakness pills */
.strength-pill {
    display: inline-block;
    padding: 3px 10px;
    font-size: 0.75rem;
    font-weight: 500;
    background: #e8f5e9;
    color: #2e7d32;
    border-radius: 12px;
    margin: 2px 4px 2px 0;
}
.weakness-pill {
    display: inline-block;
    padding: 3px 10px;
    font-size: 0.75rem;
    font-weight: 500;
    background: #fce4ec;
    color: #c62828;
    border-radius: 12px;
    margin: 2px 4px 2px 0;
}

/* Progress label */
.progress-label {
    font-size: 0.8rem;
    font-weight: 600;
    color: #5a6a85;
    margin-bottom: 4px;
}

/* Saved toast */
.saved-toast {
    font-size: 0.8rem;
    color: #2e7d32;
    font-weight: 600;
}

/* Transcript card in review */
.transcript-turn {
    background: #fff;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 16px 20px;
    margin-bottom: 16px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
}
.transcript-q {
    font-weight: 600;
    color: #1a2744;
    font-size: 0.9rem;
    margin-bottom: 6px;
}
.transcript-a {
    color: #333;
    font-size: 0.88rem;
    line-height: 1.6;
    padding: 8px 0;
    border-top: 1px solid #f0f0f0;
}
</style>
""",
    unsafe_allow_html=True,
)

from shared.auth import require_auth, render_logout

require_auth()

# ── Navigation bar ──────────────────────────────────────────────────────────

st.markdown(
    """
<div class="nav-bar">
    <a href="http://localhost:8502" class="nav-back">&#8592; Staff Dashboard</a>
    <div class="nav-title">Hearing Prep<span class="nav-firm">&mdash; O'Brien Immigration Law</span></div>
    <div class="nav-spacer"></div>
</div>
""",
    unsafe_allow_html=True,
)
render_logout()

client_record = render_client_banner()
if render_tool_help:
    render_tool_help("hearing-prep")
if render_feedback_button:
    render_feedback_button("hearing-prep")

# ── Session state defaults ──────────────────────────────────────────────────

_DEFAULTS: dict = {
    "hp_session_id": None,
    "hp_phase": "setup",        # setup | active | review
    "hp_turn_number": 0,
    "hp_current_question_idx": 0,
    "hp_use_follow_up": False,
    "hp_follow_up_text": "",
    "hp_last_eval": None,
    "hp_show_tips": True,
    "hp_show_eval_to_client": True,
    "hp_auto_follow_up": True,
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ── Helpers ─────────────────────────────────────────────────────────────────

def _score_stars(score: int) -> str:
    """Render a score as filled/empty stars."""
    return "★" * score + "☆" * (5 - score)


def _render_evaluation(evaluation: dict) -> None:
    """Render an evaluation card with score, strengths, weaknesses."""
    score = evaluation.get("score", 0)
    st.markdown(
        f'<div class="score-stars">{_score_stars(score)}</div>'
        f'<div class="score-label">Score: {score}/5</div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        f'<div class="eval-card"><h4>Evaluation</h4>'
        f'<p>{html_mod.escape(evaluation.get("evaluation", ""))}</p></div>',
        unsafe_allow_html=True,
    )

    strengths = evaluation.get("strengths", [])
    if strengths:
        pills = "".join(
            f'<span class="strength-pill">{html_mod.escape(s)}</span>'
            for s in strengths
        )
        st.markdown(f"**Strengths:** {pills}", unsafe_allow_html=True)

    weaknesses = evaluation.get("weaknesses", [])
    if weaknesses:
        pills = "".join(
            f'<span class="weakness-pill">{html_mod.escape(w)}</span>'
            for w in weaknesses
        )
        st.markdown(f"**Areas to improve:** {pills}", unsafe_allow_html=True)

    follow_up = evaluation.get("follow_up_question", "")
    if follow_up:
        st.info(f"**ICE follow-up:** {follow_up}")


def _build_session_docx(session: dict, transcript: list[dict]) -> bytes:
    """Build a Word document with the full session transcript."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _run(para, text, size=12, bold=False, italic=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return r

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "HEARING PREPARATION TRANSCRIPT", size=14, bold=True)

    # Session info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_parts = []
    if session.get("client_name"):
        info_parts.append(f"Client: {session['client_name']}")
    info_parts.append(f"Case Type: {session.get('case_type', '')}")
    info_parts.append(f"Date: {session.get('created_at', '')[:10]}")
    info_parts.append(f"Mode: {session.get('mode', 'attorney').title()}")
    _run(p, " | ".join(info_parts), size=10)

    doc.add_paragraph()

    # Calculate average score
    scores = [
        t["evaluation"]["score"]
        for t in transcript
        if t.get("evaluation") and t["evaluation"].get("score")
    ]
    if scores:
        p = doc.add_paragraph()
        avg = sum(scores) / len(scores)
        _run(p, f"Average Score: {avg:.1f}/5 ({len(scores)} questions evaluated)", size=11, bold=True)
        doc.add_paragraph()

    # Turns
    for i, turn in enumerate(transcript, 1):
        # Question
        p = doc.add_paragraph()
        _run(p, f"Q{i}: ", size=11, bold=True)
        _run(p, turn.get("question_text", ""), size=11)

        # Answer
        p = doc.add_paragraph()
        _run(p, "Answer: ", size=11, bold=True, italic=True)
        _run(p, turn.get("transcript", ""), size=11)
        fmt = p.paragraph_format
        fmt.space_after = Pt(4)

        # Evaluation
        evaluation = turn.get("evaluation")
        if evaluation:
            p = doc.add_paragraph()
            score = evaluation.get("score", 0)
            _run(p, f"Score: {score}/5", size=10, bold=True)

            p = doc.add_paragraph()
            _run(p, evaluation.get("evaluation_text", ""), size=10, italic=True)
            fmt = p.paragraph_format
            fmt.space_after = Pt(2)

            strengths = evaluation.get("strengths", [])
            if strengths:
                p = doc.add_paragraph()
                _run(p, "Strengths: ", size=10, bold=True)
                _run(p, "; ".join(strengths), size=10)

            weaknesses = evaluation.get("weaknesses", [])
            if weaknesses:
                p = doc.add_paragraph()
                _run(p, "Areas to improve: ", size=10, bold=True)
                _run(p, "; ".join(weaknesses), size=10)

            follow_up = evaluation.get("follow_up_question", "")
            if follow_up:
                p = doc.add_paragraph()
                _run(p, "ICE follow-up: ", size=10, bold=True)
                _run(p, follow_up, size=10)

        # Separator
        p = doc.add_paragraph()
        _run(p, "—" * 50, size=8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Sidebar ─────────────────────────────────────────────────────────────────

with st.sidebar:
    # Session management
    st.markdown("#### Sessions")
    btn_cols = st.columns(2)
    with btn_cols[0]:
        if st.button("New Session", use_container_width=True):
            st.session_state.hp_session_id = None
            st.session_state.hp_phase = "setup"
            st.session_state.hp_turn_number = 0
            st.session_state.hp_current_question_idx = 0
            st.session_state.hp_use_follow_up = False
            st.session_state.hp_follow_up_text = ""
            st.session_state.hp_last_eval = None
            st.rerun()
    with btn_cols[1]:
        if st.session_state.hp_session_id and st.button(
            "End Session", use_container_width=True, type="primary"
        ):
            update_session(st.session_state.hp_session_id, status="completed")
            st.session_state.hp_phase = "review"
            st.rerun()

    # Session history
    saved_sessions = list_sessions(limit=20)
    if saved_sessions:
        st.divider()
        st.markdown("#### History")
        labels_map = {
            s["id"]: f"{s['client_name'] or 'Unnamed'} — {s['case_type']}"
            for s in saved_sessions
        }
        session_ids = list(labels_map.keys())
        selected_session = st.selectbox(
            "Load a session",
            options=[""] + session_ids,
            format_func=lambda x: labels_map.get(x, "Select..."),
            label_visibility="collapsed",
        )
        load_cols = st.columns(2)
        with load_cols[0]:
            if selected_session and st.button("Load", use_container_width=True):
                st.session_state.hp_session_id = selected_session
                session = get_session(selected_session)
                if session and session["status"] == "completed":
                    st.session_state.hp_phase = "review"
                else:
                    st.session_state.hp_phase = "active"
                    turns = get_turns(selected_session)
                    st.session_state.hp_turn_number = len(turns)
                    # Find where we are in the question bank
                    if session:
                        all_qs = get_all_questions(session["case_type"])
                        answered_qids = {t["question_id"] for t in turns}
                        st.session_state.hp_current_question_idx = 0
                        for idx, q in enumerate(all_qs):
                            if q["id"] not in answered_qids:
                                st.session_state.hp_current_question_idx = idx
                                break
                        else:
                            st.session_state.hp_current_question_idx = len(all_qs)
                    st.session_state.hp_use_follow_up = False
                    st.session_state.hp_follow_up_text = ""
                    st.session_state.hp_last_eval = None
                st.rerun()
        with load_cols[1]:
            if selected_session and st.button("Delete", use_container_width=True):
                delete_session(selected_session)
                if st.session_state.hp_session_id == selected_session:
                    st.session_state.hp_session_id = None
                    st.session_state.hp_phase = "setup"
                st.rerun()

    st.divider()

    # Settings
    st.markdown("#### Settings")
    st.session_state.hp_show_tips = st.checkbox(
        "Show coaching tips", value=st.session_state.hp_show_tips
    )
    st.session_state.hp_show_eval_to_client = st.checkbox(
        "Show evaluations to client", value=st.session_state.hp_show_eval_to_client
    )
    st.session_state.hp_auto_follow_up = st.checkbox(
        "Auto follow-up questions", value=st.session_state.hp_auto_follow_up
    )

    render_tool_notes("hearing-prep")


# ── Main area ───────────────────────────────────────────────────────────────

# ── SETUP PHASE ─────────────────────────────────────────────────────────────
if st.session_state.hp_phase == "setup":
    st.markdown("### Start a New Session")

    setup_col1, setup_col2 = st.columns(2)

    with setup_col1:
        mode = st.radio(
            "Mode",
            options=["Attorney-Guided", "Solo Practice"],
            help="Attorney-guided: attorney controls the flow. Solo: client practices independently.",
        )

        case_type = st.selectbox("Case Type", options=get_case_types(), key="hp_case_type")

        # Show question count
        all_qs = get_all_questions(case_type)
        banks = get_question_banks()
        sections = banks.get(case_type, [])
        st.caption(
            f"{len(all_qs)} questions across {len(sections)} sections"
        )

    with setup_col2:
        # Pre-fill client name from banner if available
        default_name = ""
        default_id = ""
        default_lang = "English"
        if client_record:
            default_name = client_record.get("Name", "")
            default_id = client_record.get("Id", "")
            best_lang = client_record.get("Best_Language__c", "")
            if best_lang and best_lang in SUPPORTED_LANGUAGES:
                default_lang = best_lang

        client_name = st.text_input("Client Name", value=default_name, key="hp_client_name")

        lang_options = list(SUPPORTED_LANGUAGES.keys())
        lang_idx = lang_options.index(default_lang) if default_lang in lang_options else 0
        language = st.selectbox(
            "Language", options=lang_options, index=lang_idx, key="hp_language"
        )

        notes = st.text_area(
            "Session Notes (optional)",
            placeholder="Any notes for this session...",
            height=80,
            key="hp_notes",
        )

    # Section overview
    if sections:
        st.markdown("---")
        st.markdown("**Question Sections**")
        for section in sections:
            n_qs = len(section["questions"])
            st.markdown(f"- **{section['title']}** — {n_qs} questions")

    st.markdown("---")
    if st.button("Start Session", type="primary", use_container_width=True):
        session_id = create_session(
            client_name=client_name,
            client_id=default_id,
            case_type=case_type,
            mode="attorney" if mode == "Attorney-Guided" else "solo",
            language=language,
            notes=notes,
        )
        st.session_state.hp_session_id = session_id
        st.session_state.hp_phase = "active"
        st.session_state.hp_turn_number = 0
        st.session_state.hp_current_question_idx = 0
        st.session_state.hp_use_follow_up = False
        st.session_state.hp_follow_up_text = ""
        st.session_state.hp_last_eval = None
        st.rerun()


# ── ACTIVE PHASE ────────────────────────────────────────────────────────────
elif st.session_state.hp_phase == "active":
    session = get_session(st.session_state.hp_session_id)
    if not session:
        st.error("Session not found.")
        st.stop()

    case_type = session["case_type"]
    language = session["language"]
    lang_code = SUPPORTED_LANGUAGES.get(language, "en-US")
    all_questions = get_all_questions(case_type)
    current_idx = st.session_state.hp_current_question_idx

    # Progress bar
    total_qs = len(all_questions)
    answered = st.session_state.hp_turn_number
    if total_qs > 0:
        pct = min(answered / total_qs, 1.0)
        st.markdown(
            f'<div class="progress-label">{answered} of {total_qs} questions answered</div>',
            unsafe_allow_html=True,
        )
        st.progress(pct)

    # Check if we're done with the question bank
    if current_idx >= total_qs and not st.session_state.hp_use_follow_up:
        st.success("All questions in this section have been answered!")
        if st.button("View Session Review", type="primary"):
            update_session(st.session_state.hp_session_id, status="completed")
            st.session_state.hp_phase = "review"
            st.rerun()
        st.stop()

    # Determine the current question
    if st.session_state.hp_use_follow_up and st.session_state.hp_follow_up_text:
        current_question_text = st.session_state.hp_follow_up_text
        current_question_id = f"follow_up_{st.session_state.hp_turn_number}"
        current_tip = "This is a follow-up question generated by the ICE evaluator based on your previous answer."
        current_section = "Follow-Up"
    else:
        if current_idx < total_qs:
            q = all_questions[current_idx]
            current_question_text = q["label"]
            current_question_id = q["id"]
            current_tip = q.get("tip", "")
            current_section = q.get("section", "")
        else:
            st.success("All questions answered!")
            if st.button("View Session Review", type="primary"):
                update_session(st.session_state.hp_session_id, status="completed")
                st.session_state.hp_phase = "review"
                st.rerun()
            st.stop()

    # Two-column layout
    q_col, eval_col = st.columns([3, 2], gap="large")

    with q_col:
        # Section label
        if current_section:
            st.caption(current_section)

        # Question
        st.markdown(f"**Q{st.session_state.hp_turn_number + 1}:** {current_question_text}")

        # Coaching tip
        if current_tip and st.session_state.hp_show_tips:
            with st.expander("Coaching tip", icon=":material/lightbulb:"):
                st.markdown(
                    f'<div class="coaching-tip">{html_mod.escape(current_tip)}</div>',
                    unsafe_allow_html=True,
                )

        # Audio recording
        st.markdown("---")
        st.markdown("**Record your answer:**")
        audio_value = st.audio_input(
            "Record answer",
            key=f"audio_{st.session_state.hp_turn_number}",
            label_visibility="collapsed",
        )

        # Transcription
        transcript_text = ""
        confidence = 0.0
        audio_bytes = None

        if audio_value is not None:
            audio_bytes = audio_value.getvalue()
            with st.spinner("Transcribing audio..."):
                try:
                    result = transcribe_audio(audio_bytes, lang_code)
                    transcript_text = result["transcript"]
                    confidence = result["confidence"]
                except Exception as e:
                    st.error(f"Transcription failed: {e}")
                    st.info("You can type your answer manually below.")

        # Editable transcript
        transcript_text = st.text_area(
            "Transcript (edit if needed)",
            value=transcript_text,
            height=150,
            key=f"transcript_{st.session_state.hp_turn_number}",
            placeholder="Record audio above or type your answer here...",
        )

        if confidence > 0:
            st.caption(f"Transcription confidence: {confidence:.0%}")

        # Submit button
        if st.button(
            "Submit Answer",
            type="primary",
            use_container_width=True,
            disabled=not transcript_text.strip(),
        ):
            # Save turn to database
            turn_id = add_turn(
                session_id=st.session_state.hp_session_id,
                turn_number=st.session_state.hp_turn_number,
                question_id=current_question_id,
                question_text=current_question_text,
                audio_blob=audio_bytes,
                transcript=transcript_text.strip(),
                language_code=lang_code,
                confidence=confidence,
            )

            # Get conversation history for context
            prior_transcript = get_session_transcript(st.session_state.hp_session_id)
            # Remove the turn we just added (it doesn't have an eval yet)
            prior_transcript = [t for t in prior_transcript if t["id"] != turn_id]
            conversation_history = build_conversation_history(prior_transcript)

            # Evaluate with Claude
            with st.spinner("ICE attorney is evaluating your answer..."):
                try:
                    eval_result = evaluate_answer(
                        system_prompt=ICE_EVALUATOR_SYSTEM_PROMPT,
                        conversation_history=conversation_history,
                        current_question=current_question_text,
                        answer_transcript=transcript_text.strip(),
                        case_type=case_type,
                    )

                    # Save evaluation
                    add_evaluation(
                        turn_id=turn_id,
                        session_id=st.session_state.hp_session_id,
                        evaluation_text=eval_result.get("evaluation", ""),
                        score=eval_result.get("score", 0),
                        strengths=eval_result.get("strengths", []),
                        weaknesses=eval_result.get("weaknesses", []),
                        follow_up_question=eval_result.get("follow_up_question", ""),
                    )

                    st.session_state.hp_last_eval = eval_result

                    # Advance to next question
                    st.session_state.hp_turn_number += 1

                    # Determine next question
                    follow_up = eval_result.get("follow_up_question", "")
                    if follow_up and st.session_state.hp_auto_follow_up:
                        st.session_state.hp_use_follow_up = True
                        st.session_state.hp_follow_up_text = follow_up
                    else:
                        st.session_state.hp_use_follow_up = False
                        st.session_state.hp_follow_up_text = ""
                        if not st.session_state.hp_use_follow_up:
                            st.session_state.hp_current_question_idx = current_idx + 1

                    st.rerun()

                except Exception as e:
                    st.error(f"Evaluation failed: {e}")
                    # Still advance even if evaluation fails
                    st.session_state.hp_turn_number += 1
                    st.session_state.hp_use_follow_up = False
                    st.session_state.hp_current_question_idx = current_idx + 1
                    st.rerun()

    with eval_col:
        st.markdown("**Last Evaluation**")

        if st.session_state.hp_last_eval:
            _render_evaluation(st.session_state.hp_last_eval)
        else:
            # Show scores from previous turns
            turns = get_turns(st.session_state.hp_session_id)
            if turns:
                transcript_data = get_session_transcript(st.session_state.hp_session_id)
                scores = [
                    t["evaluation"]["score"]
                    for t in transcript_data
                    if t.get("evaluation") and t["evaluation"].get("score")
                ]
                if scores:
                    avg = sum(scores) / len(scores)
                    st.metric("Average Score", f"{avg:.1f}/5")
                    st.caption(f"{len(scores)} questions evaluated")

                # Show the most recent evaluation
                for t in reversed(transcript_data):
                    if t.get("evaluation"):
                        eval_data = t["evaluation"]
                        eval_display = {
                            "evaluation": eval_data.get("evaluation_text", ""),
                            "score": eval_data.get("score", 0),
                            "strengths": eval_data.get("strengths", []),
                            "weaknesses": eval_data.get("weaknesses", []),
                            "follow_up_question": eval_data.get("follow_up_question", ""),
                        }
                        _render_evaluation(eval_display)
                        break
            else:
                st.caption("Submit your first answer to see the evaluation here.")

        # Session info in sidebar area
        st.markdown("---")
        st.caption(f"**Session:** {session.get('client_name', 'Unnamed')}")
        st.caption(f"**Case type:** {case_type}")
        st.caption(f"**Language:** {language}")
        st.caption(f"**Mode:** {session.get('mode', 'attorney').title()}")


# ── REVIEW PHASE ────────────────────────────────────────────────────────────
elif st.session_state.hp_phase == "review":
    session = get_session(st.session_state.hp_session_id)
    if not session:
        st.error("Session not found.")
        st.stop()

    transcript = get_session_transcript(st.session_state.hp_session_id)

    st.markdown("### Session Review")

    # Session info bar
    info_parts = []
    if session.get("client_name"):
        info_parts.append(f"**Client:** {session['client_name']}")
    info_parts.append(f"**Case Type:** {session.get('case_type', '')}")
    info_parts.append(f"**Date:** {session.get('created_at', '')[:10]}")
    info_parts.append(f"**Mode:** {session.get('mode', 'attorney').title()}")
    info_parts.append(f"**Questions:** {len(transcript)}")
    st.markdown(" | ".join(info_parts))

    # Average score
    scores = [
        t["evaluation"]["score"]
        for t in transcript
        if t.get("evaluation") and t["evaluation"].get("score")
    ]
    if scores:
        avg = sum(scores) / len(scores)
        score_cols = st.columns(3)
        with score_cols[0]:
            st.metric("Average Score", f"{avg:.1f}/5")
        with score_cols[1]:
            st.metric("Questions Answered", len(transcript))
        with score_cols[2]:
            high = sum(1 for s in scores if s >= 4)
            st.metric("Strong Answers (4+)", f"{high}/{len(scores)}")

    st.markdown("---")

    # Full transcript
    for i, turn in enumerate(transcript, 1):
        question_text = turn.get("question_text", "")
        answer_text = turn.get("transcript", "")
        evaluation = turn.get("evaluation")

        st.markdown(
            f'<div class="transcript-turn">'
            f'<div class="transcript-q">Q{i}: {html_mod.escape(question_text)}</div>'
            f'<div class="transcript-a">{html_mod.escape(answer_text)}</div>'
            f"</div>",
            unsafe_allow_html=True,
        )

        if evaluation and st.session_state.hp_show_eval_to_client:
            eval_display = {
                "evaluation": evaluation.get("evaluation_text", ""),
                "score": evaluation.get("score", 0),
                "strengths": evaluation.get("strengths", []),
                "weaknesses": evaluation.get("weaknesses", []),
                "follow_up_question": evaluation.get("follow_up_question", ""),
            }
            _render_evaluation(eval_display)

    # Export
    st.markdown("---")
    st.markdown("**Export**")

    if transcript:
        # Build export data
        export_name = session.get("client_name", "Session").replace(" ", "_")

        # Plain text export
        plain_lines = [
            "HEARING PREPARATION TRANSCRIPT",
            f"Client: {session.get('client_name', '')}",
            f"Case Type: {session.get('case_type', '')}",
            f"Date: {session.get('created_at', '')[:10]}",
            f"Mode: {session.get('mode', 'attorney').title()}",
            "",
        ]
        if scores:
            plain_lines.append(f"Average Score: {avg:.1f}/5")
            plain_lines.append("")

        for i, turn in enumerate(transcript, 1):
            plain_lines.append(f"Q{i}: {turn.get('question_text', '')}")
            plain_lines.append(f"Answer: {turn.get('transcript', '')}")
            evaluation = turn.get("evaluation")
            if evaluation:
                plain_lines.append(f"Score: {evaluation.get('score', 0)}/5")
                plain_lines.append(f"Evaluation: {evaluation.get('evaluation_text', '')}")
                strengths = evaluation.get("strengths", [])
                if strengths:
                    plain_lines.append(f"Strengths: {'; '.join(strengths)}")
                weaknesses = evaluation.get("weaknesses", [])
                if weaknesses:
                    plain_lines.append(f"Areas to improve: {'; '.join(weaknesses)}")
            plain_lines.append("")
            plain_lines.append("—" * 50)
            plain_lines.append("")

        plain_text = "\n".join(plain_lines)

        exp_cols = st.columns(3)
        with exp_cols[0]:
            st.download_button(
                "Download .txt",
                data=plain_text,
                file_name=f"HearingPrep_{export_name}.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with exp_cols[1]:
            docx_bytes = _build_session_docx(session, transcript)
            st.download_button(
                "Download .docx",
                data=docx_bytes,
                file_name=f"HearingPrep_{export_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with exp_cols[2]:
            if st.button("Upload to Google Docs", use_container_width=True):
                with st.spinner("Uploading to Google Docs..."):
                    try:
                        docx_bytes = _build_session_docx(session, transcript)
                        url = upload_to_google_docs(
                            docx_bytes,
                            f"Hearing Prep - {session.get('client_name', 'Session')}",
                        )
                        st.session_state.hp_google_doc_url = url
                    except Exception as e:
                        st.error(f"Upload failed: {e}")
            if st.session_state.get("hp_google_doc_url"):
                st.markdown(f"[Open Google Doc]({st.session_state.hp_google_doc_url})")
    else:
        st.caption("No questions were answered in this session.")

    # Back to active session
    if session.get("status") != "completed":
        if st.button("Continue Session"):
            st.session_state.hp_phase = "active"
            st.rerun()
